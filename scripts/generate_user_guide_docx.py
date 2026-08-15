"""Generate the DX-OS end-user and testing guide as a DOCX document.

Run from the repository root:
    python scripts/generate_user_guide_docx.py
"""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
OUTPUT = ROOT / "docs" / "generated" / "Huong_dan_su_dung_DX_OS_kiem_thu_luong.docx"

BLUE = "0F6CBD"
DARK_BLUE = "083B66"
LIGHT_BLUE = "EAF3FF"
LIGHT_GRAY = "F3F4F6"
LIGHT_YELLOW = "FFF7D6"
LIGHT_GREEN = "E8F5E9"
WHITE = "FFFFFF"
TEXT = "1F2937"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, *, bold: bool = False, color: str = TEXT) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    header = table.rows[0]
    set_repeat_table_header(header)
    for index, label in enumerate(headers):
        set_cell_shading(header.cells[index], BLUE)
        set_cell_text(header.cells[index], label, bold=True, color=WHITE)
        if widths:
            header.cells[index].width = Cm(widths[index])

    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            set_cell_text(cells[index], value)
            if widths:
                cells[index].width = Cm(widths[index])
        if len(table.rows) % 2 == 1:
            for cell in cells:
                set_cell_shading(cell, "F8FAFC")
    document.add_paragraph()
    return table


def add_bullet(document: Document, text: str, level: int = 0) -> None:
    style = "List Bullet" if level == 0 else "List Bullet 2"
    paragraph = document.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.add_run(text)


def add_number(document: Document, text: str, level: int = 0) -> None:
    style = "List Number" if level == 0 else "List Number 2"
    paragraph = document.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.add_run(text)


def add_code(document: Document, code: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.left_indent = Cm(0.35)
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "111827")
    p_pr.append(shd)
    run = paragraph.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(242, 242, 242)


def add_callout(document: Document, title: str, text: str, fill: str = LIGHT_BLUE) -> None:
    table = document.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.text = ""
    title_paragraph = cell.paragraphs[0]
    title_paragraph.paragraph_format.space_after = Pt(2)
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    body = cell.add_paragraph(text)
    body.paragraph_format.space_after = Pt(2)
    document.add_paragraph()


def add_heading(document: Document, text: str, level: int) -> None:
    heading = document.add_heading(text, level=level)
    heading.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    heading.paragraph_format.space_after = Pt(6)
    return heading


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, color in (("Title", 28, DARK_BLUE), ("Heading 1", 18, DARK_BLUE), ("Heading 2", 13, BLUE), ("Heading 3", 11, DARK_BLUE)):
        style = styles[name]
        style.font.name = "Aptos Display" if name != "Normal" else "Aptos"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)

    if "Code Block" not in styles:
        code_style = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = "Consolas"
        code_style.font.size = Pt(8.5)

    header = section.header.paragraphs[0]
    header.text = "DX-OS LAB  |  HƯỚNG DẪN SỬ DỤNG VÀ KIỂM THỬ"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.name = "Aptos"
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor.from_string("64748B")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("DX-OS Lab  •  Tài liệu nội bộ  •  Trang ")
    add_page_number(footer)
    for run in footer.runs:
        run.font.name = "Aptos"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string("64748B")


def add_cover(document: Document) -> None:
    document.add_paragraph()
    document.add_paragraph()
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("DX-OS LAB")
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(31)
    run.font.color.rgb = RGBColor.from_string(BLUE)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("HƯỚNG DẪN SỬ DỤNG VÀ KIỂM THỬ LUỒNG NGHIỆP VỤ")
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    document.add_paragraph()
    intro = document.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro.add_run(
        "Dành cho nhóm 4 người chuẩn bị kiểm thử, luyện demo và báo cáo dự án."
    ).italic = True

    document.add_paragraph()
    add_callout(
        document,
        "Phạm vi tài liệu",
        "Hướng dẫn này bám theo phiên bản đã cập nhật: dashboard theo vai trò, lối tắt nghiệp vụ và xuất CSV báo cáo/ngân sách. Tài liệu không mô tả RAG hoặc AI Agent vì hai mô-đun này chưa được triển khai.",
        LIGHT_BLUE,
    )
    document.add_paragraph()
    info = document.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f"Phiên bản tài liệu: 1.0  |  Ngày tạo: {date.today().strftime('%d/%m/%Y')}")
    document.add_page_break()


def add_contents(document: Document) -> None:
    add_heading(document, "Mục lục nhanh", 1)
    entries = [
        "1. Mục tiêu và phạm vi kiểm thử",
        "2. Kiến trúc, URL và điều kiện bắt đầu",
        "3. Tài khoản, role và quyền",
        "4. Những chức năng mới sau cập nhật",
        "5. Luồng mua sắm đến thanh toán Employee → Manager → Finance → Auditor",
        "6. Hướng dẫn thao tác theo từng role",
        "7. Kiểm thử dashboard theo role và xuất CSV",
        "8. Các tình huống âm (negative test)",
        "9. Checklist kiểm thử và checklist demo 10–15 phút",
        "10. Xử lý sự cố thường gặp",
    ]
    for entry in entries:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.add_run(entry)
    add_callout(
        document,
        "Cách dùng tài liệu",
        "Nếu chỉ cần luyện trình bày, làm lần lượt mục 2 → 5 → 7 → 9. Nếu cần kiểm thử đầy đủ, hoàn thành toàn bộ checklist ở mục 9 và lưu ảnh chụp màn hình làm bằng chứng.",
        LIGHT_GREEN,
    )
    document.add_page_break()


def add_system_setup(document: Document) -> None:
    add_heading(document, "1. Mục tiêu và phạm vi kiểm thử", 1)
    document.add_paragraph(
        "Mục tiêu là chứng minh một yêu cầu mua sắm đi từ lúc nhân viên tạo phiếu đến đặt hàng, nhận hàng, "
        "đối soát hóa đơn và thanh toán; ngân sách thay đổi đúng; auditor chỉ đọc; admin quản lý policy có audit."
    )
    add_bullet(document, "Không dùng tài khoản thật hoặc tài liệu có dữ liệu nhạy cảm trong môi trường lab.")
    add_bullet(document, "Không chia sẻ mật khẩu trong data/runtime; các file này được Git-ignore.")
    add_bullet(document, "Chỉ mở frontend bằng localhost:4200. Không dùng 127.0.0.1 vì redirect URI Keycloak sẽ lỗi.")

    add_heading(document, "2. Kiến trúc, URL và điều kiện bắt đầu", 1)
    document.add_paragraph("Luồng kỹ thuật chính: Angular Web → Keycloak đăng nhập → Go API → PostgreSQL/Nextcloud; Metabase đọc dữ liệu báo cáo.")
    add_table(
        document,
        ["Thành phần", "Địa chỉ", "Dùng khi nào"],
        [
            ["DX-OS Web", "http://localhost:4200", "Người dùng nghiệp vụ đăng nhập và thao tác."],
            ["Keycloak", "http://localhost:8080", "Trang đăng nhập và quản lý identity local."],
            ["Go API", "http://localhost:8081", "Kiểm tra /health/live và /health/ready; không phải màn hình nghiệp vụ."],
            ["Metabase", "http://localhost:3000", "Phân tích BI và đối chiếu dashboard."],
            ["Nextcloud", "http://localhost:8082", "Kho lưu file nội bộ, đi qua Go API trong nghiệp vụ."],
            ["DX-OS Docs", "http://localhost:4300", "Đọc tài liệu web."],
        ],
        [3.0, 4.0, 9.0],
    )
    add_heading(document, "2.1 Kiểm tra stack đang chạy", 2)
    document.add_paragraph("Từ thư mục gốc DX-OS, chạy các lệnh sau. Chỉ cần build lại khi có thay đổi code.")
    add_code(document, "docker compose --profile foundation --profile application --profile reporting ps\n"
                       "docker compose --profile foundation --profile application --profile reporting up -d --build\n"
                       "docker compose --profile documentation up -d --build docs")
    document.add_paragraph("Kết quả mong đợi: postgres, api, metabase và docs hiển thị healthy; web và keycloak ở trạng thái Up. Sau đó mở /health/ready để xác nhận API trả HTTP 200.")
    add_callout(document, "Lưu ý về cổng", "Cổng 8080 là Keycloak. API là 8081. Nếu gõ localhost:8080 để tìm API, bạn sẽ nhìn thấy trang Keycloak chứ không phải endpoint API.", LIGHT_YELLOW)


def add_accounts_and_roles(document: Document) -> None:
    add_heading(document, "3. Tài khoản, role và quyền", 1)
    document.add_paragraph("Mật khẩu không được ghi trong tài liệu này. Đọc từ file tương ứng trong data/runtime hoặc tạo lại bằng script. Mỗi lần chạy Initialize-DevUser.ps1, mật khẩu của user đó sẽ đổi.")
    add_table(
        document,
        ["Tài khoản", "Role", "File credential local", "Mục tiêu test"],
        [
            ["employee.demo", "employee", "data/runtime/employee-demo.txt", "Tạo, sửa, đính kèm và gửi phiếu."],
            ["manager.demo", "department_manager", "data/runtime/manager-demo.txt", "Duyệt cấp phòng ban hoặc yêu cầu chỉnh sửa."],
            ["finance.demo", "finance", "data/runtime/finance-demo.txt", "Duyệt, đặt hàng, hóa đơn, ngân sách, báo cáo."],
            ["auditor.demo", "auditor", "data/runtime/auditor-demo.txt", "Kiểm tra chỉ đọc toàn bộ bằng chứng vận hành."],
            ["admin.demo", "dx_admin", "data/runtime/admin-demo.txt", "Quản lý SLA/ngưỡng chứng từ; xem báo cáo."],
            ["ai.operator.demo", "ai_operator", "data/runtime/ai-operator-demo.txt", "Xem dashboard nền tảng; AI/RAG chưa có thao tác."],
        ],
        [3.0, 3.4, 5.3, 4.3],
    )
    add_heading(document, "3.1 Tạo lại tài khoản demo khi cần", 2)
    add_code(document, ".\\scripts\\Initialize-DevUser.ps1 -Username employee.demo -Role employee -CredentialsPath data\\runtime\\employee-demo.txt\n"
                       ".\\scripts\\Initialize-DevUser.ps1 -Username manager.demo -Role department_manager -CredentialsPath data\\runtime\\manager-demo.txt\n"
                       ".\\scripts\\Initialize-DevUser.ps1 -Username finance.demo -Role finance -CredentialsPath data\\runtime\\finance-demo.txt\n"
                       ".\\scripts\\Initialize-DevUser.ps1 -Username auditor.demo -Role auditor -CredentialsPath data\\runtime\\auditor-demo.txt")
    add_callout(document, "Sau khi đổi role", "Đăng xuất rồi đăng nhập lại để Keycloak cấp access token mới có role mới. Nếu không, menu có thể vẫn hiển thị theo quyền cũ.", LIGHT_YELLOW)

    add_heading(document, "4. Những chức năng mới sau cập nhật", 1)
    add_table(
        document,
        ["Role", "Thông tin mới trên Dashboard", "Lối tắt / thao tác mới"],
        [
            ["Employee", "Số phiếu trong phạm vi, số phiếu cần bổ sung, 5 phiếu cập nhật gần đây.", "Tạo phiếu mới; mở danh sách phiếu."],
            ["Manager", "Số phiếu chờ duyệt cấp bộ phận, phiếu gần đây theo phạm vi.", "Mở hàng đợi phê duyệt; tạo/mở danh sách phiếu."],
            ["Finance", "Chờ duyệt, giao hàng, công nợ, ngân sách, KPI/SLA.", "Phê duyệt, Giao nhận, Hóa đơn, Ngân sách, Báo cáo."],
            ["Auditor", "Dữ liệu và bằng chứng chỉ đọc.", "Mở Giao nhận/Hóa đơn/Chính sách/Audit; không ghi."],
            ["DX Admin", "KPI và policy toàn tổ chức.", "Sửa SLA/ngưỡng chứng từ có version và audit."],
            ["AI Operator", "Trạng thái nền tảng và mô tả phạm vi AI.", "Không có Agent/RAG giả lập."],
        ],
        [2.8, 8.2, 5.0],
    )
    add_bullet(document, "Dashboard gọi API theo đúng role. Card hiển thị “—” nghĩa là role không có quyền đọc chỉ số hoặc dịch vụ thành phần chưa sẵn sàng.")
    add_bullet(document, "Báo cáo và Ngân sách có nút Xuất CSV. CSV dùng UTF-8 BOM để Excel hiển thị tiếng Việt đúng.")


def add_workflow(document: Document) -> None:
    document.add_page_break()
    add_heading(document, "5. Luồng mua sắm đến thanh toán: Employee → Manager → Finance → Auditor", 1)
    document.add_paragraph("Đây là luồng nên dùng khi demo. Chuẩn bị bốn cửa sổ trình duyệt ẩn danh hoặc đăng xuất/đăng nhập giữa mỗi role.")
    add_code(document, "DRAFT  →  SUBMITTED  →  MANAGER_APPROVED  →  APPROVED\n"
                       "  ↑          ↓                 ↓\n"
                       "  └── CHANGES_REQUESTED     REJECTED\n"
                       "DRAFT hoặc CHANGES_REQUESTED có thể → CANCELLED")

    add_heading(document, "5.1 Bước A — Employee tạo và gửi phiếu", 2)
    add_number(document, "Đăng nhập employee.demo tại http://localhost:4200.")
    add_number(document, "Ở Dashboard, kiểm tra tên role Employee; bấm Tạo phiếu mới.")
    add_number(document, "Nhập tiêu đề dễ nhận biết, ví dụ: “Mua laptop demo báo cáo”. Chọn cost center hợp lệ, currency VND và lý do mua.")
    add_number(document, "Thêm ít nhất một dòng hàng: mô tả, số lượng, đơn vị và đơn giá. Ghi lại mã phiếu sau khi lưu.")
    add_number(document, "Nếu tổng giá trị từ 20.000.000 VND trở lên, tải một file Báo giá hợp lệ trước khi gửi. Dưới ngưỡng này không cần báo giá.")
    add_number(document, "Mở chi tiết phiếu, kiểm tra ngân sách/tài liệu, rồi bấm Gửi duyệt.")
    add_callout(document, "Kỳ vọng", "Trạng thái chuyển DRAFT → SUBMITTED. Dashboard Employee giảm hoặc cập nhật số phiếu cần bổ sung khi dữ liệu tải lại. Phiếu xuất hiện trong hàng đợi Manager cùng department.", LIGHT_GREEN)

    add_heading(document, "5.2 Bước B — Department Manager duyệt cấp bộ phận", 2)
    add_number(document, "Đăng xuất Employee. Đăng nhập manager.demo.")
    add_number(document, "Tại Dashboard, kiểm tra card Chờ trưởng bộ phận và bấm Phê duyệt.")
    add_number(document, "Mở đúng phiếu vừa tạo; đối chiếu tiêu đề, cost center, số tiền, lý do và báo giá nếu có.")
    add_number(document, "Chọn Phê duyệt. Không được duyệt phiếu do chính manager tạo.")
    add_callout(document, "Kỳ vọng", "Trạng thái chuyển SUBMITTED → MANAGER_APPROVED. Hệ thống tạo reservation: reserved tăng, available giảm. Phiếu rời hàng đợi Manager và xuất hiện ở Finance.", LIGHT_GREEN)

    add_heading(document, "5.3 Bước C — Finance duyệt cuối", 2)
    add_number(document, "Đăng nhập finance.demo.")
    add_number(document, "Dashboard phải hiển thị số Chờ duyệt tài chính. Bấm Phê duyệt.")
    add_number(document, "Mở phiếu, kiểm tra budget check, tài liệu, lịch sử và comment của Manager.")
    add_number(document, "Chọn Phê duyệt để kết thúc quy trình.")
    add_callout(document, "Kỳ vọng", "Trạng thái chuyển MANAGER_APPROVED → APPROVED. Reservation giảm; committed tăng. KPI đã duyệt và dữ liệu báo cáo cập nhật theo khoảng thời gian lọc.", LIGHT_GREEN)

    add_heading(document, "5.4 Bước D — Đặt hàng, nhận hàng và thanh toán", 2)
    add_number(document, "Finance mở Giao nhận, chọn nhà cung cấp hoạt động và phát hành đơn đặt hàng.")
    add_number(document, "Finance mở Hóa đơn, ghi số hóa đơn/ngày/hạn trả/amount/currency. Trước khi nhận hàng, nhãn phải là Chờ nhận hàng và lệnh Xác minh bị chặn.")
    add_number(document, "Employee hoặc Manager cùng phòng mở Giao nhận và xác nhận đã nhận. Finance không được tự xác nhận nhận hàng.")
    add_number(document, "Finance tải lại Hóa đơn. Khi amount/currency khớp, chọn Xác minh, sau đó nhập mã tham chiếu và ngày thanh toán.")
    add_number(document, "Employee mở Thông báo và kiểm tra thông báo hóa đơn đã thanh toán.")
    add_callout(document, "Kỳ vọng", "Luồng order → receipt → invoice verified → paid hoàn tất; mọi mutation có audit và optimistic locking. Hóa đơn lệch tiền/tệ không thể xác minh.", LIGHT_GREEN)

    add_heading(document, "5.5 Bước E — Auditor đối chiếu", 2)
    add_number(document, "Đăng nhập auditor.demo.")
    add_number(document, "Mở danh sách/chi tiết phiếu; kiểm tra trạng thái APPROVED và Timeline có các mốc tạo, gửi, manager duyệt, finance duyệt.")
    add_number(document, "Mở Ngân sách; kiểm tra allocation, reservation/commitment và lịch sử điều chỉnh. Auditor không có nút Điều chỉnh.")
    add_number(document, "Mở Báo cáo, đặt khoảng ngày bao phủ thời điểm vừa tạo phiếu; đối chiếu tổng phiếu/KPI/SLA.")
    add_number(document, "Bấm Xuất CSV ở Báo cáo và Ngân sách; mở file bằng Excel để kiểm tra tiếng Việt và dữ liệu số.")
    add_callout(document, "Kỳ vọng", "Auditor đọc được dữ liệu trong phạm vi nhưng không thể ghi. Nếu cố truy cập thao tác ghi qua UI/API, hệ thống phải trả 403 hoặc không hiển thị nút.", LIGHT_GREEN)


def add_role_guides(document: Document) -> None:
    add_heading(document, "6. Hướng dẫn nhanh theo từng role", 1)
    role_steps = [
        ["Employee", "Dashboard → Tạo phiếu mới → lưu DRAFT → đính kèm (nếu cần) → gửi duyệt.", "Chỉ thấy phiếu thuộc phạm vi của mình; có thể sửa DRAFT/CHANGES_REQUESTED."],
        ["Department Manager", "Dashboard → Phê duyệt → mở phiếu SUBMITTED → Approve / Request changes / Reject.", "Chỉ xử lý phiếu cùng department; không tự duyệt phiếu của mình."],
        ["Finance", "Phê duyệt → Giao nhận → Hóa đơn → Ngân sách/Báo cáo.", "Chỉ xác minh invoice khi đã nhận và khớp; lưu payment reference."],
        ["Auditor", "Mở phiếu, Giao nhận, Hóa đơn, Chính sách, Audit, Báo cáo.", "Chỉ đọc; mọi lệnh ghi phải bị 403."],
        ["DX Admin", "Chính sách → sửa SLA/ngưỡng chứng từ → yêu cầu kiểm tra audit.", "Không tự động có quyền duyệt, ngân sách hay hóa đơn."],
        ["AI Operator", "Dashboard → kiểm tra trạng thái nền tảng và phạm vi chuẩn bị AI.", "Không có menu chat, recommendation, tool execution hay Agent trong bản này."],
    ]
    add_table(document, ["Role", "Thao tác chính", "Giới hạn cần nhớ"], role_steps, [3.1, 8.0, 4.9])

    add_heading(document, "6.1 Khi phiếu bị yêu cầu chỉnh sửa", 2)
    add_number(document, "Manager hoặc Finance chọn Yêu cầu chỉnh sửa và nhập comment rõ ràng.")
    add_number(document, "Employee đăng nhập, dashboard hiển thị số Cần tôi bổ sung; mở phiếu có trạng thái CHANGES_REQUESTED.")
    add_number(document, "Employee sửa nội dung/tài liệu rồi gửi lại. Trạng thái quay về SUBMITTED và Manager xử lý lại.")
    add_callout(document, "Phân biệt", "Yêu cầu chỉnh sửa là quay lại người tạo để hoàn thiện. Từ chối là kết thúc phiếu ở REJECTED. Hủy là do người tạo thực hiện khi phiếu ở DRAFT hoặc CHANGES_REQUESTED.", LIGHT_YELLOW)


def add_dashboard_csv_tests(document: Document) -> None:
    add_heading(document, "7. Kiểm thử Dashboard theo role và Xuất CSV", 1)
    add_heading(document, "7.1 Dashboard", 2)
    add_table(
        document,
        ["Đăng nhập bằng", "Cần thấy", "Không được thấy / không được gọi"],
        [
            ["employee.demo", "Tạo phiếu mới, danh sách phiếu, số cần bổ sung.", "Hàng đợi duyệt, ngân sách, báo cáo."],
            ["manager.demo", "Chờ trưởng bộ phận, Phê duyệt, danh sách phiếu.", "Ngân sách, báo cáo."],
            ["finance.demo", "Chờ tài chính, cảnh báo ngân sách, KPI 30 ngày, Phê duyệt/Ngân sách/Báo cáo.", "Thao tác ngoài scope organization."],
            ["auditor.demo", "Ngân sách/Báo cáo, số liệu chỉ đọc, phiếu gần đây.", "Nút điều chỉnh ngân sách hoặc quyết định duyệt."],
            ["admin.demo", "KPI/SLA và link Báo cáo.", "Danh sách phiếu, ngân sách và thao tác workflow."],
            ["ai.operator.demo", "Trạng thái nền tảng và giới hạn AI.", "AI Agent hoặc quyền nghiệp vụ giả lập."],
        ],
        [3.2, 7.3, 5.5],
    )
    add_bullet(document, "Nếu có cảnh báo màu vàng ở dashboard, một API thành phần chưa tải được; kiểm tra API health và log, nhưng các lối tắt còn lại vẫn có thể dùng.")
    add_bullet(document, "Nếu card là “—”, trước hết kiểm tra role; đây có thể là hành vi đúng do role không có quyền đọc chỉ số đó.")

    add_heading(document, "7.2 Xuất CSV", 2)
    add_number(document, "Đăng nhập Finance, Auditor hoặc DX Admin rồi vào Báo cáo.")
    add_number(document, "Chọn từ ngày, đến ngày; có thể lọc cost center và currency. Bấm Áp dụng.")
    add_number(document, "Bấm Xuất CSV. File chứa KPI, trạng thái, phòng ban và sử dụng ngân sách theo dữ liệu API hiện có.")
    add_number(document, "Ở trang Ngân sách, bấm Xuất CSV để nhận allocation, reservation và lịch sử điều chỉnh.")
    add_callout(document, "Kỳ vọng CSV", "Mở bằng Excel không lỗi dấu tiếng Việt. Số tiền là giá trị thô để tính tiếp, không phải chuỗi trang trí. Phạm vi dữ liệu xuất đúng bằng quyền của role đang đăng nhập.", LIGHT_GREEN)


def add_negative_tests(document: Document) -> None:
    document.add_page_break()
    add_heading(document, "8. Các tình huống âm (negative test)", 1)
    document.add_paragraph("Các test này rất hữu ích khi báo cáo vì chứng minh hệ thống có kiểm soát quyền, ngân sách và dữ liệu.")
    add_table(
        document,
        ["Tình huống", "Cách thử", "Kết quả đúng"],
        [
            ["Sai redirect URI", "Mở 127.0.0.1:4200 thay vì localhost:4200.", "Keycloak báo redirect_uri không hợp lệ; quay về localhost."],
            ["Employee duyệt", "Dùng employee.demo tìm đường dẫn /approvals hoặc gọi API duyệt.", "Không có menu hoặc API trả 403."],
            ["Auditor điều chỉnh budget", "Mở Ngân sách bằng auditor.demo.", "Không có nút Điều chỉnh; lệnh ghi phải bị từ chối."],
            ["Thiếu báo giá", "Tạo phiếu từ 20 triệu VND, không tải QUOTATION, gửi duyệt.", "Submit/resubmit bị chặn và có thông báo yêu cầu tài liệu."],
            ["Không đủ ngân sách", "Tạo giá trị vượt allocation rồi cho Manager duyệt.", "Không tạo reservation; hệ thống báo thiếu ngân sách."],
            ["Version cũ", "Mở cùng một allocation ở hai phiên Finance, lưu phiên thứ nhất rồi lưu phiên còn lại.", "Lần lưu cũ bị conflict; tải lại dashboard rồi thử lại."],
            ["Truy cập ngoài scope", "Dùng user employee khác mở mã phiếu không thuộc mình.", "API/UI không trả dữ liệu ngoài phạm vi."],
        ],
        [3.5, 7.3, 5.2],
    )


def add_checklists(document: Document) -> None:
    add_heading(document, "9. Checklist kiểm thử và demo", 1)
    add_heading(document, "9.1 Checklist trước khi test", 2)
    for item in [
        "Docker Desktop đang chạy; docker compose ps không có service quan trọng ở trạng thái lỗi.",
        "http://localhost:4200 và http://localhost:8081/health/ready trả HTTP 200.",
        "Có credential cho employee.demo, manager.demo, finance.demo và auditor.demo trong data/runtime.",
        "Mở bốn cửa sổ ẩn danh hoặc sẵn sàng logout/login giữa các role.",
        "Có một file báo giá PDF/tệp hợp lệ nếu test phiếu từ 20 triệu VND.",
        "Chuẩn bị tên phiếu có chữ “demo” và ghi lại mã phiếu để dễ tìm trong báo cáo.",
    ]:
        document.add_paragraph(f"☐  {item}")

    add_heading(document, "9.2 Checklist bằng chứng cần lưu", 2)
    for item in [
        "Ảnh Dashboard của Finance có KPI, cảnh báo ngân sách và lối tắt nghiệp vụ.",
        "Ảnh chi tiết phiếu ở các trạng thái SUBMITTED, MANAGER_APPROVED và APPROVED.",
        "Ảnh Timeline cho thấy actor và thời điểm từng bước.",
        "Ảnh Ngân sách trước/sau khi duyệt cuối hoặc reservation/commitment.",
        "Ảnh Auditor ở chế độ chỉ đọc (không có nút Điều chỉnh).",
        "Hai file CSV xuất từ Báo cáo và Ngân sách, mở được bằng Excel.",
        "Kết quả các smoke test: Foundation, Application, OIDC, Workflow, Budget, Attachments, Reporting, Documentation.",
    ]:
        document.add_paragraph(f"☐  {item}")

    add_heading(document, "9.3 Kịch bản demo 10–15 phút cho 4 người", 2)
    add_table(
        document,
        ["Người", "Thời lượng", "Nội dung nói và thao tác"],
        [
            ["Người 1", "2–3 phút", "Mở kiến trúc/URL, đăng nhập Employee, tạo phiếu và gửi duyệt."],
            ["Người 2", "2–3 phút", "Đăng nhập Manager, duyệt cấp bộ phận; giải thích reservation ngân sách."],
            ["Người 3", "3–4 phút", "Đăng nhập Finance, duyệt cuối, mở dashboard ngân sách và xuất CSV."],
            ["Người 4", "3–4 phút", "Đăng nhập Auditor, đối chiếu Timeline/Báo cáo/CSV; nêu kiểm soát quyền và test âm."],
        ],
        [2.5, 2.5, 11.0],
    )
    add_callout(document, "Mẹo trình bày", "Không cố demo toàn bộ menu. Kể một câu chuyện xuyên suốt bằng cùng một mã phiếu và cho thấy vai trò sau chỉ nhìn/thao tác phần họ được phép. Đây là điểm mạnh nhất của DX-OS hiện tại.", LIGHT_BLUE)


def add_troubleshooting(document: Document) -> None:
    add_heading(document, "10. Xử lý sự cố thường gặp", 1)
    add_table(
        document,
        ["Hiện tượng", "Nguyên nhân thường gặp", "Cách xử lý nhanh"],
        [
            ["Không vào được web", "Container web chưa chạy hoặc port 4200 bị chiếm.", "docker compose ps; mở lại localhost:4200; xem docker compose logs web."],
            ["Đăng nhập báo redirect_uri", "Dùng IP/port khác localhost:4200.", "Mở đúng http://localhost:4200; không dùng 127.0.0.1."],
            ["Sai username/password", "Credential đã đổi khi script tạo user chạy lại.", "Mở file data/runtime tương ứng hoặc chạy Initialize-DevUser.ps1 để tạo lại."],
            ["API trả 401", "Chưa đăng nhập hoặc token hết hạn.", "Logout/login lại; kiểm tra http://localhost:8081/health/ready."],
            ["API trả 403", "Role hoặc data scope không cho phép.", "Đây thường là hành vi đúng; kiểm tra role và phạm vi department/organization."],
            ["Không gửi được phiếu", "Thiếu báo giá, thiếu dữ liệu, sai trạng thái hoặc ngân sách.", "Mở chi tiết phiếu để đọc message; kiểm tra tổng tiền, attachment và status."],
            ["Không thấy menu Báo cáo", "Không phải finance/auditor/dx_admin.", "Dùng đúng account và đăng nhập lại sau khi đổi role."],
            ["Metabase hiển thị 0", "Khoảng ngày/currency không bao phủ dữ liệu.", "Mở rộng khoảng ngày, bỏ filter hoặc tạo dữ liệu đúng trạng thái."],
            ["Dashboard có cảnh báo vàng", "Một API thành phần chưa tải được.", "Mở API health, kiểm tra docker compose logs api; thử tải lại trang."],
        ],
        [3.6, 5.6, 6.8],
    )
    add_heading(document, "10.1 Lệnh chẩn đoán an toàn", 2)
    add_code(document, "docker compose --profile foundation --profile application --profile reporting ps\n"
                       "docker compose --profile foundation --profile application --profile reporting logs --tail 120 api web\n"
                       "curl.exe http://localhost:8081/health/ready\n"
                       ".\\scripts\\Test-Application.ps1")
    add_callout(document, "Không làm", "Không sửa trực tiếp database để đổi trạng thái phiếu. Không gửi access token, .env, file credential hoặc mật khẩu vào nhóm chat/slide báo cáo.", LIGHT_YELLOW)

    add_heading(document, "Tài liệu nguồn trong repository", 1)
    document.add_paragraph("Khi cần chi tiết kỹ thuật hơn, tham khảo các tài liệu gốc sau:")
    for item in [
        "README.md — cài đặt, vận hành và URL.",
        "docs/GETTING_STARTED.md — dựng môi trường local.",
        "docs/USER_GUIDE.md — quyền role, dashboard và CSV.",
        "docs/runbooks/PROCUREMENT_MVP.md — state machine, authorization và budget rule.",
        "scripts/Test-*.ps1 — smoke test có thể chạy lại để tạo bằng chứng.",
    ]:
        add_bullet(document, item)


def build_document() -> Path:
    document = Document()
    configure_document(document)
    add_cover(document)
    add_contents(document)
    add_system_setup(document)
    add_accounts_and_roles(document)
    add_workflow(document)
    add_role_guides(document)
    add_dashboard_csv_tests(document)
    add_negative_tests(document)
    add_checklists(document)
    add_troubleshooting(document)

    for section in document.sections:
        section.start_type = WD_SECTION.NEW_PAGE
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    output = build_document()
    print(f"Created {output}")
