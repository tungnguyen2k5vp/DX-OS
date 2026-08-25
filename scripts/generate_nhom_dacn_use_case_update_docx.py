from __future__ import annotations

"""Generate the updated Chapter 2 use-case document for DX-OS.

The source document is intentionally kept unchanged.  This file contains a
clean, UTF-8 replacement/supplement for section 2.2 onward so it can be
reviewed before being merged into the graduation-project document.
"""

from datetime import date
from pathlib import Path
import subprocess

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "generated" / "Nhom_DACN_bo_sung_use_case_cap_nhat.docx"
SOURCE_DOC = ROOT / "docs" / "generated" / "Nhom_DACN.docx"

NAVY = "12304A"
TEAL = "007F86"
TEAL_DARK = "00636A"
SLATE = "536B80"
MUTED = "EDF4F7"
LIGHT_BLUE = "EAF4F8"
LIGHT_GREEN = "E8F5EF"
LIGHT_YELLOW = "FFF8E5"
LIGHT_RED = "FFF0F0"
GRID = "C8D8E1"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = GRID, size: str = "6") -> None:
    properties = cell._tc.get_or_add_tcPr()
    borders = properties.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn("w:" + name))
        if node is None:
            node = OxmlElement("w:" + name)
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def repeat_table_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    properties.append(header)


def keep_with_next(paragraph) -> None:
    properties = paragraph._p.get_or_add_pPr()
    element = OxmlElement("w:keepNext")
    properties.append(element)


def set_run_font(run, name="Aptos", size=10.5, color=NAVY, bold=False, italic=False) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(1.8)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.12

    for style_name, size, color, space_before, space_after in (
        ("Title", 26, TEAL_DARK, 0, 10),
        ("Heading 1", 18, NAVY, 15, 7),
        ("Heading 2", 14, TEAL_DARK, 12, 5),
        ("Heading 3", 12, NAVY, 9, 4),
        ("Heading 4", 11, TEAL_DARK, 7, 3),
        ("Heading 5", 10.5, TEAL_DARK, 6, 2),
    ):
        style = doc.styles[style_name]
        style.font.name = "Aptos Display" if style_name in {"Title", "Heading 1"} else "Aptos"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), style.font.name)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(space_before)
        style.paragraph_format.space_after = Pt(space_after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Paragraph", "Body Text"):
        style = doc.styles[style_name]
        style.font.name = "Aptos"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
        style.font.size = Pt(10.5)
        style.font.color.rgb = RGBColor.from_string(NAVY)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.1

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.add_run("DX-OS LAB  |  BỔ SUNG CHƯƠNG 2 – USE CASE")
    set_run_font(header_run, size=8, color=SLATE, bold=True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("DX-OS Lab  •  Tài liệu đối chiếu mã nguồn  •  Trang ")
    set_run_font(footer_run, size=8, color=SLATE)
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    footer._p.append(field_begin)
    footer._p.append(instr)
    footer._p.append(field_end)


def heading(doc: Document, text: str, level: int = 1):
    paragraph = doc.add_heading(text, level=level)
    keep_with_next(paragraph)
    return paragraph


def paragraph(doc: Document, text: str = "", style: str | None = None, bold_prefix: str | None = None):
    item = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        run = item.add_run(bold_prefix)
        set_run_font(run, bold=True)
        run = item.add_run(text[len(bold_prefix) :])
        set_run_font(run)
    else:
        run = item.add_run(text)
        set_run_font(run)
    return item


def bullet(doc: Document, text: str) -> None:
    item = doc.add_paragraph(style="List Bullet")
    run = item.add_run(text)
    set_run_font(run)


def numbered(doc: Document, text: str) -> None:
    item = doc.add_paragraph(style="List Number")
    run = item.add_run(text)
    set_run_font(run)


def code_block(doc: Document, text: str) -> None:
    item = doc.add_paragraph()
    item.paragraph_format.left_indent = Cm(0.35)
    item.paragraph_format.right_indent = Cm(0.35)
    item.paragraph_format.space_before = Pt(4)
    item.paragraph_format.space_after = Pt(7)
    item.paragraph_format.line_spacing = 1.0
    set_cell_shading_for_paragraph(item, LIGHT_BLUE)
    run = item.add_run(text)
    set_run_font(run, name="Cascadia Mono", size=9, color=NAVY)


def set_cell_shading_for_paragraph(paragraph, fill: str) -> None:
    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def callout(doc: Document, title: str, body: str, fill: str = LIGHT_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, fill, "0")
    set_cell_margins(cell, 150, 180, 150, 180)
    p = cell.paragraphs[0]
    run = p.add_run(title)
    set_run_font(run, size=10.5, color=TEAL_DARK, bold=True)
    p.add_run("\n")
    run = p.add_run(body)
    set_run_font(run, size=10, color=NAVY)
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    header_row = table.rows[0]
    repeat_table_header(header_row)
    for index, value in enumerate(headers):
        cell = header_row.cells[index]
        if widths:
            cell.width = Inches(widths[index])
        set_cell_shading(cell, TEAL)
        set_cell_border(cell, TEAL)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(value)
        set_run_font(run, size=9.2, color=WHITE, bold=True)
    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            cell = cells[index]
            if widths:
                cell.width = Inches(widths[index])
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(str(value))
            set_run_font(run, size=9.1, color=NAVY)
    doc.add_paragraph()


def source_snapshot() -> str:
    try:
        result = subprocess.run(
            ["git", "rev-parse", "--short", "HEAD"],
            cwd=ROOT,
            capture_output=True,
            text=True,
            check=True,
        )
        return result.stdout.strip()
    except (OSError, subprocess.CalledProcessError):
        return "không xác định"


ACTORS = [
    ["Nhân viên", "employee", "Tạo/sửa/gửi/hủy phiếu của mình; trao đổi; tải chứng từ; xác nhận nhận hàng của đơn do mình yêu cầu."],
    ["Trưởng bộ phận", "department_manager", "Xem và xử lý phiếu trong phạm vi phòng ban; có thể tạo phiếu; ủy quyền xử lý trong thời gian xác định; xác nhận giao nhận theo phạm vi được phép."],
    ["Tài chính", "finance", "Duyệt cấp tài chính; ngân sách; nhà cung cấp; báo giá; đơn hàng; hóa đơn; thanh toán; báo cáo."],
    ["Kiểm toán nội bộ", "auditor", "Đọc dữ liệu theo phạm vi tổ chức; xem nhật ký, hồ sơ kiểm toán, gói bằng chứng; không thay đổi trạng thái nghiệp vụ."],
    ["Điều phối khuyến nghị", "ai_operator", "Quét dữ liệu để tạo khuyến nghị có bằng chứng; chấp nhận, bác bỏ hoặc bỏ qua khuyến nghị; không tự động thay đổi phiếu."],
    ["Quản trị DX-OS", "dx_admin", "Quản lý hồ sơ người dùng, phòng ban, quy tắc phê duyệt, chính sách vận hành; xem báo cáo và vận hành khuyến nghị."],
]


USE_CASE_SUMMARY = [
    ["UC-01", "Đăng nhập hệ thống", "Tất cả", "Xác thực SSO qua Keycloak; tạo phiên làm việc."],
    ["UC-02", "Tạo và lưu phiếu mua sắm", "Nhân viên, Trưởng bộ phận", "Tạo bản nháp với dòng hàng, trung tâm chi phí và tiền tệ."],
    ["UC-03", "Xem danh sách và chi tiết phiếu", "Theo phạm vi vai trò", "Xem thông tin, dòng hàng, tệp, dòng thời gian và ngân sách."],
    ["UC-04", "Sửa phiếu mua sắm", "Chủ phiếu", "Chỉ sửa DRAFT/CHANGES_REQUESTED, có kiểm soát phiên bản."],
    ["UC-05", "Gửi hoặc gửi lại phiếu", "Chủ phiếu", "Kiểm tra chứng từ bắt buộc rồi chuyển SUBMITTED."],
    ["UC-06", "Phê duyệt phiếu hai cấp", "Trưởng bộ phận, Tài chính", "Approve, yêu cầu chỉnh sửa hoặc từ chối; giữ/commit/giải phóng ngân sách."],
    ["UC-07", "Hủy phiếu", "Chủ phiếu", "Hủy DRAFT/CHANGES_REQUESTED; phiếu vẫn còn lịch sử."],
    ["UC-08", "Quản lý tệp và trao đổi", "Người có quyền xem; chủ phiếu khi sửa", "Tải/xóa/tải xuống tệp; bình luận được lưu cùng phiếu."],
    ["UC-09", "Kiểm tra và điều chỉnh ngân sách", "Tài chính; Kiểm toán đọc", "Theo dõi cấp phát/đang giữ/đã cam kết/còn lại; điều chỉnh có lý do."],
    ["UC-10", "Xem báo cáo vận hành", "Tài chính, Kiểm toán, Quản trị", "Báo cáo mua sắm, ngân sách, thời gian xử lý; có liên kết Metabase."],
    ["UC-11", "Theo dõi công việc và thông báo", "Tất cả", "Hàng đợi công việc, hạn SLA, thông báo chưa đọc và đánh dấu đã đọc."],
    ["UC-12", "Approval Inbox – hàng đợi phê duyệt", "Trưởng bộ phận, Tài chính", "Chọn đúng tầng duyệt từ trạng thái phiếu; xử lý liền mạch tại một nơi."],
    ["UC-13", "Danh mục mua sắm và kiểm tra trùng", "Nhân viên, Trưởng bộ phận", "Gợi ý hàng hóa chuẩn; cảnh báo phiếu tương tự trong phòng ban 90 ngày."],
    ["UC-14", "Quản lý hồ sơ nhà cung cấp", "Tài chính; Kiểm toán đọc", "Hồ sơ pháp lý, liên hệ, ngân hàng, hợp đồng, tuân thủ và mức rủi ro."],
    ["UC-15", "Nhập, chấm điểm và chọn báo giá", "Tài chính; Kiểm toán đọc", "So sánh giá, giao hàng, chất lượng, tuân thủ; chọn một báo giá hợp lệ."],
    ["UC-16", "Tạo và quản lý đơn hàng", "Tài chính", "Tạo từ phiếu/báo giá đã duyệt; cập nhật hoặc hủy trước khi phát sinh nhận hàng/hóa đơn."],
    ["UC-17", "Ghi nhận giao nhận và ngoại lệ", "Nhân viên, Trưởng bộ phận; Tài chính đọc", "Nhận đủ/một phần/hỏng/sai hàng/từ chối; lưu biên bản và lịch sử."],
    ["UC-18", "Hóa đơn, đối soát và thanh toán", "Tài chính; Kiểm toán đọc", "Ghi hóa đơn, kiểm tra nhận hàng–tổng tiền–tiền tệ, tranh chấp, thanh toán từng phần."],
    ["UC-19", "Ủy quyền và quy tắc phê duyệt", "Quản trị; Trưởng bộ phận; Tài chính/Kiểm toán đọc", "Cấu hình tuyến duyệt theo phòng ban/giá trị và ủy quyền có thời hạn."],
    ["UC-20", "Chính sách SLA và chứng từ", "Quản trị; Kiểm toán đọc", "Thiết lập thời hạn xử lý và ngưỡng chứng từ bắt buộc."],
    ["UC-21", "Kiểm toán, hồ sơ vụ việc và gói bằng chứng", "Kiểm toán; Quản trị đọc", "Lọc nhật ký, tạo hồ sơ, theo dõi khắc phục, tải gói JSON theo một phiếu."],
    ["UC-22", "Khuyến nghị có giải thích", "Điều phối AI, Quản trị vận hành; Tài chính/Kiểm toán đọc", "Sinh cảnh báo SLA/giá trị/trùng lặp/rủi ro; con người quyết định."],
    ["UC-23", "Quản trị người dùng và phòng ban", "Quản trị DX-OS", "Cập nhật hồ sơ, phòng ban, trạng thái truy cập; phát hiện xung đột vai trò."],
]


DETAIL_CASES = [
    {
        "id": "UC-11",
        "title": "Theo dõi công việc và thông báo",
        "actors": "Tất cả người dùng đã đăng nhập",
        "summary": "Use case này gom các việc mà người dùng cần xử lý và các sự kiện phát sinh từ cùng một phiếu, giúp người dùng đi thẳng đến đúng màn hình thay vì tự tìm lại dữ liệu.",
        "basic": [
            "Người dùng mở “Việc của tôi” hoặc “Thông báo”. Hệ thống gọi danh sách việc theo tài khoản, phòng ban và vai trò hiện tại.",
            "Màn hình công việc hiển thị tổng việc cần chú ý, việc sắp đến hạn, việc quá hạn và từng loại việc: hoàn thiện phiếu, trưởng bộ phận xem xét, tài chính thẩm định hoặc theo dõi hạn xử lý.",
            "Người dùng chọn một việc. Hệ thống mở đúng phiếu mua sắm và giữ nguyên mã phiếu để tiếp tục luồng liên vai trò.",
            "Từ trung tâm thông báo, người dùng có thể lọc chỉ thông báo chưa đọc, mở liên kết nguồn, đánh dấu từng thông báo hoặc đánh dấu tất cả đã đọc.",
            "Sau các sự kiện như yêu cầu sửa, duyệt, đặt hàng, nhận hàng, ghi hóa đơn hoặc thanh toán, máy chủ xếp thông báo vào hàng đợi gửi cho người nhận phù hợp. Use case kết thúc.",
        ],
        "alternate": [
            "Nếu không có việc hoặc thông báo, hệ thống hiển thị trạng thái rỗng, không coi đó là lỗi.",
            "Nếu thông báo đang chờ gửi hoặc gửi lỗi, quản trị có thể theo dõi bộ đếm hàng đợi; việc này không tự thay đổi trạng thái phiếu.",
            "Nếu phiên đăng nhập hết hạn hoặc API không phản hồi, hệ thống báo lỗi và cho phép tải lại.",
        ],
        "special": "Dữ liệu công việc là dữ liệu hỗ trợ điều hướng; quyền thao tác vẫn phải được kiểm tra lại tại API nghiệp vụ. Thông báo dùng phân trang và có trạng thái đã đọc/chưa đọc.",
        "pre": "Người dùng đã đăng nhập; phiếu hoặc sự kiện liên quan tồn tại trong cùng tổ chức.",
        "post": "Thành công: công việc/thông báo được hiển thị và trạng thái đọc được cập nhật nếu người dùng chọn. Thất bại: dữ liệu nghiệp vụ không bị thay đổi.",
        "extension": "Từ một thông báo có resourceType là purchase_request, mở thẳng chi tiết phiếu; loại resource khác quay về trang tổng quan.",
        "evidence": "GET /api/v1/me/tasks-summary; GET/POST /api/v1/me/notifications; frontend: features/procurement/pages/work-center và features/dashboard/pages/notification-center.",
    },
    {
        "id": "UC-12",
        "title": "Approval Inbox – hàng đợi phê duyệt",
        "actors": "Trưởng bộ phận, Tài chính; người được ủy quyền hợp lệ",
        "summary": "Use case này cung cấp một hàng đợi phê duyệt duy nhất nhưng tự chọn đúng tầng xử lý theo vai trò và trạng thái phiếu.",
        "basic": [
            "Trưởng bộ phận mở “Phê duyệt”; hệ thống lọc phiếu SUBMITTED trong phòng ban mà người đó quản lý. Tài chính mở cùng màn hình; hệ thống lọc phiếu MANAGER_APPROVED trong tổ chức.",
            "Người duyệt chọn “Xem và xử lý”. Hệ thống hiển thị lý do, dòng hàng, tệp, dòng thời gian, bình luận và kiểm tra ngân sách.",
            "Người duyệt chọn Phê duyệt. Hệ thống khóa phiếu, kiểm tra phiên bản, phạm vi phòng ban/tổ chức, chống tự duyệt và tuyến duyệt hiện hành.",
            "Nếu là bước trưởng bộ phận, phiếu chuyển sang MANAGER_APPROVED và số tiền được giữ ngân sách. Nếu là bước tài chính, phiếu chuyển sang APPROVED và tiền được cam kết chính thức.",
            "Hệ thống ghi process event/audit, tạo thông báo cho vai trò tiếp theo và làm mới hàng đợi. Use case kết thúc.",
        ],
        "alternate": [
            "Từ chối hoặc yêu cầu chỉnh sửa bắt buộc có nội dung giải thích; phiếu chuyển REJECTED hoặc CHANGES_REQUESTED và thông báo cho chủ phiếu.",
            "Nếu không đủ ngân sách ở bước giữ chỗ, API trả lỗi xung đột và phiếu cùng ngân sách giữ nguyên.",
            "Nếu người duyệt là chính người yêu cầu hoặc không thuộc phạm vi, hệ thống từ chối thao tác.",
            "Nếu phiên bản đã thay đổi hoặc nút được bấm lặp, hệ thống dùng ExpectedVersion và Idempotency-Key để không ghi trùng.",
        ],
        "special": "Mỗi quyết định phải gắn người thực hiện, vai trò, thời điểm, mã tương quan và lý do khi cần. Quy tắc phê duyệt có thể làm tuyến duyệt khác mặc định.",
        "pre": "Đã đăng nhập; có phiếu ở đúng trạng thái và người dùng có vai trò/phạm vi xử lý phù hợp.",
        "post": "Thành công: trạng thái phiếu, số giữ chỗ/cam kết ngân sách, audit và thông báo được cập nhật trong cùng giao dịch. Thất bại: toàn bộ thay đổi được hoàn tác.",
        "extension": "Mở chi tiết phiếu để xem timeline, thêm bình luận hoặc chuyển sang nhánh yêu cầu chỉnh sửa; không tạo phiếu mới.",
        "evidence": "GET /api/v1/purchase-requests; POST /api/v1/purchase-requests/{requestID}/transitions; frontend: approval-inbox.ts/html; backend: store.go Transition và approval_governance.go.",
    },
    {
        "id": "UC-13",
        "title": "Danh mục mua sắm và kiểm tra trùng phiếu",
        "actors": "Nhân viên, Trưởng bộ phận",
        "summary": "Use case này hỗ trợ người lập phiếu chọn hàng hóa chuẩn và phát hiện các phiếu tương tự trước khi gửi, nhưng chỉ cảnh báo để người dùng quyết định.",
        "basic": [
            "Người dùng mở màn hình tạo phiếu. Hệ thống tải các mặt hàng đang hoạt động trong danh mục theo tổ chức, gồm mã, tên, nhóm, đơn vị, đơn giá tham khảo và tiền tệ.",
            "Người dùng chọn mặt hàng hoặc dùng thông tin đó để điền dòng hàng; tổng tiền phiếu được tính lại như luồng tạo phiếu thông thường.",
            "Trước khi lưu/gửi, hệ thống gửi tiêu đề, trung tâm chi phí, tổng tiền và mã phiếu đang sửa tới chức năng kiểm tra trùng.",
            "Máy chủ chỉ tìm các phiếu cùng phòng ban, chưa bị từ chối/hủy, tạo trong 90 ngày; tính độ giống tiêu đề, độ gần giá trị trong 10% và cùng trung tâm chi phí.",
            "Nếu có kết quả, hệ thống hiển thị tối đa 5 ứng viên cùng lý do; người dùng đối chiếu rồi tiếp tục hoặc quay lại sửa. Use case kết thúc.",
        ],
        "alternate": [
            "Không có ứng viên phù hợp thì hiển thị “không phát hiện trùng”, không chặn việc tạo phiếu.",
            "Nếu tiêu đề ngắn hơn 3 ký tự hoặc tổng tiền sai định dạng, hệ thống trả lỗi kiểm tra dữ liệu.",
            "Nếu người dùng đang sửa một phiếu, mã phiếu đó được loại khỏi danh sách so sánh.",
        ],
        "special": "Ngưỡng tương đồng tiêu đề chính là từ 60%; trường hợp từ 45% kết hợp giá gần nhau và cùng trung tâm chi phí cũng được cảnh báo. Đây là gợi ý kiểm soát, không phải kết luận gian lận.",
        "pre": "Đã đăng nhập và thuộc phạm vi tạo phiếu; danh mục có thể rỗng nhưng API vẫn trả danh sách hợp lệ.",
        "post": "Kết quả cảnh báo được hiển thị; không có bản ghi nghiệp vụ nào bị đổi trạng thái chỉ vì kiểm tra trùng.",
        "extension": "Từ ứng viên trùng, người dùng mở chi tiết phiếu để so sánh; sau đó quay lại phiếu hiện tại bằng mã PR.",
        "evidence": "GET /api/v1/procurement-catalog; POST /api/v1/purchase-requests/duplicate-check; backend: guided_requests.go.",
    },
    {
        "id": "UC-14",
        "title": "Quản lý hồ sơ nhà cung cấp",
        "actors": "Tài chính; Kiểm toán xem không sửa",
        "summary": "Use case này duy trì một hồ sơ nhà cung cấp dùng chung cho báo giá, đơn hàng, đối soát và kiểm tra rủi ro.",
        "basic": [
            "Tài chính mở “Nhà cung cấp”. Hệ thống liệt kê nhà cung cấp trong cùng tổ chức, ưu tiên nhà cung cấp đang hoạt động và hiển thị mức rủi ro.",
            "Người dùng có quyền quản lý tạo hồ sơ với mã, tên, mã số thuế, liên hệ, địa chỉ, ngân hàng, số tài khoản, hợp đồng, ngày hết hạn, tình trạng tuân thủ, điểm hiệu suất và ghi chú.",
            "Hệ thống kiểm tra dữ liệu, mã nhà cung cấp/mã số thuế không trùng, lưu hồ sơ và audit SUPPLIER_CREATED.",
            "Khi có thay đổi, Tài chính chọn chỉnh sửa; hệ thống khóa bản ghi, kiểm tra ExpectedVersion, lưu phiên bản mới và audit SUPPLIER_UPDATED.",
            "Kiểm toán chỉ đọc hồ sơ để đối chiếu. Hồ sơ nhà cung cấp được dùng tiếp ở luồng báo giá và tạo đơn hàng. Use case kết thúc.",
        ],
        "alternate": [
            "Mã hoặc mã số thuế đã tồn tại thì từ chối lưu và giữ nguyên biểu mẫu.",
            "Nếu phiên bản đã cũ, hệ thống báo xung đột; người dùng phải tải lại rồi đối chiếu trước khi sửa tiếp.",
            "Nhà cung cấp BLOCKED hoặc không ACTIVE không được dùng để ghi báo giá hợp lệ/đặt hàng.",
        ],
        "special": "Tài chính mới có quyền tạo/sửa; người có đồng thời finance và auditor bị coi là xung đột và không được dùng quyền thay đổi. Thông tin ngân hàng là dữ liệu nhạy cảm, chỉ hiển thị trong phạm vi được cấp.",
        "pre": "Đã đăng nhập; Tài chính có quyền quản lý hoặc Kiểm toán có quyền đọc.",
        "post": "Hồ sơ được tạo/cập nhật, tăng version và có resource audit; thất bại không làm thay đổi hồ sơ hiện tại.",
        "extension": "Chuyển từ hồ sơ nhà cung cấp sang so sánh báo giá; báo giá giữ supplierId để liên kết ngược.",
        "evidence": "GET/POST/PATCH /api/v1/suppliers; backend: store.go ListSuppliers/CreateSupplier/UpdateSupplier; model Supplier.",
    },
    {
        "id": "UC-15",
        "title": "Nhập, chấm điểm và chọn báo giá",
        "actors": "Tài chính; Kiểm toán xem kết quả",
        "summary": "Use case này biến một phiếu đã được duyệt thành hồ sơ so sánh báo giá, giúp Tài chính chọn nhà cung cấp có căn cứ và bàn giao trực tiếp sang tạo đơn hàng.",
        "basic": [
            "Tài chính mở “So sánh báo giá”. Hệ thống chỉ đưa vào các phiếu APPROVED chưa có đơn hàng không bị hủy.",
            "Tài chính nhập báo giá cho nhà cung cấp ACTIVE, không BLOCKED: mã báo giá, tổng tiền, tiền tệ, ngày giao dự kiến, bảo hành, điều khoản thanh toán và ghi chú.",
            "Hệ thống tính điểm theo dữ liệu hiện có: giá 40%, tiến độ giao 25%, chất lượng/hiệu suất nhà cung cấp 20%, tuân thủ và rủi ro 15%. Màn hình hiển thị điểm từng thành phần và điểm tổng hợp trên thang 100.",
            "Tài chính đối chiếu các báo giá, ghi lý do tối thiểu 10 ký tự và chọn một báo giá. Hệ thống đánh dấu báo giá được chọn SELECTED, các báo giá còn lại REJECTED, hồ sơ so sánh thành AWARDED.",
            "Người dùng bấm “Tạo đơn hàng” ngay cạnh báo giá được chọn. Hệ thống chuyển sang màn hình đơn hàng và điền sẵn phiếu, nhà cung cấp, mã tham chiếu, ngày giao dự kiến và ghi chú; người dùng chỉ cần kiểm tra rồi phát hành. Use case kết thúc.",
        ],
        "alternate": [
            "Phiếu chưa APPROVED, nhà cung cấp không hoạt động/đã bị chặn hoặc tiền tệ không khớp thì không được ghi báo giá.",
            "Nếu thiếu lý do chọn, ExpectedVersion hoặc Idempotency-Key, hệ thống không chốt lựa chọn.",
            "Nếu báo giá/hồ sơ đã bị người khác cập nhật, hệ thống báo xung đột version và yêu cầu tải lại.",
        ],
        "special": "Điểm tổng hợp là điểm hỗ trợ quyết định, không phải tự động chọn nhà cung cấp. Điểm giá được chuẩn hóa theo giá thấp nhất; điểm tuân thủ bị trừ theo rủi ro nhà cung cấp.",
        "pre": "Phiếu đã được duyệt cấp cuối; Tài chính có quyền quản lý sourcing; nhà cung cấp đã có hồ sơ hợp lệ.",
        "post": "Hồ sơ sourcing AWARDED, một báo giá SELECTED và các báo giá khác REJECTED; sự kiện và audit được lưu.",
        "extension": "Tạo đơn hàng từ báo giá được chọn là điểm mở rộng trực tiếp; không nhập lại dữ liệu nhà cung cấp thủ công.",
        "evidence": "GET /api/v1/sourcing; POST/PATCH /api/v1/sourcing/quotes; POST /api/v1/sourcing/quotes/{quoteID}/selection; frontend: sourcing-board; backend: sourcing.go.",
    },
    {
        "id": "UC-16",
        "title": "Tạo và quản lý đơn hàng",
        "actors": "Tài chính; Nhân viên/Trưởng bộ phận xem kết quả",
        "summary": "Use case này tạo PO từ phiếu đã duyệt hoặc báo giá đã chọn, sau đó cho phép chỉnh thông tin trước khi phát sinh giao nhận/hóa đơn.",
        "basic": [
            "Tài chính mở “Đặt hàng và giao nhận” hoặc bấm “Tạo đơn hàng” từ báo giá đã chọn.",
            "Biểu mẫu nhận sẵn phiếu, nhà cung cấp, mã tham chiếu ngoài, ngày giao dự kiến và ghi chú; người dùng kiểm tra lại rồi phát hành đơn hàng.",
            "API khóa phiếu, yêu cầu phiếu APPROVED, nhà cung cấp ACTIVE và kiểm tra điều kiện báo giá. Với phiếu VND từ 50.000.000 trở lên, phải có hồ sơ AWARDED nếu chưa có đơn hàng.",
            "Hệ thống tạo mã PO, ghi sự kiện ORDER_PLACED, audit và thông báo cho người yêu cầu. Bảng giao nhận hiển thị trạng thái ORDERED.",
            "Trước khi có giao nhận hoặc hóa đơn, Tài chính có thể cập nhật nhà cung cấp, ngày giao, mã ngoài và ghi chú; hoặc hủy với lý do. Use case kết thúc.",
        ],
        "alternate": [
            "Phiếu chưa APPROVED, nhà cung cấp không ACTIVE hoặc báo giá chưa AWARDED ở ngưỡng bắt buộc thì tạo PO bị từ chối.",
            "Gửi lại cùng Idempotency-Key với đúng dữ liệu sẽ trả kết quả cũ; cùng khóa nhưng khác dữ liệu là xung đột.",
            "Không cho cập nhật/hủy PO khi đã có receipt hoặc invoice; hủy bắt buộc lý do 10–2.000 ký tự.",
            "Nếu version PO đã thay đổi, hệ thống từ chối cập nhật để tránh ghi đè.",
        ],
        "special": "Tạo, cập nhật và hủy đơn hàng chỉ do Tài chính không đồng thời là Kiểm toán thực hiện. Tạo PO không đổi trạng thái phiếu APPROVED, nhưng tạo dấu vết ORDER_PLACED.",
        "pre": "Đã có phiếu APPROVED; nếu thuộc ngưỡng cần sourcing thì đã chọn báo giá; có nhà cung cấp hợp lệ.",
        "post": "Thành công: PO tồn tại và liên kết PR–supplier; thất bại: không tạo PO một phần và dữ liệu phiếu không đổi.",
        "extension": "Từ PO chuyển sang ghi nhận giao nhận; sau khi nhận hàng, PO trở thành đầu vào cho luồng hóa đơn.",
        "evidence": "POST/PATCH /api/v1/procurement-operations/orders; POST /orders/{requestID}/transitions; backend: store.go CreatePurchaseOrder và operations_extensions.go.",
    },
    {
        "id": "UC-17",
        "title": "Ghi nhận giao nhận và xử lý ngoại lệ",
        "actors": "Nhân viên là người yêu cầu; Trưởng bộ phận trong phạm vi; Tài chính/Kiểm toán xem",
        "summary": "Use case này ghi nhận số lượng thực nhận theo từng dòng, tình trạng hàng và các trường hợp giao thiếu, hỏng, sai hàng hoặc từ chối nhận.",
        "basic": [
            "Người yêu cầu mở PO ở trạng thái ORDERED/PARTIALLY_RECEIVED/RECEIPT_EXCEPTION và chọn ghi nhận giao nhận.",
            "Người dùng nhập ngày nhận, kết quả (nhận một phần, nhận đủ, hỏng, sai hàng hoặc từ chối), ghi chú tối thiểu 5 ký tự và số lượng/tình trạng cho từng dòng hàng.",
            "Hệ thống kiểm tra dòng hàng thuộc PO, không nhận vượt số lượng đặt đối với dòng ACCEPTED, kiểm tra version và Idempotency-Key.",
            "Hệ thống lưu biên bản GR, các dòng nhận hàng, audit và process event. Nếu chưa nhận đủ thì PO là PARTIALLY_RECEIVED; nhận đủ là RECEIVED; có lỗi là RECEIPT_EXCEPTION.",
            "Thông báo cập nhật giao nhận được gửi cho Tài chính và người yêu cầu; lịch sử biên bản có thể mở lại từ chi tiết phiếu. Use case kết thúc.",
        ],
        "alternate": [
            "Kết quả COMPLETE nhưng còn dòng chưa nhận đủ sẽ bị từ chối; người dùng phải chọn PARTIAL hoặc bổ sung số lượng.",
            "Số lượng nhận vượt số lượng đặt, dòng không thuộc PO, ngày sai hoặc điều kiện dòng không hợp lệ thì không lưu biên bản.",
            "Kiểm toán chỉ đọc; người không phải chủ phiếu hoặc quản lý đúng phòng ban không được ghi nhận.",
        ],
        "special": "Hệ thống hỗ trợ nhận nhiều lần; tổng ACCEPTED của các biên bản không được vượt số lượng đặt. Đây là ghi nhận thực tế, không tự động tạo hóa đơn.",
        "pre": "PO tồn tại và đang ở trạng thái cho phép nhận; người dùng thuộc đúng requester/department scope.",
        "post": "Biên bản và từng dòng nhận được lưu; trạng thái PO, audit, timeline và thông báo được cập nhật trong một giao dịch.",
        "extension": "Sau khi PO RECEIVED, Tài chính có thể ghi hóa đơn; nếu có ngoại lệ, Tài chính xử lý trước khi đối soát.",
        "evidence": "POST /api/v1/procurement-operations/orders/{requestID}/receipts; GET cùng đường dẫn; backend: operations_extensions.go RecordReceipt/ListReceipts.",
    },
    {
        "id": "UC-18",
        "title": "Hóa đơn, đối soát và thanh toán",
        "actors": "Tài chính; Kiểm toán đọc và kiểm tra lịch sử",
        "summary": "Use case này nối PO đã nhận hàng với hóa đơn và các lần thanh toán, đồng thời ghi rõ lý do khi số tiền hoặc tiền tệ không khớp.",
        "basic": [
            "Tài chính mở “Hóa đơn và thanh toán”. Hệ thống hiển thị PO, nhà cung cấp, trạng thái nhận hàng, hóa đơn, số đã trả, số còn lại và hạn thanh toán.",
            "Tài chính ghi hóa đơn với số hóa đơn, ngày phát hành, hạn trả, số tiền, tiền tệ và ghi chú. Hệ thống không cho ghi cho PO đã hủy và lưu trạng thái RECORDED.",
            "Hệ thống đối soát ở mức hiện đang triển khai: PO phải RECEIVED; tiền tệ hóa đơn phải khớp; tổng hóa đơn không vượt giá trị phiếu; thấp hơn là PARTIAL_MATCH, bằng là MATCHED.",
            "Nếu phù hợp, Tài chính chọn xác minh để chuyển VERIFIED. Nếu có chênh lệch, chọn tranh chấp DISPUTED và ghi lý do; có thể mở lại về RECORDED sau khi xử lý.",
            "Với hóa đơn VERIFIED, Tài chính ghi một lần thanh toán toàn bộ hoặc một phần. Hệ thống không cho trả vượt số dư; đủ tiền thì chuyển PAID, chưa đủ vẫn VERIFIED. Use case kết thúc.",
        ],
        "alternate": [
            "PO chưa RECEIVED, sai tiền tệ hoặc tổng tiền vượt giá trị phiếu thì không được VERIFY; trạng thái vẫn RECORDED để xử lý.",
            "DISPUTE bắt buộc có ghi chú; UPDATE chỉ được thực hiện với RECORDED/DISPUTED.",
            "Số tiền thanh toán lớn hơn số còn lại, ngày thanh toán ở tương lai, version sai hoặc khóa gửi trùng không hợp lệ thì bị từ chối.",
            "Kiểm toán xem được bảng và lịch sử thanh toán nhưng không tạo/sửa/verify/pay.",
        ],
        "special": "Đối soát hiện tại chưa kiểm tra chi tiết từng dòng hàng của hóa đơn; nó dùng trạng thái nhận PO, tổng tiền và tiền tệ. Tài liệu không gọi đây là đối soát dòng hàng đầy đủ.",
        "pre": "PO không bị hủy; hóa đơn có dữ liệu hợp lệ; Tài chính có quyền quản lý.",
        "post": "Hóa đơn, event, audit, số đã trả/còn lại và thông báo được cập nhật nhất quán; lịch sử thanh toán không bị xóa.",
        "extension": "Từ bảng hóa đơn mở danh sách các lần thanh toán hoặc chuyển sang trung tâm kiểm toán để truy vết.",
        "evidence": "GET/POST/PATCH /api/v1/invoices; POST /invoices/{invoiceID}/transitions; GET/POST /invoices/{invoiceID}/payments; backend: invoices.go, payments.go.",
    },
    {
        "id": "UC-19",
        "title": "Ủy quyền và quy tắc phê duyệt",
        "actors": "Quản trị DX-OS; Trưởng bộ phận; Tài chính/Kiểm toán xem",
        "summary": "Use case này bảo đảm phiếu vẫn có tuyến xử lý khi người duyệt vắng mặt và tự chọn số vòng duyệt theo phòng ban, tiền tệ và giá trị.",
        "basic": [
            "Người dùng mở “Ủy quyền và quy tắc”. Hệ thống tải quy tắc của tổ chức, các ủy quyền và danh sách người có thể nhận ủy quyền theo phạm vi.",
            "Quản trị tạo/cập nhật quy tắc gồm phòng ban tùy chọn, tên, tiền tệ, khoảng giá trị, yêu cầu trưởng bộ phận, yêu cầu tài chính, độ ưu tiên và trạng thái hoạt động.",
            "Trưởng bộ phận tạo ủy quyền cho một người dùng đang hoạt động, nhập ngày bắt đầu/kết thúc và lý do; hệ thống không cho tự ủy quyền cho chính mình.",
            "Khi gửi phiếu, hệ thống ưu tiên quy tắc phòng ban hơn quy tắc chung, sau đó ưu tiên nhỏ hơn để resolve route; khi duyệt, ủy quyền chỉ có hiệu lực trong ngày đã cấu hình.",
            "Người tạo ủy quyền hoặc Quản trị có thể bật/tắt ủy quyền, có ExpectedVersion và audit. Use case kết thúc.",
        ],
        "alternate": [
            "Khoảng ngày kết thúc trước ngày bắt đầu, lý do quá ngắn hoặc người nhận không hoạt động/cùng tổ chức thì không lưu.",
            "Quy tắc không có ít nhất một cấp duyệt hoặc khoảng giá trị sai thì bị từ chối.",
            "Kiểm toán chỉ đọc; Tài chính được xem nhưng không sửa quy tắc/ủy quyền.",
        ],
        "special": "Vai trò nghiệp vụ vẫn lấy từ Keycloak; màn hình này chỉ cấu hình tuyến và trạng thái ủy quyền, không cấp role mới. Mọi thay đổi quy tắc/ủy quyền có resource audit.",
        "pre": "Đã đăng nhập và có quyền xem/quản lý tương ứng; phòng ban và người dùng mục tiêu tồn tại.",
        "post": "Quy tắc/ủy quyền được lưu với version mới; các phiếu gửi sau đó dùng tuyến phù hợp, phiếu đang xử lý không bị tự ý đổi ngược.",
        "extension": "Từ quy tắc chuyển về Approval Inbox để kiểm thử bằng một phiếu có giá trị/phòng ban tương ứng.",
        "evidence": "GET /api/v1/approval-governance; POST/PATCH /api/v1/approval-governance/delegations; POST/PATCH /api/v1/approval-governance/rules; backend: approval_governance.go.",
    },
    {
        "id": "UC-20",
        "title": "Chính sách SLA và chứng từ",
        "actors": "Quản trị DX-OS; Kiểm toán xem",
        "summary": "Use case này cấu hình mục tiêu thời gian xử lý và ngưỡng chứng từ bắt buộc để các luồng gửi duyệt và khuyến nghị có cùng một nguồn quy tắc.",
        "basic": [
            "Quản trị hoặc Kiểm toán mở “Chính sách vận hành”. Hệ thống hiển thị danh sách chính sách SLA theo process và quy tắc chứng từ theo tiền tệ.",
            "Quản trị cập nhật số giờ mục tiêu và trạng thái hoạt động của một SLA; hoặc cập nhật ngưỡng tiền, loại tài liệu bắt buộc và trạng thái của quy tắc chứng từ.",
            "Hệ thống kiểm tra giới hạn giờ 1–720, ngưỡng tiền không âm, loại chứng từ thuộc QUOTATION/SPECIFICATION/CONTRACT/OTHER và ExpectedVersion.",
            "Nếu hợp lệ, hệ thống lưu version mới, ghi audit rồi dùng chính sách đó cho kiểm tra gửi duyệt và theo dõi hạn. Use case kết thúc.",
        ],
        "alternate": [
            "Kiểm toán nhìn thấy dữ liệu nhưng không có nút chỉnh sửa và API cũng từ chối PATCH.",
            "Nếu version đã cũ hoặc mã process/rule không tồn tại, hệ thống báo xung đột/không tìm thấy.",
            "Tắt một quy tắc chỉ làm quy tắc không còn áp dụng cho lần kiểm tra sau; không xóa lịch sử audit.",
        ],
        "special": "SLA là mục tiêu thời gian xử lý, không phải cam kết giao hàng. Quy tắc chứng từ chỉ kiểm soát việc gửi duyệt theo ngưỡng; tệp thực tế vẫn được quản lý trong phiếu.",
        "pre": "Đã đăng nhập; Quản trị có quyền quản lý hoặc Kiểm toán có quyền đọc.",
        "post": "Chính sách được cập nhật có version/audit; phiếu mới hoặc lần gửi lại sau đó dùng giá trị chính sách hiện hành.",
        "extension": "Từ chính sách chứng từ mở phiếu có giá trị vượt ngưỡng để kiểm tra lỗi thiếu báo giá trước khi gửi duyệt.",
        "evidence": "GET /api/v1/admin/policies; PATCH /api/v1/admin/policies/sla/{processName}; PATCH /api/v1/admin/policies/attachments/{ruleID}; backend: policies.go.",
    },
    {
        "id": "UC-21",
        "title": "Kiểm toán, hồ sơ vụ việc và gói bằng chứng",
        "actors": "Kiểm toán nội bộ; Quản trị xem dữ liệu quản trị",
        "summary": "Use case này biến các dấu vết rải rác của một phiếu thành nhật ký có thể lọc, hồ sơ kiểm toán có người phụ trách và gói bằng chứng tải xuống được.",
        "basic": [
            "Kiểm toán mở “Trung tâm kiểm toán”, lọc nhật ký theo thời gian, tài nguyên, hành động, người thực hiện hoặc mã tương quan.",
            "Kiểm toán tạo hồ sơ vụ việc với tiêu đề, mô tả, mức độ, tài nguyên liên quan, người phụ trách, hạn xử lý và trạng thái OPEN.",
            "Trong quá trình xử lý, Kiểm toán cập nhật OPEN → IN_REMEDIATION → RESOLVED/CLOSED; khi RESOLVED/CLOSED bắt buộc có nội dung khắc phục.",
            "Kiểm toán nhập mã định danh phiếu và chọn tải gói bằng chứng. API tập hợp phiếu, timeline, tệp đính kèm, PO, biên bản nhận, hóa đơn và resource audit event thành JSON.",
            "Hệ thống tải file với tên theo mã phiếu; dữ liệu gốc vẫn ở hệ thống và không bị xóa. Use case kết thúc.",
        ],
        "alternate": [
            "Người không có auditor/dx_admin không truy cập nhật ký hoặc hồ sơ.",
            "Hồ sơ đóng mà thiếu resolution, tài nguyên/owner không thuộc tổ chức hoặc version sai thì bị từ chối.",
            "Nếu phiếu không tồn tại hoặc gói bằng chứng không sinh được đầy đủ thành phần, hệ thống báo lỗi thay vì tải file thiếu.",
        ],
        "special": "Gói bằng chứng hiện được tải ở định dạng JSON; đây là bản chụp có thời điểm sinh và người sinh, không phải file PDF ký số.",
        "pre": "Đã đăng nhập với quyền Kiểm toán; phiếu hoặc sự kiện cần kiểm tra thuộc cùng tổ chức.",
        "post": "Hồ sơ/audit event được lưu; gói bằng chứng tải về chứa các dữ liệu liên quan và có thể dùng để review.",
        "extension": "Từ một audit event mở resource gốc; từ gói bằng chứng quay lại timeline để đối chiếu trạng thái.",
        "evidence": "GET /api/v1/audit/events; GET/POST/PATCH /api/v1/audit/cases; GET /api/v1/audit/evidence/{requestID}; backend: audit_cases.go và reporting/audit.go.",
    },
    {
        "id": "UC-22",
        "title": "Khuyến nghị có giải thích",
        "actors": "Điều phối khuyến nghị, Quản trị; Tài chính/Kiểm toán xem",
        "summary": "Use case này quét dữ liệu để tạo các cảnh báo có bằng chứng, sau đó yêu cầu con người ghi quyết định; hệ thống không tự động sửa phiếu, ngân sách hay hóa đơn.",
        "basic": [
            "Điều phối khuyến nghị hoặc Quản trị mở “Trung tâm khuyến nghị” và chọn quét dữ liệu mới.",
            "Máy chủ tạo các khuyến nghị theo quy tắc có thể giải thích như quá hạn SLA, phiếu giá trị lớn, trùng phiếu, chia nhỏ mua sắm, bất thường đơn giá, quá hạn thanh toán, rủi ro nhà cung cấp, thay đổi hồ sơ nhà cung cấp hoặc xung đột vai trò.",
            "Mỗi khuyến nghị liên kết mã PR nếu có, mức rủi ro, tóm tắt, loại khuyến nghị, thời điểm sinh và evidence JSON; người xem mở “Bằng chứng có thể giải thích” để đối chiếu dữ liệu gốc.",
            "Người có quyền vận hành nhập lý do tối thiểu 5 ký tự và chọn Chấp nhận, Bác bỏ hoặc Bỏ qua. Hệ thống lưu người quyết định, thời điểm, version và bình luận.",
            "Tài chính/Kiểm toán có thể xem để đưa quyết định nghiệp vụ ở màn hình chính; khuyến nghị không tự chuyển trạng thái PR/PO/invoice. Use case kết thúc.",
        ],
        "alternate": [
            "Người chỉ có quyền xem không thấy nút quét/quyết định và API từ chối thao tác thay đổi.",
            "Nếu lý do ngắn, trạng thái không hợp lệ, version cũ hoặc khuyến nghị không tồn tại, hệ thống không lưu quyết định.",
            "Khuyến nghị trùng fingerprint không tạo thêm bản ghi; quét lại có thể chỉ làm mới danh sách hiện có.",
        ],
        "special": "AI trong hệ thống hiện là bộ quy tắc kiểm soát có thể giải thích, không phải mô hình tự trị. Evidence phải hiển thị đơn vị tiền, ngày, ngưỡng và mã phiếu để người đọc hiểu.",
        "pre": "Đã đăng nhập; người dùng thuộc nhóm được xem AI, còn quyền vận hành chỉ dành cho ai_operator/dx_admin.",
        "post": "Khuyến nghị và quyết định được lưu audit; dữ liệu nghiệp vụ chỉ đổi nếu một vai trò nghiệp vụ thực hiện bước riêng.",
        "extension": "Từ một khuyến nghị mở chi tiết PR để xử lý luồng phê duyệt, sourcing, thanh toán hoặc kiểm toán tương ứng.",
        "evidence": "GET /api/v1/ai/recommendations; POST /generate; POST /{recommendationID}/decisions; backend: ai_recommendations.go; frontend: recommendation-center.",
    },
    {
        "id": "UC-23",
        "title": "Quản trị người dùng và phòng ban",
        "actors": "Quản trị DX-OS",
        "summary": "Use case này đồng bộ hồ sơ mà DX-OS cần để tính phạm vi dữ liệu, phòng ban, trạng thái truy cập và cảnh báo xung đột vai trò; việc cấp role nghiệp vụ vẫn do Keycloak quản lý.",
        "basic": [
            "Quản trị mở “Quản trị”. Hệ thống hiển thị tổng quan tổ chức: người dùng hoạt động/ngừng hoạt động, phòng ban hoạt động, phiếu mở, hàng đợi thông báo và số người có xung đột role.",
            "Quản trị chọn một người dùng và chỉnh tên hiển thị, email, phòng ban hoặc trạng thái hoạt động. Hệ thống kiểm tra email, phòng ban đang hoạt động, version và không cho tự vô hiệu hóa tài khoản đang đăng nhập.",
            "Quản trị tạo phòng ban với mã, tên, trung tâm chi phí, phòng ban cha và trạng thái hoạt động; hoặc chỉnh sửa phòng ban với kiểm tra không tạo vòng lặp cây.",
            "Nếu tắt phòng ban, hệ thống không cho thực hiện khi còn người dùng hoạt động thuộc phòng ban đó; mọi thay đổi được ghi resource audit.",
            "Màn hình hiển thị role snapshot và nhãn xung đột như finance đồng thời auditor, dx_admin đồng thời auditor hoặc ai_operator đồng thời auditor. Use case kết thúc.",
        ],
        "alternate": [
            "Người dùng/phòng ban không thuộc cùng tổ chức, mã trùng, version cũ hoặc dữ liệu sai thì không lưu.",
            "Quản trị không được tự tắt tài khoản của mình.",
            "Màn hình không cấp mới hoặc xóa role Keycloak; cần thao tác ở hệ thống định danh rồi đồng bộ snapshot.",
        ],
        "special": "Phạm vi dữ liệu nghiệp vụ dựa vào tổ chức/phòng ban trong database và role từ identity provider. Quản trị hồ sơ không đồng nghĩa với có quyền duyệt phiếu hay sửa dữ liệu tài chính.",
        "pre": "Đã đăng nhập với dx_admin; tổ chức hiện tại tồn tại.",
        "post": "Hồ sơ/phòng ban được cập nhật nhất quán; role conflict được hiển thị để nhóm vận hành xử lý ở Keycloak.",
        "extension": "Sau khi thay đổi phòng ban, đăng nhập lại vai trò liên quan để kiểm tra phạm vi xem phiếu, hàng đợi phê duyệt và báo cáo.",
        "evidence": "GET /api/v1/admin/center; PATCH /api/v1/admin/users/{userID}; POST/PATCH /api/v1/admin/departments; backend: admin.go; frontend: admin-center.",
    },
]


def add_cover(doc: Document) -> None:
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DX-OS LAB")
    set_run_font(run, name="Aptos Display", size=34, color=TEAL, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("BỔ SUNG VÀ CHUẨN HÓA USE CASE")
    set_run_font(run, name="Aptos Display", size=21, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Chương 2 – từ mục 2.2 trở đi, đối chiếu mã nguồn hiện tại")
    set_run_font(run, size=12, color=SLATE, italic=True)

    doc.add_paragraph()
    callout(
        doc,
        "Mục đích của file",
        "File này giữ nguyên Nhom_DACN.docx và trình bày bản cập nhật để nhóm xem thử trước khi ghép vào báo cáo. Các use case 01–10 của tài liệu gốc được chuẩn hóa trong bảng tổng quan; các use case từ 11 trở đi được mô tả chi tiết theo đúng form: mô tả, luồng cơ bản, luồng rẽ nhánh, yêu cầu đặc biệt, tiền điều kiện, hậu điều kiện và điểm mở rộng.",
        LIGHT_GREEN,
    )
    callout(
        doc,
        "Cách đọc",
        "Hãy đọc luồng liên kết ở mục 2.3 trước, sau đó đọc từng use case bổ sung. Các tên endpoint và file mã nguồn ở cuối mỗi use case là căn cứ đối chiếu, không phải bước người dùng phải tự gọi bằng tay.",
        LIGHT_BLUE,
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Phiên bản tài liệu: 1.0  •  Ngày tạo: {date.today().strftime('%d/%m/%Y')}  •  Mã nguồn đối chiếu: {source_snapshot()}")
    set_run_font(run, size=9.5, color=SLATE)
    doc.add_page_break()


def add_scope(doc: Document) -> None:
    heading(doc, "2.2. Tổng quan các Use Case cập nhật", 1)
    paragraph(doc, "Tài liệu gốc mô tả các vòng đời cốt lõi của phiếu mua sắm nhưng chưa bao quát các mô-đun được triển khai sau đó. Bảng dưới đây giữ các nghiệp vụ cũ và bổ sung những nghiệp vụ đã có route/API trong mã nguồn hiện tại.")
    callout(
        doc,
        "Phạm vi đối chiếu",
        f"Bản này được đối chiếu từ mã nguồn trong workspace hiện tại, commit HEAD {source_snapshot()}, cùng các route frontend và endpoint backend đang có. Các thay đổi giao diện cục bộ chưa làm thay đổi nghiệp vụ cũng không được coi là use case mới.",
        LIGHT_YELLOW,
    )
    heading(doc, "2.2.1. Actor và phạm vi trách nhiệm", 2)
    add_table(doc, ["Actor hiển thị", "Mã role", "Trách nhiệm thực tế trong phiên bản hiện tại"], ACTORS, [1.7, 1.8, 3.8])
    paragraph(doc, "Lưu ý: một tài khoản có thể mang nhiều role, nhưng các tổ hợp xung đột (ví dụ finance + auditor) bị hệ thống nhận diện và quyền ghi thường bị chặn. Vai trò nghiệp vụ được cấp tập trung trong Keycloak; DX-OS quản lý hồ sơ, phòng ban và trạng thái truy cập.")

    heading(doc, "2.2.2. Danh sách Use Case tổng hợp", 2)
    add_table(doc, ["Mã", "Tên Use Case", "Actor chính", "Phạm vi/cập nhật"], USE_CASE_SUMMARY, [0.8, 2.2, 1.8, 2.5])

    heading(doc, "2.2.3. Phân biệt phần kế thừa và phần bổ sung", 2)
    add_table(
        doc,
        ["Nhóm", "Use Case", "Tình trạng trong Nhom_DACN.docx", "Cách dùng file này"],
        [
            ["Luồng cốt lõi", "UC-01 đến UC-10", "Đã có mô tả chi tiết nhưng một số tên/actor chưa đồng nhất với code mới.", "Giữ nội dung nền; dùng ma trận và business rule cập nhật để sửa chỗ lệch."],
            ["Đã liệt kê nhưng thiếu chi tiết", "UC-11, UC-12", "Có trong bảng tóm tắt hoặc giao diện nhưng chưa có phần mô tả đầy đủ.", "Dùng phần 2.3.3.11–2.3.3.12 của file này."],
            ["Bổ sung sau cập nhật", "UC-13 đến UC-23", "Chưa có trong tài liệu gốc hoặc mới chỉ xuất hiện rời rạc ở chương khác.", "Dùng toàn bộ mô tả chi tiết và bảng truy vết ở cuối file."],
        ],
        [1.4, 1.4, 2.8, 1.7],
    )


def add_model(doc: Document) -> None:
    heading(doc, "2.3. Mô hình hóa chức năng cập nhật", 1)
    heading(doc, "2.3.1. Bản đồ luồng nghiệp vụ liên vai trò", 2)
    paragraph(doc, "Một nghiệp vụ hoàn chỉnh không dừng ở việc tạo phiếu. Các actor bàn giao cùng một mã PR, sau đó sinh PO, biên bản nhận, hóa đơn và dấu vết kiểm toán.")
    code_block(
        doc,
        "QUẢN TRỊ DX-OS: người dùng/phòng ban/quy tắc/chính sách\n"
        "          │\n"
        "NHÂN VIÊN hoặc TRƯỞNG BỘ PHẬN: tạo PR → bản nháp → kiểm tra trùng → gửi duyệt\n"
        "          │                                  ▲\n"
        "          ▼                                  │ yêu cầu chỉnh sửa\n"
        "TRƯỞNG BỘ PHẬN: duyệt tầng 1 / giữ ngân sách ─┘\n"
        "          │\n"
        "TÀI CHÍNH: duyệt tầng 2 / cam kết ngân sách\n"
        "          │\n"
        "TÀI CHÍNH: nhập & chấm điểm báo giá → chọn NCC → tạo PO → phát hành\n"
        "          │\n"
        "NHÂN VIÊN/TRƯỞNG BỘ PHẬN: nhận đủ/một phần/ngoại lệ → biên bản GR\n"
        "          │\n"
        "TÀI CHÍNH: ghi hóa đơn → đối soát → tranh chấp/xác minh → thanh toán\n"
        "          │\n"
        "KIỂM TOÁN: nhật ký → hồ sơ vụ việc → gói bằng chứng JSON\n"
        "ĐIỀU PHỐI AI: quét và giải thích cảnh báo; không tự đổi dữ liệu nghiệp vụ."
    )
    add_table(
        doc,
        ["Chặng", "Actor thực hiện", "Dữ liệu bàn giao", "Actor tiếp theo"],
        [
            ["1. Nhu cầu", "Nhân viên/Trưởng bộ phận", "PR + dòng hàng + chứng từ + mã PR", "Trưởng bộ phận"],
            ["2. Phê duyệt", "Trưởng bộ phận → Tài chính", "Trạng thái + budget reservation/commit + lý do", "Tài chính"],
            ["3. Nguồn cung", "Tài chính", "Supplier + các quote + điểm + quote được chọn", "Tài chính tạo PO"],
            ["4. Đặt hàng", "Tài chính", "PO liên kết PR và supplier", "Người yêu cầu nhận hàng"],
            ["5. Giao nhận", "Người yêu cầu/Quản lý", "GR theo dòng, số lượng, tình trạng", "Tài chính"],
            ["6. Hóa đơn", "Tài chính", "Invoice + match status + payments", "Kiểm toán"],
            ["7. Kiểm soát", "Kiểm toán/AI", "Audit event, case, evidence, recommendation", "Người xử lý nghiệp vụ"],
        ],
        [1.1, 1.7, 2.8, 1.7],
    )

    heading(doc, "2.3.2. Ma trận quyền theo chức năng", 2)
    add_table(
        doc,
        ["Chức năng", "Nhân viên", "Trưởng bộ phận", "Tài chính", "Kiểm toán", "Điều phối AI", "Quản trị"],
        [
            ["Tạo/sửa/gửi/hủy PR của mình", "Ghi", "Ghi", "Xem", "Xem", "Không", "Không mặc định"],
            ["Phê duyệt PR", "Không", "Tầng phòng ban", "Tầng tài chính", "Không", "Không", "Không mặc định"],
            ["Bình luận/timeline/tệp", "Phạm vi PR", "Phạm vi phòng ban", "Phạm vi tổ chức", "Đọc", "Không", "Đọc theo quyền"],
            ["Nhà cung cấp/báo giá", "Không", "Không", "Ghi", "Đọc", "Không", "Đọc sourcing"],
            ["Đơn hàng/giao nhận", "Nhận đơn của mình", "Nhận trong phạm vi", "Ghi PO", "Đọc", "Không", "Không mặc định"],
            ["Hóa đơn/thanh toán", "Không", "Không", "Ghi", "Đọc", "Không", "Không mặc định"],
            ["Ngân sách/báo cáo", "Không", "Không", "Ghi/xem", "Đọc", "Không", "Đọc báo cáo"],
            ["Ủy quyền/quy tắc", "Không", "Ủy quyền", "Đọc", "Đọc", "Không", "Ghi quy tắc"],
            ["Chính sách SLA/chứng từ", "Không", "Không", "Không", "Đọc", "Không", "Ghi"],
            ["Kiểm toán/gói bằng chứng", "Không", "Không", "Không", "Ghi/xuất", "Đọc", "Đọc"],
            ["Khuyến nghị", "Không", "Không", "Đọc", "Đọc", "Ghi quyết định", "Ghi quyết định"],
            ["Người dùng/phòng ban", "Không", "Không", "Không", "Không", "Không", "Ghi"],
        ],
        [1.45, 0.9, 1.05, 1.0, 0.9, 0.95, 0.9],
    )

    heading(doc, "2.3.3. Trạng thái và điểm bàn giao", 2)
    add_table(
        doc,
        ["Đối tượng", "Trạng thái/giá trị", "Ai tạo/chuyển", "Điều kiện chính"],
        [
            ["Phiếu mua sắm", "DRAFT → SUBMITTED", "Chủ phiếu", "Đủ dữ liệu và chứng từ bắt buộc."],
            ["Phiếu mua sắm", "SUBMITTED → MANAGER_APPROVED", "Trưởng bộ phận", "Đúng phòng ban, không tự duyệt; giữ ngân sách."],
            ["Phiếu mua sắm", "MANAGER_APPROVED → APPROVED", "Tài chính", "Đúng tổ chức; cam kết ngân sách."],
            ["Phiếu mua sắm", "→ CHANGES_REQUESTED/REJECTED/CANCELLED", "Vai trò tương ứng/chủ phiếu", "Yêu cầu chỉnh sửa/từ chối cần lý do; hủy chỉ DRAFT/CHANGES_REQUESTED."],
            ["Hồ sơ sourcing", "NOT_STARTED → OPEN → AWARDED", "Tài chính", "Phiếu APPROVED; chọn một quote hợp lệ."],
            ["Đơn hàng", "ORDERED → PARTIALLY_RECEIVED → RECEIVED", "Tài chính/Người nhận", "Nhận theo số lượng từng dòng."],
            ["Đơn hàng", "ORDERED/PARTIAL → RECEIPT_EXCEPTION", "Người nhận", "Hỏng, sai hàng hoặc từ chối nhận."],
            ["Hóa đơn", "RECORDED → VERIFIED/DISPUTED → PAID", "Tài chính", "Verify khi match; thanh toán không vượt số dư."],
            ["Khuyến nghị", "PENDING → APPROVED/REJECTED/DISMISSED", "Điều phối AI/Quản trị", "Lý do tối thiểu 5 ký tự; không tự đổi nghiệp vụ."],
            ["Hồ sơ kiểm toán", "OPEN → IN_REMEDIATION → RESOLVED/CLOSED", "Kiểm toán", "RESOLVED/CLOSED cần nội dung khắc phục."],
        ],
        [1.35, 2.0, 1.35, 2.35],
    )

    heading(doc, "2.3.4. Business Rules cập nhật cần ghi vào báo cáo", 2)
    add_table(
        doc,
        ["Mã", "Quy tắc", "Ý nghĩa khi kiểm thử"],
        [
            ["BR-01", "Phạm vi tổ chức/phòng ban", "Không thể dùng URL/API để xem hoặc sửa dữ liệu ngoài tổ chức; manager bị giới hạn phòng ban."],
            ["BR-02", "Optimistic Locking", "Update/approve/receive/pay dùng version; version cũ phải trả lỗi và không ghi đè."],
            ["BR-03", "Idempotency", "Gửi lặp create/transition/quote/PO/receipt/payment không sinh bản ghi trùng."],
            ["BR-04", "Chống tự phê duyệt", "Người yêu cầu không thể phê duyệt chính phiếu của mình; xung đột role cũng bị cảnh báo/chặn."],
            ["BR-05", "Ngân sách theo vòng đời", "Manager approve giữ chỗ; finance approve cam kết; từ chối/hủy phù hợp giải phóng."],
            ["BR-06", "Chứng từ theo ngưỡng", "Phiếu vượt ngưỡng phải có loại tài liệu được cấu hình trước khi gửi duyệt."],
            ["BR-07", "Sourcing bắt buộc ở ngưỡng VND", "PO từ 50.000.000 VND trở lên không có sourcing award bị từ chối."],
            ["BR-08", "Nhận hàng theo dòng", "Không nhận vượt số lượng đặt; COMPLETE chỉ hợp lệ khi tất cả dòng đã nhận đủ."],
            ["BR-09", "Đối soát hóa đơn hiện hành", "Chỉ verify khi PO RECEIVED, tiền tệ khớp và tổng hóa đơn không vượt tổng phiếu."],
            ["BR-10", "Audit và correlation", "Mọi thay đổi chính có actor, role, thời điểm, action, from/to và correlation ID."],
            ["BR-11", "Khuyến nghị chỉ hỗ trợ", "AI sinh evidence và lưu quyết định; không tự thay đổi PR/PO/invoice/budget."],
            ["BR-12", "Tách quyền", "Auditor đọc/xuất; Finance ghi nghiệp vụ tài chính; Admin ghi cấu hình; AI Operator ghi quyết định AI."],
        ],
        [0.8, 2.1, 4.15],
    )


def add_detail_case(doc: Document, case: dict, index: int) -> None:
    heading(doc, f"2.3.3.{index}. Mô tả use case {case['title']} ({case['id']})", 3)
    paragraph(doc, f"Tên use case: {case['title']}", style="List Paragraph", bold_prefix="Tên use case:")
    paragraph(doc, f"Actor chính: {case['actors']}", style="List Paragraph", bold_prefix="Actor chính:")
    paragraph(doc, f"Mô tả vắn tắt: {case['summary']}", bold_prefix="Mô tả vắn tắt:")
    heading(doc, "Các luồng sự kiện", 4)
    heading(doc, "3.1. Luồng cơ bản", 5)
    for item in case["basic"]:
        numbered(doc, item)
    heading(doc, "3.2. Luồng rẽ nhánh", 5)
    for item in case["alternate"]:
        bullet(doc, item)
    paragraph(doc, f"Các yêu cầu đặc biệt: {case['special']}", bold_prefix="Các yêu cầu đặc biệt:")
    paragraph(doc, f"Tiền điều kiện: {case['pre']}", bold_prefix="Tiền điều kiện:")
    paragraph(doc, f"Hậu điều kiện: {case['post']}", bold_prefix="Hậu điều kiện:")
    paragraph(doc, f"Điểm mở rộng: {case['extension']}", bold_prefix="Điểm mở rộng:")
    callout(doc, "Căn cứ mã nguồn", case["evidence"], LIGHT_BLUE)


def add_analysis(doc: Document) -> None:
    heading(doc, "2.4. Phân tích use case cập nhật", 1)
    paragraph(doc, "Phần này thay cho việc vẽ các sơ đồ rời rạc không còn khớp với chức năng mới. Bảng phân tích dưới đây mô tả đúng luồng giữa UI, API, cơ sở dữ liệu/sự kiện và actor tiếp theo; khi cần nộp bản UML, nhóm có thể chuyển từng dòng thành sequence diagram.")
    heading(doc, "2.4.1. Phân tích chuỗi PR → PO → nhận hàng → thanh toán", 2)
    add_table(
        doc,
        ["Bước", "Actor/UI", "API/service chính", "Dữ liệu và sự kiện", "Đầu ra"],
        [
            ["1", "Nhân viên/Manager – Tạo phiếu", "POST /purchase-requests", "purchase_requests + purchase_request_items; audit tạo phiếu", "PR DRAFT"],
            ["2", "Chủ phiếu – Gửi duyệt", "POST /purchase-requests/{id}/transitions", "Kiểm tra policy/tệp; process event SUBMITTED", "PR SUBMITTED"],
            ["3", "Manager – Duyệt tầng 1", "POST transitions", "Khóa PR + budget reservation + notification", "PR MANAGER_APPROVED"],
            ["4", "Finance – Duyệt tầng 2", "POST transitions", "Budget commit + audit + notification", "PR APPROVED"],
            ["5", "Finance – Nhập/chọn quote", "POST sourcing/quotes + POST selection", "sourcing_case + supplier_quotes + sourcing_events", "Case AWARDED"],
            ["6", "Finance – Tạo PO", "POST procurement-operations/orders", "purchase_orders + ORDER_PLACED + thông báo", "PO ORDERED"],
            ["7", "Requester/Manager – Ghi nhận GR", "POST orders/{requestID}/receipts", "purchase_order_receipts + receipt_items + event giao nhận", "PO RECEIVED/PARTIAL/EXCEPTION"],
            ["8", "Finance – Ghi/verify invoice", "POST invoices + POST transitions", "purchase_invoices + invoice_events; match status", "Invoice VERIFIED/DISPUTED"],
            ["9", "Finance – Thanh toán", "POST invoices/{id}/payments", "invoice_payments + paid_amount + event", "Invoice VERIFIED hoặc PAID"],
            ["10", "Auditor – Truy vết", "GET audit/events + audit/evidence", "timeline + attachments + PO + receipts + invoices + audit events", "Gói bằng chứng JSON"],
        ],
        [0.45, 1.65, 1.75, 2.3, 1.0],
    )

    heading(doc, "2.4.2. Phân tích các điểm kiểm soát xuyên suốt", 2)
    add_table(
        doc,
        ["Điểm kiểm soát", "Nằm ở đâu", "Nếu lỗi thì chuyện gì xảy ra", "Cách kiểm thử"],
        [
            ["Phạm vi dữ liệu", "ScopeFor/organization_id/department_id", "403 hoặc không trả dữ liệu ngoài phạm vi", "Đổi sang manager phòng ban khác và mở cùng PR."],
            ["Phiên bản", "ExpectedVersion ở PR, supplier, quote, PO, invoice, case", "409/xung đột; không ghi đè", "Mở hai cửa sổ, sửa cùng bản ghi rồi lưu cửa sổ thứ hai."],
            ["Gửi trùng", "Idempotency-Key trong create/transition/PO/receipt/payment", "Trả kết quả cũ hoặc 409 nếu dữ liệu khác", "Bấm gửi hai lần hoặc gửi lại cùng khóa."],
            ["Dấu vết", "process_events, resource audit, invoice/sourcing/audit events", "Giao dịch lỗi thì rollback; thành công có actor/action", "Sau mỗi bước mở timeline/audit và đối chiếu mã tương quan."],
            ["Thông báo", "Outbox notification queue", "Sự kiện vẫn nằm trong hàng đợi để xử lý; admin nhìn pending/dead", "Duyệt/đặt hàng/nhận hàng rồi kiểm tra Thông báo."],
            ["AI có kiểm soát", "Evidence + decision comment", "Không tự đổi PR/PO/invoice/budget", "Chấp nhận cảnh báo rồi kiểm tra trạng thái PR trước/sau."],
        ],
        [1.25, 1.7, 2.0, 2.0],
    )

    heading(doc, "2.4.3. Ánh xạ use case mới với màn hình và mã nguồn", 2)
    add_table(
        doc,
        ["Use Case", "Route/màn hình frontend", "Endpoint chính", "File backend/đối tượng"],
        [
            ["UC-11", "/work-center; /notifications", "GET /me/tasks-summary; /me/notifications", "store.go TaskSummary; notification service"],
            ["UC-12", "/approvals; /purchase-requests/{id}", "GET purchase-requests; POST transitions", "store.go Transition; approval-inbox"],
            ["UC-13", "/purchase-requests/new", "GET procurement-catalog; POST duplicate-check", "guided_requests.go; purchase-request-create"],
            ["UC-14", "/suppliers", "GET/POST/PATCH suppliers", "store.go supplier methods; Supplier model"],
            ["UC-15", "/sourcing", "GET sourcing; POST/PATCH quotes; POST selection", "sourcing.go; sourcing-board"],
            ["UC-16", "/operations", "POST/PATCH order; POST order transition", "store.go/operations_extensions.go"],
            ["UC-17", "/operations; chi tiết PR", "POST/GET orders/{requestID}/receipts", "operations_extensions.go; ReceiptRecord"],
            ["UC-18", "/invoices", "POST/PATCH invoices; transitions; payments", "invoices.go; payments.go"],
            ["UC-19", "/approval-governance", "GET governance; delegations; rules", "approval_governance.go"],
            ["UC-20", "/policies", "GET admin/policies; PATCH SLA/attachments", "policies.go"],
            ["UC-21", "/audit", "GET audit/events/cases/evidence", "audit_cases.go; reporting/audit.go"],
            ["UC-22", "/ai-center", "GET/POST ai/recommendations", "ai_recommendations.go"],
            ["UC-23", "/admin", "GET admin/center; PATCH user; POST/PATCH department", "admin.go"],
        ],
        [0.65, 1.55, 2.4, 2.35],
    )


def add_terms(doc: Document) -> None:
    heading(doc, "Phụ lục A. Thuật ngữ cần giải thích trong báo cáo", 1)
    add_table(
        doc,
        ["Thuật ngữ", "Cách hiểu dễ đọc", "Nó xuất hiện ở đâu trong DX-OS"],
        [
            ["SLA", "Mục tiêu thời gian xử lý một quy trình. Ví dụ: phiếu phải được xem xét trong 24 giờ.", "Chính sách SLA, hàng đợi công việc, khuyến nghị quá hạn."],
            ["Correlation ID", "Mã liên kết một yêu cầu qua API, audit và log để truy ngược cùng một lần thao tác.", "PR transition, PO, receipt, invoice, audit event."],
            ["Idempotency-Key", "Khóa chống gửi trùng. Gửi lại cùng thao tác và cùng dữ liệu không tạo thêm bản ghi.", "Tạo phiếu nghiệp vụ, quote, PO, receipt, payment, budget."],
            ["Optimistic Locking", "Kiểm tra version trước khi lưu; ai lưu sau phải biết dữ liệu đã bị người khác đổi.", "Sửa PR, supplier, quote, PO, invoice, rule và audit case."],
            ["Reservation", "Khoản ngân sách tạm giữ khi cấp quản lý đã duyệt nhưng chưa cam kết cuối.", "Manager approve và dashboard ngân sách."],
            ["Commitment", "Khoản ngân sách đã cam kết chính thức sau phê duyệt tài chính.", "Finance approve và budget dashboard."],
            ["Match status", "Kết quả đối chiếu hiện tại giữa trạng thái nhận hàng, tổng tiền và tiền tệ của hóa đơn.", "Bảng Hóa đơn và thanh toán."],
            ["Evidence", "Dữ liệu giải thích vì sao hệ thống đưa ra cảnh báo hoặc quyết định, không phải bằng chứng kết luận tự động.", "Trung tâm khuyến nghị và gói kiểm toán."],
            ["Correlation khác với Evidence", "Correlation là mã liên kết kỹ thuật; Evidence là dữ liệu nghiệp vụ dùng để giải thích.", "Audit/log so với AI recommendation."],
        ],
        [1.35, 3.0, 2.6],
    )
    heading(doc, "Phụ lục B. Checklist ghép vào Nhom_DACN.docx", 1)
    for item in (
        "Giữ Chương 1 của tài liệu gốc; thay bảng actor và ma trận chức năng bằng bản ở mục 2.3.1–2.3.2.",
        "Giữ mô tả chi tiết UC-01 đến UC-10 nhưng rà lại actor Finance/Auditor/Admin theo ma trận mới.",
        "Chèn các mô tả 2.3.3.11–2.3.3.23 sau use case 2.3.3.10 của tài liệu gốc.",
        "Bổ sung bảng trạng thái và business rule cập nhật trước phần 2.4 để các biểu đồ không mô tả nhầm trạng thái cũ.",
        "Khi vẽ UML, dùng mã PR → quote → PO → receipt → invoice → audit làm khóa liên kết; không tạo actor rời rạc cho mỗi màn hình.",
        "Nếu sử dụng thuật ngữ tiếng Anh trong báo cáo, thêm giải thích tiếng Việt ở Phụ lục A lần đầu xuất hiện.",
    ):
        bullet(doc, item)


def main() -> None:
    if not SOURCE_DOC.exists():
        raise FileNotFoundError(f"Không tìm thấy tài liệu gốc: {SOURCE_DOC}")
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_scope(doc)
    add_model(doc)
    heading(doc, "2.3.3. Mô tả chi tiết các Use Case bổ sung", 2)
    paragraph(doc, "Các mục dưới đây tiếp nối sau 2.3.3.10 của Nhom_DACN.docx. Nội dung được viết theo cùng biểu mẫu nhưng mô tả đúng các endpoint và trạng thái đang có trong mã nguồn.")
    for index, case in enumerate(DETAIL_CASES, start=11):
        add_detail_case(doc, case, index)
    add_analysis(doc)
    add_terms(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
