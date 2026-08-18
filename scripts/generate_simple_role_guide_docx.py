"""Generate a simple DX-OS guide that mirrors the /employee-guide page structure."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from generate_role_user_guide_docx import (
    BLUE,
    BODY_FONT,
    HEADING_FONT,
    NAVY,
    PALE_BLUE,
    PALE_GREEN,
    PALE_RED,
    PALE_YELLOW,
    TEXT,
    add_table,
    apply_font,
    bullet,
    callout,
    configure,
    heading,
    number,
)


ROOT = Path(__file__).resolve().parent.parent
OUTPUT = ROOT / "docs" / "generated" / "Huong_dan_DX_OS_don_gian_day_du_theo_vai_tro.docx"


ROLES = [
    {
        "name": "Employee",
        "vietnamese": "Nhân viên",
        "account": "employee.demo",
        "intro": "Tạo phiếu đúng, chuyển được sang Manager, bổ sung khi bị trả lại và xác nhận hàng thực tế.",
        "flow": [
            "Tạo và kiểm tra bản nháp",
            "Đính kèm tài liệu cần thiết",
            "Gửi phiếu cho Trưởng bộ phận",
            "Bổ sung nếu bị yêu cầu chỉnh sửa",
            "Theo dõi duyệt và đặt hàng",
            "Xác nhận hàng đã nhận",
        ],
        "sections": [
            ("01", "Tổng quan", "/dashboard", ["Xem ô Cần tôi bổ sung.", "Xem phiếu cập nhật gần đây.", "Bấm Tạo phiếu mới khi có nhu cầu."]),
            ("02", "Phiếu mua sắm", "/purchase-requests", ["Nhập tiêu đề, lý do, cost center và ít nhất một dòng hàng.", "Kiểm tra số lượng, đơn giá và tổng tiền.", "Tải báo giá nếu phiếu từ 20.000.000 VND.", "Lưu bản nháp rồi mở chi tiết để bấm Gửi phê duyệt."]),
            ("03", "Việc của tôi", "/work-center", ["Mở bản nháp chưa hoàn tất.", "Đọc lý do của phiếu bị yêu cầu chỉnh sửa.", "Sửa nội dung, bổ sung tệp và bấm Gửi lại."]),
            ("04", "Giao nhận", "/operations", ["Theo dõi mã đơn và ngày giao dự kiến.", "Nhập số lượng thực nhận theo từng dòng.", "Chọn nhận đủ, nhận một phần hoặc ghi ngoại lệ hỏng/sai hàng."]),
            ("05", "Thông báo", "/notifications", ["Mở thông báo chưa đọc.", "Đi thẳng tới phiếu liên quan.", "Đánh dấu đã đọc sau khi xử lý."]),
            ("06", "Hướng dẫn", "/employee-guide", ["Xem lại luồng nhanh.", "Đọc quyền được làm và không được làm.", "Kiểm tra trạng thái trước khi hỏi Manager."]),
        ],
        "statuses": [
            ("DRAFT", "Bản nháp", "Hoàn thiện nội dung, tải tệp và gửi phê duyệt."),
            ("SUBMITTED", "Đã gửi", "Chờ Manager duyệt; theo dõi và trao đổi trên phiếu."),
            ("CHANGES_REQUESTED", "Yêu cầu chỉnh sửa", "Đọc lý do, sửa, bổ sung tệp và bấm Gửi lại."),
            ("MANAGER_APPROVED", "Manager đã duyệt", "Chờ Finance kiểm tra và duyệt cuối."),
            ("APPROVED", "Đã phê duyệt", "Theo dõi đặt hàng và giao nhận."),
            ("REJECTED / CANCELLED", "Từ chối / Đã hủy", "Đọc lý do; phiếu đã kết thúc."),
        ],
        "can": ["Tạo, sửa và gửi phiếu của mình", "Tải lên và tải xuống tài liệu", "Trao đổi và xem Timeline", "Xác nhận hàng trong phạm vi được phép", "Xem thông báo"],
        "cannot": ["Phê duyệt thay Manager/Finance", "Điều chỉnh ngân sách", "Chọn nhà cung cấp hoặc phát hành đơn", "Xử lý hóa đơn/thanh toán", "Sửa chính sách hoặc audit"],
        "help_title": "Không thấy phiếu ở phía Manager?",
        "help": "Mở chi tiết phiếu, kiểm tra tệp bắt buộc và chắc chắn trạng thái đã chuyển từ Bản nháp sang Đã gửi.",
    },
    {
        "name": "Department Manager",
        "vietnamese": "Trưởng bộ phận",
        "account": "manager.demo",
        "intro": "Kiểm tra nhu cầu của phòng ban, phản hồi rõ ràng và chuyển phiếu hợp lệ sang Finance.",
        "flow": ["Mở hàng đợi", "Tìm phiếu Đã gửi", "Kiểm tra nhu cầu và chứng từ", "Phê duyệt hoặc phản hồi", "Theo dõi Finance", "Hỗ trợ nhận hàng"],
        "sections": [
            ("01", "Tổng quan", "/dashboard", ["Xem số phiếu đang chờ cấp bộ phận.", "Mở lối tắt Phê duyệt.", "Ưu tiên phiếu sắp quá SLA."]),
            ("02", "Phê duyệt", "/approvals", ["Tìm phiếu SUBMITTED.", "Kiểm tra lý do, dòng hàng, tổng tiền và báo giá.", "Chọn Phê duyệt, Yêu cầu chỉnh sửa hoặc Từ chối.", "Nhập nhận xét rõ ràng trước khi xác nhận."]),
            ("03", "Việc của tôi", "/work-center", ["Mở công việc ưu tiên.", "Xử lý phiếu sắp đến hạn.", "Kiểm tra phiếu đã rời hàng đợi sau khi quyết định."]),
            ("04", "Phiếu mua sắm", "/purchase-requests", ["Tìm kiếm theo mã hoặc tiêu đề.", "Đọc Timeline và trao đổi.", "Theo dõi trạng thái sau khi chuyển Finance."]),
            ("05", "Giao nhận", "/operations", ["Tìm đơn thuộc phòng ban.", "Đối chiếu từng dòng hàng.", "Xác nhận nhận đủ/một phần hoặc ghi ngoại lệ khi được giao nhận thực tế."]),
            ("06", "Thông báo", "/notifications", ["Xem việc mới được chuyển đến.", "Mở phiếu từ thông báo.", "Đánh dấu đã đọc sau khi xử lý."]),
        ],
        "statuses": [
            ("SUBMITTED", "Chờ Manager", "Kiểm tra và ra quyết định."),
            ("CHANGES_REQUESTED", "Đã trả lại", "Chờ Employee sửa và gửi lại."),
            ("MANAGER_APPROVED", "Đã duyệt cấp bộ phận", "Phiếu đã chuyển sang Finance."),
            ("REJECTED", "Đã từ chối", "Luồng kết thúc; lý do phải có trong Timeline."),
            ("ORDERED", "Đã đặt hàng", "Theo dõi ngày giao dự kiến."),
            ("PARTIALLY_RECEIVED", "Nhận một phần", "Theo dõi và ghi lần nhận tiếp theo."),
        ],
        "can": ["Duyệt phiếu cùng phòng ban", "Yêu cầu chỉnh sửa hoặc từ chối", "Xem chứng từ và Timeline", "Trao đổi với Employee/Finance", "Xác nhận nhận hàng trong phạm vi"],
        "cannot": ["Tự duyệt phiếu do mình tạo", "Duyệt phiếu phòng ban khác", "Điều chỉnh hạn mức ngân sách", "Phát hành đơn hàng", "Xác minh hoặc thanh toán hóa đơn"],
        "help_title": "Không thấy phiếu Employee vừa gửi?",
        "help": "Kiểm tra Employee đã bấm Gửi phê duyệt, phiếu đang SUBMITTED và hai tài khoản thuộc cùng phòng ban.",
    },
    {
        "name": "Finance",
        "vietnamese": "Tài chính",
        "account": "finance.demo",
        "intro": "Duyệt cuối, kiểm soát ngân sách, quản lý nhà cung cấp, đơn hàng, hóa đơn và thanh toán.",
        "flow": ["Duyệt cuối", "Kiểm tra ngân sách", "Chọn nhà cung cấp", "Phát hành đơn hàng", "Đối soát hóa đơn", "Thanh toán và báo cáo"],
        "sections": [
            ("01", "Tổng quan", "/dashboard", ["Xem phiếu chờ Finance.", "Xem cảnh báo ngân sách và công nợ.", "Mở nhanh Phê duyệt, Giao nhận hoặc Hóa đơn."]),
            ("02", "Phê duyệt", "/approvals", ["Mở phiếu MANAGER_APPROVED.", "Kiểm tra ngân sách và chứng từ.", "Phê duyệt, yêu cầu chỉnh sửa hoặc từ chối."]),
            ("03", "Ngân sách", "/budgets", ["Xem allocation, available, reserved và committed.", "Lọc theo cost center và tiền tệ.", "Điều chỉnh khi có lý do hợp lệ và kiểm tra lịch sử."]),
            ("04", "Nhà cung cấp", "/suppliers", ["Tạo hoặc sửa hồ sơ nhà cung cấp.", "Kiểm tra ngân hàng, hợp đồng, compliance và risk.", "Không dùng nhà cung cấp BLOCKED hoặc hết hạn."]),
            ("05", "Giao nhận", "/operations", ["Chọn phiếu APPROVED và phát hành order.", "Cập nhật lịch giao/tham chiếu trước khi có receipt.", "Hủy order chưa có receipt/hóa đơn và phải nhập lý do."]),
            ("06", "Hóa đơn", "/invoices", ["Tạo một hoặc nhiều hóa đơn cho order.", "Kiểm tra trạng thái đối soát.", "Xác minh hóa đơn hợp lệ.", "Ghi thanh toán một phần hoặc đủ; kiểm tra số dư còn lại."]),
            ("07", "Báo cáo", "/reports", ["Chọn khoảng ngày và bộ lọc.", "Đối chiếu KPI.", "Xuất CSV khi cần báo cáo."]),
            ("08", "Khuyến nghị", "/ai-center", ["Đọc cảnh báo SLA, giá trị lớn và rủi ro nhà cung cấp.", "Kiểm tra evidence.", "Chuyển vấn đề tới người có trách nhiệm."]),
        ],
        "statuses": [
            ("MANAGER_APPROVED", "Chờ Finance", "Kiểm tra và duyệt cuối."),
            ("APPROVED", "Đã phê duyệt", "Chọn nhà cung cấp và phát hành order."),
            ("ORDERED", "Chờ giao", "Theo dõi giao nhận; Finance không tự nhận hàng."),
            ("PARTIALLY_RECEIVED", "Nhận một phần", "Chờ Employee/Manager nhận phần còn lại."),
            ("VERIFIED", "Hóa đơn đã xác minh", "Được phép ghi nhận thanh toán."),
            ("PAID", "Đã thanh toán", "Chỉ đọc và đối chiếu lịch sử."),
        ],
        "can": ["Duyệt cuối", "Quản lý ngân sách và nhà cung cấp", "Phát hành/sửa/hủy order hợp lệ", "Tạo và đối soát nhiều hóa đơn", "Thanh toán nhiều đợt và xuất báo cáo"],
        "cannot": ["Tự xác nhận hàng đã nhận", "Thanh toán hóa đơn chưa VERIFIED", "Thanh toán vượt số dư", "Dùng nhà cung cấp bị chặn", "Xóa audit trail hoặc lịch sử thanh toán"],
        "help_title": "Không xác minh hoặc thanh toán được hóa đơn?",
        "help": "Kiểm tra receipt, match status, tiền tệ, số tiền, trạng thái VERIFIED và tải lại trang để lấy version mới nhất.",
    },
    {
        "name": "Auditor",
        "vietnamese": "Kiểm toán",
        "account": "auditor.demo",
        "intro": "Đọc toàn bộ chuỗi bằng chứng, lập hồ sơ phát hiện và theo dõi việc khắc phục.",
        "flow": ["Chọn phiếu cần kiểm tra", "Đọc Timeline", "Đối chiếu order/receipt", "Đối chiếu invoice/payment", "Xuất evidence", "Mở và đóng audit case"],
        "sections": [
            ("01", "Kiểm toán", "/audit", ["Tìm audit event và hồ sơ phát hiện.", "Tạo audit case với mức độ, owner và hạn xử lý.", "Cập nhật OPEN → IN_REMEDIATION → RESOLVED/CLOSED.", "Xuất evidence package cho phiếu cần kiểm tra."]),
            ("02", "Phiếu mua sắm", "/purchase-requests", ["Tìm theo mã phiếu.", "Đọc dòng hàng, attachment, comment và Timeline.", "Đối chiếu actor và thời gian."]),
            ("03", "Giao nhận và hóa đơn", "/operations và /invoices", ["Đọc order, receipt và ngoại lệ.", "Đọc hóa đơn, đối soát và payment history.", "Không thao tác ghi dữ liệu."]),
            ("04", "Ngân sách", "/budgets", ["Xem hạn mức và mức sử dụng.", "Đối chiếu reserved/committed.", "Đọc lịch sử điều chỉnh."]),
            ("05", "Chính sách", "/policies", ["Xem SLA và quy tắc chứng từ.", "Đối chiếu policy áp dụng tại thời điểm phiếu được gửi.", "Không sửa policy."]),
            ("06", "Báo cáo và khuyến nghị", "/reports và /ai-center", ["Lọc và xuất báo cáo.", "Đọc khuyến nghị/evidence ở chế độ chỉ đọc.", "Dùng dữ liệu để chọn mẫu kiểm toán."]),
        ],
        "statuses": [
            ("OPEN", "Mới mở", "Xác định phát hiện, owner và hạn xử lý."),
            ("IN_REMEDIATION", "Đang khắc phục", "Theo dõi tiến độ và bằng chứng bổ sung."),
            ("RESOLVED", "Đã xử lý", "Kiểm tra resolution và bằng chứng."),
            ("CLOSED", "Đã đóng", "Hồ sơ hoàn tất; giữ nguyên lịch sử."),
            ("HIGH / CRITICAL", "Mức độ cao", "Ưu tiên kiểm tra và theo dõi sát hạn."),
        ],
        "can": ["Đọc dữ liệu toàn tổ chức theo phạm vi", "Tạo và cập nhật audit case", "Xuất evidence package", "Xuất báo cáo", "Đọc khuyến nghị AI"],
        "cannot": ["Sửa phiếu mua sắm", "Duyệt phiếu", "Điều chỉnh ngân sách/policy", "Sửa order, supplier hoặc invoice", "Thanh toán hoặc quyết định AI recommendation"],
        "help_title": "Không xuất được evidence package?",
        "help": "Kiểm tra đang đăng nhập đúng Auditor, dùng đúng mã phiếu và phiếu thuộc cùng tổ chức. DX Admin chỉ xem case, không thay Auditor xuất evidence.",
    },
    {
        "name": "DX Admin",
        "vietnamese": "Quản trị hệ thống",
        "account": "admin.demo",
        "intro": "Quản lý hồ sơ người dùng, phòng ban, chính sách và theo dõi hoạt động toàn tổ chức.",
        "flow": ["Xem tổng quan quản trị", "Kiểm tra người dùng", "Cập nhật phòng ban", "Quản lý chính sách", "Theo dõi audit/report", "Xử lý khuyến nghị"],
        "sections": [
            ("01", "Quản trị", "/admin", ["Xem số người dùng/phòng ban và backlog thông báo.", "Sửa tên hiển thị, email, phòng ban và active của user.", "Tạo/sửa phòng ban và phòng ban cha.", "Không tự vô hiệu hóa tài khoản đang đăng nhập."]),
            ("02", "Chính sách", "/policies", ["Xem SLA và quy tắc chứng từ.", "Sửa giá trị có lý do.", "Tải lại để kiểm tra version tăng."]),
            ("03", "Kiểm toán", "/audit", ["Xem danh sách audit case.", "Theo dõi mức độ và trạng thái.", "Không tạo case hoặc xuất evidence thay Auditor."]),
            ("04", "Báo cáo", "/reports", ["Xem KPI toàn tổ chức.", "Lọc theo thời gian/phòng ban.", "Xuất dữ liệu khi được phép."]),
            ("05", "Khuyến nghị", "/ai-center", ["Chạy bộ quy tắc khuyến nghị.", "Đọc evidence.", "Chọn Approved/Rejected/Dismissed và nhập comment."]),
            ("06", "Keycloak", "http://localhost:8080", ["Cấp hoặc thu hồi realm role theo quy trình quản trị.", "Không sửa password/role trực tiếp trong database DX-OS.", "Kiểm tra lại session sau khi đổi role."]),
        ],
        "statuses": [
            ("ACTIVE", "Đang hoạt động", "Người dùng/phòng ban có thể tiếp tục sử dụng."),
            ("INACTIVE", "Ngừng hoạt động", "Chỉ vô hiệu hóa khi không còn ràng buộc đang sử dụng."),
            ("PENDING", "Khuyến nghị chờ quyết định", "Đọc evidence và ra quyết định có comment."),
            ("APPROVED / REJECTED", "Đã quyết định", "Kiểm tra người quyết định và audit trail."),
        ],
        "can": ["Quản lý hồ sơ user và phòng ban", "Quản lý SLA/quy tắc chứng từ", "Xem audit case và báo cáo", "Chạy và quyết định khuyến nghị", "Theo dõi backlog vận hành"],
        "cannot": ["Tự có quyền Manager/Finance", "Duyệt hoặc thanh toán nếu chưa được cấp role", "Tự vô hiệu hóa tài khoản đang dùng", "Tạo vòng lặp cây phòng ban", "Xóa lịch sử audit"],
        "help_title": "Đổi thông tin nhưng người dùng chưa thấy?",
        "help": "Tải lại trang. Nếu vừa đổi realm role trong Keycloak, yêu cầu người dùng đăng xuất và đăng nhập lại để token nhận quyền mới.",
    },
    {
        "name": "AI Operator",
        "vietnamese": "Điều phối khuyến nghị",
        "account": "ai.operator.demo",
        "intro": "Quét rủi ro có giải thích, đọc bằng chứng và ghi nhận quyết định của con người.",
        "flow": ["Mở trung tâm khuyến nghị", "Chạy quét dữ liệu", "Ưu tiên rủi ro cao", "Đọc evidence", "Ra quyết định", "Chuyển việc tới đúng vai trò"],
        "sections": [
            ("01", "Khuyến nghị", "/ai-center", ["Đọc Methodology trước khi dùng.", "Bấm tạo/quét khuyến nghị.", "Lọc PENDING hoặc HIGH/CRITICAL.", "Mở từng khuyến nghị để đọc evidence.", "Chọn Approved, Rejected hoặc Dismissed và nhập comment."]),
            ("02", "Tổng quan", "/dashboard", ["Xác nhận đúng role AI Operator.", "Xem tình trạng nền tảng.", "Mở nhanh trung tâm khuyến nghị."]),
            ("03", "Phối hợp xử lý", "Theo mã phiếu trong khuyến nghị", ["Ghi lại mã phiếu và loại rủi ro.", "Chuyển rủi ro SLA cho Manager/Finance.", "Chuyển supplier risk cho Finance/Auditor.", "Theo dõi kết quả qua trạng thái khuyến nghị."]),
        ],
        "statuses": [
            ("PENDING", "Chờ quyết định", "Đọc evidence và đánh giá."),
            ("APPROVED", "Chấp nhận khuyến nghị", "Chuyển việc tới vai trò nghiệp vụ phù hợp."),
            ("REJECTED", "Không chấp nhận", "Comment phải nêu lý do."),
            ("DISMISSED", "Bỏ qua", "Dùng khi không còn phù hợp hoặc trùng lặp."),
            ("HIGH / CRITICAL", "Rủi ro cao", "Ưu tiên đọc và phản hồi trước."),
        ],
        "can": ["Tạo bộ khuyến nghị theo luật", "Xem evidence và mức rủi ro", "Approved/Rejected/Dismissed có comment", "Xem lịch sử quyết định", "Chuyển thông tin tới người xử lý"],
        "cannot": ["Tự duyệt phiếu mua sắm", "Tự đặt hàng hoặc nhận hàng", "Tự sửa supplier/ngân sách", "Tự xác minh hoặc thanh toán hóa đơn", "Coi khuyến nghị là quyết định nghiệp vụ cuối cùng"],
        "help_title": "Không thấy khuyến nghị mới?",
        "help": "Bấm quét lại, kiểm tra dữ liệu có phiếu quá SLA, phiếu từ 50 triệu VND hoặc supplier rủi ro. Hệ thống không tạo khuyến nghị trùng cho cùng phiên bản dữ liệu.",
    },
]


def add_cover(document: Document) -> None:
    document.add_paragraph()
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("DX-OS")
    apply_font(run, HEADING_FONT, 28, NAVY, True)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("HƯỚNG DẪN ĐƠN GIẢN, ĐẦY ĐỦ THEO TỪNG VAI TRÒ")
    apply_font(run, HEADING_FONT, 17, BLUE, True)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Trình bày theo cùng form với trang /employee-guide")
    apply_font(run, BODY_FONT, 11, TEXT)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run(f"Cập nhật ngày {date.today().strftime('%d/%m/%Y')}")
    callout(document, "Cách đọc nhanh", "Chọn đúng vai trò của bạn. Làm theo Luồng nhanh, sau đó mở đúng khu vực trong bảng. Cuối trang luôn có danh sách Được làm, Không được làm và cách tự kiểm tra khi gặp lỗi.", PALE_BLUE)
    document.add_page_break()


def add_login(document: Document) -> None:
    heading(document, "1. Đăng nhập và chọn đúng vai trò", 1)
    for item in [
        "Mở http://localhost:4200.",
        "Đăng nhập bằng tài khoản được cấp. Có thể dùng tài khoản demo trong bảng dưới.",
        "Nhìn góc phải trên để kiểm tra username và nhãn vai trò.",
        "Khi test nhiều vai trò, dùng mỗi vai trò một cửa sổ ẩn danh riêng.",
    ]:
        number(document, item)
    add_table(document, ["Vai trò", "Tài khoản demo", "Mở phần nào trong tài liệu"], [[role["vietnamese"], role["account"], f"Hướng dẫn {role['name']}"] for role in ROLES], [4.7, 5.0, 7.5])
    callout(document, "Bảo mật", "Mật khẩu nằm trong data/runtime/*.txt của máy local. Không đưa password, token hoặc file .env vào DOCX, GitHub hay nhóm chat.", PALE_YELLOW)


def add_quick_flow(document: Document, role: dict) -> None:
    heading(document, "Luồng nhanh cần nhớ", 2)
    rows = []
    for index, item in enumerate(role["flow"], start=1):
        rows.append([str(index), item])
    add_table(document, ["Bước", "Việc cần làm"], rows, [2.0, 15.2])


def add_role(document: Document, role: dict, index: int) -> None:
    document.add_page_break()
    heading(document, f"{index}. Hướng dẫn công việc {role['name']} ({role['vietnamese']})", 1)
    document.add_paragraph(role["intro"])
    callout(document, "Tài khoản demo", role["account"], PALE_GREEN)
    add_quick_flow(document, role)

    heading(document, "Các phần cần dùng", 2)
    rows = []
    for number_code, title, route, tasks in role["sections"]:
        task_text = "• " + "\n• ".join(tasks)
        rows.append([number_code, f"{title}\n{route}", task_text])
    add_table(document, ["STT", "Mở ở đâu", "Làm gì tại đây"], rows, [1.4, 4.9, 10.9])

    heading(document, "Đọc trạng thái và biết phải làm gì", 2)
    add_table(document, ["Mã trạng thái", "Hiển thị", f"{role['name']} cần làm"], [list(item) for item in role["statuses"]], [4.1, 4.6, 8.5])

    heading(document, f"{role['name']} được làm và không được làm", 2)
    max_rows = max(len(role["can"]), len(role["cannot"]))
    permission_rows = []
    for row_index in range(max_rows):
        allowed = f"✓ {role['can'][row_index]}" if row_index < len(role["can"]) else ""
        denied = f"— {role['cannot'][row_index]}" if row_index < len(role["cannot"]) else ""
        permission_rows.append([allowed, denied])
    add_table(document, ["Có thể làm", "Không thực hiện"], permission_rows, [8.6, 8.6])
    callout(document, role["help_title"], role["help"], PALE_YELLOW)


def add_end_to_end(document: Document) -> None:
    document.add_page_break()
    heading(document, "8. Một luồng đầy đủ để cả nhóm làm theo", 1)
    add_table(
        document,
        ["Bước", "Ai làm", "Thao tác", "Kết quả"],
        [
            ["1", "Employee", "Tạo phiếu, thêm hàng/tài liệu và gửi.", "SUBMITTED"],
            ["2", "Manager", "Kiểm tra và phê duyệt.", "MANAGER_APPROVED"],
            ["3", "Finance", "Duyệt cuối.", "APPROVED"],
            ["4", "Finance", "Chọn supplier và phát hành order.", "ORDERED"],
            ["5", "Employee/Manager", "Nhận hàng theo từng dòng.", "PARTIALLY_RECEIVED hoặc RECEIVED"],
            ["6", "Finance", "Tạo/xác minh invoice.", "VERIFIED"],
            ["7", "Finance", "Thanh toán một hoặc nhiều đợt.", "PAID"],
            ["8", "Auditor", "Xuất evidence và mở case nếu cần.", "Audit trail đầy đủ"],
            ["9", "AI Operator/Admin", "Quét và xử lý khuyến nghị.", "Decision có comment/audit"],
        ],
        [1.2, 3.1, 8.4, 4.5],
    )
    callout(document, "Luồng đạt", "Mỗi bước đúng vai trò, trạng thái chuyển đúng, Employee nhận thông báo và Auditor đối chiếu được toàn bộ Timeline/order/receipt/invoice/payment.", PALE_GREEN)


def add_troubleshooting(document: Document) -> None:
    heading(document, "9. Khi gặp lỗi", 1)
    add_table(
        document,
        ["Hiện tượng", "Cách xử lý nhanh"],
        [
            ["Không vào được web", "Kiểm tra http://localhost:4200 và http://localhost:8081/health/ready."],
            ["401", "Đăng xuất rồi đăng nhập lại."],
            ["403", "Kiểm tra đúng role và phạm vi. Trong test phân quyền, 403 thường là đúng."],
            ["Phiếu không tới người tiếp theo", "Kiểm tra trạng thái và nút xác nhận cuối cùng đã được bấm."],
            ["Không xác minh được invoice", "Kiểm tra receipt, số tiền, tiền tệ và match status."],
            ["Không thanh toán được", "Invoice phải VERIFIED; số tiền không được vượt remaining amount."],
            ["Dữ liệu vừa sửa bị conflict", "Tải lại trang để lấy version mới rồi thao tác lại."],
        ],
        [5.3, 11.9],
    )
    callout(document, "Khi báo lỗi cho nhóm phát triển", "Gửi thời điểm, username (không gửi password), vai trò, mã phiếu, màn hình đang mở, thao tác vừa bấm và correlation ID nếu có.", PALE_RED)


def build() -> Path:
    document = Document()
    configure(document)
    add_cover(document)
    add_login(document)
    for index, role in enumerate(ROLES, start=2):
        add_role(document, role, index)
    add_end_to_end(document)
    add_troubleshooting(document)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(f"Created {build()}")
