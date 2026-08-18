"""Create the current DX-OS role-based user guide.

Run from the repository root:
    python scripts/generate_role_user_guide_docx.py
"""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
OUTPUT = ROOT / "docs" / "generated" / "Huong_dan_su_dung_DX_OS_theo_tung_vai_tro_de_doc.docx"

NAVY = "083B66"
BLUE = "0F6CBD"
TEAL = "008C95"
TEXT = "1F2937"
MUTED = "64748B"
WHITE = "FFFFFF"
PALE_BLUE = "EAF3FF"
PALE_GREEN = "E8F5E9"
PALE_YELLOW = "FFF7D6"
PALE_RED = "FDECEC"
BODY_FONT = "Arial"
HEADING_FONT = "Arial"


def apply_font(run, name: str, size: float, color: str = TEXT, bold: bool = False) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    r_fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:eastAsia"), name)


def set_cell_margins(cell, top: int = 90, start: int = 110, bottom: int = 90, end: int = 110) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = margins.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def shade(cell, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    element = OxmlElement("w:shd")
    element.set(qn("w:fill"), color)
    properties.append(element)


def set_cell(cell, value: str, *, bold: bool = False, color: str = TEXT) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(value)
    apply_font(run, BODY_FONT, 9.5, color, bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)


def repeat_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    element = OxmlElement("w:tblHeader")
    element.set(qn("w:val"), "true")
    properties.append(element)


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    header = table.rows[0]
    repeat_header(header)
    for index, value in enumerate(headers):
        shade(header.cells[index], BLUE)
        set_cell(header.cells[index], value, bold=True, color=WHITE)
        if widths:
            header.cells[index].width = Cm(widths[index])
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            set_cell(cells[index], value)
            if widths:
                cells[index].width = Cm(widths[index])
        if row_index % 2:
            for cell in cells:
                shade(cell, "F8FAFC")
    document.add_paragraph()
    return table


def heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_heading(text, level)
    paragraph.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    paragraph.paragraph_format.space_after = Pt(6)


def bullet(document: Document, text: str, level: int = 0) -> None:
    paragraph = document.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    apply_font(run, BODY_FONT, 11)


def number(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    apply_font(run, BODY_FONT, 11)


def callout(document: Document, title: str, body: str, fill: str = PALE_BLUE) -> None:
    table = document.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    shade(cell, fill)
    cell.text = ""
    first = cell.paragraphs[0]
    first.paragraph_format.space_after = Pt(2)
    run = first.add_run(title)
    run.bold = True
    apply_font(run, BODY_FONT, 10.5, NAVY, bold=True)
    paragraph = cell.add_paragraph(body)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.12
    for run in paragraph.runs:
        apply_font(run, BODY_FONT, 10.5)
    set_cell_margins(cell, top=130, start=150, bottom=130, end=150)
    document.add_paragraph()


def page_number(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instruction)
    run._r.append(end)


def configure(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    normal = document.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.line_spacing = 1.2
    normal.paragraph_format.space_after = Pt(7)
    normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    for name, size, color in (("Title", 26, NAVY), ("Heading 1", 18, NAVY), ("Heading 2", 14, BLUE), ("Heading 3", 12, NAVY)):
        style = document.styles[name]
        style.font.name = HEADING_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), HEADING_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), HEADING_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), HEADING_FONT)
    if "DX Tip" not in document.styles:
        style = document.styles.add_style("DX Tip", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = BODY_FONT
        style.font.size = Pt(10)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("DX-OS | HƯỚNG DẪN SỬ DỤNG THEO VAI TRÒ")
    apply_font(run, BODY_FONT, 8.5, MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("DX-OS Lab • Tài liệu hướng dẫn nội bộ • Trang ")
    apply_font(run, BODY_FONT, 8.5, MUTED)
    page_number(footer)


def cover(document: Document) -> None:
    document.add_paragraph()
    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("DX-OS")
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("HƯỚNG DẪN SỬ DỤNG THEO TỪNG VAI TRÒ")
    run.bold = True
    apply_font(run, HEADING_FONT, 18, BLUE, bold=True)
    document.add_paragraph()
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run("Phiên bản vận hành doanh nghiệp: procure-to-pay, kiểm toán, quản trị và khuyến nghị kiểm soát").italic = True
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run(f"Ngày cập nhật: {date.today().strftime('%d/%m/%Y')}")
    document.add_paragraph()
    callout(
        document,
        "Mục đích tài liệu",
        "Giúp từng người dùng biết mình được làm gì, thao tác theo thứ tự nào, dữ liệu nào phải kiểm tra và khi nào cần chuyển việc sang vai trò tiếp theo.",
        PALE_BLUE,
    )
    callout(
        document,
        "Nguyên tắc quan trọng",
        "Không chia sẻ password hoặc access token. Không sửa trực tiếp database để đổi trạng thái phiếu. Mỗi bước phải thực hiện trên giao diện/API để giữ timeline và audit trail.",
        PALE_YELLOW,
    )
    document.add_page_break()


def getting_started(document: Document) -> None:
    heading(document, "1. Bắt đầu sử dụng", 1)
    heading(document, "1.1 Địa chỉ hệ thống", 2)
    add_table(
        document,
        ["Thành phần", "Địa chỉ", "Dùng để làm gì"],
        [
            ["DX-OS Web", "http://localhost:4200", "Đăng nhập và thao tác nghiệp vụ hằng ngày."],
            ["Đăng nhập Keycloak", "http://localhost:8080", "Tự mở khi DX-OS chuyển sang màn hình đăng nhập."],
            ["API health", "http://localhost:8081/health/ready", "Kiểm tra nhanh API khi web báo không tải được dữ liệu."],
        ],
        [4.0, 5.2, 8.0],
    )
    heading(document, "1.2 Cách đăng nhập", 2)
    for item in [
        "Mở http://localhost:4200 bằng cửa sổ ẩn danh hoặc một profile trình duyệt riêng cho mỗi vai trò.",
        "Khi chuyển đến trang DX-OS LAB, nhập username và password của tài khoản được cấp.",
        "Sau khi vào web, nhìn góc phải trên để kiểm tra username và nhãn vai trò. Nếu không đúng, đăng xuất ngay.",
        "Mật khẩu local nằm trong data/runtime/*.txt và không được commit hoặc gửi qua chat. Nếu chưa có tài khoản demo, chạy scripts/Initialize-DevUser.ps1 theo hướng dẫn dự án.",
    ]:
        number(document, item)
    callout(document, "Mẹo demo", "Mở ít nhất bốn cửa sổ ẩn danh: Employee, Manager, Finance và Auditor. Nhờ vậy có thể chuyển vai trò mà không nhầm session đăng nhập.", PALE_GREEN)
    heading(document, "1.3 Vai trò và phạm vi", 2)
    add_table(
        document,
        ["Vai trò", "Tài khoản demo thường dùng", "Trách nhiệm chính"],
        [
            ["Employee", "employee.demo", "Tạo yêu cầu, bổ sung hồ sơ, nhận hàng và xem thông báo."],
            ["Department Manager", "manager.demo", "Duyệt cấp phòng ban, trả yêu cầu chỉnh sửa, hỗ trợ xác nhận nhận hàng."],
            ["Finance", "finance.demo", "Duyệt cuối, ngân sách, nhà cung cấp, đơn hàng, hóa đơn và thanh toán."],
            ["Auditor", "auditor.demo", "Kiểm tra dấu vết, lập hồ sơ phát hiện và xuất bằng chứng."],
            ["DX Admin", "admin.demo", "Quản trị người dùng/phòng ban, chính sách và giám sát vận hành."],
            ["AI Operator", "ai.operator.demo", "Chạy và quyết định khuyến nghị kiểm soát có giải thích."],
        ],
        [3.7, 4.4, 9.1],
    )


def flow(document: Document) -> None:
    document.add_page_break()
    heading(document, "2. Luồng tổng quát từ yêu cầu đến thanh toán", 1)
    document.add_paragraph("Một phiếu mua sắm đi theo chuỗi trách nhiệm dưới đây. Không có vai trò nào được bỏ qua bước kiểm soát của vai trò khác.")
    add_table(
        document,
        ["Bước", "Vai trò thực hiện", "Kết quả cần có"],
        [
            ["1", "Employee", "Tạo phiếu DRAFT, thêm dòng hàng/lý do/chứng từ, rồi gửi duyệt thành SUBMITTED."],
            ["2", "Department Manager", "Duyệt, yêu cầu chỉnh sửa hoặc từ chối. Duyệt thành MANAGER_APPROVED và giữ ngân sách."],
            ["3", "Finance", "Duyệt cuối thành APPROVED; ngân sách chuyển sang committed."],
            ["4", "Finance", "Chọn nhà cung cấp và phát hành purchase order (ORDERED)."],
            ["5", "Employee hoặc Manager", "Ghi nhận từng dòng hàng: đủ, một phần, hỏng/sai hoặc từ chối."],
            ["6", "Finance", "Tạo một hoặc nhiều hóa đơn; đối soát order – receipt – invoice."],
            ["7", "Finance", "Xác minh hóa đơn và ghi một hoặc nhiều lần thanh toán đến khi PAID."],
            ["8", "Auditor", "Đọc bằng chứng, lập audit case nếu phát hiện và theo dõi khắc phục."],
        ],
        [1.1, 4.0, 12.1],
    )
    callout(document, "Quy tắc tách nhiệm vụ", "Finance phát hành đơn và thanh toán nhưng không tự xác nhận hàng đã nhận. Auditor đọc/kiểm tra nhưng không sửa dữ liệu mua sắm. Manager không tự duyệt phiếu của mình.", PALE_RED)


def step_table(document: Document, rows: list[list[str]]) -> None:
    add_table(document, ["Bước", "Thao tác", "Kết quả cần kiểm tra"], rows, [1.0, 8.8, 7.4])


def employee(document: Document) -> None:
    document.add_page_break()
    heading(document, "3. Employee — Nhân viên yêu cầu và nhận hàng", 1)
    document.add_paragraph("Employee khởi tạo nhu cầu mua sắm, hoàn thiện hồ sơ khi bị trả về và xác nhận hàng thực tế đã nhận. Employee không có quyền phê duyệt, điều chỉnh ngân sách hay thanh toán.")
    heading(document, "3.1 Menu nên dùng", 2)
    for item in ["Tổng quan: xem số phiếu và việc cần chú ý.", "Phiếu mua sắm: tạo, tìm kiếm, lọc, mở và sửa phiếu của mình.", "Việc của tôi: xử lý DRAFT hoặc CHANGES_REQUESTED.", "Giao nhận: xác nhận tình trạng hàng cho order thuộc phạm vi của mình.", "Hướng dẫn và Thông báo: xem hướng dẫn nội bộ, phản hồi/nhắc việc và thông báo thanh toán."]:
        bullet(document, item)
    heading(document, "3.2 Tạo và gửi phiếu mua sắm", 2)
    step_table(document, [
        ["1", "Vào Phiếu mua sắm → Tạo phiếu mua sắm. Nhập tiêu đề và lý do mua sắm rõ ràng (ít nhất 10 ký tự).", "Phiếu ở trạng thái DRAFT; người yêu cầu và phòng ban được lấy từ tài khoản đăng nhập."],
        ["2", "Chọn cost center, tiền tệ; thêm từng dòng hàng: mô tả, đơn vị, số lượng, đơn giá.", "Tổng tiền phải đúng với nhu cầu; kiểm tra phần Tóm tắt bản nháp."],
        ["3", "Nếu phiếu đạt ngưỡng yêu cầu chứng từ, tải báo giá/tài liệu cần thiết trước khi gửi.", "Không còn cảnh báo thiếu attachment bắt buộc."],
        ["4", "Bấm Lưu bản nháp nếu chưa xong; khi đủ thông tin, mở phiếu và bấm Gửi duyệt.", "Trạng thái đổi thành SUBMITTED; phiếu được chuyển sang Manager cùng phòng."],
    ])
    heading(document, "3.3 Khi bị yêu cầu chỉnh sửa", 2)
    step_table(document, [
        ["1", "Mở Việc của tôi hoặc lọc Phiếu mua sắm theo CHANGES_REQUESTED.", "Thấy đúng mã phiếu cần bổ sung."],
        ["2", "Đọc Timeline và phần trao đổi để biết lý do trả về.", "Hiểu rõ nội dung cần sửa, ví dụ báo giá hoặc mô tả cấu hình."],
        ["3", "Chỉnh nội dung/dòng hàng/đính kèm, lưu lại rồi Gửi duyệt lại.", "Trạng thái quay về SUBMITTED; lịch sử cũ vẫn giữ nguyên."],
    ])
    heading(document, "3.4 Xác nhận nhận hàng theo từng dòng", 2)
    step_table(document, [
        ["1", "Vào Giao nhận, tìm order theo mã phiếu hoặc tên hàng; mở chi tiết order.", "Order đang ORDERED, PARTIALLY_RECEIVED hoặc RECEIPT_EXCEPTION."],
        ["2", "Đối chiếu hàng thực tế với từng dòng: số lượng nhận, tình trạng và ghi chú.", "Không nhập số lượng vượt số đã đặt hoặc vượt số còn thiếu."],
        ["3", "Chọn kết quả phù hợp: nhận đủ, nhận một phần, hỏng/sai hoặc từ chối; bấm ghi nhận.", "Order đổi PARTIALLY_RECEIVED, RECEIPT_EXCEPTION hoặc RECEIVED; lịch sử receipt xuất hiện."],
        ["4", "Mở Thông báo sau khi Finance thanh toán hóa đơn.", "Có thông báo cập nhật thanh toán/hóa đơn theo phiếu của mình."],
    ])
    callout(document, "Không làm", "Không xác nhận hàng khi chưa nhận thực tế; không dùng tài khoản Employee để thử duyệt, đặt hàng hoặc thanh toán. API/UI phải chặn các thao tác này.", PALE_RED)


def manager(document: Document) -> None:
    document.add_page_break()
    heading(document, "4. Department Manager — Trưởng bộ phận", 1)
    document.add_paragraph("Manager là cấp duyệt đầu tiên trong phòng ban. Trọng tâm là kiểm tra nhu cầu, hồ sơ, số tiền và phản hồi cụ thể cho Employee.")
    heading(document, "4.1 Duyệt phiếu cấp phòng ban", 2)
    step_table(document, [
        ["1", "Vào Phê duyệt hoặc Việc của tôi; lọc trạng thái SUBMITTED và tìm theo mã/tên phiếu.", "Chỉ thấy phiếu trong phòng ban được phép xem."],
        ["2", "Mở chi tiết; kiểm tra lý do mua, dòng hàng, tổng tiền, cost center, attachment, comment và Timeline.", "Thông tin đủ để ra quyết định; không có bất thường về hồ sơ."],
        ["3", "Chọn Phê duyệt khi phù hợp; nhập nhận xét nếu cần.", "Trạng thái thành MANAGER_APPROVED, ngân sách được reserve và Finance nhận việc."],
        ["4", "Chọn Yêu cầu chỉnh sửa khi thiếu thông tin; mô tả rõ cần bổ sung gì.", "Trạng thái thành CHANGES_REQUESTED; Employee nhận việc và xem được comment."],
        ["5", "Chọn Từ chối khi không còn nhu cầu hoặc sai chính sách; nêu lý do cụ thể.", "Trạng thái thành REJECTED; luồng dừng và có audit/timeline."],
    ])
    heading(document, "4.2 Hỗ trợ giao nhận", 2)
    document.add_paragraph("Manager cùng phòng có thể xác nhận hàng thay Employee khi được giao thực tế. Cách thao tác giống mục 3.4; cần đối chiếu từng dòng và ghi chú rõ ngoại lệ.")
    callout(document, "Kiểm soát", "Manager không tự duyệt phiếu do chính mình tạo và không điều chỉnh ngân sách hoặc thanh toán. Nếu không thấy nút là đúng phân quyền.", PALE_YELLOW)


def finance(document: Document) -> None:
    document.add_page_break()
    heading(document, "5. Finance — Tài chính và mua sắm", 1)
    document.add_paragraph("Finance hoàn tất duyệt cuối, điều phối nhà cung cấp/đơn hàng, kiểm soát hóa đơn–thanh toán, ngân sách và báo cáo. Đây là vai trò có ảnh hưởng tiền tệ lớn nhất.")
    heading(document, "5.1 Duyệt cuối và kiểm tra ngân sách", 2)
    step_table(document, [
        ["1", "Mở Phê duyệt, lọc MANAGER_APPROVED, kiểm tra hồ sơ và budget check.", "Phiếu có đủ tài liệu theo chính sách, số tiền và cost center hợp lý."],
        ["2", "Bấm Phê duyệt.", "Trạng thái APPROVED; reserved giảm và committed tăng theo ngân sách."],
        ["3", "Vào Ngân sách để xem allocation, available, reserved, committed và lịch sử điều chỉnh.", "Số liệu phản ánh đúng phiếu đã duyệt; thay đổi dùng expected version để tránh ghi đè."],
    ])
    heading(document, "5.2 Quản lý nhà cung cấp", 2)
    step_table(document, [
        ["1", "Vào Nhà cung cấp → tạo mới hoặc mở hồ sơ có sẵn.", "Có mã, tên, mã số thuế, liên hệ, trạng thái và mức rủi ro."],
        ["2", "Hoàn thiện địa chỉ, ngân hàng, số tài khoản, hợp đồng/hạn hợp đồng, compliance, điểm hiệu suất và ghi chú.", "Không chọn nhà cung cấp BLOCKED, compliance không đạt hoặc hợp đồng hết hạn cho order mới."],
        ["3", "Lưu thay đổi và kiểm tra audit/timeline khi cần.", "Dữ liệu được cập nhật theo version, không mất thay đổi của người khác."],
    ])
    heading(document, "5.3 Phát hành, sửa hoặc hủy purchase order", 2)
    step_table(document, [
        ["1", "Vào Giao nhận; chọn phiếu APPROVED chưa có order, chọn supplier, mã tham chiếu và ngày giao dự kiến.", "Order được phát hành với trạng thái ORDERED."],
        ["2", "Trước khi có receipt/hóa đơn, dùng thao tác sửa để cập nhật lịch giao, tham chiếu hoặc ghi chú khi cần.", "Version tăng, thay đổi có lịch sử."],
        ["3", "Chỉ hủy khi order chưa phát sinh receipt/hóa đơn; nhập lý do hủy rõ ràng.", "Order thành CANCELLED; không được xóa lịch sử."],
    ])
    heading(document, "5.4 Hóa đơn, đối soát và thanh toán nhiều đợt", 2)
    step_table(document, [
        ["1", "Vào Hóa đơn; chọn order đã phát hành và tạo hóa đơn với số hóa đơn, ngày, hạn trả, tiền tệ và số tiền.", "Có thể ghi nhiều hóa đơn trên cùng một order."],
        ["2", "Kiểm tra match status. Khi receipt chưa đủ, hóa đơn chỉ chờ đối soát; khi tổng hóa đơn khớp order/receipt, trạng thái MATCHED hoặc PARTIAL_MATCH phù hợp.", "Không xác minh khi chưa đáp ứng đối soát ba bên hoặc sai tiền tệ/số tiền."],
        ["3", "Bấm Xác minh để chuyển RECORDED thành VERIFIED.", "Hóa đơn đủ điều kiện thanh toán."],
        ["4", "Ghi nhận thanh toán: nhập số tiền dương, ngày không ở tương lai, mã tham chiếu ngân hàng và ghi chú. Có thể trả một phần.", "Hệ thống tăng paid amount, giảm remaining amount và lưu payment history."],
        ["5", "Ghi các đợt còn lại. Khi tổng paid amount bằng invoice amount, hóa đơn tự thành PAID.", "Hệ thống chặn số tiền vượt dư nợ; gửi lại cùng Idempotency-Key không tạo dòng thanh toán trùng."],
    ])
    heading(document, "5.5 Báo cáo và khuyến nghị", 2)
    for item in [
        "Báo cáo: chọn khoảng ngày/phòng ban/cost center/tiền tệ rồi xuất CSV khi cần đối chiếu.",
        "Khuyến nghị: Finance có thể xem rủi ro SLA, giá trị lớn và supplier risk để hỗ trợ ra quyết định, nhưng không thay thế quy trình phê duyệt.",
    ]:
        bullet(document, item)
    callout(document, "Không làm", "Finance không xác nhận receipt của chính order đang điều phối. Hãy để Employee hoặc Manager cùng phòng ghi nhận hàng thực tế.", PALE_RED)


def auditor(document: Document) -> None:
    document.add_page_break()
    heading(document, "6. Auditor — Kiểm toán", 1)
    document.add_paragraph("Auditor có quyền đọc rộng để kiểm tra, nhưng không thay đổi dữ liệu mua sắm, ngân sách, order, hóa đơn hoặc thanh toán. Riêng Audit Center cho phép lập và quản lý hồ sơ phát hiện.")
    heading(document, "6.1 Đối chiếu hồ sơ và bằng chứng", 2)
    step_table(document, [
        ["1", "Mở Phiếu mua sắm và tìm theo mã phiếu cần kiểm tra.", "Đọc được thông tin, dòng hàng, tài liệu, comment và Timeline."],
        ["2", "Mở Giao nhận và Hóa đơn để đối chiếu purchase order, receipt, invoice, payment reference và lịch sử thanh toán.", "Chuỗi thời gian logic: duyệt → order → receipt → verify → payment."],
        ["3", "Vào Kiểm toán; dùng xuất evidence package tại phiếu cần kiểm tra.", "Tải được JSON gồm phiếu, timeline, tài liệu, order, receipt, invoice và audit events."],
    ])
    heading(document, "6.2 Lập và theo dõi audit case", 2)
    step_table(document, [
        ["1", "Trong Kiểm toán, tạo Audit case; nhập tiêu đề, mô tả, mức độ, đối tượng liên quan, người phụ trách và hạn xử lý.", "Case mới ở OPEN; có mã case và audit event."],
        ["2", "Khi đơn vị xử lý, cập nhật IN_REMEDIATION; khi đủ bằng chứng, cập nhật RESOLVED hoặc CLOSED kèm resolution.", "Version tăng và lịch sử case được giữ lại."],
        ["3", "Mở Ngân sách, Chính sách, Báo cáo, Nhà cung cấp/Hóa đơn/Giao nhận để kiểm tra chéo ở chế độ đọc.", "Không có nút ghi dữ liệu hoặc thao tác ghi trả 403."],
    ])
    callout(document, "Quyền đặc biệt", "Auditor xem được AI recommendations nhưng không tạo/chạy/ra quyết định thay AI Operator hoặc DX Admin.", PALE_YELLOW)


def admin(document: Document) -> None:
    document.add_page_break()
    heading(document, "7. DX Admin — Quản trị DX-OS", 1)
    document.add_paragraph("DX Admin quản trị cấu trúc tổ chức và chính sách vận hành. Vai trò nghiệp vụ như Finance/Manager vẫn được Keycloak cấp riêng, không tự động có chỉ vì là Admin.")
    heading(document, "7.1 Quản lý người dùng và phòng ban", 2)
    step_table(document, [
        ["1", "Mở Quản trị để xem tổng quan tổ chức, số người dùng, số phòng ban và backlog thông báo.", "Xác nhận dữ liệu tổng quan tải thành công."],
        ["2", "Mở người dùng; chỉnh display name, email, phòng ban hoặc trạng thái active khi có quyết định quản trị hợp lệ.", "Không tự vô hiệu hóa chính tài khoản đang đăng nhập; thay đổi có version/audit."],
        ["3", "Mở phòng ban; tạo hoặc cập nhật mã, tên, phòng ban cha và trạng thái active.", "Không tạo vòng lặp cây phòng ban, không vô hiệu hóa phòng ban đang được người dùng sử dụng."],
        ["4", "Nếu cần đổi quyền nghiệp vụ, thực hiện tại Keycloak theo quy trình quản trị danh tính.", "DX-OS không cho sửa trực tiếp realm role từ màn hình quản trị."],
    ])
    heading(document, "7.2 Chính sách, giám sát và AI", 2)
    for item in [
        "Chính sách: quản lý SLA và quy tắc chứng từ bằng expected version; kiểm tra audit sau mỗi thay đổi.",
        "Kiểm toán/Báo cáo: xem audit case và báo cáo toàn tổ chức; DX Admin không xuất evidence package thay Auditor.",
        "Khuyến nghị: DX Admin có thể chạy bộ quy tắc và quyết định APPROVED/REJECTED/DISMISSED kèm nhận xét, nhưng quyết định không tự thay đổi phiếu hay thanh toán.",
    ]:
        bullet(document, item)


def ai_operator(document: Document) -> None:
    document.add_page_break()
    heading(document, "8. AI Operator — Điều phối khuyến nghị kiểm soát", 1)
    document.add_paragraph("AI Operator làm việc với trung tâm khuyến nghị có giải thích. Đây không phải chatbot tự thực thi; mọi thay đổi nghiệp vụ vẫn do người có thẩm quyền thực hiện qua quy trình chuẩn.")
    heading(document, "8.1 Chạy và quyết định khuyến nghị", 2)
    step_table(document, [
        ["1", "Mở Khuyến nghị; đọc phần Methodology để hiểu dữ liệu và giới hạn của hệ thống.", "Hiểu rõ khuyến nghị dùng luật kiểm soát, không phải quyết định tự động."],
        ["2", "Bấm tạo/quét khuyến nghị khi cần rà soát dữ liệu mới.", "Danh sách sinh các nhóm: rủi ro quá SLA, phiếu giá trị lớn, rủi ro nhà cung cấp."],
        ["3", "Mở từng khuyến nghị, đọc evidence (mã phiếu, trạng thái, giá trị, SLA hoặc supplier risk).", "Đủ căn cứ để đánh giá, không chỉ dựa vào nhãn rủi ro."],
        ["4", "Chọn APPROVED, REJECTED hoặc DISMISSED; nhập comment tối thiểu 5 ký tự.", "Quyết định có người thực hiện, thời điểm, version và audit trail."],
        ["5", "Chuyển thông tin cần xử lý cho Manager/Finance/Auditor qua quy trình phù hợp.", "AI không tự duyệt phiếu, không tự tạo order và không tự trả tiền."],
    ])
    callout(document, "Cách dùng đúng", "Dùng khuyến nghị để ưu tiên kiểm tra, không coi đây là bằng chứng duy nhất hoặc kết quả phê duyệt cuối cùng.", PALE_GREEN)


def scenarios(document: Document) -> None:
    document.add_page_break()
    heading(document, "9. Ba kịch bản nên thực hành", 1)
    heading(document, "9.1 Luồng chuẩn", 2)
    for item in [
        "Employee tạo phiếu, đính kèm báo giá nếu cần và gửi duyệt.",
        "Manager phê duyệt; Finance phê duyệt cuối và phát hành order.",
        "Employee/Manager nhận đủ hàng; Finance tạo hóa đơn, xác minh và thanh toán đủ một lần.",
        "Auditor kiểm tra timeline/evidence package; Employee kiểm tra thông báo thanh toán.",
    ]:
        number(document, item)
    heading(document, "9.2 Nhận hàng và thanh toán nhiều đợt", 2)
    for item in [
        "Finance phát hành order gồm ít nhất một dòng hàng.",
        "Employee nhận một phần, ghi đúng số lượng và note; order thành PARTIALLY_RECEIVED.",
        "Employee nhận phần còn lại; order thành RECEIVED.",
        "Finance tạo hóa đơn, xác minh, thanh toán một phần; kiểm tra remaining amount.",
        "Finance thanh toán phần còn lại; kiểm tra PAID và payment history có đúng số đợt.",
    ]:
        number(document, item)
    heading(document, "9.3 Phát hiện kiểm toán", 2)
    for item in [
        "Auditor mở một phiếu có chứng từ/hóa đơn để xuất evidence package.",
        "Auditor tạo audit case HIGH, gán owner và due date.",
        "Sau khi có biện pháp xử lý, Auditor cập nhật IN_REMEDIATION rồi RESOLVED kèm resolution.",
        "DX Admin/Auditor kiểm tra audit log và báo cáo; không sửa lịch sử mua sắm để khép case.",
    ]:
        number(document, item)


def troubleshooting(document: Document) -> None:
    document.add_page_break()
    heading(document, "10. Lỗi thường gặp và checklist bàn giao", 1)
    add_table(
        document,
        ["Hiện tượng", "Nguyên nhân thường gặp", "Cách xử lý"],
        [
            ["Không vào được web", "Web/API chưa chạy hoặc port bị chiếm.", "Mở localhost:4200; kiểm tra API health và docker compose ps."],
            ["Sai username/password", "Password demo đã được tạo lại.", "Mở đúng file data/runtime của tài khoản hoặc tạo lại user local."],
            ["401", "Token hết hạn hoặc session cũ.", "Đăng xuất, đóng session cũ và đăng nhập lại."],
            ["403", "Sai role hoặc vượt phạm vi phòng ban/tổ chức.", "Kiểm tra nhãn role. Với test phân quyền, 403 là kết quả đúng."],
            ["Không xác minh được hóa đơn", "Receipt chưa đủ hoặc đối soát số tiền/tiền tệ chưa khớp.", "Kiểm tra receipt, match status, amount và currency."],
            ["Không thanh toán được", "Hóa đơn chưa VERIFIED, version cũ hoặc số tiền vượt remaining amount.", "Tải lại hóa đơn, dùng version mới và nhập số tiền không vượt dư nợ."],
        ],
        [3.6, 5.8, 7.8],
    )
    heading(document, "10.1 Checklist trước khi demo/bàn giao", 2)
    for item in [
        "☐ Web localhost:4200 và API health/ready đều mở được.",
        "☐ Chuẩn bị các cửa sổ ẩn danh và credential cho Employee, Manager, Finance, Auditor, Admin, AI Operator.",
        "☐ Có ít nhất một phiếu demo để đi qua APPROVED → ORDERED → RECEIVED → VERIFIED → PAID.",
        "☐ Có một ví dụ nhận hàng một phần và một ví dụ payment nhiều đợt nếu cần demo chức năng mới.",
        "☐ Chụp/ghi mã phiếu, payment reference, audit case và evidence package làm bằng chứng báo cáo.",
        "☐ Không đưa password, access token hoặc file .env vào slide, DOCX, GitHub hay nhóm chat.",
    ]:
        document.add_paragraph(item)
    callout(document, "Tài liệu liên quan", "Xem docs/USER_GUIDE.md để tra cứu nhanh; docs/features/ENTERPRISE_OPERATIONS.md để hiểu quy tắc nghiệp vụ và scripts/Test-ProcurementWorkflow.ps1 để chạy smoke test toàn luồng.", PALE_BLUE)


def build() -> Path:
    document = Document()
    configure(document)
    cover(document)
    getting_started(document)
    flow(document)
    employee(document)
    manager(document)
    finance(document)
    auditor(document)
    admin(document)
    ai_operator(document)
    scenarios(document)
    troubleshooting(document)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(f"Created {build()}")
