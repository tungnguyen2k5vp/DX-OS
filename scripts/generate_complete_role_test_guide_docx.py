from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "generated" / "Huong_dan_kiem_thu_toan_bo_DX_OS_theo_6_vai_tro.docx"

NAVY = "123047"
TEAL = "007C83"
LIGHT_TEAL = "E6F5F4"
LIGHT_BLUE = "EDF5FA"
LIGHT_YELLOW = "FFF6D8"
LIGHT_GREEN = "EAF7EE"
LIGHT_RED = "FDECEC"
SLATE = "526577"
WHITE = "FFFFFF"
GRID = "D8E2E8"


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell(cell, text: str, *, bold: bool = False, color: str = NAVY) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(color)


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    result = doc.add_table(rows=1, cols=len(headers))
    result.style = "Table Grid"
    result.autofit = False
    repeat_header(result.rows[0])
    for index, value in enumerate(headers):
        shade(result.rows[0].cells[index], TEAL)
        set_cell(result.rows[0].cells[index], value, bold=True, color=WHITE)
        if widths:
            result.rows[0].cells[index].width = Cm(widths[index])
    for row_index, values in enumerate(rows):
        cells = result.add_row().cells
        for index, value in enumerate(values):
            set_cell(cells[index], value)
            if widths:
                cells[index].width = Cm(widths[index])
        if row_index % 2:
            for cell in cells:
                shade(cell, "F7FAFC")
    doc.add_paragraph()
    return result


def heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)


def bullet(doc: Document, text: str, level: int = 0) -> None:
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text)


def numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text)


def callout(doc: Document, title: str, body: str, fill: str = LIGHT_BLUE) -> None:
    box = doc.add_table(rows=1, cols=1)
    cell = box.cell(0, 0)
    shade(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)
    p = cell.add_paragraph(body)
    p.paragraph_format.space_after = Pt(2)
    doc.add_paragraph()


def code(doc: Document, text: str) -> None:
    box = doc.add_table(rows=1, cols=1)
    cell = box.cell(0, 0)
    shade(cell, NAVY)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for index, line in enumerate(text.splitlines()):
        if index:
            p.add_run().add_break()
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor.from_string(WHITE)
    doc.add_paragraph()


def page_number(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def add_toc(doc: Document) -> None:
    heading(doc, "Mục lục", 1)
    p = doc.add_paragraph()
    run = p.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.extend([begin, instr, separate])
    p.add_run("Mở bằng Microsoft Word, nhấn Ctrl+A rồi F9 để cập nhật mục lục và số trang.")
    end_run = p.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)
    doc.add_page_break()


def configure(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].font.color.rgb = RGBColor.from_string(NAVY)
    styles["Normal"].paragraph_format.line_spacing = 1.12
    for name, size, color in (
        ("Title", 28, NAVY),
        ("Heading 1", 18, NAVY),
        ("Heading 2", 14, TEAL),
        ("Heading 3", 11.5, NAVY),
    ):
        styles[name].font.name = "Aptos Display"
        styles[name].font.size = Pt(size)
        styles[name].font.color.rgb = RGBColor.from_string(color)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("DX-OS LAB | HƯỚNG DẪN KIỂM THỬ TOÀN BỘ HỆ THỐNG")
    r.font.name = "Aptos"
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string(SLATE)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("DX-OS Lab • Tài liệu nội bộ • Trang ")
    r.font.name = "Aptos"
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string(SLATE)
    page_number(footer)


def cover(doc: Document) -> None:
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DX-OS LAB")
    r.bold = True
    r.font.name = "Aptos Display"
    r.font.size = Pt(34)
    r.font.color.rgb = RGBColor.from_string(TEAL)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("HƯỚNG DẪN KIỂM THỬ TOÀN BỘ CHỨC NĂNG THEO 6 VAI TRÒ")
    r.bold = True
    r.font.name = "Aptos Display"
    r.font.size = Pt(17)
    r.font.color.rgb = RGBColor.from_string(NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Từ tạo yêu cầu mua sắm đến phê duyệt, ngân sách, chọn nhà cung cấp, giao nhận, hóa đơn, kiểm toán và khuyến nghị kiểm soát.").italic = True
    doc.add_paragraph()
    callout(
        doc,
        "Dành cho người mới",
        "Mỗi ca kiểm thử đều ghi rõ: đăng nhập bằng tài khoản nào, chuẩn bị gì, bấm ở đâu, kết quả đúng phải thấy và bằng chứng cần chụp. Có thể dùng tài liệu này làm kịch bản demo khi báo cáo dự án.",
        LIGHT_GREEN,
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Phiên bản 3.0 • Ngày tạo {date.today().strftime('%d/%m/%Y')}")
    doc.add_page_break()


def test_case(
    doc: Document,
    code_id: str,
    title: str,
    account: str,
    preparation: str,
    steps: list[str],
    expected: str,
    evidence: str,
    negative: str = "",
) -> None:
    heading(doc, f"{code_id}. {title}", 3)
    rows = [
        ["Tài khoản", account],
        ["Chuẩn bị", preparation],
        ["Kết quả mong đợi", expected],
        ["Bằng chứng", evidence],
    ]
    if negative:
        rows.append(["Kiểm tra chặn sai quyền", negative])
    table(doc, ["Mục", "Nội dung"], rows, [3.4, 13.9])
    doc.add_paragraph("Các bước thực hiện:")
    for step in steps:
        numbered(doc, step)
    p = doc.add_paragraph()
    p.add_run("Kết luận:  ☐ Đạt    ☐ Không đạt    Ghi chú: __________________________________________").bold = True


def add_getting_started(doc: Document) -> None:
    heading(doc, "1. Chuẩn bị trước khi kiểm thử", 1)
    doc.add_paragraph("Không cần biết code để làm theo. Trước mỗi ca kiểm thử, hãy kiểm tra hệ thống, mở đúng địa chỉ và tách phiên đăng nhập của từng vai trò.")
    heading(doc, "1.1 Khởi động và kiểm tra dịch vụ", 2)
    code(doc, "docker compose --profile foundation --profile application --profile reporting up -d --build\ndocker compose --profile foundation --profile application --profile reporting ps")
    table(
        doc,
        ["Thành phần", "Địa chỉ", "Kết quả đúng"],
        [
            ["Ứng dụng DX-OS", "http://localhost:4200", "Tự chuyển đến trang đăng nhập Keycloak."],
            ["API", "http://localhost:8081/health/ready", "HTTP 200 và trạng thái ready."],
            ["Keycloak", "http://localhost:8080", "Trang đăng nhập mở được."],
            ["Metabase", "http://localhost:3000", "Trang báo cáo mở được; chờ thêm 1–3 phút ở lần khởi động đầu."],
            ["Nextcloud", "http://localhost:8082", "Trang lưu trữ tài liệu mở được nếu dùng chức năng này."],
        ],
        [3.3, 5.3, 8.7],
    )
    callout(doc, "Địa chỉ bắt buộc", "Dùng http://localhost:4200, không thay bằng 127.0.0.1:4200 vì Keycloak có thể từ chối redirect URI.", LIGHT_YELLOW)
    heading(doc, "1.2 Tài khoản demo", 2)
    doc.add_paragraph("Mật khẩu không in trong tài liệu. Mở file tương ứng trong data/runtime để lấy thông tin đăng nhập; tuyệt đối không đưa các file này lên GitHub.")
    table(
        doc,
        ["Vai trò", "Tên đăng nhập", "File chứa mật khẩu", "Phạm vi chính"],
        [
            ["Nhân viên", "employee.demo", "data/runtime/employee-demo.txt", "Tạo/sửa/gửi phiếu, phản hồi yêu cầu sửa, theo dõi và xác nhận nhận hàng."],
            ["Trưởng bộ phận", "manager.demo", "data/runtime/manager-demo.txt", "Duyệt theo phòng ban, duyệt hàng loạt, ủy quyền và theo dõi giao nhận."],
            ["Tài chính", "finance.demo", "data/runtime/finance-demo.txt", "Duyệt cuối, ngân sách, nhà cung cấp, báo giá, đơn hàng, hóa đơn và thanh toán."],
            ["Kiểm toán", "auditor.demo", "data/runtime/auditor-demo.txt", "Đọc dữ liệu, mở hồ sơ kiểm toán, kiểm tra bằng chứng và báo cáo."],
            ["Quản trị", "admin.demo", "data/runtime/admin-demo.txt", "Hồ sơ người dùng, phòng ban, chính sách, quy tắc duyệt và xung đột vai trò."],
            ["AI Operator", "ai.operator.demo", "data/runtime/ai-operator-demo.txt", "Tạo/đánh giá khuyến nghị kiểm soát; không tự thay đổi nghiệp vụ."],
        ],
        [2.6, 3.3, 5.0, 6.4],
    )
    bullet(doc, "Nên mở 4 cửa sổ ẩn danh riêng cho Employee, Manager, Finance và Auditor để tránh nhầm phiên.")
    bullet(doc, "Kiểm tra Employee và Manager cùng phòng ban tại Quản trị > Người dùng nghiệp vụ. Nếu khác phòng ban, Manager sẽ không thấy phiếu của Employee.")
    heading(doc, "1.3 Bộ dữ liệu test nên tạo", 2)
    table(
        doc,
        ["Bộ dữ liệu", "Giá trị", "Mục đích"],
        [
            ["Luồng nhanh", "10.000.000 VND", "Không bắt buộc báo giá; dùng đi nhanh hết quy trình."],
            ["Luồng chứng từ", "25.000.000 VND", "Kiểm tra bắt buộc đính kèm báo giá từ 20 triệu."],
            ["Luồng giá trị lớn", "60.000.000 VND", "Kiểm tra chọn nhà cung cấp từ ít nhất 2 báo giá trước khi phát hành đơn."],
            ["Luồng trùng lặp", "Cùng tiêu đề, số tiền, phòng ban", "Kiểm tra cảnh báo yêu cầu mua sắm có nguy cơ trùng."],
        ],
        [3.4, 4.7, 8.8],
    )
    callout(doc, "Quy tắc đặt tên", "Thêm ngày giờ vào tiêu đề, ví dụ “Mua màn hình demo 20-08 14h30”, để dễ tìm và không vô tình kích hoạt cảnh báo trùng lặp.", LIGHT_BLUE)


def add_overview(doc: Document) -> None:
    doc.add_page_break()
    heading(doc, "2. Hiểu luồng liên kết giữa các vai trò", 1)
    code(doc, "NHÂN VIÊN: Bản nháp → Gửi duyệt\nTRƯỞNG BỘ PHẬN: Duyệt / Yêu cầu sửa / Từ chối\nTÀI CHÍNH: Duyệt cuối → Giữ/cam kết ngân sách → Chọn nhà cung cấp → Đơn hàng\nNHÂN VIÊN: Xác nhận nhận hàng\nTÀI CHÍNH: Hóa đơn → Đối soát 3 bên → Thanh toán\nKIỂM TOÁN: Đọc hồ sơ và bằng chứng\nQUẢN TRỊ: Quản lý chính sách, quy tắc và quyền\nAI OPERATOR: Tạo khuyến nghị có bằng chứng, con người quyết định")
    heading(doc, "2.1 Bản đồ trạng thái", 2)
    table(
        doc,
        ["Mã trạng thái", "Tên dễ hiểu", "Ai xử lý tiếp"],
        [
            ["DRAFT", "Bản nháp", "Nhân viên hoàn thiện rồi gửi."],
            ["SUBMITTED", "Đã gửi", "Trưởng bộ phận cùng phòng ban."],
            ["CHANGES_REQUESTED", "Yêu cầu chỉnh sửa", "Quay về Nhân viên sửa và gửi lại."],
            ["MANAGER_APPROVED", "Trưởng bộ phận đã duyệt", "Tài chính duyệt cuối."],
            ["APPROVED", "Đã phê duyệt", "Tài chính chọn nhà cung cấp/phát hành đơn."],
            ["REJECTED / CANCELLED", "Bị từ chối / Đã hủy", "Luồng kết thúc."],
            ["AWAITING_ORDER / ORDERED", "Chờ đặt hàng / Đã đặt", "Tài chính và bộ phận nhận hàng."],
            ["PARTIALLY_RECEIVED / RECEIVED", "Nhận một phần / Đã nhận đủ", "Nhân viên xác nhận từng đợt."],
            ["RECEIPT_EXCEPTION", "Có sai lệch khi nhận", "Tài chính/đơn vị mua xử lý ngoại lệ."],
            ["RECORDED / VERIFIED / DISPUTED / PAID", "Hóa đơn đã ghi nhận / xác minh / tranh chấp / thanh toán", "Tài chính."],
        ],
        [4.6, 5.3, 7.0],
    )
    heading(doc, "2.2 Quyền quan trọng phải nhớ", 2)
    bullet(doc, "Nhân viên chỉ sửa phiếu của mình khi phiếu còn ở Bản nháp hoặc Yêu cầu chỉnh sửa.")
    bullet(doc, "Trưởng bộ phận chỉ thấy phiếu trong phạm vi phòng ban/ủy quyền và không được tự duyệt phiếu của chính mình.")
    bullet(doc, "Tài chính không được tự xác nhận đã nhận hàng thay bộ phận yêu cầu.")
    bullet(doc, "Kiểm toán là vai trò chỉ đọc đối với nghiệp vụ mua sắm; thao tác sửa phải không xuất hiện hoặc trả 403.")
    bullet(doc, "Quản trị không phải siêu người dùng nghiệp vụ: không dùng Admin để duyệt phiếu hay thanh toán.")
    bullet(doc, "Khuyến nghị chỉ hỗ trợ quyết định; không tự động duyệt, sửa ngân sách hay thanh toán.")


def add_employee(doc: Document) -> None:
    doc.add_page_break()
    heading(doc, "3. Kiểm thử vai trò Nhân viên (Employee)", 1)
    doc.add_paragraph("Mục tiêu: chứng minh Nhân viên tạo nhu cầu đúng biểu mẫu, cung cấp đủ chứng từ, theo dõi phản hồi và xác nhận hàng thực nhận.")
    cases = [
        ("E01", "Đăng nhập, Tổng quan và Hướng dẫn nhân viên", "employee.demo", "API ready; tài khoản hoạt động.", ["Mở ứng dụng và đăng nhập.", "Kiểm tra tên/nhãn vai trò ở góc phải.", "Mở Tổng quan, Việc của tôi, Thông báo và trang Hướng dẫn nhân viên.", "Thử truy cập trực tiếp /admin hoặc /budgets."], "Các trang dành cho Nhân viên mở được; trang quản trị/ngân sách bị điều hướng hoặc từ chối.", "Ảnh menu Nhân viên và trang Hướng dẫn.", "Không thấy nút Phê duyệt, quản trị ngân sách hoặc thanh toán."),
        ("E02", "Tạo và sửa bản nháp", "employee.demo", "Dùng bộ dữ liệu 10 triệu.", ["Vào Phiếu mua sắm > Tạo phiếu.", "Chọn hạng mục gợi ý hoặc nhập dòng hàng: mô tả, số lượng, đơn vị, đơn giá.", "Nhập tiêu đề, lý do tối thiểu 10 ký tự, trung tâm chi phí và tiền tệ.", "Lưu bản nháp, mở lại và chỉnh một trường."], "Tổng tiền tự tính đúng; người yêu cầu và phòng ban lấy từ tài khoản; trạng thái vẫn là Bản nháp.", "Ảnh chi tiết phiếu có mã PR và tổng tiền."),
        ("E03", "Gửi phiếu nhỏ dưới 20 triệu", "employee.demo", "Phiếu E02 hợp lệ.", ["Mở chi tiết phiếu.", "Kiểm tra dòng hàng và tổng tiền.", "Bấm Gửi duyệt và xác nhận.", "Mở Timeline/Thông báo."], "Trạng thái chuyển Đã gửi; Manager cùng phòng ban thấy công việc; có lịch sử và thông báo.", "Ảnh trạng thái Đã gửi và Timeline."),
        ("E04", "Bắt buộc báo giá từ 20 triệu", "employee.demo", "Tạo phiếu 25 triệu; có sẵn tệp PDF báo giá thử nghiệm.", ["Lưu phiếu 25 triệu rồi bấm Gửi khi chưa có tệp.", "Ghi nhận thông báo chặn.", "Tải tệp lên loại Báo giá.", "Gửi lại."], "Lần đầu bị chặn rõ lý do; sau khi có tệp hợp lệ thì gửi thành công.", "Ảnh cảnh báo và danh sách tệp đính kèm."),
        ("E05", "Cảnh báo yêu cầu có nguy cơ trùng", "employee.demo", "Đã có một phiếu; tạo phiếu mới cùng tiêu đề/nội dung/số tiền.", ["Nhập dữ liệu gần giống phiếu đã có.", "Quan sát phần kiểm tra trùng lặp.", "Mở phiếu được hệ thống gợi ý để đối chiếu.", "Chỉ tiếp tục khi đã xác nhận đây là nhu cầu khác."], "Hiển thị cảnh báo và phiếu tương tự; hệ thống không âm thầm tạo trùng.", "Ảnh cảnh báo, mã phiếu đối chiếu và quyết định của người dùng."),
        ("E06", "Nhận yêu cầu chỉnh sửa và gửi lại", "employee.demo", "Manager đã chọn Yêu cầu chỉnh sửa ở ca M03.", ["Mở Thông báo hoặc Việc của tôi.", "Mở phiếu Yêu cầu chỉnh sửa và đọc lý do.", "Sửa đúng nội dung/chứng từ được yêu cầu.", "Thêm bình luận phản hồi rồi Gửi lại."], "Phiếu quay lại Đã gửi; lý do cũ, thay đổi mới và bình luận còn trong Timeline.", "Ảnh lý do sửa, nội dung sau sửa và trạng thái mới."),
        ("E07", "Bình luận, thông báo và nhật ký", "employee.demo", "Có ít nhất một phiếu đang xử lý.", ["Thêm bình luận nghiệp vụ.", "Mở chuông Thông báo, đọc một mục.", "Quay lại chi tiết phiếu và kiểm tra Timeline."], "Bình luận hiển thị đúng người/thời gian; thông báo đã đọc thay đổi số đếm; Timeline không mất dữ liệu.", "Ảnh bình luận và thông báo."),
        ("E08", "Hủy phiếu đúng trạng thái", "employee.demo", "Tạo một bản nháp riêng để hủy.", ["Mở phiếu bản nháp của mình.", "Bấm Hủy và xác nhận.", "Thử chỉnh sửa/gửi lại sau khi hủy."], "Trạng thái Đã hủy; không thể tiếp tục quy trình.", "Ảnh trạng thái Đã hủy."),
        ("E09", "Xác nhận giao nhận từng phần và đủ", "employee.demo", "Finance đã phát hành đơn hàng từ ca F08.", ["Mở Giao nhận hoặc chi tiết đơn.", "Ghi nhận số lượng nhỏ hơn số đặt và lưu.", "Kiểm tra trạng thái Nhận một phần.", "Ghi nhận phần còn lại và xác nhận đủ.", "Nếu hàng lỗi, thử tạo ngoại lệ nhận hàng."], "Số lượng không vượt quá số đặt; trạng thái lần lượt Nhận một phần/Đã nhận đủ hoặc Có sai lệch; lịch sử đầy đủ.", "Ảnh từng trạng thái và số lượng thực nhận."),
        ("E10", "Kiểm tra riêng tư giữa người dùng", "employee.demo", "Có phiếu của một người dùng khác.", ["Dùng URL trực tiếp mở phiếu không thuộc phạm vi.", "Thử sửa hoặc hủy phiếu đó."], "Không xem/sửa được dữ liệu ngoài phạm vi; API trả 403/404 phù hợp.", "Ảnh thông báo từ chối, không chụp token hay mật khẩu."),
    ]
    for item in cases:
        test_case(doc, *item)


def add_manager(doc: Document) -> None:
    doc.add_page_break()
    heading(doc, "4. Kiểm thử vai trò Trưởng bộ phận (Manager)", 1)
    doc.add_paragraph("Mục tiêu: kiểm tra hàng đợi theo phòng ban, quyết định phê duyệt có lý do, duyệt hàng loạt và ủy quyền có thời hạn.")
    cases = [
        ("M01", "Kiểm tra hàng đợi đúng phòng ban", "manager.demo", "Employee đã gửi phiếu E03; hai tài khoản cùng phòng ban.", ["Đăng nhập Manager.", "Mở Phê duyệt hoặc Việc của tôi.", "Tìm mã PR của E03.", "Đối chiếu người yêu cầu, phòng ban, tổng tiền và hạn xử lý."], "Phiếu Đã gửi xuất hiện đúng một lần; phiếu phòng ban khác không xuất hiện.", "Ảnh hàng đợi có mã phiếu và phòng ban."),
        ("M02", "Phê duyệt cấp phòng ban", "manager.demo", "Phiếu hợp lệ ở trạng thái Đã gửi.", ["Mở chi tiết, đọc lý do, dòng hàng, ngân sách và tệp.", "Nhập lý do quyết định nếu form yêu cầu.", "Bấm Phê duyệt."], "Trạng thái thành Trưởng bộ phận đã duyệt; Finance nhận việc; ngân sách được giữ nếu áp dụng.", "Ảnh quyết định và Timeline."),
        ("M03", "Yêu cầu chỉnh sửa", "manager.demo", "Một phiếu riêng đang Đã gửi.", ["Mở phiếu.", "Bấm Yêu cầu chỉnh sửa.", "Nhập lý do cụ thể từ 5 ký tự, ví dụ thiếu thông số kỹ thuật.", "Xác nhận."], "Phiếu thành Yêu cầu chỉnh sửa và quay về Nhân viên; Manager không duyệt tiếp cho đến khi gửi lại.", "Ảnh lý do và trạng thái."),
        ("M04", "Từ chối phiếu", "manager.demo", "Một phiếu không phù hợp đang Đã gửi.", ["Mở phiếu.", "Bấm Từ chối.", "Nhập lý do rõ ràng và xác nhận."], "Phiếu thành Bị từ chối, luồng kết thúc; người yêu cầu nhận thông báo.", "Ảnh trạng thái và lý do."),
        ("M05", "Duyệt hàng loạt", "manager.demo", "Có từ 2 phiếu hợp lệ cùng phòng ban đang chờ.", ["Mở danh sách Phê duyệt.", "Chọn nhiều phiếu.", "Thực hiện duyệt hàng loạt.", "Mở từng phiếu kiểm tra kết quả."], "Chỉ phiếu đủ điều kiện được duyệt; kết quả từng phiếu rõ ràng; không duyệt nhầm phiếu ngoài phạm vi.", "Ảnh danh sách trước/sau và thông báo tổng hợp."),
        ("M06", "Tạo ủy quyền phê duyệt có thời hạn", "manager.demo", "Có tài khoản Manager thay thế hợp lệ.", ["Mở Ủy quyền và quy tắc.", "Chọn người nhận, ngày bắt đầu/kết thúc và phạm vi.", "Lưu, sau đó đăng nhập người nhận để kiểm tra hàng đợi.", "Quay lại và vô hiệu hóa ủy quyền sau test."], "Trong thời hạn, người nhận thấy đúng phạm vi; hết hạn/vô hiệu hóa thì mất quyền; lịch sử vẫn được lưu.", "Ảnh cấu hình, hàng đợi người nhận và trạng thái vô hiệu hóa."),
        ("M07", "Chặn tự phê duyệt và chặn quyền Tài chính", "manager.demo", "Có phiếu do chính Manager tạo hoặc URL thao tác Finance.", ["Thử duyệt phiếu do chính tài khoản tạo.", "Thử truy cập ngân sách, thanh toán hoặc API tương ứng."], "Tự phê duyệt bị chặn; Manager không thể sửa ngân sách, nhà cung cấp hay thanh toán.", "Ảnh lỗi 403/thông báo nghiệp vụ."),
        ("M08", "Theo dõi đặt hàng và giao nhận", "manager.demo", "Có phiếu đã duyệt và đơn hàng.", ["Mở Giao nhận.", "Đối chiếu nhà cung cấp, mã đơn, lịch giao và trạng thái.", "Nếu được phân quyền nhận hàng, ghi nhận đúng phạm vi; không sửa hóa đơn."], "Manager theo dõi được đơn trong phạm vi; không thực hiện nghiệp vụ Tài chính trái quyền.", "Ảnh bảng giao nhận căn cột và trạng thái."),
    ]
    for item in cases:
        test_case(doc, *item)


def add_finance(doc: Document) -> None:
    doc.add_page_break()
    heading(doc, "5. Kiểm thử vai trò Tài chính (Finance)", 1)
    doc.add_paragraph("Mục tiêu: kiểm soát ngân sách, phê duyệt cuối, nhà cung cấp, so sánh báo giá, đơn hàng, đối soát hóa đơn và thanh toán.")
    cases = [
        ("F01", "Phê duyệt cuối", "finance.demo", "Phiếu đã được Manager duyệt.", ["Mở Phê duyệt và tìm mã PR.", "Kiểm tra ngân sách, chứng từ và quyết định Manager.", "Bấm Phê duyệt cuối."], "Phiếu thành Đã phê duyệt; số tiền giữ chuyển sang cam kết theo quy tắc ngân sách.", "Ảnh trạng thái và số liệu ngân sách."),
        ("F02", "Từ chối hoặc yêu cầu sửa ở cấp Tài chính", "finance.demo", "Một phiếu riêng đã qua Manager.", ["Mở phiếu.", "Chọn hành động không phê duyệt và nhập lý do.", "Kiểm tra thông báo tới các bên."], "Trạng thái, lý do và tác động ngân sách đúng; tiền giữ được giải phóng khi luồng kết thúc.", "Ảnh Timeline và ngân sách trước/sau."),
        ("F03", "Kiểm soát hạn mức ngân sách", "finance.demo", "Có trung tâm chi phí và kỳ ngân sách.", ["Mở Ngân sách.", "Đối chiếu tổng hạn mức, khả dụng, đang giữ và đã cam kết.", "Thử duyệt phiếu vượt số khả dụng.", "Kiểm tra release khi phiếu bị hủy/từ chối."], "Không chi vượt hạn mức; reservation/commitment/release cân đối, không âm và không cộng trùng.", "Ảnh số liệu trước/sau cùng mã PR."),
        ("F04", "Điều chỉnh ngân sách an toàn", "finance.demo", "Ghi lại phiên bản/số liệu hiện tại.", ["Tạo điều chỉnh với lý do rõ ràng.", "Lưu và kiểm tra số mới.", "Mở nhật ký thay đổi.", "Thử gửi lại bằng phiên bản cũ nếu có thể."], "Điều chỉnh có người/thời gian/lý do; cập nhật đồng thời bằng phiên bản cũ bị chặn.", "Ảnh bản ghi điều chỉnh và audit."),
        ("F05", "Quản lý nhà cung cấp và tuân thủ", "finance.demo", "Chuẩn bị dữ liệu nhà cung cấp thử nghiệm không nhạy cảm.", ["Tạo/cập nhật hồ sơ nhà cung cấp.", "Nhập mã, tên, thông tin liên hệ, trạng thái tuân thủ và rủi ro.", "Lưu và kiểm tra điểm/hiệu suất.", "Kiểm tra cảnh báo khi chọn nhà cung cấp rủi ro."], "Không trùng mã; trạng thái tuân thủ và lịch sử thay đổi rõ; nhà cung cấp không đạt bị cảnh báo/chặn đúng quy tắc.", "Ảnh hồ sơ và lịch sử."),
        ("F06", "So sánh ít nhất hai báo giá", "finance.demo", "Phiếu 60 triệu đã Được phê duyệt; có 2 nhà cung cấp.", ["Mở So sánh báo giá của phiếu.", "Nhập ít nhất 2 báo giá gồm giá, thời gian giao, chất lượng và tuân thủ.", "Kiểm tra điểm tổng hợp.", "Chọn một nhà cung cấp và ghi lý do."], "Bảng so sánh thẳng cột; điểm minh bạch; chỉ một báo giá được trao; lưu người quyết định/thời gian.", "Ảnh bảng so sánh và nhà cung cấp được chọn."),
        ("F07", "Bắt buộc sourcing với phiếu từ 50 triệu", "finance.demo", "Phiếu 60 triệu chưa chọn báo giá.", ["Thử phát hành đơn hàng ngay.", "Ghi nhận thông báo chặn.", "Hoàn thành F06.", "Phát hành đơn lại với đúng nhà cung cấp được chọn."], "Lần đầu bị chặn; chỉ nhà cung cấp trúng lựa chọn mới được dùng phát hành đơn.", "Ảnh cảnh báo và đơn hàng hợp lệ."),
        ("F08", "Phát hành, cập nhật và hủy đơn hàng", "finance.demo", "Phiếu đã Được phê duyệt; sourcing hoàn tất nếu >=50 triệu.", ["Tạo đơn hàng, chọn nhà cung cấp, ngày giao và mã ngoài nếu có.", "Lưu và đối chiếu tổng tiền với phiếu.", "Cập nhật thông tin cho phép.", "Tạo một đơn thử riêng để kiểm tra hủy."], "Đơn có mã PO duy nhất; không vượt tổng phiếu; hủy đúng trạng thái và để lại audit.", "Ảnh mã PO và trạng thái."),
        ("F09", "Theo dõi giao trễ và ngoại lệ", "finance.demo", "Có đơn với ngày giao quá hạn hoặc receipt exception.", ["Mở Giao nhận.", "Kiểm tra Chờ đặt hàng, Đang giao, Giao trễ, Đã nhận.", "Mở đơn có ngoại lệ và ghi nhận phương án xử lý."], "Giao trễ được ưu tiên; ngoại lệ không bị coi là đã nhận đủ; quyết định có lịch sử.", "Ảnh KPI và dòng giao nhận."),
        ("F10", "Ghi nhận hóa đơn", "finance.demo", "Đơn đã phát hành; có số hóa đơn thử nghiệm.", ["Mở Hóa đơn & thanh toán.", "Tạo hóa đơn theo PO, nhập số, ngày, số tiền, hạn trả.", "Lưu và mở lại."], "Số hóa đơn không trùng; liên kết đúng PR/PO/nhà cung cấp; trạng thái Đã ghi nhận.", "Ảnh hóa đơn và liên kết chứng từ."),
        ("F11", "Đối soát ba bên", "finance.demo", "Có PO, biên nhận và hóa đơn.", ["Mở hóa đơn.", "Đối chiếu số lượng đặt, số lượng nhận và số tiền hóa đơn.", "Xác minh hóa đơn khớp.", "Tạo một hóa đơn lệch để kiểm tra tranh chấp."], "Hóa đơn chỉ Xác minh khi PR/PO/receipt/invoice khớp; sai lệch thành Tranh chấp và không thanh toán được.", "Ảnh kết quả đối soát đúng và sai."),
        ("F12", "Thanh toán từng phần và đầy đủ", "finance.demo", "Hóa đơn đã Xác minh.", ["Ghi nhận một khoản thanh toán nhỏ hơn số còn lại.", "Kiểm tra số dư và trạng thái.", "Ghi nhận phần còn lại.", "Thử nhập vượt số còn lại."], "Số dư giảm đúng; đủ tiền thành Đã thanh toán; vượt tiền bị chặn; có người/thời gian/tham chiếu.", "Ảnh lịch sử thanh toán và số dư."),
        ("F13", "Theo dõi hóa đơn đến hạn/quá hạn", "finance.demo", "Có hóa đơn chưa thanh toán với hạn trả khác nhau.", ["Mở hàng đợi đối soát/thanh toán.", "Kiểm tra sắp xếp ưu tiên quá hạn.", "Đối chiếu hạn trả và nhãn trạng thái."], "Hóa đơn quá hạn được đánh dấu rõ, không bị lẫn với chưa có hóa đơn.", "Ảnh bảng hàng đợi đầy đủ cột."),
        ("F14", "Báo cáo và Metabase", "finance.demo", "Stack reporting hoạt động.", ["Mở Báo cáo trong DX-OS.", "Đối chiếu KPI với các phiếu vừa tạo.", "Mở Metabase từ menu hoặc http://localhost:3000.", "Nếu chưa lên, kiểm tra container/log và chờ khởi tạo."], "Dashboard/Metabase mở được; số liệu không lẫn đơn vị/trạng thái; liên kết không trỏ sai cổng.", "Ảnh báo cáo và Metabase."),
    ]
    for item in cases:
        test_case(doc, *item)


def add_auditor(doc: Document) -> None:
    doc.add_page_break()
    heading(doc, "6. Kiểm thử vai trò Kiểm toán (Auditor)", 1)
    cases = [
        ("A01", "Kiểm tra phạm vi chỉ đọc", "auditor.demo", "Có dữ liệu PR, ngân sách, PO, hóa đơn.", ["Đăng nhập Auditor.", "Mở từng khu vực mua sắm, nhà cung cấp, ngân sách, hóa đơn và sourcing.", "Tìm nút tạo/sửa/xóa/phê duyệt."], "Xem được dữ liệu phục vụ kiểm toán; hành động thay đổi không xuất hiện.", "Ảnh menu và một trang chi tiết.", "Gọi thao tác sửa bằng URL/API phải trả 403; dữ liệu không đổi."),
        ("A02", "Đối chiếu dấu vết một giao dịch đầu-cuối", "auditor.demo", "Biết mã PR của luồng nhanh.", ["Mở PR và Timeline.", "Theo liên kết sang phê duyệt, ngân sách, PO, receipt, invoice và payment.", "Đối chiếu người thực hiện và thời điểm."], "Mỗi bước truy vết được về cùng mã PR; không có khoảng trống hoặc thay đổi không tên người.", "Ảnh chuỗi bằng chứng."),
        ("A03", "Mở hồ sơ kiểm toán", "auditor.demo", "Chọn một giao dịch có ngoại lệ hoặc giá trị lớn.", ["Vào Kiểm toán > Tạo hồ sơ.", "Ghi phạm vi, lý do, mức rủi ro và người phụ trách.", "Liên kết giao dịch/bằng chứng liên quan."], "Hồ sơ có mã, trạng thái, phạm vi và liên kết bằng chứng; không sửa dữ liệu nguồn.", "Ảnh hồ sơ audit."),
        ("A04", "Xuất gói bằng chứng", "auditor.demo", "Hồ sơ A03 có dữ liệu.", ["Mở hồ sơ.", "Chọn xuất/tải gói bằng chứng.", "Kiểm tra tệp tải về và danh sách nội dung."], "Gói có dữ liệu nhất quán, thời điểm tạo và tham chiếu PR/PO/invoice; không chứa mật khẩu/token.", "Ảnh màn hình và tên tệp xuất."),
        ("A05", "Kiểm tra chính sách và quy tắc duyệt", "auditor.demo", "Admin đã có chính sách và quy tắc.", ["Mở Chính sách.", "Đối chiếu ngưỡng báo giá, SLA và quy tắc phê duyệt.", "Thử sửa một giá trị."], "Đọc được phiên bản/chính sách; không lưu thay đổi được.", "Ảnh chính sách và phản hồi chặn sửa."),
        ("A06", "Kiểm tra phân tách nhiệm vụ", "auditor.demo", "Có dữ liệu role snapshot/conflict.", ["Mở báo cáo kiểm soát hoặc audit.", "Tìm trường hợp cùng người có vai trò xung đột.", "Đối chiếu với Admin center."], "Xung đột được chỉ ra rõ tài khoản, vai trò và mức rủi ro.", "Ảnh bản ghi xung đột."),
        ("A07", "Báo cáo và Metabase", "auditor.demo", "Metabase đang chạy.", ["Mở Báo cáo.", "Lọc theo khoảng thời gian/phòng ban nếu có.", "Mở Metabase và đối chiếu tổng số giao dịch."], "Auditor xem được báo cáo nhưng không sửa dữ liệu nguồn; KPI khớp dữ liệu nghiệp vụ.", "Ảnh bộ lọc và kết quả."),
        ("A08", "Đọc khuyến nghị và bằng chứng", "auditor.demo", "AI Operator đã tạo khuyến nghị.", ["Mở Khuyến nghị.", "Mở Bằng chứng có thể giải thích.", "Đối chiếu mã PR/số tiền/rủi ro với dữ liệu nguồn."], "Bằng chứng hiển thị tiếng Việt dễ hiểu và truy vết được; Auditor không có nút quyết định nếu không được cấp.", "Ảnh thẻ khuyến nghị và bằng chứng."),
    ]
    for item in cases:
        test_case(doc, *item)


def add_admin(doc: Document) -> None:
    doc.add_page_break()
    heading(doc, "7. Kiểm thử vai trò Quản trị DX-OS (Admin)", 1)
    cases = [
        ("D01", "Hồ sơ người dùng và phòng ban", "admin.demo", "Danh sách tài khoản demo đã đồng bộ.", ["Mở Quản trị.", "Tìm employee.demo và manager.demo.", "Đối chiếu tên hiển thị, email, phòng ban và trạng thái.", "Chọn Chỉnh sửa, đổi một trường không nguy hiểm rồi hoàn tác."], "Danh sách căn cột, không bị cắt nút; lưu có xác nhận và audit; hai tài khoản dùng test duyệt cùng phòng ban.", "Ảnh trước/sau và phòng ban."),
        ("D02", "Khóa và mở lại người dùng", "admin.demo", "Chọn tài khoản test không phải tài khoản đang dùng.", ["Đặt trạng thái không hoạt động.", "Thử đăng nhập tài khoản đó.", "Mở lại trạng thái và đăng nhập lại."], "Tài khoản bị khóa không truy cập; mở lại hoạt động; lịch sử thay đổi được lưu.", "Ảnh trạng thái và kết quả đăng nhập."),
        ("D03", "Ảnh chụp vai trò và xung đột quyền", "admin.demo", "Có role assignments trong Keycloak.", ["Mở phần vai trò/xung đột.", "Đồng bộ hoặc tải lại snapshot.", "Kiểm tra cặp quyền xung đột như yêu cầu–duyệt hoặc nhà cung cấp–thanh toán."], "Hiển thị tài khoản, vai trò, lý do xung đột; không âm thầm tự gỡ quyền.", "Ảnh báo cáo xung đột."),
        ("D04", "Cấu hình ngưỡng chứng từ và SLA", "admin.demo", "Ghi lại giá trị hiện tại để hoàn tác.", ["Mở Chính sách.", "Đổi ngưỡng báo giá hoặc SLA bằng giá trị thử nghiệm.", "Tạo phiếu kiểm tra tác động.", "Khôi phục giá trị ban đầu."], "Chính sách mới có hiệu lực, có phiên bản/người sửa/thời gian; hoàn tác thành công.", "Ảnh chính sách trước/sau và kết quả phiếu."),
        ("D05", "Tạo quy tắc phê duyệt", "admin.demo", "Chuẩn bị phòng ban, tiền tệ, khoảng tiền và mức duyệt.", ["Mở Ủy quyền và quy tắc hoặc Chính sách > Quy tắc duyệt.", "Tạo rule theo phòng ban, VND, min/max, số cấp Manager/Finance và ưu tiên.", "Tạo phiếu nằm trong khoảng để kiểm tra.", "Vô hiệu hóa rule test sau khi xong."], "Đúng rule ưu tiên được áp dụng; khoảng tiền không mâu thuẫn âm thầm; có audit.", "Ảnh rule và luồng phiếu áp dụng."),
        ("D06", "Kiểm tra chồng lấn quy tắc", "admin.demo", "Đã có ít nhất một rule.", ["Thử tạo rule trùng phạm vi và cùng ưu tiên.", "Đọc cảnh báo/validation.", "Điều chỉnh ưu tiên hoặc phạm vi rồi lưu."], "Hệ thống cảnh báo hoặc giải quyết ưu tiên rõ; không chọn rule ngẫu nhiên.", "Ảnh thông báo và rule sau sửa."),
        ("D07", "Nhật ký quản trị", "admin.demo", "Đã thực hiện D01–D06.", ["Mở Kiểm toán/audit log.", "Lọc theo admin.demo, thời gian và loại hành động.", "Đối chiếu giá trị trước/sau."], "Mọi thay đổi quản trị quan trọng có actor, thời gian, đối tượng và trước/sau.", "Ảnh bộ lọc và bản ghi."),
        ("D08", "Chặn Admin làm thay nghiệp vụ", "admin.demo", "Có phiếu chờ duyệt và hóa đơn.", ["Thử mở/phê duyệt phiếu bằng URL trực tiếp.", "Thử ghi thanh toán hoặc xác nhận nhận hàng."], "Admin không được dùng quyền quản trị để duyệt mua sắm, nhận hàng hay thanh toán nếu không có role nghiệp vụ tương ứng.", "Ảnh 403/thông báo từ chối."),
    ]
    for item in cases:
        test_case(doc, *item)


def add_ai(doc: Document) -> None:
    doc.add_page_break()
    heading(doc, "8. Kiểm thử vai trò AI Operator / Khuyến nghị kiểm soát", 1)
    callout(doc, "Hiểu đúng chức năng", "Đây là bộ quy tắc kiểm soát có thể giải thích, không phải chatbot tự ra quyết định. Nó đọc dấu hiệu rủi ro, tạo khuyến nghị và bằng chứng; người có quyền vẫn phải Chấp nhận, Bác bỏ hoặc Bỏ qua.", LIGHT_TEAL)
    cases = [
        ("AI01", "Tạo khuyến nghị từ dữ liệu hiện có", "ai.operator.demo", "Có dữ liệu giá trị lớn, quá SLA, nhà cung cấp rủi ro hoặc hóa đơn quá hạn.", ["Mở Khuyến nghị.", "Chạy/tải lại quá trình sinh khuyến nghị nếu giao diện có nút.", "Kiểm tra số lượng và loại cảnh báo."], "Sinh các loại phù hợp như SLA_BREACH, HIGH_VALUE_REVIEW, SUPPLIER_RISK, DUPLICATE_REQUEST_RISK, SPLIT_PURCHASE_RISK, PRICE_ANOMALY, PAYMENT_OVERDUE, SUPPLIER_MASTER_CHANGED hoặc ROLE_CONFLICT.", "Ảnh danh sách khuyến nghị."),
        ("AI02", "Bằng chứng có thể giải thích", "ai.operator.demo", "Có ít nhất một khuyến nghị.", ["Mở một thẻ.", "Mở Bằng chứng có thể giải thích.", "Đối chiếu số tiền, mã PR, nhà cung cấp hoặc tài khoản với màn hình nguồn."], "Hiển thị nhãn tiếng Việt dễ hiểu, giá trị được định dạng; không bắt người dùng đọc JSON thô.", "Ảnh bằng chứng và dữ liệu nguồn."),
        ("AI03", "Chấp nhận khuyến nghị", "ai.operator.demo", "Một khuyến nghị đang Chờ quyết định.", ["Nhập lý do tối thiểu 5 ký tự.", "Bấm Chấp nhận.", "Tải lại trang."], "Trạng thái đã quyết định, lưu người/thời gian/lý do; không tự động duyệt hay sửa dữ liệu nghiệp vụ.", "Ảnh quyết định và audit."),
        ("AI04", "Bác bỏ khuyến nghị", "ai.operator.demo", "Một khuyến nghị khác đang chờ.", ["Nhập lý do dựa trên bằng chứng.", "Bấm Bác bỏ.", "Kiểm tra lịch sử."], "Trạng thái Bị bác bỏ; lý do còn sau tải lại.", "Ảnh trạng thái và lý do."),
        ("AI05", "Bỏ qua khuyến nghị", "ai.operator.demo", "Một cảnh báo không còn phù hợp.", ["Nhập lý do.", "Bấm Bỏ qua.", "Lọc danh sách theo trạng thái."], "Khuyến nghị được đóng/bỏ qua nhưng không bị xóa khỏi lịch sử.", "Ảnh bộ lọc và lịch sử."),
        ("AI06", "Chặn quyết định thiếu lý do", "ai.operator.demo", "Một khuyến nghị đang chờ.", ["Để trống hoặc nhập dưới 5 ký tự.", "Bấm một hành động quyết định."], "Nút bị vô hiệu hóa hoặc API từ chối; trạng thái không đổi.", "Ảnh validation."),
        ("AI07", "Chặn AI Operator làm nghiệp vụ", "ai.operator.demo", "Có URL phiếu/ngân sách/hóa đơn.", ["Thử duyệt phiếu.", "Thử sửa ngân sách, nhà cung cấp hoặc thanh toán."], "AI Operator chỉ vận hành khuyến nghị; mọi mutation nghiệp vụ trái quyền bị chặn.", "Ảnh 403/thông báo."),
    ]
    for item in cases:
        test_case(doc, *item)


def add_cross_role(doc: Document) -> None:
    doc.add_page_break()
    heading(doc, "9. Kịch bản test liên vai trò bắt buộc", 1)
    scenarios = [
        ["X01", "Luồng nhanh 10 triệu", "Employee tạo/gửi → Manager duyệt → Finance duyệt → PO → Employee nhận hàng → Finance hóa đơn/thanh toán → Auditor đối chiếu.", "Kết thúc PAID, audit đầy đủ."],
        ["X02", "Luồng chứng từ 25 triệu", "Gửi không có báo giá bị chặn → tải PDF → gửi thành công.", "Rule chứng từ hoạt động."],
        ["X03", "Luồng 60 triệu", "Duyệt xong → nhập 2 báo giá → chọn nhà cung cấp → mới phát hành PO.", "Không thể bỏ qua sourcing."],
        ["X04", "Yêu cầu chỉnh sửa", "Manager yêu cầu sửa → Employee sửa/bình luận/gửi lại → Manager duyệt.", "Không mất lý do và Timeline."],
        ["X05", "Hủy/từ chối và hoàn ngân sách", "Hủy hoặc từ chối sau khi có reservation.", "Tiền giữ được release đúng một lần."],
        ["X06", "Giao nhận một phần", "Nhận một phần → hóa đơn vượt phần đã nhận bị tranh chấp → nhận đủ → xác minh lại.", "Đối soát 3 bên chính xác."],
        ["X07", "Ủy quyền", "Manager tạo ủy quyền → người nhận duyệt đúng phạm vi → vô hiệu hóa.", "Quyền có thời hạn và audit."],
        ["X08", "Kiểm soát rủi ro", "Tạo dữ liệu trùng/giá trị lớn/quá hạn → AI sinh khuyến nghị → Auditor kiểm tra.", "Bằng chứng giải thích được; không tự sửa dữ liệu."],
    ]
    table(doc, ["Mã", "Kịch bản", "Chuỗi thao tác", "Điều kiện đạt"], scenarios, [1.4, 3.7, 8.0, 4.2])
    heading(doc, "9.1 Phiếu ghi kết quả cho mỗi kịch bản", 2)
    table(
        doc,
        ["Thông tin", "Điền khi test"],
        [
            ["Mã kịch bản / mã PR", "____________________________________________"],
            ["Người thực hiện / thời gian", "____________________________________________"],
            ["Kết quả", "☐ Đạt   ☐ Không đạt   ☐ Bị chặn do dữ liệu"],
            ["Ảnh/video/log", "____________________________________________"],
            ["Mô tả lỗi và bước tái hiện", "____________________________________________"],
        ],
        [5.2, 12.1],
    )


def add_regression(doc: Document) -> None:
    doc.add_page_break()
    heading(doc, "10. Checklist ổn định, bảo mật và hồi quy", 1)
    checks = [
        ["Hệ thống", "API health ready; Keycloak, web, Metabase và database không restart liên tục.", "☐"],
        ["Đăng nhập", "Đăng xuất xong không dùng lại được phiên; tài khoản khóa không đăng nhập.", "☐"],
        ["Phân quyền", "Mỗi role chỉ thấy menu và thao tác đúng phạm vi; URL trực tiếp vẫn bị chặn ở API.", "☐"],
        ["Phòng ban", "Manager chỉ xử lý phiếu đúng phòng ban/ủy quyền.", "☐"],
        ["Tính toàn vẹn", "Tổng dòng hàng = tổng phiếu; receipt không vượt PO; payment không vượt invoice.", "☐"],
        ["Đồng thời", "Hai người sửa cùng dữ liệu không ghi đè âm thầm; phiên bản cũ bị từ chối.", "☐"],
        ["Idempotency", "Bấm hai lần/refresh không tạo hai quyết định, hai đơn hoặc hai thanh toán.", "☐"],
        ["Tệp", "Chặn loại/kích thước không hợp lệ; tên tệp an toàn; người sai quyền không tải được.", "☐"],
        ["Ngôn ngữ", "Nhãn người dùng là tiếng Việt; mã chuẩn như PR, PO, SLA, API, VND có thể giữ nguyên.", "☐"],
        ["Giao diện bảng", "Cột tiêu đề thẳng dữ liệu, nút thao tác không bị cắt; màn hình nhỏ có dạng thẻ/ưu tiên cột.", "☐"],
        ["Audit", "Mọi quyết định và thay đổi quan trọng có actor, thời gian, lý do, trước/sau.", "☐"],
        ["Mất điện/restart", "Sau docker restart, dữ liệu đã commit vẫn còn; giao dịch dở không tạo bản ghi nửa vời.", "☐"],
        ["Sao lưu", "Có bản sao lưu database/volume và đã thử quy trình khôi phục ở môi trường test.", "☐"],
    ]
    table(doc, ["Nhóm", "Điều kiện kiểm tra", "Đạt"], checks, [3.0, 13.1, 1.2])
    callout(doc, "Không test phá hoại trên dữ liệu thật", "Kiểm tra mất điện/khôi phục chỉ nên thực hiện trên dữ liệu test. Dùng docker stop/start có kiểm soát; không xóa volume để thử nghiệm.", LIGHT_RED)


def add_troubleshooting(doc: Document) -> None:
    doc.add_page_break()
    heading(doc, "11. Xử lý lỗi thường gặp", 1)
    table(
        doc,
        ["Hiện tượng", "Kiểm tra", "Cách xử lý an toàn"],
        [
            ["Không vào được web", "API ready, container frontend và cổng 4200.", "Khởi động đúng profile; xem docker compose ps/logs."],
            ["Đăng nhập quay vòng", "URL có phải localhost; Keycloak realm/client redirect URI.", "Mở lại http://localhost:4200 trong cửa sổ ẩn danh."],
            ["Manager không thấy phiếu", "Trạng thái có phải Đã gửi; Employee/Manager có cùng phòng ban; ủy quyền còn hiệu lực.", "Sửa hồ sơ phòng ban bằng Admin rồi đăng nhập lại."],
            ["Phiếu vẫn Bản nháp", "Người dùng mới chỉ bấm Lưu bản nháp.", "Mở chi tiết và bấm Gửi duyệt sau khi đủ dữ liệu/chứng từ."],
            ["Không phát hành được PO", "Phiếu đã Được phê duyệt chưa; >=50 triệu đã chọn báo giá chưa.", "Hoàn thành sourcing và dùng đúng nhà cung cấp trúng lựa chọn."],
            ["Không xác minh hóa đơn", "Đã nhận hàng chưa; số lượng/số tiền có khớp PO và receipt không.", "Sửa dữ liệu nguồn đúng nghiệp vụ hoặc chuyển Tranh chấp."],
            ["Metabase từ chối kết nối", "Container metabase có chạy/healthy; cổng 3000; database reporting.", "Chờ 1–3 phút, xem logs rồi restart riêng dịch vụ reporting nếu cần."],
            ["Bảng bị cắt nút/cột", "Mức zoom trình duyệt và độ rộng màn hình.", "Đưa zoom về 100%; ghi ảnh, tên trang và độ phân giải để sửa responsive."],
            ["403 Forbidden", "Role và phạm vi dữ liệu.", "Nếu đang test sai quyền thì đây là kết quả đúng; không nâng role chỉ để vượt lỗi."],
            ["Dữ liệu mất sau restart", "Volume/database, migration và giao dịch đã commit.", "Dừng test, sao lưu log/volume metadata; không chạy lệnh xóa volume."],
        ],
        [4.0, 6.2, 7.1],
    )
    heading(doc, "11.1 Thông tin cần gửi khi báo lỗi", 2)
    bullet(doc, "Tên role và username (không gửi password/token).")
    bullet(doc, "URL trang, mã PR/PO/invoice và thời điểm xảy ra.")
    bullet(doc, "Các bước tái hiện ngắn gọn, kết quả mong đợi và kết quả thực tế.")
    bullet(doc, "Ảnh toàn màn hình và log liên quan đã che dữ liệu nhạy cảm.")


def add_signoff(doc: Document) -> None:
    doc.add_page_break()
    heading(doc, "12. Biên bản chốt kiểm thử", 1)
    table(
        doc,
        ["Hạng mục", "Kết quả / người xác nhận"],
        [
            ["Nhân viên E01–E10", "☐ Đạt  ☐ Không đạt   Người test: __________________"],
            ["Trưởng bộ phận M01–M08", "☐ Đạt  ☐ Không đạt   Người test: __________________"],
            ["Tài chính F01–F14", "☐ Đạt  ☐ Không đạt   Người test: __________________"],
            ["Kiểm toán A01–A08", "☐ Đạt  ☐ Không đạt   Người test: __________________"],
            ["Quản trị D01–D08", "☐ Đạt  ☐ Không đạt   Người test: __________________"],
            ["AI Operator AI01–AI07", "☐ Đạt  ☐ Không đạt   Người test: __________________"],
            ["Liên vai trò X01–X08", "☐ Đạt  ☐ Không đạt   Người test: __________________"],
            ["Hồi quy và ổn định", "☐ Đạt  ☐ Không đạt   Người test: __________________"],
            ["Lỗi nghiêm trọng còn mở", "________________________________________________________"],
            ["Kết luận sẵn sàng demo", "☐ Có   ☐ Chưa   Ngày xác nhận: ____/____/________"],
        ],
        [6.0, 11.3],
    )
    callout(doc, "Tiêu chí sẵn sàng demo", "Không còn lỗi làm mất dữ liệu, sai phân quyền, duyệt/chi tiền trùng, vượt ngân sách hoặc làm đứt luồng chính. Lỗi giao diện nhỏ phải được ghi rõ và có phương án tránh trong buổi demo.", LIGHT_GREEN)


def build() -> Path:
    doc = Document()
    configure(doc)
    cover(doc)
    add_toc(doc)
    add_getting_started(doc)
    add_overview(doc)
    add_employee(doc)
    add_manager(doc)
    add_finance(doc)
    add_auditor(doc)
    add_admin(doc)
    add_ai(doc)
    add_cross_role(doc)
    add_regression(doc)
    add_troubleshooting(doc)
    add_signoff(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = "Hướng dẫn kiểm thử toàn bộ DX-OS theo 6 vai trò"
    doc.core_properties.subject = "Kịch bản kiểm thử nghiệp vụ và phân quyền DX-OS"
    doc.core_properties.author = "DX-OS Lab"
    doc.core_properties.keywords = "DX-OS, kiểm thử, Employee, Manager, Finance, Auditor, Admin, AI Operator"
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
