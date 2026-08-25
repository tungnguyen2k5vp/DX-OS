from __future__ import annotations

"""Generate a copy-ready source-code flow guide for the DX-OS defense demo.

The document intentionally contains only a question objective and a copyable flow.
It avoids the explanatory labels used in the longer defense guide so it can be kept
open as a discreet navigation sheet during a live code demonstration.
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "generated" / "Luong_ma_nguon_demo_huy_phieu_va_tep_dinh_kem_DX_OS.docx"

NAVY = "18324A"
TEAL = "007C83"
TEAL_DARK = "005E63"
SLATE = "52677B"
MUTED = "6D7E8D"
WHITE = "FFFFFF"
PALE_TEAL = "E9F7F5"
PALE_BLUE = "F1F7FB"
PALE_GOLD = "FFF7E7"
GRID = "C9D9E3"


FLOWS = [
    (
        "01",
        "Hủy phiếu đúng quyền và đúng trạng thái",
        "Giảng viên hỏi: Ai được hủy phiếu và hệ thống kiểm tra ở đâu?",
        [
            "purchase-request-detail.html  [nút Hủy phiếu]",
            "→ purchase-request-detail.ts  availableActions()",
            "→ isPurchaseRequestOwner(request, auth.username())",
            "→ chỉ cho phép: owner + DRAFT | CHANGES_REQUESTED",
            "→ performTransition()",
            "→ ProcurementService.transition()",
            "→ POST /api/v1/purchase-requests/{requestID}/transitions",
            "→ transitionPurchaseRequest()",
            "→ Store.Transition()",
            "→ DecideTransition()",
            "→ CANCELLED",
        ],
        "frontend/src/app/features/procurement/pages/purchase-request-detail/purchase-request-detail.ts\nbackend/internal/procurement/model.go\nbackend/internal/procurement/store.go",
    ),
    (
        "02",
        "Chặn hủy sai người hoặc sai trạng thái",
        "Giảng viên hỏi: Nếu người khác cố hủy phiếu hoặc phiếu đã gửi duyệt thì sao?",
        [
            "Store.Transition()",
            "→ lockRequest() lấy trạng thái mới nhất của phiếu",
            "→ DecideTransition(actor, request, ActionCancel)",
            "→ actor.UserID ≠ request.RequesterID  → từ chối",
            "→ request.Status = SUBMITTED | MANAGER_APPROVED | APPROVED  → từ chối",
            "→ chỉ DRAFT | CHANGES_REQUESTED + đúng requester mới đi tiếp",
            "→ không có UPDATE status",
        ],
        "backend/internal/procurement/model.go  — DecideTransition()\nbackend/internal/procurement/store.go  — Transition()",
    ),
    (
        "03",
        "Chống bấm Hủy hai lần và chống ghi đè",
        "Giảng viên hỏi: Mạng chậm, bấm hai lần hoặc mở hai tab thì xử lý thế nào?",
        [
            "performTransition()",
            "→ gửi expectedVersion + Idempotency-Key",
            "→ Store.Transition()",
            "→ tra process_events bằng idempotency_key",
            "→ cùng request + cùng action  → trả lại kết quả cũ",
            "→ expectedVersion khác version hiện tại  → ErrVersionConflict",
            "→ không tạo thêm event CANCELLED",
        ],
        "frontend/src/app/features/procurement/pages/purchase-request-detail/purchase-request-detail.ts\nbackend/internal/procurement/store.go\nbackend/migrations/000002_procurement_mvp.sql",
    ),
    (
        "04",
        "Lưu trạng thái và lịch sử hủy phiếu",
        "Giảng viên hỏi: Hủy rồi dữ liệu mất hay còn lịch sử ở đâu?",
        [
            "Store.Transition()",
            "→ BEGIN database transaction",
            "→ UPDATE purchase_requests.status = CANCELLED",
            "→ INSERT process_events  [Timeline]",
            "→ insertAudit()",
            "→ INSERT audit_logs  [Trung tâm kiểm toán]",
            "→ COMMIT",
            "→ nếu có lỗi  → ROLLBACK",
        ],
        "backend/internal/procurement/store.go  — Transition(), insertAudit()\nbackend/migrations/000002_procurement_mvp.sql  — purchase_requests, process_events\nbackend/migrations/000003_procurement_audit.sql  — audit_logs",
    ),
    (
        "05",
        "Tải lên tệp đính kèm an toàn",
        "Giảng viên hỏi: File đi từ giao diện vào hệ thống như thế nào và kiểm tra gì?",
        [
            "purchase-request-detail.html  [Chọn tệp / Tải lên]",
            "→ uploadAttachment()",
            "→ ProcurementService.uploadAttachment()",
            "→ POST /api/v1/purchase-requests/{requestID}/attachments",
            "→ uploadPurchaseRequestAttachment()",
            "→ ValidateAttachment()",
            "→ PDF | DOCX | XLSX | JPG | PNG; tối đa 10 MB",
            "→ kiểm tra tên tệp + loại tệp + nội dung thực",
            "→ Store.UploadAttachment()",
        ],
        "frontend/src/app/features/procurement/pages/purchase-request-detail/purchase-request-detail.ts\nfrontend/src/app/features/procurement/data-access/procurement.service.ts\nbackend/internal/platform/httpapi/purchase_requests.go\nbackend/internal/procurement/model.go\nbackend/internal/procurement/attachments.go",
    ),
    (
        "06",
        "Lưu tệp vào Nextcloud và đánh dấu hợp lệ",
        "Giảng viên hỏi: File lưu ở đâu? Nếu Nextcloud lỗi giữa chừng thì sao?",
        [
            "Store.UploadAttachment()",
            "→ tính checksum SHA-256",
            "→ INSERT purchase_request_attachments  [status = UPLOADING]",
            "→ Nextcloud.Put(storagePath, content)",
            "→ UPDATE purchase_request_attachments  [status = ACTIVE]",
            "→ INSERT audit_logs  [ATTACHMENT_UPLOADED]",
            "→ nếu Put/finalize lỗi  → cleanupUploadingAttachment()",
            "→ không để tệp nửa chừng xuất hiện như đã thành công",
        ],
        "backend/internal/procurement/attachments.go  — UploadAttachment()\nbackend/internal/platform/documentstore/nextcloud.go  — Put()\nbackend/migrations/000006_purchase_request_attachments.sql",
    ),
    (
        "07",
        "Bắt buộc Báo giá với phiếu từ 20 triệu VND",
        "Giảng viên hỏi: Quy tắc 20 triệu được kiểm tra lúc nào và ở đâu?",
        [
            "người dùng bấm Gửi duyệt | Gửi duyệt lại",
            "→ Store.Transition()",
            "→ ActionSubmit | ActionResubmit",
            "→ requireAttachmentForSubmission()",
            "→ đọc attachment_rules theo organization + currency + threshold_amount",
            "→ kiểm tra có attachment ACTIVE loại QUOTATION hay chưa",
            "→ thiếu Báo giá  → ErrAttachmentRequired",
            "→ đủ Báo giá  → tiếp tục SUBMITTED",
        ],
        "backend/internal/procurement/store.go  — Transition()\nbackend/internal/procurement/attachments.go  — requireAttachmentForSubmission()\nbackend/migrations/000006_purchase_request_attachments.sql  — attachment_rules",
    ),
    (
        "08",
        "Tải xuống tệp và kiểm tra file nguyên vẹn",
        "Giảng viên hỏi: Khi tải xuống, hệ thống làm sao biết file không bị thay đổi?",
        [
            "purchase-request-detail.ts  → downloadAttachment()",
            "→ GET /api/v1/purchase-requests/{requestID}/attachments/{attachmentID}/content",
            "→ downloadPurchaseRequestAttachment()",
            "→ Store.DownloadAttachment()",
            "→ kiểm tra quyền xem phiếu",
            "→ Nextcloud.Get(storagePath)",
            "→ SHA-256(file tải về) = checksum đã lưu",
            "→ trả Blob cho trình duyệt tải file",
        ],
        "frontend/src/app/features/procurement/pages/purchase-request-detail/purchase-request-detail.ts\nbackend/internal/platform/httpapi/purchase_requests.go\nbackend/internal/procurement/attachments.go  — DownloadAttachment()\nbackend/internal/platform/documentstore/nextcloud.go  — Get()",
    ),
    (
        "09",
        "Xóa tệp có kiểm soát và vẫn có lịch sử",
        "Giảng viên hỏi: Xóa file có xóa luôn dấu vết không? Nếu kho file lỗi thì sao?",
        [
            "purchase-request-detail.ts  → deleteAttachment()",
            "→ DELETE /api/v1/purchase-requests/{requestID}/attachments/{attachmentID}",
            "→ Store.DeleteAttachment()",
            "→ owner + DRAFT | CHANGES_REQUESTED",
            "→ UPDATE attachment status = DELETING",
            "→ Nextcloud.Delete(storagePath)",
            "→ UPDATE attachment status = DELETED",
            "→ INSERT audit_logs  [ATTACHMENT_DELETED]",
            "→ storage lỗi  → trả status về ACTIVE",
        ],
        "backend/internal/procurement/attachments.go  — DeleteAttachment()\nbackend/internal/platform/documentstore/nextcloud.go  — Delete()\nbackend/migrations/000006_purchase_request_attachments.sql",
    ),
    (
        "10",
        "Chỉ ra kiểm thử mã nguồn",
        "Giảng viên hỏi: Em kiểm thử luật nghiệp vụ và an toàn tệp ở đâu?",
        [
            "backend/internal/procurement/model_test.go",
            "→ TestDecideTransition()",
            "→ TestValidateAttachmentAcceptsQuotationPDF()",
            "→ TestValidateAttachmentRejectsUnsafeFile()",
            "→ TestValidateAttachmentRejectsSpoofedContentType()",
            "→ TestValidateAttachmentRejectsMismatchedExtension()",
            "→ frontend purchase-request-ownership.spec.ts",
            "→ kiểm tra username chủ phiếu khác tên hiển thị vẫn hiện đúng nút thao tác",
        ],
        "backend/internal/procurement/model_test.go\nfrontend/src/app/features/procurement/utils/purchase-request-ownership.spec.ts",
    ),
]


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        tc_pr.append(node)
    node.set(qn("w:fill"), fill)


def borders(cell, color: str = GRID) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.first_child_found_in("w:tcBorders")
    if node is None:
        node = OxmlElement("w:tcBorders")
        tc_pr.append(node)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_node = node.find(qn(f"w:{edge}"))
        if edge_node is None:
            edge_node = OxmlElement(f"w:{edge}")
            node.append(edge_node)
        edge_node.set(qn("w:val"), "single")
        edge_node.set(qn("w:sz"), "5")
        edge_node.set(qn("w:color"), color)


def margins(cell, top: int = 110, start: int = 145, bottom: int = 110, end: int = 145) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.first_child_found_in("w:tcMar")
    if node is None:
        node = OxmlElement("w:tcMar")
        tc_pr.append(node)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        edge_node = node.find(qn(f"w:{side}"))
        if edge_node is None:
            edge_node = OxmlElement(f"w:{side}")
            node.append(edge_node)
        edge_node.set(qn("w:w"), str(value))
        edge_node.set(qn("w:type"), "dxa")


def keep_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:cantSplit")
    tr_pr.append(node)


def run(paragraph, text: str, *, size: float = 11, color: str = NAVY, bold: bool = False, mono: bool = False):
    result = paragraph.add_run(text)
    font_name = "Aptos Mono" if mono else "Aptos"
    result.font.name = font_name
    result._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    result.font.size = Pt(size)
    result.font.color.rgb = RGBColor.from_string(color)
    result.bold = bold
    return result


def page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run(paragraph, "DX-OS  |  Luồng mã nguồn demo  |  Trang ", size=8.5, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def configure(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.55)
    section.bottom_margin = Cm(1.45)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)
    section.header_distance = Cm(0.65)
    section.footer_distance = Cm(0.65)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.line_spacing = 1.2
    normal.paragraph_format.space_after = Pt(6)
    for name, size, color in (("Title", 26, NAVY), ("Heading 1", 18, TEAL_DARK), ("Heading 2", 14, NAVY)):
        style = doc.styles[name]
        style.font.name = "Aptos"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(14 if name != "Title" else 0)
        style.paragraph_format.space_after = Pt(7)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run(header, "DX-OS  •  LUỒNG MÃ NGUỒN DEMO", size=8.5, color=TEAL, bold=True)
    page_number(section.footer.paragraphs[0])


def add_note(doc: Document, body: str, fill: str = PALE_GOLD) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, fill)
    borders(cell)
    margins(cell, 135, 175, 135, 175)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run(p, body, size=10.5, color=NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_flow_card(doc: Document, number: str, title: str, objective: str, steps: list[str], sources: str) -> None:
    table = doc.add_table(rows=3, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    header = table.cell(0, 0)
    shade(header, TEAL_DARK)
    borders(header, TEAL_DARK)
    margins(header, 115, 150, 115, 150)
    p = header.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run(p, f"LUỒNG {number}  |  {title}", size=11.6, color=WHITE, bold=True)

    goal = table.cell(1, 0)
    shade(goal, PALE_TEAL)
    borders(goal)
    margins(goal, 115, 150, 115, 150)
    p = goal.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run(p, "Mục tiêu cần chứng minh: ", size=10.7, color=TEAL_DARK, bold=True)
    run(p, objective, size=10.7)

    flow = table.cell(2, 0)
    shade(flow, PALE_BLUE)
    borders(flow)
    margins(flow, 125, 160, 125, 160)
    p = flow.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    run(p, "LUỒNG MÃ NGUỒN", size=10.4, color=TEAL_DARK, bold=True)
    for step in steps:
        p = flow.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.15)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.08
        run(p, step, size=9.7, mono=True)
    p = flow.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(0)
    run(p, "FILE: ", size=9.2, color=TEAL_DARK, bold=True)
    run(p, sources, size=9.2, color=SLATE, mono=True)

    for row in table.rows:
        keep_together(row)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def build_document() -> Document:
    doc = Document()
    configure(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(44)
    run(p, "DX-OS", size=14, color=TEAL, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run(p, "LUỒNG MÃ NGUỒN DEMO", size=27, color=NAVY, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run(p, "Use case Hủy phiếu và Quản lý tệp đính kèm", size=16, color=TEAL_DARK, bold=True)
    add_note(
        doc,
        "Chỉ dùng như tờ điều hướng khi trình diễn mã nguồn: chọn đúng mục tiêu hỏi, sau đó đi lần lượt theo các mũi tên. Tài liệu không có phần diễn giải mẫu hay nhắc “tìm/nói”.",
        PALE_TEAL,
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(25)
    run(p, "Mã nguồn đối chiếu theo phiên bản DX-OS hiện tại.", size=10.5, color=SLATE)

    doc.add_page_break()
    doc.add_heading("Cách dùng nhanh", level=1)
    add_note(
        doc,
        "Mỗi thẻ là một đường demo hoàn chỉnh. Khối xanh là mục tiêu giảng viên đang hỏi; khối xanh nhạt là chuỗi file/hàm cần mở theo thứ tự. Các dòng FILE ở cuối thẻ là đường dẫn để dán vào Ctrl + P trong VS Code.",
        PALE_GOLD,
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run(p, "Thứ tự tổng quát: ", size=11, color=TEAL_DARK, bold=True)
    run(p, "Giao diện → Service → API → Quy tắc nghiệp vụ → Cơ sở dữ liệu/Kho tệp → Nhật ký", size=11)

    for index, flow in enumerate(FLOWS):
        if index in (0, 4, 8):
            doc.add_page_break()
            heading = "Hủy phiếu" if index == 0 else "Tệp đính kèm" if index == 4 else "Độ tin cậy và kiểm thử"
            doc.add_heading(heading, level=1)
        add_flow_card(doc, *flow)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14)
    run(p, "Không mở .env, data/runtime, mật khẩu hoặc access token trong buổi bảo vệ.", size=9.2, color=MUTED)
    return doc


if __name__ == "__main__":
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    build_document().save(OUTPUT)
    print(OUTPUT)
