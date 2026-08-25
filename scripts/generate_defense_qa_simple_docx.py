from __future__ import annotations

"""Create a plain-language defense guide for cancellation and attachments.

The document is deliberately tied to the currently implemented DX-OS flow:
* an employee can cancel only their own DRAFT / CHANGES_REQUESTED request;
* attachment management is available only while that request is editable;
* requests from 20,000,000 VND require an ACTIVE QUOTATION before submission.
"""

from pathlib import Path
import sys

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "generated" / "Bo_cau_hoi_bao_ve_UC_huy_phieu_va_quan_ly_tep_dinh_kem.docx"

NAVY = "18324A"
TEAL = "007C83"
TEAL_DARK = "005E63"
SLATE = "52677B"
MUTED = "6D7E8D"
WHITE = "FFFFFF"
PALE_TEAL = "E9F7F5"
PALE_BLUE = "F1F7FB"
PALE_GOLD = "FFF7E7"
PALE_RED = "FCEEF0"
GRID = "C9D9E3"


CANCEL_QUESTIONS = [
    (
        "Hủy phiếu để làm gì?",
        "Dùng khi nhu cầu mua không còn cần nữa. Hệ thống đóng phiếu, không xóa phiếu.",
        "Tôi chọn đổi trạng thái sang Đã hủy thay vì xóa dữ liệu. Vì vậy sau này vẫn xem được ai tạo, ai hủy, hủy lúc nào và lý do hủy.",
        "Trạng thái Đã hủy trên phiếu và dòng Hủy phiếu trong phần Tiến trình.",
    ),
    (
        "Ai được hủy phiếu?",
        "Chỉ người đã tạo phiếu mới được hủy phiếu của mình.",
        "Hệ thống kiểm tra ở máy chủ, không chỉ ẩn nút trên giao diện. Vì thế đăng nhập bằng tài khoản khác, kể cả Kiểm toán, cũng không thể hủy thay.",
        "Đăng nhập employee.demo, mở phiếu do chính tài khoản này tạo để thấy nút Hủy phiếu.",
    ),
    (
        "Phiếu ở trạng thái nào thì được hủy?",
        "Chỉ Bản nháp hoặc Đang yêu cầu chỉnh sửa.",
        "Khi phiếu đã gửi quản lý duyệt, người tạo không tự ý hủy nữa để không làm đứt công việc của người đang duyệt.",
        "Một phiếu Bản nháp có nút Hủy; sau khi Gửi duyệt thì nút này không còn.",
    ),
    (
        "Hủy phiếu khác gì Từ chối phiếu?",
        "Hủy là người tạo chủ động dừng nhu cầu. Từ chối là quản lý hoặc Tài chính không chấp nhận phiếu.",
        "Hai việc có người thực hiện và ý nghĩa khác nhau nên hệ thống lưu hai lịch sử riêng: Hủy phiếu và Từ chối phiếu.",
        "So sánh nhãn trạng thái và dòng Tiến trình của hai phiếu khác nhau.",
    ),
    (
        "Hủy rồi dữ liệu có mất không?",
        "Không mất. Phiếu vẫn còn trong hệ thống với trạng thái Đã hủy.",
        "Các dòng hàng, người tạo và lịch sử vẫn được giữ để báo cáo và kiểm toán. Người dùng chỉ không thể sửa hoặc gửi tiếp phiếu đã hủy.",
        "Mở lại phiếu sau khi hủy: vẫn thấy mã PR, nội dung và tổng tiền.",
    ),
    (
        "Nếu bấm Hủy hai lần hoặc mạng chậm thì sao?",
        "Hệ thống chống xử lý trùng, nên không tạo hai lần hủy trong lịch sử.",
        "Mỗi lần gửi có mã chống bấm trùng. Đồng thời phiếu có số phiên bản; nếu một tab khác đã sửa trước, thao tác cũ sẽ bị báo cần tải lại thay vì ghi đè.",
        "Không cần cố tạo lỗi khi demo; chỉ cần nêu đây là cơ chế chống bấm đôi và chống hai người cùng sửa.",
    ),
    (
        "Nếu mất điện hoặc lỗi giữa lúc hủy thì sao?",
        "Phiếu không rơi vào tình trạng nửa hủy nửa chưa hủy.",
        "Việc đổi trạng thái, ghi Tiến trình và ghi nhật ký được lưu trong một lần cập nhật của cơ sở dữ liệu. Hoặc tất cả thành công, hoặc hệ thống quay về trạng thái cũ.",
        "Nói ngắn: “Hệ thống lưu theo nguyên tắc làm trọn gói hoặc không làm gì cả.”",
    ),
    (
        "Có khôi phục phiếu đã hủy không?",
        "Bản hiện tại chưa có nút khôi phục. Nếu cần mua lại thì tạo phiếu mới.",
        "Đây là giới hạn có chủ đích để lịch sử rõ ràng. Một bản nâng cấp hợp lý là thêm yêu cầu mở lại phiếu, có người duyệt và lịch sử riêng.",
        "Nêu rõ đây là hướng phát triển, không nói là hệ thống đã có.",
    ),
    (
        "Em kiểm thử chức năng hủy như thế nào?",
        "Tôi kiểm tra đúng quyền, đúng trạng thái, lịch sử sau hủy và trường hợp thao tác bị từ chối.",
        "Các ca quan trọng là: chủ phiếu hủy Bản nháp thành công; tài khoản khác không hủy được; phiếu đã gửi không có nút hủy; sau khi hủy còn dấu vết trong Tiến trình và Kiểm toán.",
        "Chụp màn hình trước hủy, sau hủy và phần Tiến trình.",
    ),
]


ATTACHMENT_QUESTIONS = [
    (
        "Tệp đính kèm dùng để làm gì?",
        "Để đính kèm báo giá, đặc tả kỹ thuật, hợp đồng hoặc tài liệu liên quan đến phiếu mua sắm.",
        "Tệp giúp người duyệt có chứng cứ để quyết định. Ví dụ phiếu giá trị lớn phải có Báo giá trước khi gửi duyệt.",
        "Phần Tài liệu đính kèm hiển thị loại tệp, tên tệp, dung lượng và người tải lên.",
    ),
    (
        "Tệp được lưu ở đâu?",
        "Nội dung tệp lưu ở Nextcloud; thông tin mô tả tệp lưu trong cơ sở dữ liệu DX-OS.",
        "Tách như vậy giúp cơ sở dữ liệu không phải chứa tệp lớn, còn DX-OS vẫn kiểm soát được tệp thuộc phiếu nào, do ai tải và có còn hiệu lực không.",
        "Nói theo cách đơn giản: “Một nơi giữ file, một nơi giữ sổ theo dõi file.”",
    ),
    (
        "Ai được tải lên hoặc xóa tệp?",
        "Chỉ người tạo phiếu, khi phiếu còn Bản nháp hoặc đang được yêu cầu chỉnh sửa.",
        "Quyền được kiểm tra ở máy chủ. Kiểm toán chỉ xem/đối chiếu, không thể tự tải lên hay xóa chứng từ của người khác.",
        "Đăng nhập employee.demo và dùng một phiếu Bản nháp của chính mình.",
    ),
    (
        "Hệ thống nhận những loại tệp nào?",
        "PDF, DOCX, XLSX, JPG và PNG; mỗi tệp tối đa 10 MB.",
        "Hệ thống kiểm tra cả đuôi tệp và dấu hiệu nội dung bên trong, nên không thể chỉ đổi tên một file lạ thành .pdf để qua kiểm tra.",
        "Chọn một PDF nhỏ hơn 10 MB để demo thành công. Không cần dùng tệp nguy hiểm khi bảo vệ.",
    ),
    (
        "Làm sao chứng minh tệp tải lên đúng?",
        "Sau khi tải lên, hệ thống hiện tên, loại, dung lượng, người tải và thời gian tải.",
        "Máy chủ tạo thêm mã kiểm tra cho nội dung tệp. Khi tải xuống, hệ thống so lại mã này để phát hiện tệp bị thay đổi ngoài ý muốn.",
        "Tải xuống chính tệp vừa tải lên và mở để chứng minh tệp dùng được.",
    ),
    (
        "Nếu đang tải lên mà Nextcloud lỗi thì sao?",
        "Hệ thống không để lại một tệp nửa chừng như thể đã tải thành công.",
        "Tệp chỉ được xem là hợp lệ khi cả phần lưu file và phần ghi thông tin tệp hoàn tất. Nếu bước sau lỗi, hệ thống dọn phần tạm đã tạo và báo lỗi cho người dùng.",
        "Nêu đây là cách xử lý để hai nơi lưu trữ không lệch nhau; không cần cố tắt Nextcloud khi demo.",
    ),
    (
        "Xóa tệp thì có mất dấu vết không?",
        "Tệp không còn được dùng trong phiếu, nhưng hệ thống vẫn lưu lịch sử đã xóa để kiểm toán.",
        "Hệ thống đánh dấu đang xóa, xóa file ở kho lưu trữ, sau đó chốt trạng thái đã xóa và ghi nhật ký. Nếu kho lưu trữ báo lỗi, tệp được trả về trạng thái dùng được thay vì biến mất mập mờ.",
        "Xóa một tệp thử khi phiếu còn Bản nháp, rồi tải lại danh sách và mở Tiến trình.",
    ),
    (
        "Tại sao phiếu từ 20 triệu VND phải có báo giá?",
        "Đó là quy tắc chứng từ của hệ thống: từ 20.000.000 VND phải có ít nhất một tệp loại Báo giá trước khi gửi duyệt.",
        "Quy tắc được kiểm tra ở máy chủ khi Gửi duyệt và Gửi duyệt lại. Chỉ tệp đang còn hiệu lực và đúng loại Báo giá mới được tính.",
        "Tạo phiếu 25 triệu, thử Gửi duyệt khi chưa có Báo giá để thấy cảnh báo; tải PDF loại Báo giá lên rồi gửi lại.",
    ),
    (
        "Người dùng có thể bỏ qua kiểm tra bằng cách sửa giao diện không?",
        "Không. Giao diện chỉ hỗ trợ thao tác; máy chủ mới là nơi quyết định có cho tải và gửi phiếu hay không.",
        "Ngay cả khi cố gọi API trực tiếp, máy chủ vẫn kiểm tra quyền, trạng thái phiếu, loại tệp, dung lượng và quy tắc Báo giá.",
        "Có thể nói: “Tôi không đặt bảo mật chỉ ở nút bấm trên màn hình.”",
    ),
    (
        "Điểm hạn chế hiện tại là gì?",
        "Chưa có quét virus, phiên bản hóa tệp hoặc khôi phục tệp đã xóa trên giao diện.",
        "Các phần này là hướng nâng cấp thực tế: tích hợp phần mềm quét virus, lưu nhiều phiên bản và có quy trình phục hồi có kiểm soát.",
        "Nêu trung thực đây là phạm vi chưa triển khai, sau đó đưa ra hướng cải tiến.",
    ),
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_borders(cell, color: str = GRID) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "5")
        node.set(qn("w:color"), color)


def set_cell_margins(cell, top: int = 110, start: int = 145, bottom: int = 110, end: int = 145) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def keep_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:cantSplit")
    tr_pr.append(node)


def add_run(paragraph, value: str, *, size: float = 11, color: str = NAVY, bold: bool = False, italic: bool = False):
    run = paragraph.add_run(value)
    run.font.name = "Aptos"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic
    return run


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(paragraph, "DX-OS  |  Hướng dẫn bảo vệ  |  Trang ", size=8.5, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.55)
    section.bottom_margin = Cm(1.45)
    section.left_margin = Cm(1.65)
    section.right_margin = Cm(1.65)
    section.header_distance = Cm(0.65)
    section.footer_distance = Cm(0.65)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.line_spacing = 1.22
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color in (("Title", 26, NAVY), ("Heading 1", 18, TEAL_DARK), ("Heading 2", 14, NAVY)):
        style = doc.styles[name]
        style.font.name = "Aptos"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(16 if name != "Title" else 0)
        style.paragraph_format.space_after = Pt(7)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(header, "DX-OS  •  HỦY PHIẾU & TỆP ĐÍNH KÈM", size=8.5, color=TEAL, bold=True)
    add_page_number(section.footer.paragraphs[0])


def add_note(doc: Document, heading: str, body: str, fill: str = PALE_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_borders(cell)
    set_cell_margins(cell, 135, 180, 135, 180)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    add_run(p, heading, size=11.2, color=TEAL_DARK, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    add_run(p2, body, size=10.8)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_section_intro(doc: Document, title: str, subtitle: str) -> None:
    doc.add_heading(title, level=1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    add_run(p, subtitle, size=11, color=SLATE, italic=True)


def add_simple_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for cell, label in zip(table.rows[0].cells, ("Nội dung", "Ghi nhớ thật nhanh")):
        set_cell_shading(cell, TEAL_DARK)
        set_cell_borders(cell, TEAL_DARK)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, label, size=10.5, color=WHITE, bold=True)
    keep_row_together(table.rows[0])
    for left, right in rows:
        cells = table.add_row().cells
        for cell, value, fill, bold in ((cells[0], left, PALE_TEAL, True), (cells[1], right, WHITE, False)):
            set_cell_shading(cell, fill)
            set_cell_borders(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.12
            add_run(p, value, size=10.4, bold=bold)
        cells[0].width = Cm(4.4)
        cells[1].width = Cm(12.8)
        keep_row_together(table.rows[-1])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_demo_step(doc: Document, number: int, action: str, expected: str, say: str) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    badge, body = table.rows[0].cells
    badge.width = Cm(1.15)
    body.width = Cm(16.05)
    set_cell_shading(badge, TEAL_DARK)
    set_cell_borders(badge, TEAL_DARK)
    set_cell_margins(badge, 150, 60, 150, 60)
    badge.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = badge.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, f"{number}", size=15, color=WHITE, bold=True)

    set_cell_shading(body, WHITE)
    set_cell_borders(body)
    set_cell_margins(body, 120, 170, 120, 170)
    body.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = body.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    add_run(p, "Thao tác: ", size=10.8, color=TEAL_DARK, bold=True)
    add_run(p, action, size=10.8)
    p2 = body.add_paragraph()
    p2.paragraph_format.space_after = Pt(3)
    add_run(p2, "Cần thấy: ", size=10.8, color=TEAL_DARK, bold=True)
    add_run(p2, expected, size=10.8)
    p3 = body.add_paragraph()
    p3.paragraph_format.space_after = Pt(0)
    add_run(p3, "Nói khi demo: ", size=10.8, color=TEAL_DARK, bold=True)
    add_run(p3, say, size=10.8, italic=True)
    keep_row_together(table.rows[0])
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_question(doc: Document, number: int, question: str, short: str, detail: str, demo: str) -> None:
    table = doc.add_table(rows=4, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    values = [
        (f"{number:02d}. {question}", TEAL_DARK, WHITE, 11.4, True),
        (short, PALE_TEAL, NAVY, 10.8, False),
        (detail, PALE_BLUE, NAVY, 10.5, False),
        (demo, PALE_GOLD, NAVY, 10.2, False),
    ]
    labels = ["Câu trả lời ngắn", "Nói thêm nếu thầy cô hỏi sâu", "Minh chứng khi demo"]
    for row_index, (value, fill, color, size, bold) in enumerate(values):
        cell = table.cell(row_index, 0)
        set_cell_shading(cell, fill)
        set_cell_borders(cell, TEAL_DARK if row_index == 0 else GRID)
        set_cell_margins(cell, 105, 155, 105, 155)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.12
        if row_index == 0:
            add_run(p, value, size=size, color=color, bold=bold)
        else:
            add_run(p, labels[row_index - 1] + ": ", size=size, color=TEAL_DARK, bold=True)
            add_run(p, value, size=size, color=color)
        keep_row_together(table.rows[row_index])
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(3)


def add_question_section(doc: Document, title: str, subtitle: str, questions: list[tuple[str, str, str, str]]) -> None:
    doc.add_page_break()
    add_section_intro(doc, title, subtitle)
    for index, item in enumerate(questions, start=1):
        add_question(doc, index, *item)


def add_source_reference_table(doc: Document, rows: list[tuple[str, str, str, str]]) -> None:
    """Add an easy-to-scan source-code map for the defense presentation."""
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    headers = ("Nếu được hỏi về", "Mở file nào và chỉ vào đâu")
    for cell, label in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, TEAL_DARK)
        set_cell_borders(cell, TEAL_DARK)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, label, size=10.5, color=WHITE, bold=True)
    keep_row_together(table.rows[0])

    for topic, source, find, explain in rows:
        left, right = table.add_row().cells
        left.width = Cm(4.25)
        right.width = Cm(12.95)
        for cell, fill in ((left, PALE_TEAL), (right, WHITE)):
            set_cell_shading(cell, fill)
            set_cell_borders(cell)
            set_cell_margins(cell, 110, 145, 110, 145)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = left.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.12
        add_run(p, topic, size=10.25, bold=True)

        p = right.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        add_run(p, "Mở: ", size=10.1, color=TEAL_DARK, bold=True)
        add_run(p, source, size=10.1, color=NAVY)
        p2 = right.add_paragraph()
        p2.paragraph_format.space_after = Pt(2)
        add_run(p2, "Tìm: ", size=10.1, color=TEAL_DARK, bold=True)
        add_run(p2, find, size=10.1)
        p3 = right.add_paragraph()
        p3.paragraph_format.space_after = Pt(0)
        add_run(p3, "Nói: ", size=10.1, color=TEAL_DARK, bold=True)
        add_run(p3, explain, size=10.1, italic=True)
        keep_row_together(table.rows[-1])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_source_code_section(doc: Document) -> None:
    """Append this section without changing the defense content that precedes it."""
    doc.add_page_break()
    add_section_intro(
        doc,
        "8. Khi giảng viên hỏi mã nguồn: mở file nào?",
        "Đây là bản đồ tra nhanh. Trong VS Code bấm Ctrl + P, dán đường dẫn; sau đó Ctrl + F để tìm đúng tên hàm. Số dòng có thể thay đổi khi code được cập nhật, tên hàm là mốc đáng tin hơn.",
    )
    add_note(
        doc,
        "Cách trình bày mã nguồn an toàn",
        "Đừng cuộn cả file cho giảng viên xem. Hãy đi theo một đường: giao diện → API → quy tắc nghiệp vụ → lưu dữ liệu/lịch sử. Khi chỉ code, đọc điều kiện bằng lời Việt; không cần đọc từng dòng lệnh.",
        PALE_GOLD,
    )

    doc.add_heading("8.1. Luồng hủy phiếu", level=2)
    add_source_reference_table(
        doc,
        [
            (
                "Vì sao employee.demo mới thấy nút Hủy?",
                "frontend/src/app/features/procurement/pages/purchase-request-detail/purchase-request-detail.ts — availableActions() (khoảng dòng 154)",
                "DRAFT và CHANGES_REQUESTED trả về CANCEL khi isOwner() là đúng.",
                "“Giao diện chỉ đưa nút hủy cho đúng chủ phiếu và đúng hai trạng thái còn chỉnh sửa được.”",
            ),
            (
                "Hệ thống biết ai là chủ phiếu bằng cách nào?",
                "frontend/src/app/features/procurement/utils/purchase-request-ownership.ts — isPurchaseRequestOwner() (dòng 3)",
                "request.requesterUsername được so với username đang đăng nhập; không so tên hiển thị.",
                "“Tôi dùng tên tài khoản ổn định, nên tên hiển thị như Nguyễn Minh Anh không làm sai quyền.”",
            ),
            (
                "Nút Xác nhận gọi gì?",
                "frontend/src/app/features/procurement/pages/purchase-request-detail/purchase-request-detail.ts — performTransition() (khoảng dòng 358)",
                "Gửi action, expectedVersion và ghi chú nếu có sang ProcurementService.",
                "“Sau khi người dùng bấm xác nhận, giao diện gửi yêu cầu chuyển trạng thái kèm phiên bản phiếu đang xem.”",
            ),
            (
                "API nào nhận yêu cầu hủy?",
                "backend/internal/platform/httpapi/router.go — dòng 184; backend/internal/platform/httpapi/purchase_requests.go — transitionPurchaseRequest() (dòng 418)",
                "POST /purchase-requests/{requestID}/transitions rồi handler gọi lớp nghiệp vụ.",
                "“Hủy không dùng API DELETE vì đây là một quyết định trong quy trình, không phải xóa bản ghi.”",
            ),
            (
                "Quy tắc ai được hủy và hủy ở trạng thái nào nằm ở đâu?",
                "backend/internal/procurement/model.go — DecideTransition() (dòng 1242)",
                "Hai nhánh StatusDraft và StatusChangesRequested: actor.UserID phải bằng request.RequesterID, action phải là ActionCancel.",
                "“Đây là nơi quy định luật nghiệp vụ; người khác hoặc trạng thái khác sẽ bị từ chối ở máy chủ.”",
            ),
            (
                "Làm sao chống bấm hai lần và ghi đè dữ liệu cũ?",
                "backend/internal/procurement/store.go — Transition() (dòng 1917), đoạn idempotency quanh dòng 1953 và kiểm tra version quanh dòng 2009",
                "Tìm idempotency_key, ErrIdempotencyConflict và ErrVersionConflict.",
                "“Nếu mạng chậm hoặc có hai tab, hệ thống không tạo thêm lần hủy hoặc ghi đè quyết định mới hơn.”",
            ),
            (
                "Lịch sử Hủy phiếu được lưu ở đâu?",
                "backend/internal/procurement/store.go — Transition() (dòng 2066–2096); backend/migrations/000002_procurement_mvp.sql — process_events (dòng 81); backend/migrations/000003_procurement_audit.sql — audit_logs (dòng 1)",
                "Tìm lệnh ghi process_events, insertAudit(...) và tx.Commit(...).",
                "“Một lần hủy ghi cả Tiến trình của phiếu và nhật ký kiểm toán trước khi chốt dữ liệu.”",
            ),
            (
                "Nếu mất điện giữa lúc hủy?",
                "backend/internal/procurement/store.go — Transition(): bắt đầu giao dịch quanh dòng 1930, chốt tại dòng 2096",
                "Tìm Begin(ctx), Rollback và Commit(ctx).",
                "“Các thay đổi quan trọng được lưu trọn gói; lỗi giữa chừng thì cơ sở dữ liệu quay lại trạng thái cũ.”",
            ),
        ],
    )

    doc.add_heading("8.2. Luồng tệp đính kèm", level=2)
    add_source_reference_table(
        doc,
        [
            (
                "Nút tải lên, tải xuống và xóa nằm ở đâu?",
                "frontend/src/app/features/procurement/pages/purchase-request-detail/purchase-request-detail.html — khoảng dòng 239, 264, 275; file .ts cùng thư mục — uploadAttachment(), downloadAttachment(), deleteAttachment()",
                "Ba hàm giao diện lần lượt gọi service để tải lên, tải xuống và xóa.",
                "“Đây là phần người dùng nhìn thấy; mọi thao tác đều gọi về máy chủ để kiểm tra lại.”",
            ),
            (
                "Các API tệp là gì?",
                "backend/internal/platform/httpapi/router.go — dòng 185–188; frontend/src/app/features/procurement/data-access/procurement.service.ts — dòng 343–366",
                "Bốn đường dẫn: xem danh sách, tải lên, tải nội dung, xóa.",
                "“Frontend chỉ là nơi gửi tệp; API giới hạn tệp trong đúng phiếu mua sắm.”",
            ),
            (
                "Loại tệp và dung lượng được kiểm tra ở đâu?",
                "backend/internal/procurement/model.go — ValidateAttachment() (dòng 662)",
                "PDF, DOCX, XLSX, JPG, PNG; 1 byte đến 10 MB; đồng thời kiểm tra tên tệp và nội dung thực.",
                "“Không chỉ tin phần đuôi file. Máy chủ còn kiểm tra dấu hiệu bên trong của PDF, ảnh và Office.”",
            ),
            (
                "Tệp được lưu như thế nào?",
                "backend/internal/procurement/attachments.go — UploadAttachment() (dòng 25); backend/internal/platform/documentstore/nextcloud.go — Put() (dòng 40)",
                "Tìm trạng thái UPLOADING, lời gọi documents.Put(...), rồi trạng thái ACTIVE và ATTACHMENT_UPLOADED.",
                "“DX-OS ghi sổ theo dõi trước, lưu file vào Nextcloud, sau đó mới đánh dấu tệp dùng được.”",
            ),
            (
                "Nếu kho file lỗi lúc tải lên?",
                "backend/internal/procurement/attachments.go — UploadAttachment(), các nhánh cleanupUploadingAttachment() quanh dòng 105–165",
                "Các nhánh lỗi gọi dọn dữ liệu tạm khi Put hoặc bước chốt thông tin thất bại.",
                "“Nếu lưu chưa xong, tệp không được hiện như đã thành công và phần tạm sẽ được dọn.”",
            ),
            (
                "Tải xuống kiểm tra tệp có nguyên vẹn không?",
                "backend/internal/procurement/attachments.go — DownloadAttachment() (dòng 285)",
                "Tìm sha256.Sum256 và so sánh ChecksumSHA256 trước khi trả file.",
                "“Trước khi gửi file về trình duyệt, hệ thống đối chiếu dấu vân tay của file với lúc đã tải lên.”",
            ),
            (
                "Xóa tệp có lưu lịch sử không?",
                "backend/internal/procurement/attachments.go — DeleteAttachment() (dòng 309)",
                "Tìm DELETING → documents.Delete(...) → DELETED và insertAudit(... ATTACHMENT_DELETED ...).",
                "“Xóa tệp có các bước rõ ràng; lỗi kho file thì trả trạng thái về dùng được thay vì mất dấu vết.”",
            ),
            (
                "Quy tắc 20 triệu phải có Báo giá ở đâu?",
                "backend/internal/procurement/attachments.go — requireAttachmentForSubmission() (dòng 443); backend/internal/procurement/store.go — gọi hàm này quanh dòng 2012",
                "Tìm threshold_amount, required_document_type và status = ACTIVE.",
                "“Khi Gửi duyệt, máy chủ kiểm tra có Báo giá đang hiệu lực theo quy tắc tiền tệ và ngưỡng tiền hay chưa.”",
            ),
            (
                "Các bảng dữ liệu liên quan?",
                "backend/migrations/000006_purchase_request_attachments.sql — attachment_rules (dòng 1) và purchase_request_attachments (dòng 14)",
                "attachment_rules là quy tắc bắt buộc; purchase_request_attachments là thông tin theo dõi từng tệp.",
                "“File nằm ở Nextcloud; hai bảng này giúp DX-OS biết tệp nào thuộc phiếu nào và có còn dùng được không.”",
            ),
        ],
    )

    doc.add_heading("8.3. Nếu giảng viên yêu cầu xem kiểm thử", level=2)
    add_source_reference_table(
        doc,
        [
            (
                "Kiểm thử luật chuyển trạng thái", 
                "backend/internal/procurement/model_test.go — TestDecideTransition() (dòng 126)",
                "Các trường hợp vai trò, trạng thái và hành động hợp lệ/không hợp lệ.",
                "“Tôi kiểm thử luật nghiệp vụ tách khỏi giao diện, nên lỗi quyền hoặc trạng thái bị phát hiện sớm.”",
            ),
            (
                "Kiểm thử an toàn tệp", 
                "backend/internal/procurement/model_test.go — các test ValidateAttachment từ dòng 293", 
                "Tìm các ca PDF Báo giá hợp lệ, file không an toàn, giả loại nội dung và sai phần mở rộng.",
                "“Tôi có test cả file đúng và các cách giả mạo phổ biến, không chỉ test đường thành công.”",
            ),
            (
                "Kiểm thử hiển thị đúng chủ phiếu", 
                "frontend/src/app/features/procurement/utils/purchase-request-ownership.spec.ts — dòng 10", 
                "Tên hiển thị Nguyễn Minh Anh không được dùng để cấp quyền; username employee.demo mới là định danh.",
                "“Đây là kiểm thử hồi quy cho lỗi giao diện từng ẩn nhầm nút chủ phiếu.”",
            ),
        ],
    )
    add_note(
        doc,
        "Không nên mở khi bảo vệ",
        "Không mở file .env, file tài khoản trong data/runtime, mật khẩu, access token hoặc dữ liệu cấu hình bí mật. Nếu cần nói về bảo mật, chỉ giải thích rằng các giá trị này lấy từ biến môi trường và không được đưa vào tài liệu/slide.",
        PALE_RED,
    )


def build_document() -> Document:
    doc = Document()
    configure_document(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(45)
    add_run(p, "DX-OS", size=14, color=TEAL, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    add_run(p, "BỘ CÂU HỎI BẢO VỆ", size=27, color=NAVY, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    add_run(p, "Hủy phiếu và quản lý tệp đính kèm", size=17, color=TEAL_DARK, bold=True)
    add_note(
        doc,
        "Bản viết lại dễ dùng khi bảo vệ",
        "Tài liệu này ưu tiên cách nói đơn giản và thao tác thật trên DX-OS. Mỗi câu có ba phần: trả lời ngắn để nói ngay, phần giải thích thêm khi bị hỏi sâu, và vị trí cần chỉ trên màn hình.",
        PALE_TEAL,
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(26)
    add_run(p, "Phạm vi: phiếu mua sắm, lịch sử xử lý, tệp đính kèm, PostgreSQL và Nextcloud.", size=10.5, color=SLATE, italic=True)

    doc.add_page_break()
    add_section_intro(doc, "1. Cách dùng tài liệu này", "Đọc trang này trước, sau đó tập trung vào phần demo. Không cần học thuộc câu chữ.")
    add_simple_table(
        doc,
        [
            ("Khi được hỏi nhanh", "Đọc ý ở dòng “Câu trả lời ngắn”. Mỗi câu chỉ cần 1–2 câu, nói rõ ràng và dừng lại."),
            ("Khi bị hỏi sâu", "Mở rộng bằng dòng “Nói thêm”. Chỉ dùng thuật ngữ kỹ thuật nếu giảng viên hỏi tiếp."),
            ("Khi demo", "Luôn vừa bấm vừa nói: bấm ở đâu, màn hình cần hiện gì, và ý nghĩa của kết quả đó."),
            ("Khi chưa có chức năng", "Nói thẳng “bản hiện tại chưa có”, rồi nêu hướng phát triển. Không nhận là đã làm nếu giao diện chưa có."),
            ("Câu mở đầu 20 giây", "“Em phụ trách việc hủy một nhu cầu mua sắm nhưng không làm mất lịch sử, và quản lý chứng từ đính kèm an toàn trước khi phiếu được gửi duyệt.”"),
        ],
    )
    add_note(
        doc,
        "Ba ý cần nhớ nhất",
        "(1) Hủy là đóng phiếu chứ không xóa phiếu. (2) Chỉ người tạo mới thao tác được khi phiếu còn sửa được. (3) Với phiếu từ 20 triệu VND, phải có tệp Báo giá trước khi gửi duyệt.",
        PALE_GOLD,
    )

    doc.add_page_break()
    add_section_intro(doc, "2. Chuẩn bị trước khi demo", "Dùng hai phiếu test riêng: một phiếu chỉ để hủy, một phiếu chỉ để thử tệp đính kèm.")
    add_simple_table(
        doc,
        [
            ("Tài khoản", "Đăng nhập employee.demo. Đây phải là người tạo hai phiếu test để thấy nút Hủy phiếu, tải lên và xóa tệp."),
            ("Phiếu A – để hủy", "Tạo phiếu giá trị nhỏ, ví dụ Chuột không dây 350.000 VND. Lưu ở trạng thái Bản nháp; không gửi duyệt."),
            ("Phiếu B – để tải tệp", "Tạo phiếu 25.000.000 VND, ví dụ Mua laptop cho nhân viên mới. Chuẩn bị một PDF Báo giá dưới 10 MB."),
            ("Vị trí nút", "Trong trang chi tiết phiếu, kéo lên phần đầu cột bên phải để thấy khung Hành động. Tại đây có Gửi duyệt và Hủy phiếu."),
            ("Nếu không thấy nút", "Kiểm tra lại: đang đăng nhập employee.demo, đúng phiếu do employee.demo tạo, và phiếu đang Bản nháp hoặc Yêu cầu chỉnh sửa. Sau khi cập nhật hệ thống, dùng Ctrl + F5 để tải lại trang."),
        ],
    )

    doc.add_page_break()
    add_section_intro(doc, "3. Demo A — Hủy phiếu trong khoảng 3 phút", "Không hủy phiếu đang cần cho luồng duyệt khác. Hãy dùng Phiếu A đã chuẩn bị riêng.")
    add_demo_step(doc, 1, "Vào Phiếu mua sắm → Tạo phiếu. Nhập tiêu đề “Demo hủy phiếu - không mua nữa”, lý do ít nhất 10 ký tự, một dòng hàng và lưu Bản nháp.", "Phiếu có mã PR và nhãn Bản nháp.", "“Em tạo một phiếu test ở trạng thái còn được phép chỉnh sửa.”")
    add_demo_step(doc, 2, "Mở trang chi tiết của chính phiếu vừa tạo. Ở khung Hành động bên phải, bấm Hủy phiếu.", "Màn hình hiện ô Lý do/ghi chú và nút Xác nhận.", "“Nút này chỉ hiện với người tạo phiếu và ở đúng trạng thái.”")
    add_demo_step(doc, 3, "Nhập “Không còn nhu cầu mua” rồi bấm Xác nhận.", "Nhãn phiếu đổi thành Đã hủy. Không còn nút sửa hay gửi duyệt.", "“Hệ thống đóng quy trình, nhưng không xóa dữ liệu của phiếu.”")
    add_demo_step(doc, 4, "Kéo đến phần Tiến trình hoặc Timeline của phiếu.", "Có dòng Hủy phiếu, người thực hiện và thời gian.", "“Đây là bằng chứng hệ thống giữ lại lịch sử để truy vết.”")
    add_demo_step(doc, 5, "Nếu còn thời gian, đăng nhập tài khoản khác và mở lại phiếu.", "Không có quyền hủy thay hoặc sửa phiếu.", "“Quyền được kiểm tra ở máy chủ, không chỉ dựa vào việc ẩn nút.”")
    add_note(doc, "Ảnh nên chụp", "Chụp một ảnh trước khi hủy có nhãn Bản nháp và nút Hủy phiếu; một ảnh sau khi hủy có nhãn Đã hủy; một ảnh phần Tiến trình có dòng Hủy phiếu.", PALE_BLUE)

    doc.add_page_break()
    add_section_intro(doc, "4. Demo B — Quản lý tệp đính kèm trong khoảng 4 phút", "Dùng Phiếu B 25.000.000 VND để vừa demo tệp, vừa chứng minh quy tắc bắt buộc Báo giá.")
    add_demo_step(doc, 1, "Tạo Phiếu B: tiêu đề “Demo tệp đính kèm - mua laptop”, một dòng hàng 1 × 25.000.000 VND, sau đó lưu Bản nháp.", "Phiếu hiển thị tổng tiền 25.000.000 VND.", "“Tổng tiền này cao hơn ngưỡng 20 triệu, nên cần Báo giá trước khi gửi.”")
    add_demo_step(doc, 2, "Trong chi tiết phiếu, đến phần Tài liệu đính kèm. Chọn loại Báo giá, chọn tệp PDF nhỏ hơn 10 MB rồi bấm Tải lên.", "Danh sách hiện tên tệp, nhãn Báo giá, dung lượng, người tải và thời gian.", "“Tệp được kiểm tra loại, dung lượng và nội dung trước khi được chấp nhận.”")
    add_demo_step(doc, 3, "Bấm Tải xuống và mở tệp vừa tải về.", "Tệp tải về mở được bình thường.", "“Khi tải xuống, hệ thống kiểm tra lại tệp có còn nguyên vẹn hay không.”")
    add_demo_step(doc, 4, "Bấm Xóa tệp, sau đó tải lại danh sách. Nếu muốn gửi duyệt, tải lại PDF Báo giá một lần nữa.", "Tệp không còn trong danh sách đang dùng; sau khi tải lại thì xuất hiện lại.", "“Xóa tệp vẫn có lịch sử để kiểm toán, không phải xóa mù.”")
    add_demo_step(doc, 5, "Khi chưa có Báo giá, bấm Gửi duyệt để quan sát cảnh báo. Sau đó tải lại Báo giá và bấm Gửi duyệt → Xác nhận.", "Lần đầu bị chặn vì thiếu Báo giá; lần sau phiếu chuyển Đã gửi.", "“Quy tắc này nằm ở máy chủ nên không thể bỏ qua chỉ bằng cách sửa giao diện.”")
    add_note(doc, "Ảnh nên chụp", "Chụp cảnh báo thiếu Báo giá, danh sách có tệp Báo giá, và kết quả Gửi duyệt thành công sau khi tệp đã được tải lên.", PALE_GOLD)

    add_question_section(doc, "5. Câu hỏi thường gặp — Hủy phiếu", "Dùng câu trả lời ngắn trước. Phần dưới chỉ mở rộng khi giảng viên hỏi tiếp.", CANCEL_QUESTIONS)
    add_question_section(doc, "6. Câu hỏi thường gặp — Tệp đính kèm", "Tập trung vào đúng luồng bạn vừa demo, không cần trình bày hết kiến trúc ngay từ đầu.", ATTACHMENT_QUESTIONS)

    doc.add_page_break()
    add_section_intro(doc, "7. Các từ kỹ thuật chỉ cần dùng khi bị hỏi sâu", "Nếu không bị hỏi sâu, hãy dùng phần giải thích bằng tiếng Việt ở cột bên phải.")
    add_simple_table(
        doc,
        [
            ("Giao dịch dữ liệu", "Nói đơn giản là: các bước lưu trong cơ sở dữ liệu hoặc thành công hết, hoặc quay lại như chưa làm."),
            ("Số phiên bản", "Số thứ tự thay đổi của phiếu. Nó giúp phát hiện bạn đang thao tác trên dữ liệu cũ."),
            ("Mã chống bấm trùng", "Giúp việc bấm lại do mạng chậm không tạo hai lần hủy hoặc hai dòng lịch sử."),
            ("Mã kiểm tra tệp", "Dấu vân tay của file. Hệ thống so lại khi tải xuống để biết file có bị thay đổi không."),
            ("Nhật ký kiểm toán", "Sổ theo dõi ai đã làm gì, vào lúc nào, trước và sau khi thay đổi ra sao."),
            ("Nextcloud", "Kho lưu tệp nội bộ mà DX-OS dùng để chứa file đính kèm."),
        ],
    )
    add_note(doc, "Câu chốt phần của bạn", "“Điểm chính của hai use case là đảm bảo người dùng có thể dừng hoặc bổ sung chứng từ đúng quyền, nhưng hệ thống vẫn giữ được dấu vết để sau này kiểm tra lại.”", PALE_TEAL)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    add_run(p, "Tài liệu bám theo DX-OS hiện tại. Không mô tả các chức năng chưa có như quét virus, khôi phục phiếu hoặc khôi phục phiên bản tệp.", size=9.2, color=MUTED, italic=True)
    add_source_code_section(doc)
    return doc


if __name__ == "__main__":
    output = Path(sys.argv[1]).resolve() if len(sys.argv) > 1 else OUTPUT
    output.parent.mkdir(parents=True, exist_ok=True)
    build_document().save(output)
    print(output)
