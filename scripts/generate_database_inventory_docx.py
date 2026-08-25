from __future__ import annotations

"""Create a DX-OS database inventory aligned with Nhom_DACN (1).docx.

The source report's original use cases are retained.  The inventory only adds
the tables and functions introduced by later DX-OS migrations.
"""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "generated" / "Nhom_DACN (1).docx"
OUTPUT = ROOT / "docs" / "generated" / "Thong_ke_du_lieu_DX_OS_doi_chieu_Nhom_DACN_1.docx"

NAVY = "15324B"
TEAL = "007C83"
SLATE = "526A7D"
PALE_TEAL = "E8F5F5"
PALE_BLUE = "EEF5FB"
PALE_GREEN = "EAF6EE"
PALE_YELLOW = "FFF7E6"
PALE_PURPLE = "F3EEFC"
GRID = "C8D8E0"
WHITE = "FFFFFF"


GROUPS = [
    ("Nền tảng, tổ chức và danh tính", "5", "Cấu hình chung, tổ chức, phòng ban, tài khoản và ảnh chụp quyền."),
    ("Quy trình yêu cầu mua sắm", "5", "Phiếu mua sắm, dòng hàng, lịch sử xử lý và tệp đính kèm."),
    ("Ngân sách", "4", "Kỳ ngân sách, hạn mức, giữ chỗ/cam kết và điều chỉnh."),
    ("Nhà cung cấp, danh mục và báo giá", "5", "Danh mục chuẩn, nhà cung cấp, hồ sơ và báo giá so sánh."),
    ("Quy tắc phê duyệt", "2", "Quy tắc theo giá trị/phòng ban và ủy quyền người duyệt."),
    ("Đặt hàng, giao nhận, hóa đơn và thanh toán", "6", "Phát hành đơn hàng, nhận hàng, hóa đơn và thanh toán."),
    ("Thông báo", "3", "Hàng đợi thông báo, thông báo theo người nhận và trạng thái đã đọc."),
    ("Kiểm toán và hỗ trợ quyết định", "4", "Nhật ký kiểm toán, vụ việc kiểm toán và khuyến nghị hỗ trợ."),
    ("Báo cáo", "4", "Chính sách thời hạn xử lý cùng ba view chỉ đọc cho Metabase/báo cáo."),
]


OBJECTS = [
    ("1", "Nền tảng, tổ chức và danh tính", "app_metadata", "Bảng", "Giữ — nền tảng", "Lưu phiên bản/cấu hình nghiệp vụ chung của DX-OS."),
    ("2", "Nền tảng, tổ chức và danh tính", "organizations", "Bảng", "Giữ — phạm vi gốc", "Đơn vị tổ chức; là phạm vi dữ liệu của các phòng ban, ngân sách, nhà cung cấp và báo cáo."),
    ("3", "Nền tảng, tổ chức và danh tính", "departments", "Bảng", "Giữ — UC cũ", "Phòng ban, có quan hệ cha–con; xác định phạm vi người dùng và phiếu mua sắm."),
    ("4", "Nền tảng, tổ chức và danh tính", "users", "Bảng", "Giữ — UC cũ", "Hồ sơ người dùng DX-OS, liên kết danh tính Keycloak qua keycloak_subject."),
    ("5", "Nền tảng, tổ chức và danh tính", "user_role_snapshots", "Bảng", "Bổ sung mới", "Lưu ảnh chụp quyền/role dùng để hiển thị và kiểm soát nghiệp vụ hiện tại."),
    ("6", "Quy trình yêu cầu mua sắm", "purchase_requests", "Bảng", "Giữ — UC-02…UC-08", "Hồ sơ phiếu mua sắm: người yêu cầu, phòng ban, trạng thái, tổng tiền và người đang xử lý."),
    ("7", "Quy trình yêu cầu mua sắm", "purchase_request_items", "Bảng", "Giữ — UC-02, UC-03", "Các dòng hàng/dịch vụ của một phiếu mua sắm."),
    ("8", "Quy trình yêu cầu mua sắm", "process_events", "Bảng", "Giữ + mở rộng", "Lịch sử các hành động workflow trên phiếu: gửi, duyệt, từ chối, yêu cầu sửa, hủy…"),
    ("9", "Quy trình yêu cầu mua sắm", "attachment_rules", "Bảng", "Giữ — UC-09", "Quy định loại tệp/ngưỡng tiền cần có chứng từ."),
    ("10", "Quy trình yêu cầu mua sắm", "purchase_request_attachments", "Bảng", "Giữ — UC-09", "Metadata tệp đính kèm: đường dẫn, mã băm, người tải lên; nội dung tệp nằm ở Nextcloud."),
    ("11", "Ngân sách", "budget_periods", "Bảng", "Bổ sung cấu trúc", "Kỳ ngân sách theo tổ chức, dùng để tổ chức các hạn mức."),
    ("12", "Ngân sách", "budget_allocations", "Bảng", "Giữ — UC-10", "Hạn mức theo kỳ/trung tâm chi phí; là số liệu nền của dashboard ngân sách."),
    ("13", "Ngân sách", "budget_reservations", "Bảng", "Giữ — UC-07, UC-10", "Ghi nhận tiền giữ chỗ, cam kết hoặc giải phóng theo phiếu."),
    ("14", "Ngân sách", "budget_adjustments", "Bảng", "Giữ — UC-10", "Lịch sử tăng/giảm hạn mức và người thực hiện điều chỉnh."),
    ("15", "Nhà cung cấp, danh mục và báo giá", "suppliers", "Bảng", "Bổ sung mới", "Hồ sơ nhà cung cấp, tuân thủ, thông tin hợp đồng/ngân hàng và đánh giá hiệu suất."),
    ("16", "Nhà cung cấp, danh mục và báo giá", "procurement_catalog_items", "Bảng", "Bổ sung mới", "Danh mục hàng hóa/dịch vụ chuẩn: mã, nhóm, đơn vị và giá tham chiếu."),
    ("17", "Nhà cung cấp, danh mục và báo giá", "sourcing_cases", "Bảng", "Bổ sung mới", "Hồ sơ so sánh báo giá của một phiếu; lưu báo giá được chọn và trạng thái trao thầu."),
    ("18", "Nhà cung cấp, danh mục và báo giá", "supplier_quotes", "Bảng", "Bổ sung mới", "Các báo giá của nhà cung cấp: tổng giá, ngày giao, bảo hành, điều khoản và điểm tổng hợp."),
    ("19", "Nhà cung cấp, danh mục và báo giá", "sourcing_events", "Bảng", "Bổ sung mới", "Lịch sử nhập/sửa/chọn báo giá trong một hồ sơ so sánh."),
    ("20", "Quy tắc phê duyệt", "approval_rules", "Bảng", "Bổ sung mới", "Quy tắc tuyến phê duyệt theo tổ chức, phòng ban và khoảng giá trị phiếu."),
    ("21", "Quy tắc phê duyệt", "approval_delegations", "Bảng", "Bổ sung mới", "Ủy quyền người duyệt theo phòng ban và thời hạn hiệu lực."),
    ("22", "Đặt hàng, giao nhận, hóa đơn và thanh toán", "purchase_orders", "Bảng", "Bổ sung mới", "Đơn đặt hàng phát hành từ phiếu đã được chọn nhà cung cấp."),
    ("23", "Đặt hàng, giao nhận, hóa đơn và thanh toán", "purchase_order_receipts", "Bảng", "Bổ sung mới", "Biên bản/lần nhận hàng cho một đơn đặt hàng; hỗ trợ nhận một phần, hỏng, sai hàng hoặc từ chối."),
    ("24", "Đặt hàng, giao nhận, hóa đơn và thanh toán", "purchase_order_receipt_items", "Bảng trung gian", "Bổ sung mới", "Dòng nhận hàng; ghép một biên bản nhận với các dòng hàng của phiếu."),
    ("25", "Đặt hàng, giao nhận, hóa đơn và thanh toán", "purchase_invoices", "Bảng", "Bổ sung mới", "Hóa đơn theo đơn đặt hàng; hỗ trợ hóa đơn chuẩn, tạm ứng, quyết toán và điều chỉnh."),
    ("26", "Đặt hàng, giao nhận, hóa đơn và thanh toán", "invoice_events", "Bảng", "Bổ sung mới", "Lịch sử xác minh, tranh chấp, duyệt hoặc thay đổi trạng thái hóa đơn."),
    ("27", "Đặt hàng, giao nhận, hóa đơn và thanh toán", "invoice_payments", "Bảng", "Bổ sung mới", "Các khoản thanh toán, ngày thanh toán và mã chứng từ thanh toán."),
    ("28", "Thông báo", "outbox_events", "Bảng", "Bổ sung mới", "Hàng đợi sự kiện cần phát thông báo, có người nhận/phòng ban/tổ chức."),
    ("29", "Thông báo", "user_notifications", "Bảng", "Bổ sung mới", "Thông báo đã được tạo cho người dùng từ outbox event."),
    ("30", "Thông báo", "notification_reads", "Bảng trung gian", "Bổ sung mới", "Ghi nhận người dùng đã đọc thông báo nào và thời điểm đọc."),
    ("31", "Kiểm toán và hỗ trợ quyết định", "audit_logs", "Bảng", "Giữ — UC cũ", "Dấu vết kiểm toán các thao tác nghiệp vụ; dùng resource_type/resource_id để tham chiếu đối tượng."),
    ("32", "Kiểm toán và hỗ trợ quyết định", "audit_cases", "Bảng", "Bổ sung mới", "Vụ việc kiểm toán cần theo dõi, mức độ rủi ro, người phụ trách và hạn xử lý."),
    ("33", "Kiểm toán và hỗ trợ quyết định", "audit_case_events", "Bảng", "Bổ sung mới", "Lịch sử thao tác/khắc phục của một vụ việc kiểm toán."),
    ("34", "Kiểm toán và hỗ trợ quyết định", "ai_recommendations", "Bảng", "Bổ sung mới", "Khuyến nghị dựa trên quy tắc, bằng chứng JSON, người tạo và quyết định chấp nhận/bác bỏ."),
    ("35", "Báo cáo", "reporting.sla_policies", "Bảng", "Bổ sung mới", "Chính sách thời hạn xử lý cho từng quy trình của tổ chức."),
    ("36", "Báo cáo", "reporting.purchase_request_facts", "View", "Giữ + chuẩn hóa", "Dữ liệu tổng hợp phiếu, workflow, SLA và tệp đính kèm phục vụ báo cáo."),
    ("37", "Báo cáo", "reporting.daily_procurement_metrics", "View", "Giữ + chuẩn hóa", "Chỉ số mua sắm theo ngày, được tổng hợp từ purchase_request_facts."),
    ("38", "Báo cáo", "reporting.budget_utilization", "View", "Giữ + chuẩn hóa", "Tình hình sử dụng hạn mức theo kỳ ngân sách và trung tâm chi phí."),
]


RELATIONS = [
    ("Tổ chức – phòng ban", "organizations → departments", "1 – N", "Một tổ chức có nhiều phòng ban; mỗi phòng ban thuộc một tổ chức."),
    ("Cây phòng ban", "departments → departments", "1 – N", "Một phòng ban cha có thể có nhiều phòng ban con."),
    ("Phòng ban – người dùng", "departments → users", "1 – N", "Một phòng ban có nhiều người dùng; mỗi người dùng thuộc một phòng ban."),
    ("Người dùng – ảnh chụp quyền", "users → user_role_snapshots", "1 – 0..1", "Một người dùng có tối đa một ảnh chụp role hiện thời."),
    ("Người dùng/phòng ban – phiếu", "users, departments → purchase_requests", "1 – N", "Người dùng là người yêu cầu/người được giao; phòng ban là phạm vi của phiếu."),
    ("Phiếu – dòng hàng", "purchase_requests → purchase_request_items", "1 – N", "Một phiếu có một hay nhiều dòng hàng/dịch vụ."),
    ("Phiếu – lịch sử workflow", "purchase_requests → process_events", "1 – N", "Mọi chuyển trạng thái của phiếu tạo thêm một process event."),
    ("Phiếu – tệp đính kèm", "purchase_requests → purchase_request_attachments", "1 – N", "Một phiếu có thể có nhiều tệp/chứng từ."),
    ("Tổ chức – quy tắc tệp", "organizations → attachment_rules", "1 – N", "Một tổ chức có thể cấu hình nhiều quy tắc tệp đính kèm."),
    ("Tổ chức – kỳ ngân sách", "organizations → budget_periods", "1 – N", "Một tổ chức có nhiều kỳ ngân sách."),
    ("Kỳ – hạn mức", "budget_periods → budget_allocations", "1 – N", "Một kỳ có nhiều hạn mức theo trung tâm chi phí."),
    ("Hạn mức – giữ chỗ/điều chỉnh", "budget_allocations → budget_reservations, budget_adjustments", "1 – N", "Một hạn mức có nhiều bản ghi giữ chỗ/cam kết và nhiều lần điều chỉnh."),
    ("Phiếu – giữ chỗ ngân sách", "purchase_requests → budget_reservations", "1 – N", "Một phiếu có thể phát sinh nhiều bản ghi reservation theo vòng đời xử lý."),
    ("Phiếu – hồ sơ báo giá", "purchase_requests → sourcing_cases", "1 – 0..1", "Một phiếu có tối đa một hồ sơ so sánh báo giá."),
    ("Hồ sơ – báo giá", "sourcing_cases → supplier_quotes", "1 – N", "Một hồ sơ có nhiều báo giá từ các nhà cung cấp."),
    ("Nhà cung cấp – báo giá", "suppliers → supplier_quotes", "1 – N", "Một nhà cung cấp có thể gửi nhiều báo giá."),
    ("Hồ sơ – lịch sử báo giá", "sourcing_cases → sourcing_events", "1 – N", "Các thao tác báo giá được ghi theo hồ sơ so sánh."),
    ("Phiếu – đơn đặt hàng", "purchase_requests → purchase_orders", "1 – 0..1", "Một phiếu được phát hành tối đa một đơn đặt hàng."),
    ("Nhà cung cấp – đơn đặt hàng", "suppliers → purchase_orders", "1 – N", "Một nhà cung cấp có thể có nhiều đơn đặt hàng."),
    ("Đơn hàng – nhận hàng", "purchase_orders → purchase_order_receipts", "1 – N", "Một đơn hàng có thể được nhận nhiều lần."),
    ("Nhận hàng – dòng nhận", "purchase_order_receipts → purchase_order_receipt_items", "1 – N", "Một biên bản nhận có nhiều dòng nhận hàng."),
    ("Dòng phiếu – dòng nhận", "purchase_request_items → purchase_order_receipt_items", "1 – N", "Một dòng hàng có thể được nhận qua nhiều lần."),
    ("Nhận hàng – dòng phiếu", "purchase_order_receipts ↔ purchase_request_items", "N – N", "Quan hệ nhiều–nhiều được triển khai qua bảng trung gian purchase_order_receipt_items."),
    ("Đơn hàng – hóa đơn", "purchase_orders → purchase_invoices", "1 – N", "Một đơn đặt hàng có thể có nhiều hóa đơn (tạm ứng/quyết toán/điều chỉnh)."),
    ("Hóa đơn – sự kiện/thanh toán", "purchase_invoices → invoice_events, invoice_payments", "1 – N", "Một hóa đơn có lịch sử xử lý và có thể nhận nhiều khoản thanh toán."),
    ("Quy tắc/ủy quyền", "organizations, departments, users → approval_rules, approval_delegations", "1 – N", "Quy tắc và ủy quyền thuộc tổ chức/phòng ban; ủy quyền tham chiếu người giao và người nhận."),
    ("Thông báo", "outbox_events → user_notifications", "1 – 0..1", "Một sự kiện hàng đợi tạo tối đa một thông báo đã phát; thông báo có người nhận."),
    ("Trạng thái đã đọc", "user_notifications, users → notification_reads", "1 – N / 1 – N", "Bảng notification_reads lưu liên kết thông báo – người đọc và thời điểm đọc."),
    ("Kiểm toán", "organizations, users → audit_logs, audit_cases", "1 – N", "Audit log/case thuộc tổ chức và có actor/owner; resource_type/resource_id là tham chiếu đa hình, không phải khóa ngoại vật lý."),
    ("Vụ việc kiểm toán", "audit_cases → audit_case_events", "1 – N", "Một vụ việc có nhiều sự kiện theo dõi/khắc phục."),
    ("Khuyến nghị", "organizations, purchase_requests → ai_recommendations", "1 – N", "Khuyến nghị thuộc tổ chức và có thể gắn với một phiếu mua sắm."),
    ("Báo cáo", "organizations → reporting.sla_policies", "1 – N", "Một tổ chức có nhiều chính sách thời hạn xử lý theo quy trình."),
    ("View báo cáo", "Bảng nghiệp vụ → reporting views", "Dẫn xuất", "Các view chỉ đọc lấy dữ liệu từ bảng nghiệp vụ; đây không phải quan hệ khóa ngoại."),
]


FUNCTIONS = [
    ("UC-01", "Đăng nhập hệ thống", "Giữ nguyên", "users, user_role_snapshots", "Keycloak xác thực bên ngoài; DX-OS dùng users.keycloak_subject để ánh xạ tài khoản."),
    ("UC-02", "Tạo phiếu mua sắm", "Giữ nguyên", "purchase_requests, purchase_request_items, process_events, audit_logs", "Giữ luồng tạo phiếu cũ, đồng thời lưu lịch sử nghiệp vụ."),
    ("UC-03", "Sửa phiếu mua sắm", "Giữ nguyên", "purchase_requests, purchase_request_items, process_events, audit_logs", "Giữ kiểm soát phiên bản và lịch sử thao tác."),
    ("UC-04, UC-05", "Xem danh sách/chi tiết phiếu", "Giữ nguyên", "purchase_requests, purchase_request_items, process_events, purchase_request_attachments", "Bổ sung khả năng xem tiến trình, ngân sách, báo giá, đơn hàng tùy role."),
    ("UC-06", "Gửi/gửi lại phiếu duyệt", "Giữ nguyên", "purchase_requests, attachment_rules, purchase_request_attachments, process_events, audit_logs", "Giữ kiểm tra điều kiện tệp; quy tắc phê duyệt mở rộng bằng approval_rules."),
    ("UC-07", "Duyệt phiếu hai cấp", "Giữ nguyên", "purchase_requests, process_events, budget_reservations, approval_rules, approval_delegations, audit_logs", "Giữ manager/finance approval; thêm tuyến theo giá trị và cơ chế ủy quyền."),
    ("UC-08", "Hủy phiếu", "Giữ nguyên", "purchase_requests, process_events, budget_reservations, audit_logs", "Phiếu không biến mất; trạng thái và lịch sử được lưu."),
    ("UC-09", "Quản lý tệp đính kèm", "Giữ nguyên", "attachment_rules, purchase_request_attachments, audit_logs", "Tệp thật ở Nextcloud, PostgreSQL chỉ lưu metadata và checksum."),
    ("UC-10", "Quản lý ngân sách", "Giữ nguyên", "budget_periods, budget_allocations, budget_reservations, budget_adjustments", "Bổ sung phân kỳ ngân sách và lịch sử giữ chỗ/cam kết/giải phóng."),
    ("UC-11", "Xem báo cáo vận hành", "Giữ nguyên", "reporting.sla_policies, reporting.purchase_request_facts, reporting.daily_procurement_metrics, reporting.budget_utilization", "Báo cáo/Metabase đọc từ schema reporting chỉ đọc."),
    ("UC-12", "Duyệt phiếu từ Approval Inbox", "Giữ nguyên", "purchase_requests, process_events, approval_rules, approval_delegations", "Inbox chọn đúng người xử lý theo scope, quy tắc và ủy quyền."),
    ("F-13", "Quản lý danh mục hàng hóa", "Bổ sung mới", "procurement_catalog_items", "Cung cấp hàng hóa/dịch vụ chuẩn để tạo phiếu nhanh và thống nhất dữ liệu."),
    ("F-14", "Quản lý nhà cung cấp", "Bổ sung mới", "suppliers", "Theo dõi hồ sơ, tuân thủ, điều khoản và hiệu suất nhà cung cấp."),
    ("F-15", "So sánh và chọn báo giá", "Bổ sung mới", "sourcing_cases, supplier_quotes, sourcing_events", "Nhập báo giá, chấm điểm, chọn nhà cung cấp và lưu lịch sử quyết định."),
    ("F-16", "Phát hành đơn đặt hàng", "Bổ sung mới", "purchase_orders", "Tạo đơn hàng từ phiếu đã chọn nhà cung cấp để giảm nhập tay."),
    ("F-17", "Giao nhận", "Bổ sung mới", "purchase_order_receipts, purchase_order_receipt_items", "Theo dõi nhận một phần, hỏng/sai hàng và đối chiếu với dòng hàng yêu cầu."),
    ("F-18", "Hóa đơn và thanh toán", "Bổ sung mới", "purchase_invoices, invoice_events, invoice_payments", "Theo dõi kiểm tra hóa đơn, trạng thái và khoản thanh toán."),
    ("F-19", "Thông báo công việc", "Bổ sung mới", "outbox_events, user_notifications, notification_reads", "Gửi và theo dõi thông báo theo người nhận/trạng thái đọc."),
    ("F-20", "Vụ việc kiểm toán", "Bổ sung mới", "audit_cases, audit_case_events", "Tạo và theo dõi việc khắc phục rủi ro; không thay đổi audit_logs cũ."),
    ("F-21", "Khuyến nghị hỗ trợ quyết định", "Bổ sung mới", "ai_recommendations", "Sinh khuyến nghị có bằng chứng, người dùng vẫn là người chấp nhận/bác bỏ."),
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = GRID) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        node = borders.find(tag)
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "5")
        node.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    element = OxmlElement("w:tblHeader")
    element.set(qn("w:val"), "true")
    tr_pr.append(element)


def set_font(run, size=9.0, color=NAVY, bold=False) -> None:
    run.font.name = "Aptos"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def cell_text(cell, value: str, bold=False, color=NAVY, size=8.5) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(str(value))
    set_font(run, size=size, color=color, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)
    set_cell_border(cell)


def add_table(doc: Document, headers: list[str], rows: list[tuple], widths: list[float], status_col: int | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    header = table.rows[0]
    repeat_header(header)
    for index, value in enumerate(headers):
        cell = header.cells[index]
        cell.width = Cm(widths[index])
        set_cell_shading(cell, TEAL)
        cell_text(cell, value, bold=True, color=WHITE, size=8.5)
    for row_values in rows:
        row = table.add_row()
        for index, value in enumerate(row_values):
            cell = row.cells[index]
            cell.width = Cm(widths[index])
            fill = WHITE
            if status_col is not None and index == status_col:
                text = str(value)
                if text.startswith("Giữ"):
                    fill = PALE_GREEN
                elif text.startswith("Bổ sung"):
                    fill = PALE_YELLOW
                elif text.startswith("View"):
                    fill = PALE_BLUE
            set_cell_shading(cell, fill)
            cell_text(cell, value, size=8.1)
    doc.add_paragraph()
    return table


def add_heading(doc: Document, text: str, level: int) -> None:
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True


def add_paragraph(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.1
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        set_font(run, bold=True)
        run = paragraph.add_run(text[len(bold_prefix):])
        set_font(run)
    else:
        set_font(paragraph.add_run(text))


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(2)
        set_font(paragraph.add_run(item), size=9.2)


def add_page_number(footer) -> None:
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(paragraph.add_run("DX-OS Lab · Thống kê cơ sở dữ liệu · Trang "), size=8, color=SLATE)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    paragraph._p.append(begin)
    paragraph._p.append(instr)
    paragraph._p.append(end)


def configure(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.3)
    section.bottom_margin = Cm(1.3)
    section.left_margin = Cm(1.25)
    section.right_margin = Cm(1.25)
    section.header_distance = Cm(0.6)
    section.footer_distance = Cm(0.65)
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    for style_name, size, color in (("Title", 23, TEAL), ("Heading 1", 16, NAVY), ("Heading 2", 12.5, TEAL), ("Heading 3", 10.5, NAVY)):
        style = doc.styles[style_name]
        style.font.name = "Aptos Display" if style_name in {"Title", "Heading 1"} else "Aptos"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), style.font.name)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(header.add_run("DX-OS LAB · DATABASE INVENTORY"), size=8, color=SLATE, bold=True)
    add_page_number(section.footer)


def build() -> None:
    if not SOURCE.exists():
        raise FileNotFoundError(f"Không tìm thấy tài liệu nguồn: {SOURCE}")
    doc = Document()
    configure(doc)

    title = doc.add_heading("THỐNG KÊ BẢNG DỮ LIỆU DX-OS", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(subtitle.add_run("Đối chiếu Nhom_DACN (1).docx với migrations hiện tại"), size=12, color=SLATE)
    issued = doc.add_paragraph()
    issued.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(issued.add_run(f"Cập nhật ngày {date.today().strftime('%d/%m/%Y')} · Phạm vi: CSDL nghiệp vụ DX-OS"), size=9.2, color=SLATE)

    add_heading(doc, "1. Phạm vi và nguyên tắc đối chiếu", 1)
    add_paragraph(doc, "Tài liệu này lấy Nhom_DACN (1).docx làm mốc nghiệp vụ. Các use case cũ UC-01 đến UC-12 được giữ nguyên; bảng/chức năng được triển khai thêm sau đó chỉ được bổ sung, không thay thế hoặc xóa nội dung cũ.")
    add_bullets(doc, [
        "Nguồn đối chiếu kỹ thuật: backend/migrations/000001_bootstrap.sql đến 000019_align_demo_titles_with_amounts.sql.",
        "Phạm vi thống kê: 35 bảng dữ liệu (34 bảng trong schema public và reporting.sla_policies) cùng 3 view báo cáo trong PostgreSQL DX-OS.",
        "Không liệt kê bảng nội bộ của Keycloak, Nextcloud và Metabase vì các bảng đó do từng sản phẩm bên ngoài quản lý; DX-OS chỉ tích hợp qua định danh, metadata tệp và schema reporting.",
        "Ký hiệu 1–N: một bản ghi nguồn liên kết nhiều bản ghi đích; 1–0..1: quan hệ tối đa một bản ghi đích; N–N: quan hệ nhiều–nhiều qua bảng trung gian.",
    ])

    add_heading(doc, "2. Nhóm bảng dữ liệu", 1)
    add_table(doc, ["STT", "Nhóm dữ liệu", "Số đối tượng", "Mục đích nghiệp vụ"], [(str(i + 1), *row) for i, row in enumerate(GROUPS)], [1.1, 6.4, 2.1, 17.5])

    add_heading(doc, "3. Danh mục bảng và view hiện tại", 1)
    add_paragraph(doc, "Cột “Tình trạng” phản ánh mức liên hệ với tài liệu gốc: “Giữ” là nền tảng/use case đã có trong Nhom_DACN (1); “Bổ sung mới” là phần được triển khai thêm sau đó.")
    add_table(doc, ["STT", "Nhóm", "Bảng / đối tượng", "Loại", "Tình trạng", "Mục đích và chức năng"], OBJECTS, [0.8, 4.2, 4.6, 1.6, 2.8, 13.0], status_col=4)

    add_heading(doc, "4. Thống kê quan hệ giữa các bảng", 1)
    add_paragraph(doc, "Các quan hệ dưới đây được tổng hợp từ khóa ngoại và ràng buộc hiện có trong migrations. Với quan hệ đa hình, tài liệu ghi rõ là không có khóa ngoại vật lý.")
    add_table(doc, ["Nhóm quan hệ", "Bảng liên quan", "Lực lượng", "Ý nghĩa dữ liệu"], RELATIONS, [4.0, 8.3, 2.5, 12.3])

    add_heading(doc, "5. Đối chiếu chức năng cũ và phần bổ sung", 1)
    add_paragraph(doc, "Nhóm use case UC-01 đến UC-12 được giữ nguyên theo Nhom_DACN (1). Các dòng F-13 trở đi là chức năng bổ sung của phiên bản DX-OS hiện tại và nêu rõ bảng dữ liệu phục vụ chức năng đó.")
    add_table(doc, ["Mã", "Chức năng", "Tình trạng", "Bảng dữ liệu chính", "Ghi chú cập nhật"], FUNCTIONS, [1.5, 5.2, 2.4, 7.3, 10.7], status_col=2)

    add_heading(doc, "6. Ranh giới dữ liệu với các hệ thống tích hợp", 1)
    integration_rows = [
        ("Keycloak", "CSDL Keycloak riêng", "Đăng nhập/xác thực; DX-OS lưu keycloak_subject trong users để ánh xạ hồ sơ nghiệp vụ."),
        ("Nextcloud", "Kho tệp riêng", "Lưu bytes của tệp; purchase_request_attachments lưu metadata, đường dẫn, checksum và người tải."),
        ("Metabase", "CSDL Metabase riêng", "Đọc schema reporting chỉ đọc của DX-OS để hiển thị dashboard/báo cáo."),
    ]
    add_table(doc, ["Hệ thống", "Dữ liệu quản lý bên ngoài", "Điểm liên kết với DX-OS"], integration_rows, [4.0, 6.5, 16.6])

    add_heading(doc, "7. Kết luận", 1)
    add_paragraph(doc, "CSDL DX-OS hiện kế thừa đầy đủ các chức năng nền tảng trong Nhom_DACN (1): quản lý phiếu, phê duyệt, ngân sách, tệp đính kèm, kiểm toán và báo cáo. Các phần phát triển thêm gồm danh mục, nhà cung cấp/báo giá, quy tắc phê duyệt/ủy quyền, đặt hàng–giao nhận–hóa đơn–thanh toán, thông báo, vụ việc kiểm toán và khuyến nghị hỗ trợ quyết định. Nhờ đó, luồng dữ liệu đi liền mạch từ yêu cầu mua sắm đến thanh toán và kiểm soát sau giao dịch.")
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
