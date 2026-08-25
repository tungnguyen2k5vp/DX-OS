from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from generate_complete_role_test_guide_docx import (
    LIGHT_BLUE,
    LIGHT_GREEN,
    LIGHT_RED,
    LIGHT_YELLOW,
    NAVY,
    SLATE,
    TEAL,
    add_toc,
    bullet,
    callout,
    code,
    configure,
    heading,
    numbered,
    page_number,
    table,
)


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = (
    ROOT
    / "docs"
    / "generated"
    / "Huong_dan_kiem_thu_luong_lien_vai_tro_DX_OS.docx"
)


def replace_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.clear()
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("DX-OS LAB | KỊCH BẢN LIÊN VAI TRÒ")
    run.font.name = "Aptos"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(SLATE)

    footer = section.footer.paragraphs[0]
    footer.clear()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("DX-OS Lab • Một phiếu xuyên suốt • Trang ")
    run.font.name = "Aptos"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(SLATE)
    page_number(footer)


def cover(doc: Document) -> None:
    doc.add_paragraph()
    doc.add_paragraph()
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("DX-OS LAB")
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(34)
    run.font.color.rgb = RGBColor.from_string(TEAL)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("HƯỚNG DẪN KIỂM THỬ LUỒNG LIÊN VAI TRÒ")
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(19)
    run.font.color.rgb = RGBColor.from_string(NAVY)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Một phiếu mua sắm đi xuyên suốt từ đề nghị đến thanh toán và kiểm toán")
    run.italic = True
    run.font.size = Pt(12)

    doc.add_paragraph()
    callout(
        doc,
        "Điểm khác với tài liệu cũ",
        "Tài liệu này không liệt kê rời rạc tất cả chức năng của từng vai trò. Người kiểm thử tạo đúng một phiếu, ghi lại mã phiếu, đổi tài khoản theo từng chặng và tiếp tục xử lý chính phiếu đó cho đến khi hoàn tất.",
        LIGHT_GREEN,
    )
    callout(
        doc,
        "Đối tượng sử dụng",
        "Thành viên nhóm dự án, giảng viên/người đánh giá và người mới cần hiểu cách Nhân viên, Trưởng bộ phận, Tài chính, Kiểm toán, Quản trị và Điều phối AI phối hợp trên cùng một giao dịch.",
        LIGHT_BLUE,
    )
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run(
        f"Phiên bản 2.0 • Đối chiếu mã nguồn nhánh feat/enterprise-operations-review "
        f"• Cập nhật {date.today().strftime('%d/%m/%Y')}"
    )
    doc.add_page_break()


def add_reading_guide(doc: Document) -> None:
    heading(doc, "1. Cách dùng tài liệu", 1)
    doc.add_paragraph(
        "Hãy làm các chặng theo đúng thứ tự. Không tự tạo phiếu mới khi đổi vai trò. Mọi vai trò phải tìm và xử lý cùng mã PR đã ghi ở Chặng 1."
    )
    table(
        doc,
        ["Ký hiệu", "Ý nghĩa", "Việc cần làm"],
        [
            ["ĐẦU VÀO", "Trạng thái hoặc dữ liệu do vai trò trước bàn giao", "Chỉ bắt đầu khi đầu vào đã đúng."],
            ["THAO TÁC", "Các bước thực hiện bằng vai trò hiện tại", "Bấm đúng menu/nút và nhập lý do rõ ràng."],
            ["ĐẦU RA", "Trạng thái và dữ liệu sau khi hoàn thành", "Đối chiếu trước khi đổi tài khoản."],
            ["BÀN GIAO", "Vai trò tiếp theo nhận việc", "Gửi mã PR/PO/hóa đơn và trạng thái hiện tại."],
            ["BẰNG CHỨNG", "Ảnh hoặc mã cần giữ lại", "Dùng để demo và truy vết nếu có lỗi."],
        ],
        [2.7, 6.2, 8.4],
    )
    callout(
        doc,
        "Quy tắc quan trọng nhất",
        "Sau mỗi chặng, ghi trạng thái thực tế vào Phiếu theo dõi ở mục 3.3. Nếu trạng thái không đúng, dừng lại và xử lý lỗi; không chuyển sang vai trò tiếp theo vì các màn hình sau sẽ không thấy phiếu.",
        LIGHT_YELLOW,
    )


def add_flow_map(doc: Document) -> None:
    heading(doc, "2. Bản đồ liên kết giữa các vai trò", 1)
    code(
        doc,
        "QUẢN TRỊ chuẩn bị người dùng, phòng ban và quy tắc\n"
        "        ↓\n"
        "NHÂN VIÊN tạo phiếu → TRƯỞNG BỘ PHẬN yêu cầu sửa\n"
        "        ↑                              ↓\n"
        "        └──── sửa và gửi lại ←─────────┘\n"
        "                       ↓\n"
        "TRƯỞNG BỘ PHẬN duyệt → TÀI CHÍNH duyệt cuối\n"
        "                       ↓\n"
        "TÀI CHÍNH nhập/so sánh báo giá → chọn phương án → tạo đơn từ báo giá\n"
        "                       ↓\n"
        "NHÂN VIÊN xác nhận nhận hàng\n"
        "                       ↓\n"
        "TÀI CHÍNH ghi hóa đơn → đối soát → thanh toán\n"
        "                       ↓\n"
        "KIỂM TOÁN truy vết và xuất gói bằng chứng\n"
        "\n"
        "ĐIỀU PHỐI AI quét khuyến nghị có quy tắc và ghi quyết định; không tự thay đổi phiếu.\n"
        "MỌI TÀI KHOẢN có thể hỏi Trợ lý AI nội bộ về quy trình; câu trả lời phải kèm nguồn."
    )
    table(
        doc,
        ["Vai trò", "Nhận từ ai", "Đầu việc trên cùng phiếu", "Bàn giao cho ai"],
        [
            ["Quản trị DX-OS", "Nhóm triển khai", "Kiểm tra tài khoản, phòng ban và quy tắc", "Nhân viên/Trưởng bộ phận"],
            ["Nhân viên", "Quản trị", "Tạo, bổ sung chứng từ, gửi lại và nhận hàng", "Trưởng bộ phận/Tài chính"],
            ["Trưởng bộ phận", "Nhân viên", "Yêu cầu sửa rồi phê duyệt nhu cầu phòng ban", "Nhân viên/Tài chính"],
            ["Tài chính", "Trưởng bộ phận", "Duyệt cuối, báo giá, đơn hàng, hóa đơn và thanh toán", "Nhân viên/Kiểm toán"],
            ["Kiểm toán", "Toàn bộ chuỗi", "Đối chiếu dấu vết, mở hồ sơ và xuất bằng chứng", "Người xử lý khắc phục"],
            ["Điều phối AI", "Dữ liệu đã phát sinh", "Quét khuyến nghị theo quy tắc, đọc bằng chứng và ghi quyết định", "Tài chính/Kiểm toán/Quản trị"],
            ["Trợ lý AI nội bộ", "Tài liệu Markdown trong kho tri thức", "Trả lời câu hỏi quy trình có dẫn nguồn; không ghi dữ liệu nghiệp vụ", "Mọi người dùng đã xác thực"],
        ],
        [3.2, 3.0, 7.6, 3.5],
    )


def add_preparation(doc: Document) -> None:
    doc.add_page_break()
    heading(doc, "3. Chuẩn bị một lần trước khi bắt đầu", 1)
    heading(doc, "3.0 Kiểm tra các dịch vụ cần cho kịch bản", 2)
    table(
        doc,
        ["Thành phần", "Địa chỉ kiểm tra", "Kết quả cần thấy"],
        [
            ["DX-OS Web", "http://localhost:4200", "Trang chuyển sang đăng nhập hoặc Tổng quan."],
            ["Go API", "http://localhost:8081/health/ready", "Phản hồi status = ready."],
            ["Keycloak", "http://localhost:8080", "Máy chủ xác thực phản hồi; đăng nhập DX-OS chuyển về realm dx-os."],
            ["Metabase", "http://localhost:3000", "Trang đăng nhập hoặc dashboard Metabase mở được."],
            ["Nextcloud", "http://localhost:8082", "Trang Nextcloud phản hồi; tệp DX-OS vẫn tải qua Go API."],
            ["Ollama local", "Kiểm tra trong menu Trợ lý AI nội bộ", "Nhãn AI local sẵn sàng và có tên mô hình."],
        ],
        [3.0, 7.0, 7.5],
    )
    callout(
        doc,
        "Không bắt đầu khi API hoặc Keycloak chưa sẵn sàng",
        "Nếu chỉ một thành phần phụ như Metabase hoặc Ollama lỗi, ghi nhận riêng và vẫn có thể kiểm thử nghiệp vụ mua sắm. Nếu API/Keycloak lỗi, các bước liên vai trò sẽ cho kết quả sai hoặc không đăng nhập được.",
        LIGHT_YELLOW,
    )
    heading(doc, "3.1 Mở phiên đăng nhập riêng", 2)
    doc.add_paragraph(
        "Nên dùng sáu cửa sổ ẩn danh hoặc sáu hồ sơ trình duyệt. Nếu chỉ dùng một cửa sổ, phải Đăng xuất hoàn toàn trước khi đổi tài khoản."
    )
    table(
        doc,
        ["Vai trò", "Tên đăng nhập", "Nơi lấy mật khẩu", "Tên cửa sổ nên đặt"],
        [
            ["Nhân viên", "employee.demo", "data/runtime/employee-demo.txt", "01 - Nhân viên"],
            ["Trưởng bộ phận", "manager.demo", "data/runtime/manager-demo.txt", "02 - Trưởng bộ phận"],
            ["Tài chính", "finance.demo", "data/runtime/finance-demo.txt", "03 - Tài chính"],
            ["Kiểm toán", "auditor.demo", "data/runtime/auditor-demo.txt", "04 - Kiểm toán"],
            ["Quản trị", "admin.demo", "data/runtime/admin-demo.txt", "05 - Quản trị"],
            ["Điều phối AI", "ai.operator.demo", "data/runtime/ai-operator-demo.txt", "06 - Điều phối AI"],
        ],
        [3.2, 3.5, 6.5, 4.1],
    )
    callout(
        doc,
        "Không đưa mật khẩu vào tài liệu hoặc GitHub",
        "Các file trong data/runtime chỉ dùng trên máy chạy dự án. Khi chụp ảnh báo cáo, không chụp nội dung mật khẩu, token hoặc biến môi trường.",
        LIGHT_RED,
    )

    heading(doc, "3.2 Bộ dữ liệu dùng chung cho một phiếu", 2)
    table(
        doc,
        ["Trường", "Giá trị mẫu", "Lý do chọn"],
        [
            ["Tiêu đề", "Demo liên vai trò - thiết bị văn phòng - [ngày giờ]", "Dễ tìm và tránh trùng với lần test trước."],
            ["Lý do", "Trang bị thiết bị cho nhóm dự án để làm việc và báo cáo.", "Đủ độ dài, mô tả được nhu cầu."],
            ["Trung tâm chi phí", "CC-GENERAL", "Dùng dữ liệu mặc định của môi trường demo."],
            ["Dòng 1", "Laptop văn phòng • 2 chiếc • 22.000.000 VND/chiếc", "Thành tiền 44.000.000 VND."],
            ["Dòng 2", "Bản quyền phần mềm một năm • 3 gói • 6.000.000 VND/gói", "Thành tiền 18.000.000 VND."],
            ["Tổng phiếu", "62.000.000 VND", "Kích hoạt luồng giá trị lớn và so sánh báo giá."],
            ["Chứng từ", "Một file PDF báo giá dưới 10 MB", "Đáp ứng quy tắc chứng từ của phiếu giá trị lớn."],
        ],
        [3.1, 9.2, 5.0],
    )
    bullet(doc, "Nếu chưa có PDF, tạo một văn bản ngắn có tiêu đề “Báo giá demo”, lưu thành PDF rồi tải lên với loại tài liệu Báo giá.")
    bullet(doc, "Không dùng lại chính xác tiêu đề của lần test cũ; thêm ngày giờ ở cuối tiêu đề.")

    heading(doc, "3.3 Phiếu theo dõi xuyên suốt", 2)
    table(
        doc,
        ["Thông tin phải ghi", "Giá trị khi test"],
        [
            ["Mã phiếu PR", "PR-________________________"],
            ["Mã định danh trong URL", "____________________________________________"],
            ["Mã báo giá đã chọn", "____________________________________________"],
            ["Nhà cung cấp đã chọn", "____________________________________________"],
            ["Mã đơn hàng PO", "PO-________________________"],
            ["Số hóa đơn", "INV-DEMO-__________________"],
            ["Mã hồ sơ kiểm toán", "____________________________________________"],
            ["Kết quả cuối", "☐ Đã thanh toán  ☐ Không đạt  ☐ Dừng ở chặng ____"],
        ],
        [6.3, 11.0],
    )


def stage(
    doc: Document,
    number: str,
    title: str,
    account: str,
    menu: str,
    input_state: str,
    steps: list[str],
    output_state: str,
    handoff: str,
    evidence: str,
    warning: str = "",
) -> None:
    heading(doc, f"{number}. {title}", 2)
    table(
        doc,
        ["Đăng nhập", "Mở tại", "Đầu vào", "Đầu ra"],
        [[account, menu, input_state, output_state]],
        [3.6, 4.8, 4.4, 4.5],
    )
    doc.add_paragraph("Thao tác:")
    for value in steps:
        numbered(doc, value)
    callout(doc, "Bàn giao", handoff, LIGHT_GREEN)
    callout(doc, "Bằng chứng cần giữ", evidence, LIGHT_BLUE)
    if warning:
        callout(doc, "Nếu không đúng", warning, LIGHT_YELLOW)


def add_preflight_admin(doc: Document) -> None:
    doc.add_page_break()
    heading(doc, "4. Chặng chuẩn bị — Quản trị bảo đảm hai tài khoản cùng phòng ban", 1)
    stage(
        doc,
        "4.1",
        "Kiểm tra phạm vi trước khi tạo phiếu",
        "admin.demo",
        "Quản trị",
        "Các tài khoản demo đã tồn tại.",
        [
            "Mở danh sách Người dùng nghiệp vụ trong trang Quản trị.",
            "Tìm employee.demo và manager.demo.",
            "Kiểm tra hai tài khoản cùng một phòng ban, khuyến nghị là Phòng ban chung (General Department).",
            "Kiểm tra employee.demo đang Hoạt động và có vai trò Nhân viên; manager.demo đang Hoạt động và có vai trò Trưởng bộ phận.",
            "Không thêm vai trò Tài chính hoặc Quản trị cho hai tài khoản chỉ để vượt lỗi phân quyền.",
        ],
        "Employee và Manager cùng phòng ban, đều hoạt động.",
        "Báo cho người dùng employee.demo bắt đầu Chặng 5. Không dùng Admin để tạo hoặc duyệt phiếu.",
        "Ảnh hai hồ sơ hiển thị cùng phòng ban và đúng vai trò.",
        "Nếu khác phòng ban, Manager sẽ không nhìn thấy phiếu dù Employee đã gửi. Sửa phòng ban, lưu, rồi đăng xuất/đăng nhập lại hai tài khoản để nhận token mới.",
    )
    stage(
        doc,
        "4.2",
        "Xác nhận quy tắc buộc phiếu đi qua hai cấp duyệt",
        "admin.demo",
        "Ủy quyền và quy tắc → Quy tắc theo giá trị phiếu",
        "Employee và Manager đã đúng phòng ban.",
        [
            "Tìm một quy tắc đang dùng, tiền tệ VND, có khoảng tiền bao phủ 62.000.000 VND.",
            "Kiểm tra dòng quy trình hiển thị Trưởng bộ phận → Tài chính, nghĩa là cả hai ô duyệt đều được bật.",
            "Ưu tiên quy tắc mặc định đang có. Không tạo thêm quy tắc chồng lấn nếu quy tắc hiện tại đã bao phủ số tiền test.",
            "Nếu không có quy tắc phù hợp, tạo quy tắc tạm tên “Kiểm thử liên vai trò 50–100 triệu”, từ 50.000.000 đến 100.000.000 VND, bật cả Trưởng bộ phận duyệt và Tài chính duyệt.",
            "Chụp lại tên quy tắc được dùng. Nếu đã tạo quy tắc tạm, ghi chú để tạm dừng sau khi test xong.",
        ],
        "Có đúng một quy tắc ưu tiên áp dụng cho 62.000.000 VND và yêu cầu đủ hai cấp duyệt.",
        "Báo cho employee.demo bắt đầu tạo phiếu; giữ ảnh quy tắc để giải thích nếu trạng thái chuyển khác dự kiến.",
        "Ảnh tên quy tắc, khoảng tiền, trạng thái Đang dùng và chuỗi Trưởng bộ phận → Tài chính.",
        "Nếu quy tắc chỉ có một cấp, hệ thống có thể bỏ qua Manager hoặc kết thúc ngay sau Manager; khi đó tài liệu không thể tiếp tục đúng luồng hai cấp.",
    )


def add_employee_create(doc: Document) -> None:
    doc.add_page_break()
    heading(doc, "5. Chặng tạo nhu cầu — Nhân viên tạo và gửi phiếu", 1)
    stage(
        doc,
        "5.1",
        "Tạo bản nháp và ghi mã phiếu",
        "employee.demo",
        "Phiếu mua sắm → Tạo phiếu",
        "Chưa có phiếu.",
        [
            "Nhập tiêu đề và lý do theo mục 3.2.",
            "Giữ Trung tâm chi phí là CC-GENERAL và Tiền tệ là VND.",
            "Nhập hai dòng hàng. Kiểm tra tổng tạm tính là 62.000.000 VND.",
            "Bấm Lưu bản nháp.",
            "Ở trang chi tiết, ghi Mã phiếu PR và mã định danh trong URL vào Phiếu theo dõi mục 3.3.",
        ],
        "Trạng thái Bản nháp; mã PR đã được tạo.",
        "Chưa đổi tài khoản. Employee tiếp tục tải chứng từ và gửi duyệt ở Chặng 5.2.",
        "Ảnh phần thông tin phiếu, hai dòng hàng, tổng tiền và trạng thái Bản nháp.",
        "Nếu tổng không phải 62.000.000 VND, kiểm tra lại số lượng và đơn giá trước khi gửi.",
    )
    stage(
        doc,
        "5.2",
        "Tải chứng từ và gửi cho Trưởng bộ phận",
        "employee.demo",
        "Chi tiết phiếu vừa tạo",
        "Phiếu đang ở Bản nháp.",
        [
            "Tại Tài liệu đính kèm, chọn loại Báo giá và tải file PDF dưới 10 MB.",
            "Kiểm tra tên file, dung lượng, loại Báo giá và nút Tải xuống đã xuất hiện.",
            "Bấm Tải xuống một lần để xác nhận tệp lấy về mở được và đúng nội dung; đây là kiểm tra xuyên suốt cho chức năng quản lý tệp.",
            "Tại Hành động, chọn Gửi duyệt; nhập ghi chú “Đã đủ nhu cầu và chứng từ, đề nghị xem xét”.",
            "Xác nhận trạng thái không còn là Bản nháp.",
            "Mở Thông báo hoặc Timeline để thấy sự kiện gửi duyệt.",
        ],
        "Trạng thái Đã gửi (SUBMITTED); Employee không còn sửa trực tiếp.",
        "Gửi cho manager.demo: mã PR, tổng 62.000.000 VND và lời nhắn “Vui lòng yêu cầu chỉnh sửa để test vòng phản hồi”.",
        "Ảnh trạng thái Đã gửi, file đính kèm và Timeline có hành động Gửi duyệt.",
        "Nếu vẫn Bản nháp, bạn mới chỉ lưu phiếu. Nếu hệ thống báo thiếu chứng từ, kiểm tra file đã tải xong và đúng loại Báo giá.",
    )


def add_manager_roundtrip(doc: Document) -> None:
    doc.add_page_break()
    heading(doc, "6. Chặng phản hồi — Trưởng bộ phận trả phiếu, Nhân viên sửa và gửi lại", 1)
    stage(
        doc,
        "6.1",
        "Yêu cầu chỉnh sửa trên chính phiếu vừa nhận",
        "manager.demo",
        "Phê duyệt hoặc Việc của tôi",
        "Mã PR đang Đã gửi và thuộc cùng phòng ban.",
        [
            "Tìm đúng mã PR đã ghi ở mục 3.3; không chọn một phiếu demo khác.",
            "Mở chi tiết và đối chiếu người yêu cầu là employee.demo, tổng tiền 62.000.000 VND.",
            "Trong Tài liệu đính kèm, bấm Tải xuống file Báo giá để chứng minh Manager đọc được tệp nhưng không có nút Xóa.",
            "Bấm Yêu cầu chỉnh sửa.",
            "Nhập lý do “Bổ sung thời gian cần hàng vào lý do mua sắm” rồi xác nhận.",
            "Kiểm tra Timeline có tên manager.demo và lý do vừa nhập.",
        ],
        "Trạng thái Yêu cầu chỉnh sửa (CHANGES_REQUESTED).",
        "Bàn giao lại employee.demo cùng mã PR và nội dung phải bổ sung. Đây là vòng quay ngược trong cùng một luồng, không tạo phiếu mới.",
        "Ảnh trạng thái Yêu cầu chỉnh sửa và Timeline hiển thị đầy đủ lý do.",
        "Nếu Manager không thấy phiếu, kiểm tra trạng thái Đã gửi và phòng ban ở Chặng 4; không sửa quyền tùy tiện.",
    )
    stage(
        doc,
        "6.2",
        "Nhân viên sửa rồi gửi duyệt lại",
        "employee.demo",
        "Việc của tôi → mở đúng mã PR",
        "Phiếu đang Yêu cầu chỉnh sửa.",
        [
            "Đọc lý do chỉnh sửa trong Timeline hoặc Thông báo.",
            "Bấm Sửa phiếu.",
            "Bổ sung vào Lý do: “Cần nhận hàng trước ngày báo cáo dự án”.",
            "Lưu thay đổi rồi kiểm tra hai dòng hàng và file đính kèm vẫn còn.",
            "Bấm Gửi duyệt lại và nhập ghi chú “Đã bổ sung thời gian cần hàng”.",
        ],
        "Phiếu trở lại trạng thái Đã gửi; phiên bản phiếu tăng, dữ liệu cũ không bị mất.",
        "Bàn giao lại manager.demo cùng mã PR để phê duyệt chính thức.",
        "Ảnh lý do đã bổ sung, trạng thái Đã gửi và Timeline có hành động Gửi duyệt lại.",
        "Nếu không có nút Sửa, kiểm tra bạn đang đăng nhập employee.demo và phiếu thuộc chính tài khoản này.",
    )


def add_approvals(doc: Document) -> None:
    doc.add_page_break()
    heading(doc, "7. Chặng phê duyệt — Trưởng bộ phận chuyển sang Tài chính", 1)
    stage(
        doc,
        "7.1",
        "Trưởng bộ phận phê duyệt nhu cầu",
        "manager.demo",
        "Phê duyệt → mở mã PR",
        "Phiếu đã được Employee gửi duyệt lại.",
        [
            "Đối chiếu nội dung đã sửa, chứng từ, tổng tiền và phòng ban.",
            "Bấm Phê duyệt.",
            "Nhập lý do “Nhu cầu hợp lệ, đã kiểm tra mục đích và thời gian cần hàng”.",
            "Tải lại trang và kiểm tra Timeline ghi manager.demo là người thực hiện.",
        ],
        "Trạng thái Trưởng bộ phận đã duyệt (MANAGER_APPROVED).",
        "Bàn giao finance.demo: mã PR, tổng tiền, chứng từ và thông báo phiếu đã qua cấp phòng ban.",
        "Ảnh trạng thái mới, lý do phê duyệt và Timeline.",
        "Nếu phiếu biến mất khỏi hàng chờ của Manager sau khi bấm duyệt, đó là đúng; hãy tìm bằng mã PR ở Phiếu mua sắm để xem lại.",
    )
    stage(
        doc,
        "7.2",
        "Tài chính phê duyệt cuối và kiểm tra ngân sách",
        "finance.demo",
        "Phê duyệt → mở mã PR; sau đó mở Ngân sách",
        "Phiếu đang Trưởng bộ phận đã duyệt.",
        [
            "Kiểm tra tổng tiền 62.000.000 VND, chứng từ và lịch sử yêu cầu sửa.",
            "Bấm Phê duyệt và nhập lý do “Ngân sách phù hợp, chuyển sang lấy báo giá”.",
            "Mở Ngân sách, tìm trung tâm chi phí CC-GENERAL và kiểm tra số đang giữ/đã cam kết thay đổi hợp lý.",
            "Quay lại chi tiết phiếu để xác nhận trạng thái Đã phê duyệt.",
        ],
        "Trạng thái Đã phê duyệt (APPROVED); phiếu đủ điều kiện sang mua hàng.",
        "Finance tiếp tục xử lý cùng mã PR tại So sánh báo giá; chưa bàn giao cho Employee.",
        "Ảnh trạng thái Đã phê duyệt và thẻ ngân sách CC-GENERAL.",
        "Nếu Finance không thấy nút Phê duyệt, phiếu có thể chưa qua Manager hoặc tài khoản đang sai vai trò.",
    )


def add_sourcing_order(doc: Document) -> None:
    doc.add_page_break()
    heading(doc, "8. Chặng mua hàng — Tài chính chọn báo giá và tạo đơn được điền sẵn", 1)
    stage(
        doc,
        "8.1",
        "Nhập hai báo giá và chọn một nhà cung cấp",
        "finance.demo",
        "So sánh báo giá",
        "Phiếu 62.000.000 VND đã được phê duyệt.",
        [
            "Tìm đúng mã PR trong danh sách Chờ báo giá.",
            "Bấm Nhập báo giá. Trang tự cuộn đến khu vực nhập và đặt con trỏ ở trường Nhà cung cấp; không cần tự kéo xuống.",
            "Chọn mã NCC-CNTT-ML (Công ty TNHH Công nghệ Minh Long), nhập số báo giá BG-A-[sáu số cuối mã PR], tổng giá 62.000.000 VND, ngày giao sau hôm nay 14 ngày, bảo hành 24 tháng và điều khoản thanh toán; bấm Lưu báo giá.",
            "Mở lại cùng PR, chọn NCC-CNTT-TC (Công ty Cổ phần Thiết bị Thành Công), nhập số báo giá BG-B-[sáu số cuối mã PR], tổng giá 63.200.000 VND, ngày giao sau hôm nay 16 ngày, bảo hành 12 tháng và điều khoản thanh toán; bấm Lưu báo giá.",
            "Đối chiếu Điểm giá, Điểm tiến độ, Điểm chất lượng, Điểm tuân thủ và Điểm tổng hợp. Các điểm có đơn vị /100 và do hệ thống tự tính từ giá, ngày giao cùng hồ sơ nhà cung cấp; người dùng không nhập điểm bằng tay.",
            "Bấm Chọn báo giá trên phương án Công ty TNHH Công nghệ Minh Long và xác nhận hộp thoại. Giao diện hiện tại tự ghi lý do chuẩn, không mở ô nhập lý do riêng.",
            "Kiểm tra chỉ báo giá vừa chọn có nhãn Đã chọn; báo giá còn lại có nhãn Không chọn. Ghi mã báo giá và nhà cung cấp vào Phiếu theo dõi mục 3.3.",
        ],
        "Hồ sơ báo giá ở trạng thái Đã chọn; chỉ một báo giá được chọn và xuất hiện nút Tạo đơn hàng ngay trên phương án thắng.",
        "Finance không mở form đặt hàng thủ công. Tiếp tục bấm Tạo đơn hàng tại chính báo giá đã chọn ở Chặng 8.2.",
        "Ảnh hai báo giá, điểm có đơn vị /100, nhãn Đã chọn và mã báo giá thắng.",
        "Nếu không thấy phiếu, kiểm tra phiếu đã Được phê duyệt. Nếu nút chọn bị lỗi, tải lại vì phiên bản báo giá có thể vừa thay đổi. Không sửa điểm trực tiếp trong cơ sở dữ liệu.",
    )
    stage(
        doc,
        "8.2",
        "Tạo đơn từ báo giá đã chọn và phát hành",
        "finance.demo",
        "So sánh báo giá → nút Tạo đơn hàng trên báo giá Đã chọn",
        "Phiếu đã có đúng một báo giá được chọn.",
        [
            "Bấm Tạo đơn hàng ngay cạnh nhãn Đã chọn. Hệ thống chuyển sang Đặt hàng & giao nhận và tự cuộn đến form Tạo đơn hàng.",
            "Đối chiếu các trường đã điền sẵn: đúng mã PR, nhà cung cấp Minh Long, mã tham chiếu bằng số báo giá BG-A-..., ngày giao dự kiến và ghi chú gồm nguồn báo giá, bảo hành, điều khoản thanh toán.",
            "Không đổi nhà cung cấp. Nếu muốn bổ sung, chỉ thêm ghi chú “Đã đối chiếu với phương án được chọn tại So sánh báo giá”.",
            "Bấm Phát hành đơn hàng.",
            "Tìm lại đúng mã PR trong Dòng giao nhận, kiểm tra mã PO mới, nhà cung cấp và ngày giao; ghi mã PO vào Phiếu theo dõi mục 3.3.",
        ],
        "Có mã PO; trạng thái Đang giao. Dữ liệu đơn hàng khớp báo giá được chọn mà không nhập lại thủ công.",
        "Bàn giao employee.demo: mã PR, mã PO, ngày giao dự kiến và yêu cầu xác nhận đúng số lượng thực nhận.",
        "Ảnh mã PR, mã PO, nhà cung cấp và ngày giao dự kiến trên cùng một dòng.",
        "Nếu form không tự mở hoặc không điền sẵn, quay lại So sánh báo giá và bấm đúng nút Tạo đơn hàng của phương án Đã chọn. Với phiếu từ 50.000.000 VND, backend bắt buộc phải có báo giá thắng trước khi tạo PO.",
    )


def add_receipt(doc: Document) -> None:
    doc.add_page_break()
    heading(doc, "9. Chặng giao nhận — Nhân viên xác nhận hàng cho đơn của Tài chính", 1)
    stage(
        doc,
        "9.1",
        "Lập biên bản nhận đủ hàng",
        "employee.demo",
        "Đặt hàng & giao nhận",
        "Finance đã phát hành PO cho chính mã PR.",
        [
            "Tìm dòng có đúng mã PR và mã PO đã bàn giao.",
            "Bấm Ghi nhận/Lập biên bản giao nhận.",
            "Chọn Kết quả giao nhận là Nhận đủ và Ngày nhận là ngày hiện tại.",
            "Đối với từng dòng, nhập đúng số lượng thực nhận: Laptop 2 chiếc, Bản quyền 3 gói; chọn tình trạng Chấp nhận.",
            "Nhập ghi chú “Đã kiểm tra đủ số lượng, đúng chủng loại” và bấm Lưu biên bản giao nhận.",
            "Tải lại và kiểm tra trạng thái Đã nhận.",
        ],
        "Đơn hàng có biên bản nhận đủ; trạng thái Đã nhận (RECEIVED).",
        "Bàn giao finance.demo: mã PR, mã PO, ngày nhận và xác nhận số lượng đã khớp.",
        "Ảnh biên bản giao nhận, số lượng từng dòng và trạng thái Đã nhận.",
        "Finance không được tự nhận hàng thay Employee. Nếu Employee không thấy nút, kiểm tra người yêu cầu của PR và trạng thái PO.",
    )


def add_invoice(doc: Document) -> None:
    doc.add_page_break()
    heading(doc, "10. Chặng thanh toán — Tài chính đối soát ba bên", 1)
    stage(
        doc,
        "10.1",
        "Ghi hóa đơn khớp đơn hàng và biên nhận",
        "finance.demo",
        "Hóa đơn & thanh toán",
        "PO đã được Employee xác nhận nhận đủ.",
        [
            "Tìm dòng có đúng mã PR và mã PO.",
            "Bấm Ghi hóa đơn.",
            "Nhập Số hóa đơn INV-DEMO-[mã PR], Ngày phát hành là hôm nay, Hạn thanh toán sau hôm nay và Số tiền 62.000.000 VND.",
            "Giữ Tiền tệ VND, nhập ghi chú “Hóa đơn khớp PO và biên nhận”, rồi bấm Lưu hóa đơn.",
            "Ghi số hóa đơn vào Phiếu theo dõi mục 3.3.",
            "Kiểm tra cột Đối soát hiển thị Khớp hoặc trạng thái tương đương.",
        ],
        "Hóa đơn ở Đã ghi nhận (RECORDED); đối soát đủ điều kiện xác minh.",
        "Finance tiếp tục Chặng 10.2 trên cùng dòng hóa đơn.",
        "Ảnh mã PR, PO, số hóa đơn, số tiền và kết quả đối soát.",
        "Nếu hiển thị Chờ nhận hàng hoặc Sai lệch, quay lại Chặng 9 kiểm tra biên bản và số lượng; không cố thanh toán để bỏ qua sai lệch.",
    )
    stage(
        doc,
        "10.2",
        "Xác minh và ghi nhận thanh toán",
        "finance.demo",
        "Hóa đơn & thanh toán → đúng dòng INV-DEMO",
        "Hóa đơn đã ghi nhận và đối soát khớp.",
        [
            "Bấm Xác minh; tải lại và kiểm tra trạng thái hóa đơn là Đã xác minh.",
            "Bấm Thanh toán.",
            "Nhập Số tiền đợt này 62.000.000 VND, Mã tham chiếu PAY-[mã PR], Ngày thanh toán hôm nay và ghi chú “Thanh toán đủ sau đối soát”.",
            "Bấm Ghi nhận đợt thanh toán.",
            "Kiểm tra Đã trả 62.000.000 VND, Còn lại 0 VND và trạng thái Đã thanh toán.",
        ],
        "Hóa đơn Đã thanh toán (PAID); số tiền còn lại bằng 0 VND.",
        "Bàn giao auditor.demo toàn bộ mã PR, PO, hóa đơn và mã thanh toán để kiểm tra độc lập.",
        "Ảnh trạng thái Đã xác minh, Đã thanh toán và số tiền còn lại 0 VND.",
        "Nếu nút Thanh toán chưa xuất hiện, hóa đơn chưa được Xác minh. Nếu số tiền vượt phần còn lại, hệ thống phải từ chối.",
    )


def add_audit_ai(doc: Document) -> None:
    doc.add_page_break()
    heading(doc, "11. Chặng kiểm soát — Kiểm toán, khuyến nghị và Trợ lý AI", 1)
    stage(
        doc,
        "11.1",
        "Kiểm toán truy vết từ yêu cầu đến thanh toán",
        "auditor.demo",
        "Kiểm toán; Phiếu mua sắm; Hóa đơn & thanh toán",
        "Có mã PR, PO, hóa đơn và thanh toán đã hoàn tất.",
        [
            "Mở Phiếu mua sắm, tìm mã PR và đọc Timeline từ lúc tạo đến khi phê duyệt.",
            "Đối chiếu manager.demo là người yêu cầu sửa/phê duyệt; finance.demo là người duyệt cuối, phát hành PO và thanh toán; employee.demo là người nhận hàng.",
            "Mở Trung tâm kiểm toán, lọc theo loại đối tượng hoặc thời gian của buổi test.",
            "Tại Xuất gói bằng chứng, dán mã định danh lấy từ URL của phiếu rồi bấm Tải gói bằng chứng.",
            "Mở hồ sơ kiểm toán với tiêu đề “Đối chiếu luồng liên vai trò [mã PR]”, mức Trung bình, mô tả kết quả và liên kết mã định danh phiếu.",
            "Ghi mã hồ sơ kiểm toán vào Phiếu theo dõi mục 3.3.",
        ],
        "Có chuỗi bằng chứng liên tục và hồ sơ kiểm toán gắn với cùng phiếu.",
        "Nếu có sai lệch, bàn giao đúng người chịu trách nhiệm; Auditor không sửa trực tiếp PR, PO hay hóa đơn.",
        "Gói bằng chứng tải xuống, ảnh Timeline và mã hồ sơ kiểm toán.",
        "Nếu tìm theo mã PR không ra nhật ký, dùng khoảng thời gian và loại đối tượng. Mã theo dõi thao tác dùng để gom các sự kiện cùng một lần xử lý, không phải mã phiếu.",
    )
    stage(
        doc,
        "11.2",
        "Đọc khuyến nghị AI liên quan cùng mã PR",
        "ai.operator.demo",
        "Khuyến nghị",
        "Phiếu 62.000.000 VND đã tạo dữ liệu giá trị lớn.",
        [
            "Tìm thẻ có mã PR hoặc tiêu đề của phiếu; loại thường gặp là Rà soát phiếu giá trị lớn.",
            "Mở Bằng chứng có thể giải thích và đối chiếu Mã phiếu, Giá trị phiếu 62.000.000 VND với dữ liệu nguồn.",
            "Nhập lý do “Đã đối chiếu báo giá, phê duyệt, giao nhận và thanh toán đầy đủ”.",
            "Chọn Chấp nhận nếu khuyến nghị là hợp lý, hoặc Bác bỏ/Bỏ qua nếu dữ liệu thực tế chứng minh cảnh báo không còn phù hợp.",
            "Tải lại trang để chắc chắn quyết định và lý do vẫn còn.",
        ],
        "Khuyến nghị có quyết định của con người; PR/PO/hóa đơn không bị AI tự thay đổi.",
        "Bàn giao kết quả cho Auditor hoặc Admin như một bằng chứng kiểm soát bổ sung.",
        "Ảnh thẻ khuyến nghị, bằng chứng tiếng Việt có đơn vị và lý do quyết định.",
        "Nếu chưa có khuyến nghị, chờ tiến trình nền chạy hoặc tải lại sau; không tạo dữ liệu giả trực tiếp trong database. Điều phối AI không được duyệt phiếu hay thanh toán.",
    )
    stage(
        doc,
        "11.3",
        "Hỏi Trợ lý AI nội bộ và kiểm tra nguồn",
        "auditor.demo (hoặc bất kỳ tài khoản đã đăng nhập)",
        "Trợ lý AI nội bộ",
        "Ollama local và mô hình đã cấu hình đang sẵn sàng.",
        [
            "Kiểm tra nhãn trạng thái hiển thị AI local sẵn sàng; phần Trạng thái local có tên mô hình, số tài liệu trong kho tri thức và thông báo kết nối thành công.",
            "Chọn câu hỏi mẫu “Phiếu mua sắm từ 20 triệu cần những tài liệu gì?” hoặc tự nhập câu hỏi có ít nhất 3 ký tự.",
            "Bấm Hỏi AI local và chờ trả lời. Lần đầu với mô hình 4B có thể cần khoảng 20–90 giây.",
            "Kiểm tra câu trả lời có chỉ dẫn nguồn dạng [1], [2] và phần Nguồn nội bộ liệt kê tên, đường dẫn cùng đoạn trích tài liệu.",
            "Đối chiếu nội dung với quy tắc thực tế: phiếu từ ngưỡng cấu hình cần loại tệp bắt buộc trước khi gửi. Nếu nguồn không đủ, câu trả lời phải nói rõ thay vì tự suy đoán.",
            "Xác nhận việc hỏi AI không tạo, duyệt, sửa, hủy phiếu, phát hành đơn hàng hoặc thanh toán.",
        ],
        "Có câu trả lời tiếng Việt dựa trên tài liệu nội bộ, kèm nguồn để người dùng kiểm tra; dữ liệu nghiệp vụ không đổi.",
        "Kết quả chỉ là trợ giúp tra cứu. Người có đúng vai trò vẫn phải tự thực hiện và chịu trách nhiệm cho thao tác nghiệp vụ.",
        "Ảnh trạng thái AI local, câu hỏi, câu trả lời và ít nhất một nguồn nội bộ được mở rộng.",
        "Nếu báo AI local chưa sẵn sàng, bấm Kiểm tra lại. Kiểm tra dịch vụ Ollama, biến AI_ENABLED và mô hình đã được tải; không nhầm lỗi này với Trung tâm khuyến nghị theo quy tắc.",
    )


def add_final_verification(doc: Document) -> None:
    doc.add_page_break()
    heading(doc, "12. Kiểm tra cuối — Chứng minh các vai trò thật sự liên kết", 1)
    table(
        doc,
        ["Điểm nối", "Điều kiện đạt", "Kết quả"],
        [
            ["Employee → Manager", "Manager thấy đúng PR do employee.demo tạo, cùng phòng ban.", "☐ Đạt  ☐ Không đạt"],
            ["Manager → Employee", "Lý do yêu cầu sửa xuất hiện; Employee sửa cùng PR, không tạo PR mới.", "☐ Đạt  ☐ Không đạt"],
            ["Manager → Finance", "Finance chỉ xử lý sau khi Manager duyệt; Timeline giữ đủ hai cấp.", "☐ Đạt  ☐ Không đạt"],
            ["Finance → Employee", "PO của Finance xuất hiện trong Giao nhận của Employee.", "☐ Đạt  ☐ Không đạt"],
            ["Employee → Finance", "Biên nhận của Employee làm hóa đơn đủ điều kiện đối soát/xác minh.", "☐ Đạt  ☐ Không đạt"],
            ["Finance → Auditor", "Auditor truy được PR, PO, receipt, invoice và payment bằng cùng chuỗi mã.", "☐ Đạt  ☐ Không đạt"],
            ["Dữ liệu → AI", "Khuyến nghị tham chiếu đúng PR và số tiền, nhưng không tự đổi nghiệp vụ.", "☐ Đạt  ☐ Không đạt"],
            ["Tài liệu → Trợ lý AI", "Câu trả lời có nguồn nội bộ và không tạo thao tác nghiệp vụ.", "☐ Đạt  ☐ Không đạt"],
            ["Admin → Toàn luồng", "Phòng ban/quyền đúng; Admin không trực tiếp duyệt hay thanh toán.", "☐ Đạt  ☐ Không đạt"],
        ],
        [4.1, 10.3, 2.9],
    )
    callout(
        doc,
        "Luồng đạt hoàn toàn khi",
        "Chỉ có một mã PR từ đầu đến cuối; phiếu từng quay lại Employee để sửa; mọi lần bàn giao đều hiện đúng ở tài khoản tiếp theo; đơn hàng lấy dữ liệu từ báo giá thắng; hóa đơn PAID với còn lại 0 VND; Auditor truy được bằng chứng; cả khuyến nghị và Trợ lý AI chỉ hỗ trợ con người.",
        LIGHT_GREEN,
    )

    heading(doc, "12.1 Kịch bản demo 9–11 phút", 2)
    table(
        doc,
        ["Thời lượng", "Người trình bày", "Nội dung nên mở"],
        [
            ["1 phút", "Nhân viên", "Mở PR đã chuẩn bị, chỉ hai dòng hàng, chứng từ và lần gửi duyệt."],
            ["1 phút", "Trưởng bộ phận", "Chỉ vòng Yêu cầu chỉnh sửa và lần Phê duyệt sau khi Employee bổ sung."],
            ["2 phút", "Tài chính", "Mở hai báo giá, phương án được chọn và mã PO."],
            ["1 phút", "Nhân viên", "Mở biên bản nhận đủ với đúng số lượng từng dòng."],
            ["2 phút", "Tài chính", "Mở đối soát ba bên, trạng thái Đã xác minh và Đã thanh toán."],
            ["1 phút", "Kiểm toán", "Mở Timeline/gói bằng chứng và hồ sơ kiểm toán."],
            ["1 phút", "Điều phối AI", "Mở khuyến nghị cùng mã PR và giải thích AI không tự quyết định."],
            ["1 phút", "Bất kỳ tài khoản", "Hỏi Trợ lý AI nội bộ, mở nguồn và phân biệt với khuyến nghị theo quy tắc."],
        ],
        [2.1, 3.5, 11.7],
    )


def add_troubleshooting(doc: Document) -> None:
    doc.add_page_break()
    heading(doc, "13. Khi luồng bị đứt, kiểm tra ở đâu", 1)
    table(
        doc,
        ["Hiện tượng", "Nguyên nhân thường gặp", "Quay lại chặng"],
        [
            ["Manager không thấy PR", "PR vẫn Bản nháp hoặc hai tài khoản khác phòng ban.", "4 và 5.2"],
            ["Employee không sửa được", "PR chưa ở Yêu cầu chỉnh sửa hoặc không phải người tạo.", "6.1"],
            ["Finance không thấy PR", "Manager chưa phê duyệt hoặc phiếu đã kết thúc.", "7.1"],
            ["Không chọn được báo giá", "Phiếu chưa được duyệt cuối, dữ liệu báo giá thiếu hoặc phiên bản cũ.", "7.2 và 8.1"],
            ["Không phát hành được PO", "Chưa chọn báo giá hoặc chọn sai nhà cung cấp trúng thầu.", "8.1"],
            ["Employee không nhận hàng được", "PO chưa phát hành, sai requester/phòng ban hoặc đơn đã hủy.", "8.2"],
            ["Hóa đơn không khớp", "Chưa có biên nhận, số lượng nhận thiếu hoặc số tiền khác PO.", "9.1 và 10.1"],
            ["Không có nút Thanh toán", "Hóa đơn chưa được xác minh hoặc đang tranh chấp.", "10.2"],
            ["Auditor thiếu sự kiện", "Lọc sai thời gian/đối tượng hoặc dùng mã PR thay cho mã định danh.", "11.1"],
            ["AI chưa có khuyến nghị", "Tiến trình nền chưa chạy hoặc cảnh báo đã được tạo ở lần trước.", "11.2"],
            ["Trợ lý AI chưa trả lời", "Ollama/mô hình chưa sẵn sàng, kho tri thức lỗi hoặc câu hỏi không tìm được nguồn phù hợp.", "11.3"],
        ],
        [5.0, 8.8, 3.5],
    )
    callout(
        doc,
        "Khi báo lỗi",
        "Gửi tên vai trò, URL, mã PR/PO/hóa đơn, trạng thái đầu vào, nút đã bấm, kết quả thực tế và ảnh toàn màn hình. Không gửi mật khẩu, access token hoặc dữ liệu nhạy cảm.",
        LIGHT_RED,
    )


def add_signoff(doc: Document) -> None:
    heading(doc, "14. Biên bản chốt luồng", 1)
    table(
        doc,
        ["Thông tin", "Nội dung xác nhận"],
        [
            ["Mã PR dùng kiểm thử", "____________________________________________"],
            ["Người chạy luồng", "____________________________________________"],
            ["Thời gian bắt đầu/kết thúc", "____________________________________________"],
            ["Trạng thái cuối", "☐ PAID  ☐ Dừng ở ____________________________"],
            ["Số điểm nối đạt", "____ / 9 điểm nối tại mục 12"],
            ["Lỗi còn mở", "____________________________________________"],
            ["Kết luận", "☐ Sẵn sàng demo  ☐ Cần sửa và chạy lại từ Chặng ____"],
        ],
        [6.1, 11.2],
    )


def build() -> Path:
    doc = Document()
    configure(doc)
    replace_header_footer(doc)
    cover(doc)
    add_toc(doc)
    add_reading_guide(doc)
    add_flow_map(doc)
    add_preparation(doc)
    add_preflight_admin(doc)
    add_employee_create(doc)
    add_manager_roundtrip(doc)
    add_approvals(doc)
    add_sourcing_order(doc)
    add_receipt(doc)
    add_invoice(doc)
    add_audit_ai(doc)
    add_final_verification(doc)
    add_troubleshooting(doc)
    add_signoff(doc)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = "Hướng dẫn kiểm thử luồng liên vai trò DX-OS"
    doc.core_properties.subject = "Một phiếu mua sắm xuyên suốt sáu vai trò, khuyến nghị và Trợ lý AI local"
    doc.core_properties.author = "DX-OS Lab"
    doc.core_properties.keywords = (
        "DX-OS, luồng liên vai trò, mua sắm, phê duyệt, báo giá, đơn hàng tự điền, giao nhận, hóa đơn, kiểm toán, Ollama"
    )
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
