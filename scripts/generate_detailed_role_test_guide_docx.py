"""Generate a beginner-friendly DX-OS role testing workbook.

Run from repository root:
    python scripts/generate_detailed_role_test_guide_docx.py
"""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
OUTPUT = ROOT / "docs" / "generated" / "So_tay_kiem_thu_DX_OS_theo_vai_tro_chi_tiet.docx"

NAVY = "083B66"
BLUE = "0F6CBD"
PALE_BLUE = "EAF3FF"
PALE_GREEN = "E8F5E9"
PALE_YELLOW = "FFF7D6"
PALE_RED = "FDECEC"
GRAY = "64748B"
TEXT = "1F2937"
WHITE = "FFFFFF"


def shade(cell, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    element = OxmlElement("w:shd")
    element.set(qn("w:fill"), color)
    properties.append(element)


def set_cell(cell, value: str, bold: bool = False, color: str = TEXT) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(value)
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def repeat_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    element = OxmlElement("w:tblHeader")
    element.set(qn("w:val"), "true")
    properties.append(element)


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    header = table.rows[0]
    repeat_header(header)
    for index, value in enumerate(headers):
        shade(header.cells[index], BLUE)
        set_cell(header.cells[index], value, bold=True, color=WHITE)
        if widths:
            header.cells[index].width = Cm(widths[index])
    for row_index, row_values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            set_cell(cells[index], value)
            if widths:
                cells[index].width = Cm(widths[index])
        if row_index % 2 == 1:
            for cell in cells:
                shade(cell, "F8FAFC")
    document.add_paragraph()
    return table


def add_heading(document: Document, text: str, level: int) -> None:
    heading = document.add_heading(text, level=level)
    heading.paragraph_format.space_before = Pt(13 if level == 1 else 9)
    heading.paragraph_format.space_after = Pt(6)


def bullet(document: Document, text: str, level: int = 0) -> None:
    paragraph = document.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.add_run(text)


def numbered(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.add_run(text)


def code(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="Code Block")
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.add_run(text)


def callout(document: Document, title: str, text: str, fill: str = PALE_BLUE) -> None:
    table = document.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    shade(cell, fill)
    cell.text = ""
    title_paragraph = cell.paragraphs[0]
    title_paragraph.paragraph_format.space_after = Pt(2)
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor.from_string(NAVY)
    body = cell.add_paragraph(text)
    body.paragraph_format.space_after = Pt(2)
    document.add_paragraph()


def page_number(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instruction)
    run._r.append(end)


def configure(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.line_spacing = 1.12
    normal.paragraph_format.space_after = Pt(6)
    for name, size, color in (("Title", 28, NAVY), ("Heading 1", 18, NAVY), ("Heading 2", 13, BLUE), ("Heading 3", 11, NAVY)):
        style = styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
    if "Code Block" not in styles:
        style = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Consolas"
        style.font.size = Pt(8.5)
        style.font.color.rgb = RGBColor(242, 242, 242)
    code_style = styles["Code Block"]
    code_style.paragraph_format.left_indent = Cm(0.35)
    code_style.paragraph_format.space_after = Pt(8)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("DX-OS LAB | SỔ TAY KIỂM THỬ")
    run.font.name = "Aptos"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("DX-OS Lab • Tài liệu kiểm thử nội bộ • Trang ")
    run.font.name = "Aptos"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(GRAY)
    page_number(footer)


def add_cover(document: Document) -> None:
    document.add_paragraph()
    document.add_paragraph()
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("DX-OS LAB")
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(32)
    run.font.color.rgb = RGBColor.from_string(BLUE)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("SỔ TAY KIỂM THỬ CHI TIẾT THEO TỪNG VAI TRÒ")
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor.from_string(NAVY)
    document.add_paragraph()
    intro = document.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro.add_run("Dành cho nhóm 4 người chuẩn bị test, demo và báo cáo dự án.").italic = True
    document.add_paragraph()
    callout(
        document,
        "Cách đọc tài liệu này",
        "Mỗi mục test đều có: cần đăng nhập bằng ai, phải bấm ở đâu, kết quả đúng cần thấy và ảnh nên chụp. Chỉ cần đi lần lượt từ Phần 1 đến Phần 8 là có một kịch bản demo hoàn chỉnh.",
        PALE_GREEN,
    )
    info = document.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f"Phiên bản: 2.0 | Ngày tạo: {date.today().strftime('%d/%m/%Y')}")
    document.add_page_break()


def add_getting_started(document: Document) -> None:
    add_heading(document, "1. Trước khi bắt đầu", 1)
    document.add_paragraph("Mục tiêu là kiểm thử một yêu cầu mua sắm từ lúc tạo phiếu đến thanh toán và kiểm tra quyền của từng vai trò. Không cần hiểu code để làm theo tài liệu này.")
    add_heading(document, "1.1 Mở đúng địa chỉ", 2)
    add_table(
        document,
        ["Bạn cần mở", "Địa chỉ", "Dùng để làm gì"],
        [
            ["Ứng dụng DX-OS", "http://localhost:4200", "Đăng nhập và thao tác nghiệp vụ."],
            ["Kiểm tra API", "http://localhost:8081/health/ready", "Phải trả nội dung ready/HTTP 200 trước khi test."],
            ["Keycloak", "http://localhost:8080", "Trang đăng nhập tự chuyển tới khi vào DX-OS."],
            ["Báo cáo BI", "http://localhost:3000", "Tùy chọn, dùng đối chiếu dữ liệu báo cáo."],
        ],
        [3.8, 5.5, 8.0],
    )
    callout(document, "Rất quan trọng", "Chỉ truy cập web bằng http://localhost:4200. Không dùng 127.0.0.1:4200 vì Keycloak sẽ từ chối redirect URI.", PALE_YELLOW)
    add_heading(document, "1.2 Lấy tài khoản demo", 2)
    document.add_paragraph("Mật khẩu không được in trong tài liệu. Mở file credential tương ứng trong thư mục data/runtime rồi sao chép username và password vào màn hình Keycloak.")
    add_table(
        document,
        ["Vai trò", "Username", "Mở file mật khẩu", "Dùng chính để test"],
        [
            ["Nhân viên", "employee.demo", "data/runtime/employee-demo.txt", "Tạo và gửi phiếu; xác nhận nhận hàng; xem thông báo."],
            ["Trưởng bộ phận", "manager.demo", "data/runtime/manager-demo.txt", "Duyệt cấp phòng ban; yêu cầu sửa/từ chối; xác nhận nhận hàng."],
            ["Tài chính", "finance.demo", "data/runtime/finance-demo.txt", "Duyệt cuối, ngân sách, nhà cung cấp, đơn hàng, hóa đơn, báo cáo."],
            ["Kiểm toán", "auditor.demo", "data/runtime/auditor-demo.txt", "Đọc dữ liệu, audit, báo cáo và test quyền chỉ đọc."],
            ["Quản trị DX", "admin.demo", "data/runtime/admin-demo.txt", "Sửa chính sách/SLA; xem audit và báo cáo."],
            ["AI Operator", "ai.operator.demo", "data/runtime/ai-operator-demo.txt", "Xác nhận phạm vi AI hiện tại, chưa có Agent/RAG."],
        ],
        [2.6, 3.2, 5.2, 6.3],
    )
    bullet(document, "Nên dùng cửa sổ ẩn danh riêng cho Employee, Manager, Finance và Auditor. Cách này tránh nhầm phiên đăng nhập.")
    bullet(document, "Không gửi file trong data/runtime, password, access token hoặc file .env vào nhóm chat hay GitHub.")
    add_heading(document, "1.3 Cách đọc kết quả", 2)
    add_table(
        document,
        ["Loại kết quả", "Ý nghĩa"],
        [
            ["Thành công", "Nút thao tác hoàn tất, trạng thái thay đổi đúng và có dòng lịch sử/audit tương ứng."],
            ["Bị chặn 403 hoặc không có nút", "Đây thường là kết quả đúng khi đang test quyền của role không phù hợp."],
            ["Bị chặn do điều kiện", "Ví dụ: chưa có báo giá, chưa nhận hàng, lệch tiền hóa đơn hoặc vượt ngân sách."],
        ],
        [5.3, 12.0],
    )


def add_flow(document: Document) -> None:
    document.add_page_break()
    add_heading(document, "2. Bức tranh toàn bộ luồng để không bị lạc", 1)
    document.add_paragraph("Hãy dùng cùng một mã phiếu cho toàn bộ buổi test. Tên gợi ý: Mua laptop demo - [ngày test].")
    code(document, "Employee: DRAFT → SUBMITTED\nManager: SUBMITTED → MANAGER_APPROVED\nFinance: MANAGER_APPROVED → APPROVED\nFinance: APPROVED → Purchase order\nEmployee/Manager: Purchase order → Đã nhận hàng\nFinance: Tạo hóa đơn → Xác minh → Thanh toán\nAuditor: Đọc toàn bộ bằng chứng và báo cáo")
    add_table(
        document,
        ["Bước", "Ai làm", "Kết quả phải thấy"],
        [
            ["1", "Employee", "Phiếu DRAFT có dòng hàng, lý do và cost center."],
            ["2", "Employee", "Gửi duyệt, trạng thái thành SUBMITTED."],
            ["3", "Manager", "Duyệt cấp phòng ban, trạng thái MANAGER_APPROVED; ngân sách được reserve."],
            ["4", "Finance", "Duyệt cuối, trạng thái APPROVED; ngân sách chuyển sang committed."],
            ["5", "Finance", "Phát hành purchase order với nhà cung cấp."],
            ["6", "Employee hoặc Manager", "Xác nhận đã nhận hàng; Finance không được tự xác nhận."],
            ["7", "Finance", "Hóa đơn khớp ba bên mới được xác minh và ghi nhận thanh toán."],
            ["8", "Auditor", "Đọc Timeline, budget, order, invoice, audit và báo cáo mà không sửa được."],
        ],
        [1.4, 3.7, 12.2],
    )
    callout(document, "Mẹo demo", "Nếu thời gian ngắn, chỉ demo một phiếu thành công. Các test bị chặn như thiếu báo giá, vượt ngân sách hoặc Finance tự nhận hàng có thể nói và chụp ảnh làm bằng chứng.", PALE_BLUE)


def test_case(document: Document, title: str, goal: str, preparation: str, steps: list[str], expected: str, evidence: str, warning: str | None = None) -> None:
    add_heading(document, title, 3)
    add_table(
        document,
        ["Mục", "Nội dung"],
        [["Mục tiêu", goal], ["Chuẩn bị", preparation], ["Kết quả đúng", expected], ["Bằng chứng", evidence]],
        [3.0, 14.3],
    )
    document.add_paragraph("Thao tác thực hiện:")
    for step in steps:
        numbered(document, step)
    if warning:
        callout(document, "Lưu ý", warning, PALE_YELLOW)


def add_employee(document: Document) -> None:
    document.add_page_break()
    add_heading(document, "3. Test vai trò Employee (Nhân viên)", 1)
    document.add_paragraph("Employee là người bắt đầu quy trình. Vai trò này chỉ thao tác trên phiếu do mình tạo hoặc các việc được giao trong phạm vi của mình.")
    test_case(
        document,
        "E1. Đăng nhập và kiểm tra menu",
        "Xác nhận user Employee vào đúng dashboard và không có quyền duyệt/ngân sách.",
        "Mở data/runtime/employee-demo.txt; chuẩn bị cửa sổ ẩn danh.",
        ["Mở http://localhost:4200.", "Đăng nhập bằng employee.demo.", "Ở Dashboard, kiểm tra tên/role hiển thị là Employee.", "Quan sát menu: cần có Phiếu mua sắm, Việc của tôi, Thông báo; không cần có Phê duyệt, Ngân sách hay Báo cáo."],
        "Dashboard tải được. Các menu chỉ dành cho Manager/Finance không xuất hiện hoặc bị chặn.",
        "Chụp Dashboard Employee và thanh menu.",
    )
    test_case(
        document,
        "E2. Tạo bản nháp và gửi duyệt",
        "Tạo một phiếu nhỏ để đi hết luồng nhanh.",
        "Đang đăng nhập Employee; dùng số tiền dưới 20.000.000 VND để không cần báo giá.",
        ["Bấm Tạo phiếu mới.", "Nhập tiêu đề dễ tìm, ví dụ: Mua chuột demo 16-08.", "Nhập lý do mua, cost center hợp lệ, currency VND.", "Thêm ít nhất một dòng: mô tả, số lượng, đơn vị và đơn giá; lưu DRAFT.", "Mở lại chi tiết, ghi lại mã phiếu.", "Bấm Gửi duyệt và xác nhận thao tác."],
        "Phiếu đổi DRAFT thành SUBMITTED. Timeline có mốc tạo/lưu/gửi duyệt. Phiếu biến mất khỏi Việc của tôi của Employee và xuất hiện cho Manager.",
        "Chụp trang chi tiết hiển thị mã phiếu, trạng thái SUBMITTED và Timeline.",
    )
    test_case(
        document,
        "E3. Test phiếu cần báo giá",
        "Chứng minh rule chứng từ khi giá trị từ 20 triệu VND.",
        "Tạo một phiếu mới có tổng tiền từ 20.000.000 VND trở lên; chuẩn bị một PDF/tệp báo giá giả lập, không có dữ liệu nhạy cảm.",
        ["Lưu DRAFT với tổng tiền từ 20 triệu VND.", "Thử bấm Gửi duyệt khi chưa tải báo giá.", "Đọc thông báo bị chặn.", "Tải tệp lên loại Báo giá/QUOTATION theo form.", "Gửi duyệt lại."],
        "Lần đầu bị chặn do thiếu báo giá. Sau khi có tệp hợp lệ, phiếu được SUBMITTED.",
        "Chụp thông báo bị chặn và danh sách attachment sau khi tải lên.",
    )
    test_case(
        document,
        "E4. Sửa phiếu bị yêu cầu chỉnh sửa",
        "Kiểm tra vòng lặp phản hồi giữa Employee và người duyệt.",
        "Cần một phiếu đã được Manager hoặc Finance chọn Yêu cầu chỉnh sửa.",
        ["Đăng nhập lại Employee.", "Mở Việc của tôi hoặc danh sách phiếu; tìm trạng thái CHANGES_REQUESTED.", "Mở comment/timeline để đọc lý do trả lại.", "Bấm sửa, bổ sung nội dung hoặc tệp cần thiết, lưu lại.", "Bấm Gửi duyệt lại."],
        "Phiếu quay về SUBMITTED và comment cũ vẫn còn trong Timeline.",
        "Chụp comment yêu cầu sửa và trạng thái sau khi gửi lại.",
    )
    test_case(
        document,
        "E5. Xác nhận nhận hàng và xem thông báo thanh toán",
        "Thử phần việc sau khi Finance đã phát hành order.",
        "Cần một phiếu APPROVED đã có purchase order do Finance phát hành.",
        ["Mở Giao nhận/Đặt hàng và giao nhận.", "Tìm order theo mã phiếu.", "Kiểm tra đúng thông tin hàng và bấm xác nhận đã nhận.", "Sau khi Finance ghi nhận thanh toán, mở Thông báo."],
        "Trạng thái receipt được ghi nhận. Sau thanh toán, có thông báo INVOICE_PAID hoặc thông báo hóa đơn đã thanh toán.",
        "Chụp trạng thái đã nhận và thông báo thanh toán.",
        "Employee chỉ xác nhận hàng thực tế đã nhận; không tự xác nhận khi không có order.",
    )


def add_manager(document: Document) -> None:
    document.add_page_break()
    add_heading(document, "4. Test vai trò Department Manager (Trưởng bộ phận)", 1)
    document.add_paragraph("Manager là cấp duyệt đầu tiên. Manager xử lý phiếu của phòng ban, không được tự duyệt phiếu do mình tạo.")
    test_case(
        document,
        "M1. Mở hàng đợi và duyệt cấp phòng ban",
        "Chuyển phiếu SUBMITTED sang MANAGER_APPROVED.",
        "Cần mã phiếu Employee vừa gửi; đăng nhập manager.demo.",
        ["Mở Dashboard, bấm Phê duyệt hoặc vào menu Phê duyệt.", "Tìm phiếu theo tiêu đề/mã đã ghi lại.", "Mở chi tiết; kiểm tra lý do, cost center, tổng tiền và báo giá nếu là phiếu trên 20 triệu.", "Thêm nhận xét ngắn, ví dụ: Phù hợp nhu cầu phòng ban.", "Bấm Phê duyệt."],
        "Phiếu thành MANAGER_APPROVED. Budget check thể hiện reserved tăng và available giảm. Phiếu chuyển sang hàng đợi Finance.",
        "Chụp màn hình trước/sau duyệt và budget check.",
    )
    test_case(
        document,
        "M2. Yêu cầu chỉnh sửa thay vì duyệt",
        "Kiểm tra phản hồi có comment rõ ràng.",
        "Tạo một phiếu SUBMITTED khác để không làm gián đoạn luồng chính.",
        ["Mở phiếu trong Phê duyệt.", "Chọn Yêu cầu chỉnh sửa.", "Nhập comment cụ thể, ví dụ: Bổ sung báo giá và mô tả cấu hình.", "Đăng xuất, đăng nhập Employee để kiểm tra lại."],
        "Phiếu thành CHANGES_REQUESTED; Employee nhìn thấy lý do và có thể sửa/gửi lại.",
        "Chụp comment Manager và màn hình Employee nhận việc cần bổ sung.",
    )
    test_case(
        document,
        "M3. Từ chối phiếu",
        "Kiểm tra trạng thái kết thúc REJECTED.",
        "Một phiếu SUBMITTED riêng, có lý do nên từ chối.",
        ["Mở phiếu từ hàng đợi.", "Chọn Từ chối.", "Nhập lý do rõ ràng, ví dụ: Chưa có nhu cầu trong quý này.", "Xác nhận và kiểm tra Timeline."],
        "Phiếu thành REJECTED, không xuất hiện cho Finance để duyệt tiếp.",
        "Chụp trạng thái REJECTED và comment.",
    )
    test_case(
        document,
        "M4. Test phân quyền Manager",
        "Chứng minh Manager không tự duyệt và không quản lý ngân sách.",
        "Đăng nhập manager.demo.",
        ["Nếu Manager tạo một phiếu của chính mình, gửi duyệt rồi thử xử lý trong hàng đợi.", "Quan sát hành vi bị chặn/từ chối tự duyệt.", "Thử mở trực tiếp menu Ngân sách hoặc Báo cáo."],
        "Tự duyệt bị chặn. Manager không có chức năng điều chỉnh ngân sách hoặc báo cáo toàn tổ chức.",
        "Chụp thông báo bị chặn hoặc menu không xuất hiện.",
    )


def add_finance(document: Document) -> None:
    document.add_page_break()
    add_heading(document, "5. Test vai trò Finance (Tài chính)", 1)
    document.add_paragraph("Finance là vai trò có nhiều bước nhất: duyệt cuối, quản lý ngân sách, nhà cung cấp, đơn hàng, hóa đơn, thanh toán và báo cáo.")
    test_case(
        document,
        "F1. Duyệt cuối và kiểm tra ngân sách",
        "Hoàn tất quy trình phê duyệt của phiếu MANAGER_APPROVED.",
        "Cần một phiếu do Manager vừa duyệt; đăng nhập finance.demo.",
        ["Mở Dashboard hoặc Phê duyệt.", "Mở phiếu MANAGER_APPROVED.", "Kiểm tra tổng tiền, budget check, tệp đính kèm và comment Manager.", "Bấm Phê duyệt."],
        "Phiếu thành APPROVED. Số reserved giảm, committed tăng. Phiếu xuất hiện trong nhóm có thể phát hành order.",
        "Chụp trạng thái APPROVED và trang Ngân sách trước/sau khi duyệt.",
    )
    test_case(
        document,
        "F2. Tạo hoặc kiểm tra nhà cung cấp",
        "Đảm bảo có nhà cung cấp hoạt động để phát hành order.",
        "Đang đăng nhập Finance.",
        ["Mở menu Nhà cung cấp.", "Nếu chưa có nhà cung cấp phù hợp, bấm tạo mới.", "Nhập mã, tên, mã số thuế/liên hệ (có thể là dữ liệu demo), trạng thái hoạt động và mức rủi ro thấp.", "Lưu, sau đó tìm lại nhà cung cấp trong danh sách."],
        "Nhà cung cấp xuất hiện trong danh sách; thay đổi được lưu và có audit.",
        "Chụp danh sách hoặc form nhà cung cấp vừa tạo.",
    )
    test_case(
        document,
        "F3. Phát hành purchase order",
        "Liên kết phiếu APPROVED với nhà cung cấp và lịch giao.",
        "Cần phiếu APPROVED và một nhà cung cấp đang hoạt động.",
        ["Mở Đặt hàng và giao nhận/Giao nhận.", "Chọn phiếu APPROVED đúng mã.", "Chọn nhà cung cấp, ngày giao dự kiến và mã tham chiếu demo.", "Bấm phát hành đơn hàng."],
        "Order được tạo, có trạng thái chờ nhận hàng và ghi audit.",
        "Chụp order có mã phiếu, nhà cung cấp và trạng thái chờ nhận.",
    )
    test_case(
        document,
        "F4. Tạo hóa đơn và kiểm tra chặn trước khi nhận hàng",
        "Chứng minh hệ thống bắt buộc nhận hàng trước khi xác minh hóa đơn.",
        "Cần order vừa phát hành nhưng Employee/Manager chưa xác nhận nhận hàng.",
        ["Mở Hóa đơn và thanh toán.", "Tạo hóa đơn: số hóa đơn demo, ngày hóa đơn, hạn thanh toán, amount và currency trùng với order.", "Thử bấm Xác minh ngay.", "Đọc thông báo/nhãn trạng thái."],
        "Hóa đơn được lưu nhưng xác minh bị chặn với trạng thái chờ nhận hàng.",
        "Chụp hóa đơn và thông báo bị chặn.",
    )
    test_case(
        document,
        "F5. Xác minh hóa đơn và ghi nhận thanh toán",
        "Hoàn tất đối soát ba bên: order - receipt - invoice.",
        "Employee hoặc Manager đã xác nhận receipt cho order; hóa đơn có amount/currency khớp.",
        ["Tải lại trang Hóa đơn.", "Mở hóa đơn vừa tạo; kiểm tra nhãn khớp ba bên/đủ điều kiện.", "Bấm Xác minh.", "Nhập payment reference và ngày thanh toán.", "Bấm ghi nhận thanh toán."],
        "Hóa đơn lần lượt VERIFIED rồi PAID. Requester nhận notification. Audit có dấu vết xác minh/thanh toán.",
        "Chụp trạng thái PAID, payment reference và notification của Employee.",
    )
    test_case(
        document,
        "F6. Ngân sách và báo cáo",
        "Kiểm tra dữ liệu tài chính và export CSV.",
        "Đang đăng nhập Finance; có ít nhất một phiếu APPROVED trong khoảng ngày chọn.",
        ["Mở Ngân sách, chọn cost center/currency phù hợp.", "Kiểm tra allocation, available, reserved và committed.", "Nếu cần test điều chỉnh, sửa một giá trị nhỏ và lưu; không mở đồng thời cùng allocation ở hai phiên.", "Mở Báo cáo, chọn khoảng ngày bao phủ phiếu demo và bấm Áp dụng.", "Bấm Xuất CSV, mở bằng Excel để kiểm tra dấu tiếng Việt và số tiền."],
        "Số liệu phản ánh trạng thái phiếu. CSV tải được và có dữ liệu theo quyền Finance.",
        "Chụp Budget dashboard, Report dashboard và file CSV đã mở.",
        "Nếu lưu ở hai phiên cùng lúc, phiên cũ có thể bị conflict. Đây là optimistic locking đúng thiết kế.",
    )
    test_case(
        document,
        "F7. Test phân tách nhiệm vụ",
        "Xác nhận Finance không tự xác nhận receipt.",
        "Một order đang chờ nhận hàng; đang đăng nhập Finance.",
        ["Mở order ở Giao nhận.", "Thử tìm/bấm thao tác xác nhận đã nhận hàng bằng Finance.", "Đăng xuất và xác nhận lại bằng Employee hoặc Manager cùng phòng."],
        "Finance bị chặn hoặc không có nút receipt; Employee/Manager xác nhận thành công.",
        "Chụp hành vi bị chặn bằng Finance và receipt thành công bằng Employee/Manager.",
    )


def add_auditor(document: Document) -> None:
    document.add_page_break()
    add_heading(document, "6. Test vai trò Auditor (Kiểm toán)", 1)
    document.add_paragraph("Auditor dùng hệ thống ở chế độ chỉ đọc. Đây là phần rất quan trọng trong báo cáo vì chứng minh kiểm soát quyền và audit trail.")
    test_case(
        document,
        "A1. Đọc hồ sơ mua sắm và Timeline",
        "Đối chiếu toàn bộ lịch sử của phiếu đã đi qua luồng chính.",
        "Cần một phiếu APPROVED hoặc đã thanh toán; đăng nhập auditor.demo.",
        ["Mở danh sách Phiếu mua sắm và tìm mã phiếu demo.", "Mở chi tiết; kiểm tra trạng thái, dòng hàng, attachment và Timeline.", "Xác nhận Timeline có actor, thời gian và comment của Employee/Manager/Finance."],
        "Auditor đọc được hồ sơ nhưng không có nút sửa, gửi duyệt hay phê duyệt.",
        "Chụp Timeline và menu ở chế độ chỉ đọc.",
    )
    test_case(
        document,
        "A2. Đối chiếu order, receipt, invoice và audit",
        "Kiểm tra chuỗi bằng chứng sau duyệt.",
        "Cần phiếu đã có order, receipt và hóa đơn.",
        ["Mở Giao nhận, tìm order theo mã phiếu và kiểm tra trạng thái receipt.", "Mở Hóa đơn và thanh toán, kiểm tra số hóa đơn, amount, trạng thái VERIFIED/PAID.", "Mở Trung tâm kiểm toán/Audit, lọc theo loại đối tượng như purchase order hoặc invoice nếu có bộ lọc.", "Đối chiếu mốc tạo order, receipt, xác minh hóa đơn và thanh toán."],
        "Dữ liệu nhất quán; Auditor chỉ xem, không thay đổi được supplier/order/invoice.",
        "Chụp một event audit và màn hình invoice/order.",
    )
    test_case(
        document,
        "A3. Ngân sách, chính sách và báo cáo ở chế độ chỉ đọc",
        "Kiểm tra quyền đọc rộng nhưng không có quyền ghi.",
        "Đang đăng nhập Auditor.",
        ["Mở Ngân sách, kiểm tra allocation và lịch sử điều chỉnh.", "Mở Chính sách, xem các ngưỡng/SLA hiện hành.", "Mở Báo cáo, đặt khoảng ngày và Xuất CSV.", "Thử tìm nút điều chỉnh ngân sách hoặc lưu chính sách."],
        "Auditor xem và xuất được dữ liệu nhưng nút ghi không xuất hiện hoặc API trả 403.",
        "Chụp Budget/Policy/Audit hoặc thông báo bị chặn.",
        "403 trong test này là kết quả đạt, không phải lỗi cần sửa.",
    )


def add_admin_ai(document: Document) -> None:
    document.add_page_break()
    add_heading(document, "7. Test DX Admin và AI Operator", 1)
    test_case(
        document,
        "D1. DX Admin thay đổi chính sách có version và audit",
        "Kiểm tra quản trị policy mà không cấp nhầm quyền nghiệp vụ.",
        "Đăng nhập admin.demo. Ghi lại giá trị chính sách trước khi sửa để có thể đổi lại sau test.",
        ["Mở Chính sách vận hành.", "Chọn một policy như SLA hoặc ngưỡng chứng từ.", "Thay đổi nhỏ, dễ nhận biết, rồi lưu với phiên bản hiện tại.", "Tải lại trang để xác nhận giá trị mới.", "Mở Audit để tìm event thay đổi policy.", "Sau khi chụp bằng chứng, đổi policy về giá trị ban đầu nếu đây là môi trường dùng chung."],
        "Policy được cập nhật, version tăng và có audit event. Nếu hai phiên cùng sửa, phiên cũ phải conflict thay vì ghi đè âm thầm.",
        "Chụp trước/sau policy và event audit.",
    )
    test_case(
        document,
        "D2. DX Admin không phải Finance/Manager",
        "Xác nhận admin không là superuser nghiệp vụ mặc định.",
        "Đang đăng nhập admin.demo.",
        ["Thử mở Phê duyệt, Ngân sách, Hóa đơn hoặc tạo thao tác nghiệp vụ không nằm trong menu Admin.", "Quan sát menu và phản hồi."],
        "Admin chỉ có các quyền được cấp như policy/audit/reports; không tự nhiên có quyền duyệt hoặc thanh toán.",
        "Chụp menu Admin hoặc phản hồi 403.",
    )
    test_case(
        document,
        "AI1. AI Operator và phạm vi AI hiện tại",
        "Trình bày trung thực các chức năng AI chưa triển khai.",
        "Đăng nhập ai.operator.demo.",
        ["Mở Dashboard và kiểm tra tài khoản đăng nhập thành công.", "Quan sát thông tin trạng thái nền tảng/phạm vi AI.", "Tìm menu chat AI, recommendation, tool execution hoặc Agent."],
        "User xác thực và vào dashboard được, nhưng không có AI Agent/RAG mô phỏng, không có quyền duyệt hoặc thao tác tài chính. Đây là phạm vi hiện tại của dự án.",
        "Chụp Dashboard AI Operator và ghi rõ AI/RAG là roadmap, không demo như tính năng đã có.",
        "Không dùng AI Operator để thay Manager/Finance trong bất kỳ test nghiệp vụ nào.",
    )


def add_regression(document: Document) -> None:
    document.add_page_break()
    add_heading(document, "8. Checklist cuối buổi test", 1)
    add_heading(document, "8.1 Test lỗi nên làm ít nhất một lần", 2)
    add_table(
        document,
        ["Tình huống", "Cách thử ngắn", "Kết quả đạt"],
        [
            ["Thiếu báo giá", "Tạo phiếu từ 20 triệu VND, không tải quotation rồi gửi.", "Bị chặn và nhắc bổ sung chứng từ."],
            ["Thiếu ngân sách", "Tạo phiếu lớn hơn allocation rồi để Manager duyệt.", "Không tạo reservation/hiển thị thiếu ngân sách."],
            ["Tự duyệt", "Manager/Finance thử duyệt phiếu mình tạo.", "Bị chặn."],
            ["Receipt sai vai trò", "Finance thử xác nhận đã nhận hàng.", "Bị chặn hoặc không có nút."],
            ["Invoice trước receipt", "Finance xác minh hóa đơn khi order chưa nhận.", "Bị chặn."],
            ["Auditor ghi dữ liệu", "Auditor thử sửa budget/policy/supplier.", "Không có nút hoặc 403."],
            ["Sai scope", "Employee khác mở phiếu không thuộc mình.", "Không xem được dữ liệu ngoài phạm vi."],
        ],
        [3.5, 7.2, 6.0],
    )
    add_heading(document, "8.2 Ảnh cần lưu để báo cáo", 2)
    for item in [
        "Dashboard của mỗi role (ít nhất Employee, Finance và Auditor).",
        "Một phiếu ở các trạng thái SUBMITTED, MANAGER_APPROVED và APPROVED.",
        "Budget dashboard cho thấy reserved/committed thay đổi.",
        "Order, receipt, invoice PAID và notification của Employee.",
        "Một event tại Audit Center thể hiện actor và thời gian.",
        "Một test bị chặn đúng quyền: Auditor ghi dữ liệu hoặc Finance xác nhận receipt.",
        "CSV báo cáo/ngân sách mở được trong Excel với tiếng Việt hiển thị đúng.",
    ]:
        document.add_paragraph(f"☐  {item}")
    add_heading(document, "8.3 Khi gặp lỗi", 2)
    add_table(
        document,
        ["Hiện tượng", "Xử lý nhanh"],
        [
            ["Không vào được web", "Mở http://localhost:4200; kiểm tra docker compose ps và service web."],
            ["Login báo redirect URI", "Đảm bảo dùng localhost:4200, không dùng 127.0.0.1."],
            ["Sai password", "Mở lại file credential trong data/runtime. Nếu cần tạo lại, chạy Initialize-DevUser.ps1; mật khẩu cũ sẽ đổi."],
            ["401", "Đăng xuất/đăng nhập lại; kiểm tra http://localhost:8081/health/ready."],
            ["403", "Kiểm tra role và data scope trước. Với test quyền, 403 thường là kết quả đúng."],
            ["Không có dữ liệu báo cáo", "Chọn khoảng ngày bao phủ phiếu demo, kiểm tra trạng thái phiếu và filter cost center/currency."],
        ],
        [4.5, 12.2],
    )
    callout(document, "Kết thúc buổi test", "Không sửa trực tiếp database để đổi trạng thái phiếu. Nếu cần làm lại, tạo phiếu demo mới để giữ audit trail rõ ràng. Trước khi báo cáo, kiểm tra lại ảnh, file CSV và mã phiếu đã ghi chép.", PALE_RED)


def build() -> Path:
    document = Document()
    configure(document)
    add_cover(document)
    add_getting_started(document)
    add_flow(document)
    add_employee(document)
    add_manager(document)
    add_finance(document)
    add_auditor(document)
    add_admin_ai(document)
    add_regression(document)
    for section in document.sections:
        section.start_type = WD_SECTION.NEW_PAGE
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(f"Created {build()}")
