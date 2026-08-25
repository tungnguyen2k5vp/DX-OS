from __future__ import annotations

"""Generate the DX-OS defense Q&A guide for cancellation and attachments.

Content is deliberately grounded in the current implementation:
* backend/internal/procurement/store.go and attachments.go
* backend/internal/procurement/model.go
* backend/internal/platform/httpapi/purchase_requests.go
* backend/internal/platform/documentstore/nextcloud.go
* migrations 000002, 000003 and 000006
"""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "generated" / "Bo_cau_hoi_bao_ve_UC_huy_phieu_va_quan_ly_tep_dinh_kem.docx"

NAVY = "112B46"
TEAL = "007C83"
TEAL_DARK = "00636A"
SLATE = "53677C"
MUTED = "6B7F91"
PALE_TEAL = "E8F6F5"
PALE_BLUE = "EDF5FB"
PALE_GOLD = "FFF7E8"
PALE_RED = "FCEEEE"
WHITE = "FFFFFF"
GRID = "C9D9E3"


CANCEL_QA = [
    (
        "UC Hủy phiếu dùng để giải quyết vấn đề gì?",
        "Cho phép người tạo đóng một nhu cầu mua sắm không còn cần thiết, nhưng vẫn giữ toàn bộ hồ sơ để truy vết.",
        "Đây không phải là xóa dữ liệu. Hệ thống chuyển trạng thái phiếu sang CANCELLED/Đã hủy, vì vậy người dùng không tiếp tục gửi duyệt hay xử lý phiếu đó. Cách làm này bảo đảm minh bạch: sau này vẫn biết phiếu nào đã được tạo, ai hủy, hủy từ trạng thái nào và vào thời điểm nào.",
        "Demo: mở một phiếu Bản nháp, chọn Hủy phiếu; sau khi hủy, kiểm tra nhãn Đã hủy và phần Tiến trình/Lịch sử.",
    ),
    (
        "Ai có quyền hủy phiếu?",
        "Chỉ chính người tạo phiếu mới được hủy; kiểm toán viên không có quyền này.",
        "Backend lấy người dùng từ access token, ánh xạ về bảng users rồi so sánh requester_id của phiếu với user hiện tại. Kiểm tra nằm ở domain service, không chỉ ẩn nút trên giao diện, nên gọi API thủ công bằng tài khoản khác cũng bị từ chối. Vai trò auditor bị chặn rõ ràng trước khi xử lý transition.",
        "Mã nguồn: procurement.DecideTransition và Store.Transition; thử đăng nhập tài khoản khác rồi gọi hành động CANCEL sẽ nhận lỗi không có quyền/không tìm thấy theo phạm vi dữ liệu.",
    ),
    (
        "Phiếu ở trạng thái nào thì được hủy?",
        "Chỉ hủy được khi là Bản nháp (DRAFT) hoặc đang được yêu cầu chỉnh sửa (CHANGES_REQUESTED).",
        "Đó là quy tắc trạng thái (state machine) của hệ thống. Khi phiếu đã gửi duyệt (SUBMITTED) hoặc đã đi sâu hơn vào quy trình, người tạo không thể tự ý hủy, tránh làm mất tính nhất quán với công việc của quản lý/tài chính. Nếu nghiệp vụ cần dừng một phiếu đã gửi duyệt, quy trình hiện tại phải đi qua người đang xử lý, không dùng UC này.",
        "Demo: với phiếu DRAFT sẽ thấy nút Hủy; gửi duyệt rồi quay lại sẽ không còn nút đó.",
    ),
    (
        "Tại sao không cho hủy sau khi đã gửi duyệt?",
        "Để tránh cắt ngang luồng phê duyệt và gây sai lệch trách nhiệm.",
        "Sau khi gửi, phiếu đã có thể xuất hiện trong hàng đợi của quản lý, tạo thông báo và có khả năng phát sinh giữ ngân sách ở các bước sau. Nếu người tạo tự hủy ở mọi trạng thái thì người phê duyệt có thể đang ra quyết định trên một đối tượng đã biến mất. Vì thế hệ thống giới hạn UC hủy ở giai đoạn thuộc quyền kiểm soát của người tạo.",
        "Liên hệ domain: DecideTransition chỉ cho ActionCancel ở DRAFT và CHANGES_REQUESTED.",
    ),
    (
        "Hủy phiếu khác gì với từ chối phiếu?",
        "Hủy do người tạo chủ động dừng nhu cầu; từ chối do người phê duyệt quyết định không chấp nhận.",
        "Hai hành động có actor và ý nghĩa nghiệp vụ khác nhau. CANCEL chỉ dành cho requester ở DRAFT/CHANGES_REQUESTED. REJECT dành cho trưởng bộ phận hoặc tài chính ở bước của họ và bắt buộc có nhận xét. Cả hai đều là trạng thái kết thúc nhưng lịch sử event khác nhau, do đó báo cáo và kiểm toán phân biệt được nguyên nhân kết thúc.",
        "Demo: lịch sử hiển thị sự kiện CANCELLED khác REJECTED; khi từ chối giao diện bắt buộc nhập lý do.",
    ),
    (
        "Khi hủy, dữ liệu có bị xóa khỏi database không?",
        "Không. Bản ghi purchase_requests vẫn tồn tại, chỉ đổi status thành CANCELLED.",
        "Hệ thống dùng hủy mềm theo trạng thái cho phiếu mua sắm. Bản ghi gốc, các dòng hàng và lịch sử được giữ để báo cáo và truy vết. Khi truy vấn danh sách, phiếu hủy có thể được lọc theo trạng thái, nhưng không phải là bị xóa vật lý. Đây là thiết kế phù hợp dữ liệu nghiệp vụ cần kiểm toán.",
        "Bằng chứng: UPDATE purchase_requests SET status = ... trong Store.Transition, không có DELETE purchase_requests ở UC này.",
    ),
    (
        "Lịch sử hủy được lưu ở đâu và lưu những gì?",
        "Được lưu ở process_events và audit_logs, gồm hành động, trạng thái trước/sau, người thực hiện, thời gian, correlation ID và bình luận nếu có.",
        "process_events phục vụ timeline nghiệp vụ của chính phiếu; audit_logs phục vụ tra soát ở mức toàn hệ thống. Mỗi lần chuyển trạng thái, Store.Transition chèn một process event và một audit log trong cùng transaction. Vì vậy trạng thái Đã hủy không thể xuất hiện mà thiếu dấu vết xử lý nếu transaction thành công.",
        "Demo: xem Tiến trình phiếu và trang Kiểm toán, lọc theo mã phiếu/hành động CANCELLED.",
    ),
    (
        "Có bắt buộc nhập lý do khi hủy không?",
        "Hiện tại không bắt buộc; người dùng có thể nhập bình luận để giải thích.",
        "Theo ValidateTransition, comment chỉ bắt buộc đối với REJECT và REQUEST_CHANGES; CANCEL nhận comment tùy chọn và giới hạn tối đa 2.000 ký tự. Đây là trạng thái triển khai hiện tại. Nếu đơn vị yêu cầu quản trị chặt hơn, có thể nâng quy tắc để CANCEL bắt buộc lý do mà không thay đổi mô hình dữ liệu, vì process_events đã có cột comment.",
        "Điểm trả lời trung thực: không nói hệ thống đang bắt buộc lý do hủy; hãy nêu đây là hướng cải tiến có kiểm soát.",
    ),
    (
        "API nào thực hiện việc hủy?",
        "POST /api/v1/purchase-requests/{requestID}/transitions với action = CANCEL, expectedVersion và header Idempotency-Key.",
        "Hủy là một trường hợp của API chuyển trạng thái chung, thay vì tạo API xóa riêng. Body mang action, expectedVersion và comment; Idempotency-Key nằm trong header. Cách dùng chung endpoint giúp tập trung kiểm soát state machine, kiểm toán và xử lý đồng thời cho mọi hành động workflow.",
        "Demo kỹ thuật: DevTools/Network cho thấy POST transitions và action CANCEL.",
    ),
    (
        "ExpectedVersion dùng để làm gì?",
        "Để chống ghi đè khi hai người hoặc hai tab cùng thao tác trên một phiếu.",
        "Mỗi phiếu có trường version. Giao diện gửi version đang xem; backend khóa bản ghi bằng SELECT ... FOR UPDATE, sau đó so sánh current.Version với expectedVersion. Nếu khác, trả ErrVersionConflict thay vì ghi đè quyết định mới hơn. Sau transition thành công, version tăng thêm một. Đây là optimistic locking kết hợp khóa hàng tại thời điểm cập nhật.",
        "Demo: mở cùng một phiếu ở hai tab, cập nhật/chuyển trạng thái từ tab thứ nhất, rồi thao tác từ tab còn lại để quan sát thông báo cần tải lại.",
    ),
    (
        "Nếu người dùng bấm Hủy hai lần do mạng chậm thì sao?",
        "Idempotency-Key bảo đảm cùng một yêu cầu được xử lý một lần theo ngữ nghĩa nghiệp vụ.",
        "Backend lưu idempotency_key duy nhất trong process_events. Khi nhận lại key cũ cùng phiếu, cùng người thực hiện và cùng action, hệ thống trả lại trạng thái hiện tại thay vì tạo event thứ hai. Nếu key đó bị tái sử dụng cho phiếu/actor/action khác, hệ thống báo xung đột idempotency. Đây là cơ chế chống click đúp và retry mạng tạo bản ghi trùng.",
        "Bằng chứng: chỉ mục unique có điều kiện process_events_idempotency_key_unique và nhánh kiểm tra trong Store.Transition.",
    ),
    (
        "Tại sao phải có cả Idempotency-Key và Correlation ID?",
        "Idempotency-Key chống xử lý lặp; Correlation ID dùng để liên kết dấu vết của cùng một luồng xử lý.",
        "Idempotency là chìa khóa kỹ thuật để một request retry không tạo transition trùng. Correlation ID không quyết định chống trùng; nó là mã truy vết xuyên qua event/audit/notification để điều tra sự cố hoặc theo dõi một luồng. Hai khái niệm bổ trợ chứ không thay thế nhau.",
        "Câu nói ngắn khi bị hỏi vặn: “Một cái bảo đảm đúng một lần xử lý, một cái giúp lần dấu.”",
    ),
    (
        "Hủy phiếu có nằm trong transaction không?",
        "Có. Đổi trạng thái, ghi process event, audit log và xử lý ngân sách liên quan đều nằm trong transaction PostgreSQL.",
        "Store.Transition mở transaction, khóa phiếu, kiểm tra quyền/trạng thái/version, cập nhật purchase_requests, chèn process_events, audit_logs, áp dụng budget transition rồi mới commit. Hàm queue notification cũng được gọi trong transaction chung, nhưng nhánh CANCELLED hiện không tạo thông báo mới. Nếu một bước database lỗi, defer rollback khiến các cập nhật không bị dở dang. Điều này giúp trạng thái phiếu và lịch sử luôn nhất quán trong database.",
        "Bằng chứng: Store.Transition có Begin, Rollback, Commit; các lời gọi insertAudit/applyBudgetTransition/queueTransitionNotification trước Commit.",
    ),
    (
        "Hủy có giải phóng ngân sách không?",
        "Trong luồng hiện tại, hủy chỉ xảy ra trước khi phiếu được giữ ngân sách, nên không có thao tác giải phóng trực tiếp khi CANCEL.",
        "Ngân sách được reserve khi phiếu qua bước quản lý/tài chính theo rule. Nếu phiếu bị trả về CHANGES_REQUESTED từ trạng thái MANAGER_APPROVED, reservation đã được RELEASED ngay tại thao tác yêu cầu chỉnh sửa. Sau đó người tạo có thể hủy; lúc này không còn khoản giữ chỗ cần giải phóng. Không nên nói hủy luôn giải phóng ngân sách, vì code hiện tại đúng hơn là giới hạn hủy ở trạng thái chưa có reservation hoạt động.",
        "Điểm kiểm chứng: applyBudgetTransition chỉ release khi MANAGER_APPROVED chuyển sang REJECTED hoặc CHANGES_REQUESTED.",
    ),
    (
        "Nếu hủy thất bại giữa chừng thì dữ liệu ra sao?",
        "Transaction rollback, nên database không ở trạng thái nửa hủy.",
        "Các kiểm tra quyền, version, insert event/audit và update status cùng transaction. Nếu insert audit hay xếp notification lỗi trước commit, transaction bị rollback. Client nhận lỗi và phiếu giữ nguyên trạng thái cũ. Đây là tính nguyên tử ACID trong phạm vi PostgreSQL.",
        "Demo có thể mô tả: cố ý dùng expectedVersion cũ, server trả conflict và không có event CANCELLED mới.",
    ),
    (
        "Có thể khôi phục phiếu đã hủy không?",
        "Chưa có chức năng khôi phục/restore trong state machine hiện tại.",
        "CANCELLED là trạng thái kết thúc: DecideTransition không định nghĩa transition đi ra từ trạng thái này. Dữ liệu vẫn được giữ để truy vết, nhưng không cho sửa/gửi lại để tránh làm mơ hồ lịch sử. Nếu nghiệp vụ cần tiếp tục, hướng an toàn là tạo phiếu mới hoặc bổ sung một UC “mở lại” có phê duyệt và audit riêng.",
        "Trả lời đúng phạm vi: không nhận rằng nút khôi phục đã tồn tại.",
    ),
    (
        "Hệ thống kiểm thử UC hủy như thế nào?",
        "Kiểm thử theo các nhánh: đúng quyền/đúng trạng thái, sai quyền, sai trạng thái, version conflict, request lặp và kiểm tra timeline/audit.",
        "Bộ kiểm thử tốt không chỉ click UI. Cần kiểm thử domain DecideTransition với requester/manager/finance; integration test API transitions với Idempotency-Key; kiểm tra database sau commit có status CANCELLED, event và audit; đồng thời kiểm thử retry cùng key không tạo dòng lịch sử thứ hai. Đây là các rủi ro lớn nhất của use case workflow.",
        "Checklist demo: hủy DRAFT thành công; tài khoản khác bị chặn; đã gửi duyệt không thấy hủy; timeline có CANCELLED.",
    ),
    (
        "Vì sao không dùng DELETE HTTP để hủy phiếu?",
        "Vì hủy là một quyết định nghiệp vụ, không phải xóa tài nguyên kỹ thuật.",
        "DELETE thường gợi ý xóa/ẩn tài nguyên. Ở DX-OS, hủy phải đi qua kiểm tra state machine, phân quyền, version, idempotency, audit và thông báo. POST transitions với action CANCEL biểu đạt rõ đây là event thay đổi quy trình; nó cũng đồng nhất với gửi duyệt, duyệt, từ chối và yêu cầu chỉnh sửa.",
        "Gợi ý trình bày: “Tôi mô hình hóa hủy là một business transition, không mô hình hóa như delete.”",
    ),
    (
        "Điểm hạn chế và hướng nâng cấp của UC hủy là gì?",
        "Hiện chưa có luồng yêu cầu hủy sau khi đã gửi duyệt, chưa có khôi phục và lý do hủy chưa bắt buộc.",
        "Đây là các lựa chọn phạm vi hiện tại để giảm rủi ro sai lệch workflow. Nâng cấp hợp lý là thêm CANCEL_REQUESTED cho phiếu đã gửi: requester tạo yêu cầu, người đang phụ trách xác nhận, hệ thống xử lý reservation/PO theo từng trạng thái. Có thể bắt buộc cancellation reason và thêm dashboard theo lý do hủy, nhưng phải tách thành use case mới, không phá lịch sử cũ.",
        "Nêu rõ: đây là roadmap, không phải chức năng đã triển khai.",
    ),
]


ATTACHMENT_QA = [
    (
        "UC Quản lý tệp đính kèm có phạm vi gì?",
        "Cho phép xem danh sách, tải lên, tải xuống và xóa tệp chứng từ của một phiếu mua sắm.",
        "Tệp được gắn với purchase request và có loại tài liệu: báo giá, đặc tả kỹ thuật, hợp đồng hoặc tài liệu khác. UC này không chỉ là upload file: nó có phân quyền, kiểm tra định dạng/nội dung, giới hạn dung lượng, kiểm tra tính toàn vẹn khi tải xuống, audit và quy tắc bắt buộc tệp trước khi gửi duyệt.",
        "Demo: tại chi tiết phiếu, xem danh sách tệp, chọn loại, tải PDF lên, tải xuống, xóa khi phiếu còn Bản nháp.",
    ),
    (
        "Tệp thực tế được lưu ở đâu, metadata lưu ở đâu?",
        "Nội dung file lưu ở Nextcloud; PostgreSQL lưu metadata, đường dẫn nội bộ, checksum, người tải lên và trạng thái.",
        "Không lưu binary lớn trực tiếp trong bảng nghiệp vụ giúp database gọn hơn và tận dụng Nextcloud cho lưu trữ tài liệu. Bảng purchase_request_attachments giữ original_name, content_type, size_bytes, checksum_sha256, storage_path, storage_etag, uploaded_by và status. storage_path không phải URL công khai; backend dùng nó để gọi WebDAV tới Nextcloud.",
        "Bằng chứng: migrations/000006_purchase_request_attachments.sql và platform/documentstore/nextcloud.go.",
    ),
    (
        "Ai được tải lên/xóa tệp?",
        "Chỉ người tạo phiếu, khi phiếu là DRAFT hoặc CHANGES_REQUESTED; auditor không được tải lên hay xóa.",
        "UploadAttachment và DeleteAttachment đều khóa phiếu, kiểm tra requester_id có bằng user hiện tại không, rồi kiểm tra trạng thái có phải DRAFT/CHANGES_REQUESTED không. Kiểm tra nằm trên server, vì vậy không thể vượt bằng cách sửa HTML hoặc gọi trực tiếp API. Người có quyền xem phiếu hợp lệ có thể tải xuống tệp ACTIVE, theo cơ chế Get của phiếu.",
        "Demo: tải lên/xóa với employee là chủ phiếu; chuyển trạng thái sang Submitted để thấy thao tác chỉnh tệp bị chặn.",
    ),
    (
        "Vì sao chỉ cho sửa tệp ở Bản nháp hoặc Yêu cầu chỉnh sửa?",
        "Để bộ chứng từ đã gửi duyệt không bị thay đổi âm thầm.",
        "Sau khi phiếu SUBMITTED, người phê duyệt phải thấy bộ hồ sơ ổn định. Nếu cần cập nhật chứng từ, họ trả phiếu về CHANGES_REQUESTED; lúc này người tạo được mở lại quyền sửa tệp rồi gửi lại. Cách này tạo ranh giới rõ giữa “hồ sơ đang soạn” và “hồ sơ đã trình”.",
        "Liên hệ workflow: attachment editability dùng cùng hai trạng thái mà người tạo được sửa/hủy.",
    ),
    (
        "Hệ thống hỗ trợ định dạng và dung lượng nào?",
        "PDF, DOCX, XLSX, JPG/JPEG và PNG; kích thước từ 1 byte đến tối đa 10 MB cho mỗi tệp.",
        "AllowedAttachmentContentTypes định nghĩa whitelist MIME type. MaxAttachmentSize là 10 × 1024 × 1024 bytes; migration cũng có CHECK size_bytes <= 10485760. Giao diện kiểm tra sớm để phản hồi tốt, nhưng server vẫn kiểm tra lại vì client không đáng tin cậy.",
        "Demo: thử tải .exe hoặc tệp >10 MB để nhận thông báo lỗi; không dùng tệp thật nhạy cảm khi demo.",
    ),
    (
        "Tại sao kiểm tra phần mở rộng thôi là chưa đủ?",
        "Vì có thể đổi tên file độc hại thành .pdf; hệ thống còn kiểm tra MIME, phần mở rộng và chữ ký/nội dung thực tế.",
        "ValidateAttachment đối chiếu ba lớp: content type có nằm trong whitelist, tên file có phần mở rộng khớp type, và bytes đầu/entry trong file có đúng cấu trúc. PDF phải bắt đầu %PDF-, ảnh kiểm signature, DOCX/XLSX phải là ZIP có entry word/document.xml hoặc xl/workbook.xml. Vì vậy file payload.exe đổi tên bao-gia.pdf sẽ bị từ chối.",
        "Bằng chứng: attachmentNameMatchesType và attachmentContentMatchesType; có unit test RejectsSpoofedContentType.",
    ),
    (
        "Hệ thống chống path traversal qua tên file như thế nào?",
        "Không dùng tên file do người dùng đặt làm đường dẫn lưu trữ và chặn /, \\, ký tự NUL, CR/LF trong tên.",
        "Backend tự sinh storage_path theo purchase-request ID và attachment ID: purchase-requests/{requestId}/{attachmentId}. ValidateAttachment chỉ nhận tên hiển thị an toàn, dài tối đa 255 ký tự. Adapter Nextcloud còn từ chối segment rỗng, . hoặc .. trước khi tạo WebDAV request. Vì thế chuỗi ../ hoặc ký tự xuống dòng không thể điều khiển đường dẫn hay header.",
        "Demo kỹ thuật: thử tên ../payload.pdf qua API sẽ bị validation error.",
    ),
    (
        "Luồng upload diễn ra như thế nào?",
        "Tạo metadata tạm UPLOADING trong DB, gửi file sang Nextcloud, rồi mới chuyển metadata sang ACTIVE và ghi audit.",
        "Cụ thể: server validate; mở transaction, kiểm tra quyền/trạng thái, sinh ID/path và insert metadata status UPLOADING rồi commit. Sau đó gọi documents.Put tới Nextcloud kèm SHA-256. Nếu thành công, mở transaction thứ hai để đổi ACTIVE, lưu ETag, uploaded_at và ghi ATTACHMENT_UPLOADED vào audit. Người dùng chỉ nhìn thấy tệp ACTIVE; trạng thái tạm không xuất hiện trong list.",
        "Vẽ miệng: PostgreSQL (UPLOADING) → Nextcloud PUT → PostgreSQL (ACTIVE + audit).",
    ),
    (
        "Tại sao upload không dùng một transaction duy nhất?",
        "Vì PostgreSQL và Nextcloud là hai hệ thống khác nhau, không có transaction phân tán/2PC trong thiết kế này.",
        "Không thể giữ transaction database mở trong lúc gửi file qua mạng một cách an toàn và cũng không có commit chung với WebDAV. DX-OS dùng pattern bù trừ: metadata UPLOADING là trạng thái tạm; nếu Nextcloud lỗi, xóa metadata tạm; nếu bước finalize DB lỗi, xóa file ở Nextcloud và xóa metadata tạm. Đây là eventual consistency có kiểm soát, thực tế hơn 2PC cho phạm vi đồ án.",
        "Điểm nhấn khi bị hỏi sâu: không khẳng định ACID xuyên Nextcloud; ACID chỉ trong PostgreSQL, còn giữa hai hệ dùng compensating cleanup.",
    ),
    (
        "Nếu Nextcloud mất kết nối khi đang tải lên thì sao?",
        "Upload trả lỗi, metadata UPLOADING được dọn dẹp; tệp không được hiển thị là đã tải thành công.",
        "Sau Put lỗi, code dùng context.WithoutCancel để DELETE metadata còn UPLOADING. Nếu lỗi ở finalize transaction sau khi file đã lên, cleanupUploadingAttachment gọi xóa file ở document store và xóa metadata tạm. Người dùng có thể thử lại; list chỉ truy vấn status ACTIVE nên không lẫn tệp lỗi với tệp hợp lệ.",
        "Demo mô tả: dừng Nextcloud hoặc dùng endpoint lỗi rồi thử upload; UI báo lỗi, danh sách không tăng tệp ACTIVE.",
    ),
    (
        "Checksum SHA-256 dùng để làm gì?",
        "Để phát hiện file bị thay đổi/hỏng giữa lúc lưu và lúc tải xuống.",
        "Khi upload, server tính SHA-256 trên bytes nhận được và lưu chuỗi hex 64 ký tự vào DB; Nextcloud cũng nhận header OC-Checksum. Khi download, backend lấy bytes từ storage, tính lại SHA-256 rồi so với checksum đã lưu. Nếu khác, không trả nội dung mà báo lỗi document store checksum mismatch. Đây là kiểm tra toàn vẹn, không phải mã hóa file.",
        "Câu ngắn: “Checksum trả lời file còn nguyên không; nó không dùng để che giấu nội dung.”",
    ),
    (
        "Khi tải xuống, hệ thống kiểm tra quyền thế nào?",
        "Trước khi lấy file, backend gọi quyền xem phiếu; đồng thời buộc attachment ID phải thuộc đúng request ID và có status ACTIVE.",
        "DownloadAttachment gọi s.Get với principal và requestID để áp dụng phạm vi truy cập nghiệp vụ. Sau đó getStoredAttachment truy vấn với cả attachment ID và purchase_request_id. Vì thế không thể đổi attachment ID trong URL để lấy tệp của phiếu khác; URL content không trỏ thẳng tới Nextcloud mà đi qua API đã xác thực.",
        "Demo: dùng attachment ID của phiếu A trong URL của phiếu B sẽ không tìm thấy/tải được.",
    ),
    (
        "Xóa tệp có phải xóa vĩnh viễn không?",
        "Nội dung file được yêu cầu xóa khỏi Nextcloud; metadata trong database được giữ với trạng thái DELETED để còn dấu vết.",
        "DeleteAttachment đổi ACTIVE sang DELETING, gọi DELETE WebDAV, rồi đổi metadata sang DELETED và ghi deleted_at. Danh sách chỉ hiển thị ACTIVE nên tệp biến mất khỏi giao diện, còn hàng metadata vẫn hữu ích cho audit/điều tra. Đây là khác với phiếu mua sắm: phiếu dùng hủy mềm trạng thái CANCELLED; tệp có xóa nội dung vật lý nhưng giữ lịch sử metadata.",
        "Bằng chứng: status enum UPLOADING, ACTIVE, DELETING, DELETED và UPDATE ... deleted_at = now().",
    ),
    (
        "Nếu xóa file trên Nextcloud thất bại thì sao?",
        "Metadata được trả từ DELETING về ACTIVE để người dùng không mất dấu file mà vẫn không xóa được file thật.",
        "Sau khi commit DELETING, nếu documents.Delete lỗi, code dùng context.WithoutCancel để UPDATE trạng thái trở về ACTIVE, rồi trả lỗi. Nhờ vậy DB không nói “đã xóa” trong khi file thật vẫn tồn tại. Nếu delete ngoài storage thành công mà finalize DB sau đó lỗi, đây là biên lỗi hiếm; thiết kế hiện tại báo lỗi để vận hành theo dõi, và có thể bổ sung job đối soát định kỳ trong roadmap.",
        "Điểm trung thực: cơ chế bù hiện có cho lỗi xóa storage; job reconciliation tự động chưa thấy được triển khai trong code hiện tại.",
    ),
    (
        "Tại sao trạng thái UPLOADING và DELETING không được hiển thị cho người dùng?",
        "Đó là trạng thái trung gian để bảo đảm người dùng chỉ thao tác trên tệp hoàn chỉnh, nhất quán.",
        "ListAttachments và getStoredAttachment đều lọc pa.status = ACTIVE. Trong lúc upload chưa finalize, file có thể chưa có đầy đủ metadata/ETag/audit; trong lúc delete, file đang được xử lý ở storage. Ẩn các trạng thái tạm giúp không có tình huống click tải một file nửa chừng hoặc tải file vừa bị xóa.",
        "Bằng chứng: câu query list/download đều có điều kiện status = 'ACTIVE'.",
    ),
    (
        "Tệp đính kèm liên quan gì tới việc gửi duyệt?",
        "Tùy chính sách, phiếu vượt ngưỡng có thể bắt buộc phải có một loại chứng từ trước khi gửi/gửi lại duyệt.",
        "Bảng attachment_rules cấu hình theo tổ chức và tiền tệ: ngưỡng tiền, loại tài liệu yêu cầu và active. Khi SUBMIT hoặc RESUBMIT, requireAttachmentForSubmission kiểm tra tổng tiền có đạt ngưỡng không và có ít nhất một file ACTIVE đúng document_type không. Thiếu thì trả ErrAttachmentRequired và không đổi trạng thái phiếu.",
        "Ví dụ mặc định migration: phiếu VND từ 20.000.000 trở lên cần QUOTATION/Báo giá.",
    ),
    (
        "Ai cấu hình quy tắc tệp bắt buộc?",
        "Quản trị DX-OS (role dx_admin), không phải nhân viên thông thường hay kiểm toán viên.",
        "Policy Center đọc attachment_rules theo tổ chức. UpdateAttachmentPolicy kiểm tra dx_admin và không cho auditor cập nhật; đồng thời kiểm tra version để tránh hai quản trị viên ghi đè cấu hình của nhau. Thay đổi chính sách cũng ghi audit ATTACHMENT_POLICY_UPDATED.",
        "Demo: tài khoản quản trị vào Chính sách để xem/cập nhật ngưỡng; hãy cẩn thận không đổi dữ liệu demo khi bảo vệ.",
    ),
    (
        "Làm sao chống request upload quá lớn hoặc multipart bất thường?",
        "HTTP handler giới hạn toàn body, giới hạn mỗi file và giới hạn số phần multipart.",
        "r.Body được bọc MaxBytesReader với 10 MB cộng phần overhead; parser chỉ chấp nhận đúng một part file có tên, documentType tối đa 128 bytes và tối đa 10 multipart parts. Nội dung file được đọc qua LimitReader MaxAttachmentSize + 1 rồi ValidateAttachment kiểm tra lại. Đây là lớp chống tiêu hao bộ nhớ/băng thông và multipart có cấu trúc bất thường.",
        "Bằng chứng: maxAttachmentRequestBytes, maxAttachmentMultipartParts và TestParseAttachmentUploadStreamsBoundedMultipart.",
    ),
    (
        "Có chống virus hoàn chỉnh không?",
        "Chưa có quét antivirus/anti-malware tích hợp trong code hiện tại; hệ thống mới có whitelist, kiểm chữ ký file và giới hạn dung lượng.",
        "Đây là câu cần trả lời thẳng. Việc kiểm magic bytes giảm giả mạo loại file nhưng không thay thế antivirus. Hướng production là đưa file vào quarantine, gọi ClamAV hoặc dịch vụ scan, chỉ chuyển ACTIVE sau khi PASS; đồng thời chặn macro/phân tích sandbox đối với Office. Với phạm vi đồ án, lớp kiểm tra hiện tại giảm rủi ro đầu vào cơ bản.",
        "Không nên nói “đã an toàn tuyệt đối”; hãy nêu biện pháp hiện có và roadmap.",
    ),
    (
        "Tệp có được mã hóa không?",
        "Mã nguồn hiện tại không triển khai mã hóa nội dung ở tầng ứng dụng; bảo mật dựa vào xác thực, phân quyền và hạ tầng triển khai.",
        "Nextcloud adapter dùng URL cấu hình http hoặc https; production nên bắt buộc HTTPS, mã hóa volume/backup, quản lý secret của tài khoản service và phân quyền Nextcloud tối thiểu. Mã nguồn không ghi secret vào database; username/password Nextcloud lấy từ biến môi trường. Nếu yêu cầu mã hóa end-to-end, đó là cải tiến kiến trúc cần quản lý khóa riêng.",
        "Điểm trung thực: không nhầm checksum với encryption.",
    ),
    (
        "Các loại tài liệu được phân loại để làm gì?",
        "Để người dùng và quy tắc nghiệp vụ biết chứng từ nào đang có và chứng từ nào là bắt buộc.",
        "document_type gồm QUOTATION, SPECIFICATION, CONTRACT, OTHER. Một file PDF không chỉ là “file”; nó mang ý nghĩa báo giá/đặc tả/hợp đồng. Khi submit, hệ thống kiểm tra đúng loại yêu cầu, không chỉ đếm số lượng file. Cách này hỗ trợ báo cáo, audit và giảm trường hợp tải một tệp không liên quan để vượt qua điều kiện.",
        "Demo: chọn loại Báo giá rồi upload PDF; ở rule bắt buộc QUOTATION, file OTHER sẽ không thỏa điều kiện.",
    ),
    (
        "Tại sao không dùng URL Nextcloud trực tiếp trên trình duyệt?",
        "Để không lộ credential/path nội bộ và để mọi tải xuống luôn đi qua kiểm tra quyền của DX-OS.",
        "Frontend gọi API /attachments/{id}/content sau khi đã có token DX-OS. Backend kiểm tra quyền phiếu, lấy file bằng service account WebDAV rồi trả content với Content-Disposition. Vì vậy người dùng không nhận được username/password Nextcloud hoặc URL lưu trữ nội bộ; đồng thời checksum được kiểm trước khi trả file.",
        "Bằng chứng: procurement.service downloadAttachment gọi API, còn Nextcloud adapter chỉ nằm ở backend.",
    ),
    (
        "Các test quan trọng của UC tệp đính kèm là gì?",
        "Test validation file, parser multipart có giới hạn, adapter Nextcloud WebDAV và các nhánh nghiệp vụ upload/download/delete.",
        "Trong source đã có unit test cho PDF hợp lệ, tên file nguy hiểm, MIME giả và phần mở rộng sai; HTTP test cho multipart bị giới hạn; Nextcloud adapter test auto-create folder, not found và unsafe path. Khi bảo vệ, có thể nói test được chia tầng: domain validation, HTTP boundary và adapter integration; với database cần thêm integration test quyền/trạng thái/compensation nếu mở rộng tiếp.",
        "Mã test: procurement/model_test.go, httpapi/purchase_requests_test.go, documentstore/nextcloud_test.go.",
    ),
    (
        "Điểm hạn chế và hướng nâng cấp của UC tệp đính kèm là gì?",
        "Chưa có antivirus, versioning tệp, preview và job đối soát tự động DB–Nextcloud.",
        "Roadmap hợp lý: quarantine + scan trước ACTIVE; versioning thay vì xóa rồi upload lại; preview có kiểm soát; background reconciliation phát hiện metadata/file bị lệch; retention policy cho DELETED; object storage/S3 nếu quy mô lớn. Những cải tiến này phải giữ nguyên nguyên tắc hiện có: server-side authorization, metadata tách binary và audit.",
        "Nêu roadmap như đề xuất, không mô tả là tính năng đã có.",
    ),
]


CROSS_QA = [
    (
        "Nếu mất điện hoặc service bị dừng trong khi đang xử lý thì dữ liệu có mất không?",
        "Phần thay đổi trong PostgreSQL được transaction commit hoặc rollback; file dùng trạng thái tạm và cơ chế dọn dẹp bù trừ.",
        "Với hủy phiếu, trạng thái/event/audit chỉ xuất hiện sau commit nên không có “nửa hủy” trong DB. Với upload file vì có hai hệ thống, code dùng UPLOADING/ACTIVE và cleanup. Để chịu sự cố hạ tầng tốt hơn, production cần PostgreSQL volume bền vững, backup/restore đã kiểm thử, UPS, healthcheck và job đối soát file tạm. Không hệ thống nào tránh mất dữ liệu nếu storage không có persistence/backup.",
        "Đây là câu kiến trúc: phân biệt transaction DB với nhất quán giữa DB và external storage.",
    ),
    (
        "Correlation ID là gì?",
        "Là mã theo dõi dùng để nối các log/event/audit thuộc cùng một luồng yêu cầu.",
        "Khi một thao tác có thể tạo process event, audit log, notification và gọi các service khác, correlation ID giúp truy vấn các bản ghi liên quan để debug. Nó không cấp quyền và không chống click đúp; nhiệm vụ chống xử lý lặp thuộc Idempotency-Key.",
        "Câu một dòng: “Correlation để lần dấu; Idempotency để không làm hai lần.”",
    ),
    (
        "SLA là gì và hai UC này liên quan thế nào?",
        "SLA là thời hạn/mức cam kết xử lý. Hủy và quản lý tệp không tự tạo SLA riêng, nhưng ảnh hưởng chất lượng luồng phê duyệt.",
        "Trong DX-OS, SLA due date được thiết lập khi submit/resubmit. Tệp đầy đủ giúp phiếu không bị trả về do thiếu chứng từ; hủy sớm loại bỏ nhu cầu không còn cần thiết khỏi luồng. Vì hủy chỉ ở giai đoạn sớm, nó không được coi là thao tác giải quyết SLA của phiếu đang chờ duyệt.",
        "Nếu giảng viên hỏi thuật ngữ: dùng “thời hạn xử lý cam kết” thay vì chỉ đọc chữ SLA.",
    ),
    (
        "Vì sao phải có audit log nếu đã có timeline?",
        "Timeline hướng nghiệp vụ cho một phiếu; audit log phục vụ tra soát toàn hệ thống và truy vấn theo actor/resource/correlation.",
        "Cả hai có thể ghi cùng hành động nhưng mục đích khác. process_events gắn chặt workflow purchase request để hiển thị quá trình; audit_logs dùng resource_type/resource_id tổng quát, hỗ trợ kiểm toán các tài nguyên khác như policy. Giữ hai lớp giúp UI nghiệp vụ đơn giản nhưng vẫn có bằng chứng quản trị rộng hơn.",
        "Demo: timeline trên chi tiết phiếu và màn hình Kiểm toán là hai góc nhìn khác nhau của dấu vết.",
    ),
    (
        "Bảo mật chỉ dựa vào việc ẩn nút trên giao diện có đủ không?",
        "Không. Giao diện chỉ hỗ trợ trải nghiệm; backend mới là nơi cưỡng chế quyền và trạng thái.",
        "Angular chỉ hiện nút hủy/tải lên/xóa khi phù hợp. Tuy nhiên người dùng có thể tự gửi HTTP request, nên Store.Transition, UploadAttachment và DeleteAttachment đều tự kiểm tra principal, requester, role và trạng thái. Đây là defense in depth: UX giảm thao tác sai, server ngăn vi phạm thực sự.",
        "Từ khóa nên nói: “authorization ở server side”.",
    ),
    (
        "Bạn chọn các tiêu chí đánh giá chất lượng hai UC này như thế nào?",
        "Đúng nghiệp vụ, an toàn phân quyền, toàn vẹn dữ liệu, chịu lỗi, khả năng truy vết và dễ dùng.",
        "Hủy được đánh giá qua state machine, event/audit, version/idempotency. Tệp được đánh giá qua whitelist + signature, giới hạn size, ownership, checksum, trạng thái tạm và compensation khi storage lỗi. Các tiêu chí này gắn trực tiếp với rủi ro thật của hệ thống mua sắm, không chỉ là giao diện chạy được.",
        "Gợi ý kết luận bảo vệ: “Tôi ưu tiên tính đúng và khả năng truy vết hơn là chỉ hoàn thành thao tác CRUD.”",
    ),
]


def shade(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def borders(cell, color: str = GRID) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.first_child_found_in("w:tcBorders")
    if node is None:
        node = OxmlElement("w:tcBorders")
        tc_pr.append(node)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_node = node.find(qn(f"w:{edge}"))
        if edge_node is None:
            edge_node = OxmlElement(f"w:{edge}")
            node.append(edge_node)
        edge_node.set(qn("w:val"), "single")
        edge_node.set(qn("w:sz"), "5")
        edge_node.set(qn("w:color"), color)


def margins(cell, top=110, start=150, bottom=110, end=150) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.first_child_found_in("w:tcMar")
    if node is None:
        node = OxmlElement("w:tcMar")
        tc_pr.append(node)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        child = node.find(qn(f"w:{side}"))
        if child is None:
            child = OxmlElement(f"w:{side}")
            node.append(child)
        child.set(qn("w:w"), str(value))
        child.set(qn("w:type"), "dxa")


def set_run(run, size=10.5, color=NAVY, bold=False, italic=False) -> None:
    run.font.name = "Aptos"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def text(paragraph, value: str, **kwargs):
    run = paragraph.add_run(value)
    set_run(run, **kwargs)
    return run


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def no_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    element = OxmlElement("w:cantSplit")
    tr_pr.append(element)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    text(paragraph, "DX-OS  |  Tài liệu bảo vệ  |  Trang ", size=8.5, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def setup(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.55)
    section.bottom_margin = Cm(1.35)
    section.left_margin = Cm(1.55)
    section.right_margin = Cm(1.55)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.7)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.16
    for style_name, size, color in (("Title", 25, NAVY), ("Heading 1", 16, TEAL_DARK), ("Heading 2", 13, NAVY), ("Heading 3", 11.5, TEAL_DARK)):
        style = doc.styles[style_name]
        style.font.name = "Aptos Display" if style_name == "Title" else "Aptos"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), style.font.name)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12 if style_name != "Title" else 0)
        style.paragraph_format.space_after = Pt(6)
    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    text(header_p, "DX-OS LAB  •  BỘ CÂU HỎI BẢO VỆ", size=8.5, color=TEAL, bold=True)
    add_page_number(section.footer.paragraphs[0])


def add_bullet(doc: Document, content: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    text(p, content, size=10.2)


def add_info_box(doc: Document, heading: str, body: str, fill=PALE_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, fill)
    borders(cell)
    margins(cell, 130, 180, 130, 180)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    text(p, heading, size=10.5, color=TEAL_DARK, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    text(p2, body, size=10.2, color=NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_qa(doc: Document, number: int, question: str, short: str, deep: str, evidence: str) -> None:
    table = doc.add_table(rows=4, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    labels = [
        (f"{number:02d}. {question}", WHITE, TEAL_DARK, 11.2, True),
        ("TRẢ LỜI GỌN — " + short, PALE_TEAL, NAVY, 10.2, False),
        ("TRẢ LỜI KỸ — " + deep, PALE_BLUE, NAVY, 10.0, False),
        ("DẪN CHỨNG / CÁCH DEMO — " + evidence, PALE_GOLD, NAVY, 9.6, False),
    ]
    for index, (value, fill, color, size, bold) in enumerate(labels):
        cell = table.cell(index, 0)
        shade(cell, fill)
        borders(cell)
        margins(cell, 105, 155, 105, 155)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.1
        if index == 0:
            text(p, value, size=size, color=color, bold=bold)
        else:
            prefix, content = value.split(" — ", 1)
            text(p, prefix + " — ", size=size, color=TEAL_DARK, bold=True)
            text(p, content, size=size, color=color)
    for row in table.rows:
        no_split(row)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    spacer.paragraph_format.space_before = Pt(0)


def add_two_column_table(doc: Document, rows, widths=(5.0, 11.6)) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    header = table.rows[0]
    for cell, label in zip(header.cells, ("Nội dung", "Ghi nhớ khi bảo vệ")):
        shade(cell, TEAL_DARK)
        borders(cell, TEAL_DARK)
        margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        text(p, label, size=10, color=WHITE, bold=True)
    set_repeat_header(header)
    for left, right in rows:
        cells = table.add_row().cells
        for cell, value, fill in ((cells[0], left, PALE_TEAL), (cells[1], right, WHITE)):
            shade(cell, fill)
            borders(cell)
            margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            text(p, value, size=9.7, color=NAVY, bold=(cell == cells[0]))
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        no_split(table.rows[-1])
    for row in table.rows:
        row.cells[0].width = Cm(widths[0])
        row.cells[1].width = Cm(widths[1])


def add_section_title(doc: Document, title: str, subtitle: str) -> None:
    doc.add_heading(title, level=1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    text(p, subtitle, size=10.5, color=SLATE, italic=True)


def add_qa_section(doc: Document, title: str, subtitle: str, questions) -> None:
    doc.add_page_break()
    add_section_title(doc, title, subtitle)
    for index, item in enumerate(questions, start=1):
        add_qa(doc, index, *item)


def build_document() -> Document:
    doc = Document()
    setup(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    text(p, "DX-OS", size=13, color=TEAL, bold=True)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(7)
    text(p2, "BỘ CÂU HỎI BẢO VỆ ĐỒ ÁN", size=25, color=NAVY, bold=True)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(18)
    text(p3, "Use case Hủy phiếu & Quản lý tệp đính kèm", size=16, color=TEAL_DARK, bold=True)
    add_info_box(
        doc,
        "Mục tiêu tài liệu",
        "Giúp trả lời nhanh khi bảo vệ và có sẵn phần giải thích kỹ thuật khi giảng viên hỏi sâu. Nội dung bám mã nguồn DX-OS hiện tại; các hướng nâng cấp được ghi rõ là roadmap, không nhầm với chức năng đã triển khai.",
        PALE_TEAL,
    )
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_before = Pt(22)
    text(p4, "Phạm vi thực hiện: workflow phiếu mua sắm, PostgreSQL, Nextcloud/WebDAV, API và giao diện Angular.", size=10.5, color=SLATE, italic=True)
    doc.add_page_break()

    add_section_title(doc, "1. Cách dùng khi bảo vệ", "Đừng học thuộc từng câu. Hãy nắm luồng, sau đó chọn phiên bản trả lời phù hợp mức độ câu hỏi.")
    add_two_column_table(doc, [
        ("Khi giảng viên hỏi nhanh", "Dùng phần “Trả lời gọn”: 1–2 câu, nêu mục tiêu nghiệp vụ và cơ chế chính."),
        ("Khi bị hỏi vặn", "Mở rộng theo phần “Trả lời kỹ”: quyền → trạng thái → transaction/dữ liệu → kiểm thử → giới hạn hiện tại."),
        ("Khi demo", "Luôn chỉ ra bằng chứng: trạng thái trên màn hình, Timeline, Kiểm toán hoặc tab Network. Không chỉ nói lý thuyết."),
        ("Khi chưa có chức năng", "Nói thẳng “mã hiện tại chưa triển khai”, sau đó nêu hướng nâng cấp hợp lý. Không nhận là đã làm nếu chưa có trong code."),
        ("Cấu trúc trả lời an toàn", "Vấn đề → quyết định thiết kế → cách code bảo vệ → kết quả kiểm chứng."),
    ])
    add_info_box(doc, "Câu mở đầu 30 giây", "Phần tôi phụ trách bảo vệ hai điểm dễ gây rủi ro trong hệ thống mua sắm: dừng một nhu cầu không còn hợp lệ mà không mất lịch sử, và quản lý chứng từ đính kèm an toàn. Tôi mô hình hóa hủy như một chuyển trạng thái có phân quyền, version và audit; còn tệp được tách nội dung sang Nextcloud, metadata ở PostgreSQL, kiểm tra type/size/nội dung, checksum và cơ chế xử lý lỗi giữa hai hệ thống.", PALE_GOLD)

    add_section_title(doc, "2. Bản đồ luồng cần nhớ", "Hai use case gắn trực tiếp vào vòng đời phiếu mua sắm.")
    add_two_column_table(doc, [
        ("Hủy phiếu", "Người tạo → phiếu DRAFT hoặc CHANGES_REQUESTED → POST transitions/CANCEL → kiểm tra quyền + version + idempotency → status CANCELLED + process event + audit → commit."),
        ("Tải lên tệp", "Người tạo ở trạng thái có thể chỉnh sửa → validate file → DB UPLOADING → Nextcloud PUT → DB ACTIVE + audit. Lỗi thì dọn metadata/file tạm."),
        ("Tải xuống tệp", "Token người dùng → kiểm tra quyền xem phiếu → kiểm tra file thuộc phiếu và ACTIVE → Nextcloud GET → đối chiếu SHA-256 → trả file."),
        ("Xóa tệp", "Người tạo ở trạng thái có thể chỉnh sửa → DB DELETING → Nextcloud DELETE → DB DELETED + audit; lỗi storage thì trả ACTIVE."),
        ("Gửi duyệt", "Nếu rule tệp có hiệu lực và tổng tiền đạt ngưỡng → phải có ít nhất một tệp ACTIVE đúng loại; thiếu thì chặn SUBMIT/RESUBMIT."),
    ])

    add_section_title(doc, "3. Thuật ngữ nên nói dễ hiểu", "Nói tiếng Việt trước, sau đó mới dùng thuật ngữ kỹ thuật nếu cần.")
    add_two_column_table(doc, [
        ("State machine", "Quy tắc các trạng thái nào được phép chuyển sang trạng thái nào."),
        ("Transaction", "Một gói cập nhật: hoặc thành công hết, hoặc quay lại như chưa làm trong database."),
        ("Optimistic locking / expectedVersion", "Kiểm tra phiên bản để không ghi đè thao tác mới hơn từ người/tab khác."),
        ("Idempotency-Key", "Mã chống xử lý lặp khi bấm hai lần hoặc ứng dụng gửi lại request do mạng chập chờn."),
        ("Correlation ID", "Mã liên kết log, audit, notification của cùng một luồng để dễ tra lỗi."),
        ("Checksum SHA-256", "Mã băm để kiểm tra file tải về có còn nguyên vẹn, không phải mã hóa."),
        ("SLA", "Thời hạn xử lý cam kết của quy trình."),
        ("Compensating cleanup", "Thao tác bù trừ: nếu lưu file hoặc cập nhật DB ở bước sau lỗi, hệ thống dọn phần tạm đã tạo trước đó."),
    ])

    add_qa_section(doc, "4. Câu hỏi bảo vệ — Use case Hủy phiếu", "Bao phủ nghiệp vụ, phân quyền, trạng thái, đồng thời, lịch sử và kiểm thử.", CANCEL_QA)
    add_qa_section(doc, "5. Câu hỏi bảo vệ — Use case Quản lý tệp đính kèm", "Bao phủ dữ liệu, bảo mật file, Nextcloud, consistency và kiểm thử.", ATTACHMENT_QA)
    add_qa_section(doc, "6. Câu hỏi liên kết / kiến trúc", "Các câu hỏi thường dùng để hỏi vặn về độ ổn định và lý do thiết kế.", CROSS_QA)

    doc.add_page_break()
    add_section_title(doc, "7. Kịch bản demo 5–7 phút", "Một demo ngắn nhưng có bằng chứng sẽ thuyết phục hơn việc trình bày nhiều màn hình.")
    add_two_column_table(doc, [
        ("1. Chuẩn bị", "Đăng nhập employee là người tạo phiếu. Chuẩn bị một PDF nhỏ, an toàn; tạo hoặc chọn một phiếu Bản nháp có tổng tiền vượt ngưỡng nếu muốn demo rule báo giá."),
        ("2. Upload", "Chọn loại Báo giá, tải PDF lên. Chỉ ra tên tệp, dung lượng, người tải và trạng thái. Nếu đủ ngưỡng, chỉ ra chỉ báo yêu cầu chứng từ đã được đáp ứng."),
        ("3. Validation", "Nêu rằng server kiểm tra type, tên file, signature và 10 MB. Không cần demo file độc hại thật; có thể mô tả thử .exe sẽ bị chặn."),
        ("4. Download", "Tải xuống file để minh họa quyền và kiểm tra checksum nằm ở backend."),
        ("5. Delete", "Xóa một tệp khi còn Bản nháp; tải lại danh sách để thấy file không còn ACTIVE. Nêu metadata/audit vẫn được giữ."),
        ("6. Cancel", "Trên phiếu DRAFT chọn Hủy phiếu, xác nhận status Đã hủy. Mở Timeline/kiểm toán để cho thấy phiếu không bị xóa mà có sự kiện CANCELLED."),
        ("7. Chốt", "Nhắc lại ba bảo đảm: đúng quyền/đúng trạng thái, không mất dấu vết, và chống lỗi/giả mạo dữ liệu đầu vào."),
    ])
    add_info_box(doc, "Lưu ý demo", "Không hủy phiếu bạn còn cần dùng cho các luồng role khác. Hãy tạo một phiếu test riêng. Nếu Nextcloud không chạy, nêu rõ đây là dependency của UC tệp; đừng cố trình diễn upload rồi kết luận sai về luồng nghiệp vụ.", PALE_RED)

    add_section_title(doc, "8. Bảng chốt nhanh trước khi vào bảo vệ", "Đọc lại trang này trong 2 phút trước khi trình bày.")
    add_two_column_table(doc, [
        ("Hủy phiếu", "Chủ phiếu + DRAFT/CHANGES_REQUESTED → CANCELLED, không DELETE; event + audit; version + idempotency; chưa có restore."),
        ("Tệp", "Chủ phiếu + trạng thái chỉnh sửa; PDF/DOCX/XLSX/JPG/PNG, ≤10 MB; validate server; Nextcloud chứa file, DB chứa metadata/checksum."),
        ("Toàn vẹn", "Upload: UPLOADING → Nextcloud → ACTIVE. Download: SHA-256. Delete: DELETING → storage delete → DELETED; lỗi storage thì về ACTIVE."),
        ("Chứng từ bắt buộc", "Rule theo tổ chức/tiền tệ/ngưỡng/loại file; kiểm tại SUBMIT và RESUBMIT, chỉ file ACTIVE đúng loại mới thỏa."),
        ("Điều phải nói đúng", "Không có antivirus, restore phiếu, versioning tệp hay reconciliation job tự động trong code hiện tại; đây là hướng nâng cấp."),
    ])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    text(p, "Tài liệu được tổng hợp từ mã nguồn DX-OS hiện tại, không suy diễn chức năng chưa triển khai.", size=9.2, color=MUTED, italic=True)
    return doc


if __name__ == "__main__":
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUTPUT)
    print(OUTPUT)
